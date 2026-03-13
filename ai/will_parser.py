"""Parse existing will documents and extract structured data for all wizard steps."""
import base64
import json
import anthropic
from config import ANTHROPIC_API_KEY, CLAUDE_MODEL


def parse_will_document(file_path: str) -> dict:
    """Parse a will document (PDF or DOCX) and extract structured data.

    Returns dict with keys matching session structure:
    step1_testator, step2_executors, step3_guardians, step4_beneficiaries,
    step5_gifts, step6_residuary, step7_trust, step8_other_matters
    """
    ext = file_path.rsplit('.', 1)[-1].lower()

    if ext == 'docx':
        text_content = _extract_docx_text(file_path)
        return _parse_will_text(text_content)
    elif ext == 'pdf':
        return _parse_will_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload a PDF or DOCX file.")


def _extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        return '\n'.join(paragraphs)
    except ImportError:
        # python-docx not installed, read as binary and send to Claude
        with open(file_path, 'rb') as f:
            return f"[DOCX binary content - {len(f.read())} bytes]"


def _parse_will_text(text_content: str) -> dict:
    """Send will text to Claude for structured extraction."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": f"""Parse this Malaysian will document and extract all information into structured JSON.

WILL TEXT:
{text_content}

Return ONLY a JSON object with this structure:
{_get_schema()}

Fill in as much data as you can find. Use empty strings for fields not found. For arrays, include all entries found.
Return ONLY the JSON, no explanation."""
        }]
    )

    return _parse_response(message.content[0].text)


def _parse_will_pdf(file_path: str) -> dict:
    """Send will PDF to Claude Vision for structured extraction."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    with open(file_path, 'rb') as f:
        file_data = base64.standard_b64encode(f.read()).decode('utf-8')

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": file_data,
                    }
                },
                {
                    "type": "text",
                    "text": f"""Parse this Malaysian will document and extract all information into structured JSON.

Return ONLY a JSON object with this structure:
{_get_schema()}

Fill in as much data as you can find. Use empty strings for fields not found. For arrays, include all entries found.
Return ONLY the JSON, no explanation."""
                }
            ]
        }]
    )

    return _parse_response(message.content[0].text)


def _get_schema() -> str:
    return """{
    "step1_testator": {
        "full_name": "", "nric_passport": "", "residential_address": "",
        "nationality": "Malaysian", "country_of_residence": "Malaysia",
        "date_of_birth": "", "occupation": "", "religion": "",
        "gender": "Male", "marital_status": "Single",
        "property_coverage": "Malaysia", "signing_method": "Signature",
        "contemplation_of_marriage": false, "fiance_name": "", "fiance_nric": "",
        "special_circumstances": [], "translator_name": "", "translator_nric": "",
        "translator_relationship": "", "translator_language": ""
    },
    "step2_executors": {
        "executors": [
            {"full_name": "", "nric_passport": "", "address": "", "relationship": "", "role": "Primary"}
        ]
    },
    "step3_guardians": {
        "guardians": [
            {"full_name": "", "nric_passport": "", "address": "", "relationship": "", "role": "Primary"}
        ],
        "exclude_spouse_guardian": false, "exclude_spouse_guardian_reason": "",
        "guardian_allowance": {"payment_mode": "", "amount": "", "until_age": "21", "source_of_payment": ""}
    },
    "step4_beneficiaries": {
        "beneficiaries": [
            {"full_name": "", "nric_passport_birthcert": "", "relationship": ""}
        ]
    },
    "step5_gifts": {
        "gifts": [
            {
                "description": "", "subject_to_trust": false, "subject_to_guardian_allowance": false,
                "allocations": [{"beneficiary_name": "", "share": "100%", "role": "MB"}]
            }
        ]
    },
    "step6_residuary": {
        "main_beneficiaries": [{"beneficiary_name": "", "share": ""}],
        "substitute_groups": [],
        "additional_notes": ""
    },
    "step7_trust": {
        "trust_beneficiaries": [], "purposes": [], "purposes_other_text": "",
        "immovable_property_action": "", "trust_duration": "",
        "gift_references": "", "payment_mode": "", "payment_amount": "",
        "balance_of_trust": "", "separate_trustee": false,
        "trustee_name": "", "trustee_nric": "", "trustee_address": "", "trustee_relationship": ""
    },
    "step8_other_matters": {
        "terms_of_endearment": "", "commorientes_enabled": false, "commorientes_days": "30",
        "exclusion_enabled": false, "exclusion_name": "", "exclusion_nric": "",
        "exclusion_relationship": "", "exclusion_reason": "",
        "unnamed_children_enabled": false, "unnamed_children_spouse_name": "",
        "unnamed_children_spouse_nric": "", "additional_instructions": ""
    }
}"""


def _parse_response(response_text: str) -> dict:
    """Parse Claude's response text into a dict."""
    text = response_text.strip()
    if text.startswith('```'):
        text = text.split('\n', 1)[1].rsplit('```', 1)[0].strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return {"error": "Could not parse will document", "raw": text}
