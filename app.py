"""
WillCraft AI - Malaysian AI Will Writing System
Flask application with multi-step wizard for will drafting.
"""

from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify, g, make_response, Response
from functools import wraps
import base64
import difflib
import json
import os
import re
import sys
import tempfile
import threading
import traceback
import uuid
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))

from config import FLASK_SECRET_KEY, ANTHROPIC_API_KEY, SQLALCHEMY_DATABASE_URI, DATA_DIR, UPLOAD_DIR, SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD
from database import db, Client, Will, WillEditLog, WillVersion, Person, Document, User, ROLE_PERMS, ROLE_LABELS, ProbateApplication, ProbateFormTemplate, ProbateGeneratedForm, ChatSession, ChatMessage, LegalQAGap

# Enable SQLite WAL mode + a 5s busy timeout on every new connection.
# Without this, any concurrent write attempt (e.g. inbound email webhook
# arriving while the chat polls or the wizard saves) raises
# "database is locked". WAL allows readers + a single writer concurrently;
# busy_timeout makes additional writers wait up to 5s instead of erroring.
from sqlalchemy import event
from sqlalchemy.engine import Engine
import sqlite3 as _sqlite3

@event.listens_for(Engine, 'connect')
def _enable_sqlite_wal(dbapi_connection, connection_record):
    if isinstance(dbapi_connection, _sqlite3.Connection):
        cur = dbapi_connection.cursor()
        cur.execute('PRAGMA journal_mode=WAL')
        cur.execute('PRAGMA busy_timeout=5000')
        cur.execute('PRAGMA synchronous=NORMAL')
        cur.close()

# 🔥 §10x.69 — install GLOBAL Anthropic kill switch BEFORE any other
# import that may trigger anthropic.Anthropic() construction. Patches
# Messages.create so EVERY call across the codebase (Haiku, Sonnet,
# Opus, web search) checks DISABLE_VISION_CALLS=1.
try:
    from services.anthropic_killswitch import install_global_killswitch
    install_global_killswitch()
except Exception as _e:
    print(f"WARNING: kill switch install failed: {_e}")

app = Flask(__name__)
app.secret_key = FLASK_SECRET_KEY
app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 300 * 1024 * 1024  # §10x.136 — 300MB cap (lifted from 100MB; textbooks like Gopalakrishnan can hit 150-200MB)

db.init_app(app)

# Accepted file formats for OCR scanning
OCR_ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'pdf'}

def _validate_ocr_file(file):
    """Validate uploaded file is an accepted format for OCR. Returns error message or None."""
    if not file or not file.filename:
        return 'No file selected'
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in OCR_ALLOWED_EXTENSIONS:
        return f'Unsupported file format: .{ext}. Accepted formats: JPG, PNG, PDF'
    return None


# ---------------------------------------------------------------------------
# Jinja2 filters
# ---------------------------------------------------------------------------
from fractions import Fraction
from datetime import timezone, timedelta

MYT = timezone(timedelta(hours=8))

@app.template_filter('myt')
def myt_filter(dt, fmt='%d %b %Y, %I:%M %p'):
    """Convert UTC datetime to Malaysia Time (UTC+8) and format."""
    if not dt:
        return ''
    return dt.replace(tzinfo=timezone.utc).astimezone(MYT).strftime(fmt)

@app.template_filter('oneline')
def oneline_filter(value):
    """Flatten multiline text to a single line, joining with commas."""
    if not value:
        return value
    return ', '.join(line.strip() for line in str(value).splitlines() if line.strip())

@app.template_filter('to_fraction')
def to_fraction_filter(value):
    """Convert a share value to fraction display. '40' -> '4/10', '31' -> '31/100'."""
    if not value or value == '-':
        return value
    s = str(value).strip().rstrip('%')
    if '/' in s:
        return s
    try:
        num = float(s)
        n = int(num)
        if n != num:
            return s
        if n == 100:
            return "1/1"
        # Use /10 if divisible by 10, otherwise /100
        if n % 10 == 0:
            return f"{n // 10}/10"
        else:
            return f"{n}/100"
    except (ValueError, ZeroDivisionError):
        return s


# ---------------------------------------------------------------------------
# Multi-tenant configuration
# ---------------------------------------------------------------------------
TENANT_CONFIG = {
    'will.lifa.com.my': {
        'brand': 'LIFA',
        'subtitle': 'WillCraft AI',
        'theme': 'emerald',
        'gradient_from': '#064e3b',
        'gradient_via': '#065f46',
        'gradient_to': '#0f766e',
        'accent': '#d4a745',
        'accent_light': '#fef3c7',
        'btn_bg': '#059669',
        'btn_hover': '#047857',
        'email_domain': '@lifa.com.my',
        'email_from': '',
        'email_cc': [],
        'default_users': [
            {'email': 'admin@lifa.com.my', 'password': 'Admin2026#', 'name': 'Admin', 'role': 'admin'},
            {'email': 'advisor@lifa.com.my', 'password': 'Advisor2026#', 'name': 'Advisor', 'role': 'advisor'},
            {'email': 'approver@lifa.com.my', 'password': 'Approver2026#', 'name': 'Approver', 'role': 'approver'},
        ],
    },
    'will.alantanjb.com.my': {
        'brand': 'alantanjb',
        'subtitle': 'WillCraft AI',
        'theme': 'indigo',
        'gradient_from': '#1e1b4b',
        'gradient_via': '#312e81',
        'gradient_to': '#4338ca',
        'accent': '#94a3b8',
        'accent_light': '#e2e8f0',
        'btn_bg': '#4f46e5',
        'btn_hover': '#4338ca',
        'email_domain': '@alantanjb.com',
        'email_from': 'enquiry@alantanjb.com',
        'email_cc': ['kylie.tan@alantanjb.com'],
        'firm_name': 'Tetuan Alan Tan & Associates',
        'firm_address': "24-01 & 24-02, Jln Kempas Utama 2/4, Taman Kempas Utama, 81300 Johor Bahru, Johor Darul Ta'zim",
        'firm_phone': '011-3953 2638',
        'lawyer_name': 'FAIZUL HANAFI BIN TOKIRAN',
        'lawyer_bar_number': 'BC/F/167',
        'default_users': [
            {'email': 'accounts@alantanjb.com', 'password': 'Finance88#', 'name': 'Accounts', 'role': 'admin'},
            {'email': 'enquiry@alantanjb.com', 'password': 'Enquiry88#', 'name': 'Enquiry', 'role': 'advisor'},
            {'email': 'kylie.tan@alantanjb.com', 'password': 'Aia12345#', 'name': 'Kylie Tan', 'role': 'approver'},
        ],
    },
}
# Also map without .my for Cloudflare routing
TENANT_CONFIG['will.alantanjb.com'] = TENANT_CONFIG['will.alantanjb.com.my']

DEFAULT_TENANT = {
    'brand': 'LIFA',
    'subtitle': 'WillCraft AI',
    'theme': 'emerald',
    'gradient_from': '#064e3b',
    'gradient_via': '#065f46',
    'gradient_to': '#0f766e',
    'accent': '#d4a745',
    'accent_light': '#fef3c7',
    'btn_bg': '#059669',
    'btn_hover': '#047857',
    'email_domain': '@lifa.com.my',
    'email_from': '',
    'email_cc': [],
    'default_users': [
        {'email': 'admin@lifa.com.my', 'password': 'Admin2026#', 'name': 'Admin', 'role': 'admin'},
        {'email': 'advisor@lifa.com.my', 'password': 'Advisor2026#', 'name': 'Advisor', 'role': 'advisor'},
        {'email': 'approver@lifa.com.my', 'password': 'Approver2026#', 'name': 'Approver', 'role': 'approver'},
    ],
}


def get_tenant():
    """Get tenant config based on request hostname."""
    host = request.host.split(':')[0] if request else 'localhost'
    return TENANT_CONFIG.get(host, DEFAULT_TENANT)


# ---------------------------------------------------------------------------
# Authentication decorators
# ---------------------------------------------------------------------------

def login_required(f):
    """Require login for a route."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def role_required(*roles):
    """Require specific role(s) for a route."""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session:
                return redirect(url_for('login'))
            if session.get('user_role') not in roles:
                flash('You do not have permission to access this page.', 'error')
                return redirect(url_for('index'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator


@app.before_request
def load_current_user():
    """Load current user into g for template access."""
    g.user = None
    g.perms = {}
    g.tenant = get_tenant() if request else DEFAULT_TENANT
    user_id = session.get('user_id')
    if user_id:
        g.user = db.session.get(User, user_id)
        if g.user:
            g.perms = ROLE_PERMS.get(g.user.role, {})
        else:
            # User deleted, clear session
            session.pop('user_id', None)
            session.pop('user_role', None)


@app.context_processor
def inject_global_context():
    """Make user, permissions, tenant, and testator_person_id available to all templates."""
    # Count pending approvals for approvers
    pending_count = 0
    if g.user and g.perms.get('canApprove'):
        pending_count = Will.query.filter_by(status='pending_approval').filter(Will.deleted_at.is_(None)).count()
    # Check if current will has been generated
    has_generated_will = False
    if session.get('will_id'):
        wr = db.session.get(Will, session['will_id'])
        if wr and (wr.generated_will_text or wr.status in ('generated', 'pending_approval', 'approved')):
            has_generated_will = True
    return {
        'testator_person_id': session.get('step1', {}).get('person_id', ''),
        'current_user': g.user,
        'perms': g.perms,
        'tenant': g.tenant,
        'role_labels': ROLE_LABELS,
        'pending_approval_count': pending_count,
        'has_generated_will': has_generated_will,
        # Auto cache-bust for static assets — use file mtime so any
        # change to chat.js / wizard.js etc. is picked up by the browser
        # automatically. No more manual ?v=… bumps and no more telling
        # users to hard-refresh.
        'asset_version': _asset_version,
    }


def _asset_version(filename: str) -> str:
    """Return a cache-busting query string based on file mtime.

    Usage in templates:
        <script src="{{ url_for('static', filename='js/chat.js') }}?v={{ asset_version('js/chat.js') }}"></script>
    """
    try:
        path = os.path.join(app.static_folder, filename)
        return str(int(os.path.getmtime(path)))
    except Exception:
        return '0'


with app.app_context():
    os.makedirs(DATA_DIR, exist_ok=True)
    db.create_all()
    # Migrate: add document_id column to persons if not exists
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("ALTER TABLE persons ADD COLUMN document_id VARCHAR(36)"))
            conn.commit()
    except Exception:
        pass  # Column already exists
    # Migrate: add approval columns to wills if not exists
    for col_def in [
        ("created_by", "VARCHAR(36)"),
        ("submitted_by", "VARCHAR(36)"),
        ("submitted_at", "DATETIME"),
        ("approved_by", "VARCHAR(36)"),
        ("approved_at", "DATETIME"),
        ("approval_remarks", "TEXT"),
        ("text_edited_by", "VARCHAR(36)"),
        ("text_edited_at", "DATETIME"),
        ("include_logo", "BOOLEAN DEFAULT 1"),
    ]:
        try:
            with db.engine.connect() as conn:
                conn.execute(db.text(f"ALTER TABLE wills ADD COLUMN {col_def[0]} {col_def[1]}"))
                conn.commit()
        except Exception:
            pass
    # Create new tables (WillEditLog, Probate etc.) if they don't exist
    db.create_all()
    # Migrate: add content_hash to documents for dedup-by-bytes (CLAUDE.md §10c)
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("ALTER TABLE documents ADD COLUMN content_hash VARCHAR(64)"))
            conn.execute(db.text("CREATE INDEX IF NOT EXISTS ix_documents_content_hash ON documents(content_hash)"))
            conn.commit()
    except Exception:
        pass
    # Migrate: add chat_message_id column to documents (links a doc to the chat message that uploaded it)
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("ALTER TABLE documents ADD COLUMN chat_message_id VARCHAR(36)"))
            conn.commit()
    except Exception:
        pass  # Column already exists
    # Migrate: add new probate columns if not exists
    for col_def in [
        ("probate_applications", "application_type", "VARCHAR(20) DEFAULT 'probate'"),
        ("probate_applications", "deceased_name", "VARCHAR(200)"),
        ("probate_applications", "deceased_nric", "VARCHAR(50)"),
        ("probate_applications", "deceased_address", "TEXT"),
        ("probate_applications", "applicant_name", "VARCHAR(200)"),
        ("probate_applications", "applicant_nric", "VARCHAR(50)"),
        ("probate_applications", "applicant_address", "TEXT"),
        ("probate_applications", "applicant_relationship", "VARCHAR(100)"),
        ("probate_applications", "assets_data", "TEXT DEFAULT '[]'"),
        ("probate_applications", "beneficiaries_data", "TEXT DEFAULT '[]'"),
        ("probate_applications", "will_document_id", "VARCHAR(36)"),
        ("probate_applications", "deleted_at", "DATETIME"),
        ("wills", "deleted_at", "DATETIME"),
        ("will_edit_logs", "details", "TEXT"),
        # Multi-tenant isolation — every Client owned by a User
        ("clients", "created_by", "VARCHAR(36)"),
    ]:
        try:
            with db.engine.connect() as conn:
                conn.execute(db.text(f"ALTER TABLE {col_def[0]} ADD COLUMN {col_def[1]} {col_def[2]}"))
                conn.commit()
        except Exception:
            pass
    # Index for the new clients.created_by — speeds up tenant-scoped queries
    try:
        with db.engine.connect() as conn:
            conn.execute(db.text("CREATE INDEX IF NOT EXISTS ix_clients_created_by ON clients(created_by)"))
            conn.commit()
    except Exception:
        pass
    # Seed default users if none exist — detect tenant from WILLCRAFT_DOMAIN env var
    try:
        if User.query.count() == 0:
            domain = os.environ.get('WILLCRAFT_DOMAIN', '')
            tenant = TENANT_CONFIG.get(domain, DEFAULT_TENANT)
            for u in tenant['default_users']:
                user = User(email=u['email'], name=u['name'], role=u['role'])
                user.set_password(u['password'])
                db.session.add(user)
            db.session.commit()
            print(f"[Auth] Seeded {len(tenant['default_users'])} default users for {domain or 'default'}.")
    except Exception as e:
        db.session.rollback()
        print(f"[Auth] User seeding skipped (may already exist): {e}")

    # Seed default probate form templates if table is empty
    try:
        if ProbateFormTemplate.query.count() == 0:
            PROBATE_FORM_DEFAULTS = [
                {'form_code': 'doc01', 'form_name': 'Originating Summons', 'form_name_malay': 'Saman Pemula',
                 'description': 'The main court application to start the probate process. This tells the court you want to be officially recognized as the executor of the will.',
                 'file_path': 'probate_templates/doc01_saman_pemula.docx', 'category': 'core', 'sort_order': 1},
                {'form_code': 'doc02', 'form_name': 'Affidavit under Probate Act', 'form_name_malay': 'Afidavit Menurut Akta Probet',
                 'description': "The executor's sworn statement about the deceased person, their will, and their estate. Includes exhibit references for supporting documents.",
                 'file_path': 'probate_templates/doc02_afidavit_probet.docx', 'category': 'core', 'sort_order': 2},
                {'form_code': 'doc03', 'form_name': 'Oath of Administration', 'form_name_malay': 'Sumpah Pentadbiran',
                 'description': "The executor's oath promising to honestly and faithfully manage the deceased person's estate according to the law.",
                 'file_path': 'probate_templates/doc03_sumpah_pentadbiran.docx', 'category': 'core', 'sort_order': 3},
                {'form_code': 'doc04', 'form_name': 'Witness 1 Affidavit', 'form_name_malay': 'Afidavit Saksi 1',
                 'description': 'A sworn statement from the first person who witnessed the will being signed. Confirms they saw the testator sign the will.',
                 'file_path': 'probate_templates/doc04_afidavit_saksi_1.docx', 'category': 'witness',
                 'requires_witnesses': True, 'sort_order': 4},
                {'form_code': 'doc05', 'form_name': 'Witness 2 Affidavit', 'form_name_malay': 'Afidavit Saksi 2',
                 'description': 'A sworn statement from the second person who witnessed the will being signed.',
                 'file_path': 'probate_templates/doc05_afidavit_saksi_2.docx', 'category': 'witness',
                 'requires_witnesses': True, 'sort_order': 5},
                {'form_code': 'doc06', 'form_name': 'Assets & Liabilities Schedule', 'form_name_malay': 'Jadual Aset & Liabiliti',
                 'description': "A detailed list of everything the deceased person owned (houses, cars, bank accounts, investments) and any debts they owed.",
                 'file_path': 'probate_templates/doc06_jadual_aset.docx', 'category': 'core', 'sort_order': 6},
                {'form_code': 'doc07', 'form_name': 'Beneficiary List', 'form_name_malay': 'Senarai Benefisiari',
                 'description': "A list of all people who will inherit from the deceased person's estate, including their names, ID numbers, and relationship.",
                 'file_path': 'probate_templates/doc07_senarai_benefisiari.docx', 'category': 'core', 'sort_order': 7},
                {'form_code': 'doc08', 'form_name': 'Notice of Solicitor Appointment', 'form_name_malay': 'Notis Perlantikan Peguamcara',
                 'description': 'A formal notice telling the court that a lawyer has been hired to handle this probate case.',
                 'file_path': 'probate_templates/doc08_notis_peguamcara.docx', 'category': 'core', 'sort_order': 8},
                {'form_code': 'form14a', 'form_name': 'Land Transfer (Form 14A)', 'form_name_malay': 'Borang 14A - Pindah Milik',
                 'description': 'Transfers property (land/house) from the deceased to the beneficiary named in the will. One form is needed for each property.',
                 'file_path': 'probate_templates/form14a_land_transfer.docx', 'category': 'property',
                 'requires_property': True, 'sort_order': 9},
                {'form_code': 'form346', 'form_name': 'Personal Representative (Form 346)', 'form_name_malay': 'Borang 346 - Pendaftaran Wakil Diri',
                 'description': 'Registers the executor as the legal representative at the land office so they can handle property transfers.',
                 'file_path': 'probate_templates/form346_personal_rep.docx', 'category': 'property',
                 'requires_property': True, 'sort_order': 10},
            ]
            for tpl in PROBATE_FORM_DEFAULTS:
                t = ProbateFormTemplate(**tpl)
                db.session.add(t)
            db.session.commit()
            print(f"[Probate] Seeded {len(PROBATE_FORM_DEFAULTS)} default form templates.")
    except Exception as e:
        db.session.rollback()
        print(f"[Probate] Template seeding skipped: {e}")


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def get_completed_steps():
    """Return list of completed wizard step numbers."""
    return session.get('completed_steps', [])


def mark_step_complete(step):
    """Mark a wizard step as completed in the session."""
    steps = get_completed_steps()
    if step not in steps:
        steps.append(step)
    session['completed_steps'] = steps


def ensure_client():
    """Ensure a Client record exists for the current session. Returns client_id.

    🔥 §10x.61 — NRIC is COMPULSORY for new clients.
    The inbox address is `<name5char><ic_last4>@...` per CLAUDE.md §4.
    Without an IC the inbox can't be uniquely addressed and the strict-
    match routing (also §10x.61) will reject inbound emails. Plus a
    client without an IC can't be probated — IC is on every Malaysian
    legal document.

    Raises ValueError if NRIC is missing/blank when creating a new client.
    Existing clients without IC stay accessible for backward compat but
    the chat / wizard nags the user to fill it in.
    """
    client_id = session.get('client_id')
    if client_id:
        existing = db.session.get(Client, client_id)
        if existing:
            return client_id
    step1 = session.get('step1', {})
    full_name = (step1.get('full_name') or '').strip()
    nric = (step1.get('nric_passport') or '').strip()
    if not full_name or full_name == 'New Client':
        raise ValueError(
            "§10x.61: cannot create client without a full name. "
            "Fill in Step 1 (Testator) before proceeding."
        )
    if not nric:
        raise ValueError(
            "§10x.61: NRIC/passport is COMPULSORY when creating a new client. "
            "The inbox address routing (<name><ic4>@...) requires it, and "
            "every Malaysian probate document requires the testator's IC."
        )
    client = Client(
        full_name=full_name,
        nric_passport=nric,
        email=step1.get('email'),
        phone=step1.get('phone'),
        # Multi-tenant isolation — owner is the user who created this client.
        # Defaults to None for unauthenticated background tasks; the
        # tenant-scoped query treats NULL as "legacy / globally visible"
        # so background imports still work, but every chat/UI session
        # populates this field.
        created_by=session.get('user_id'),
    )
    db.session.add(client)
    db.session.commit()
    session['client_id'] = client.id
    return client.id


def get_client_folder_name(client_id):
    """Get the friendly folder name for a client, or fall back to client_id."""
    client = db.session.get(Client, client_id)
    return client.folder_name if client else client_id


def save_will_to_db():
    """Persist current session data to the database."""
    client_id = ensure_client()
    will_id = session.get('will_id')
    step1 = session.get('step1', {})

    # Fallback: if step1 (testator) is empty but the user added someone marked
    # 'Testator' in the Identities step (Step 1), pull their details so the
    # Will title and Client name don't end up as "Will of Unknown" / "New Client".
    if not (step1 or {}).get('full_name'):
        for p in session.get('person_registry', []):
            if (p.get('relationship') or '').strip().lower() == 'testator':
                step1 = {
                    'full_name': p.get('full_name', '') or '',
                    'nric_passport': p.get('nric_passport', '') or '',
                    'residential_address': p.get('address', '') or '',
                    'date_of_birth': p.get('date_of_birth', '') or '',
                    'nationality': p.get('nationality', 'Malaysian') or 'Malaysian',
                    'gender': p.get('gender', '') or '',
                    'email': p.get('email', '') or '',
                    'phone': p.get('phone', '') or '',
                    'person_id': p.get('id', '') or '',
                }
                session['step1'] = step1
                session.modified = True
                break

    if will_id:
        will_record = db.session.get(Will, will_id)
    else:
        will_record = None

    if not will_record:
        will_record = Will(client_id=client_id, created_by=session.get('user_id'))
        db.session.add(will_record)
        db.session.flush()
        session['will_id'] = will_record.id

    will_record.identities_data = json.dumps(session.get('person_registry', []))
    will_record.step1_data = json.dumps(session.get('step1', {}))
    will_record.step2_data = json.dumps({
        'executors': session.get('step2_executors', []),
        'executor_type': session.get('step3_executor_type', 'single'),
        'trustee_data': session.get('step3_trustees', {'same_as_executor': True}),
    })
    will_record.step3_data = json.dumps({
        'guardians': session.get('step3_guardians', []),
        'guardian_allowance': session.get('step3_guardian_allowance', {}),
    })
    will_record.step4_data = json.dumps(session.get('step4_beneficiaries', []))
    will_record.step5_data = json.dumps(session.get('step5_gifts', []))
    will_record.step6_data = json.dumps(session.get('step6_residuary', {}))
    will_record.step7_data = json.dumps(session.get('step7_trust', {}))
    will_record.step8_data = json.dumps(session.get('step8_others', {}))
    will_record.completed_steps = json.dumps(session.get('completed_steps', []))
    # Only update generated_will_text when explicitly present in session
    # (it's popped after generation to keep cookie small — don't overwrite DB with None)
    if 'generated_will_text' in session:
        will_record.generated_will_text = session['generated_will_text']
        if session['generated_will_text']:
            will_record.status = 'generated'
    will_record.title = f"Will of {step1.get('full_name', 'Unknown')}"

    # Update client info from step1
    client = db.session.get(Client, client_id)
    if client and step1.get('full_name'):
        client.full_name = step1['full_name']
        client.nric_passport = step1.get('nric_passport', '')
        client.email = step1.get('email')
        client.phone = step1.get('phone')

    db.session.commit()
    return will_record


def load_will_to_session(will_record):
    """Restore a saved will into the current session."""
    session['will_id'] = will_record.id
    session['client_id'] = will_record.client_id
    session['step1'] = json.loads(will_record.step1_data or '{}')
    # Load executor data (handle both old array and new object format)
    step2_raw = json.loads(will_record.step2_data or '[]')
    if isinstance(step2_raw, list):
        # Old format: plain array of executors
        session['step2_executors'] = step2_raw
        session['step3_executor_type'] = 'joint' if len(step2_raw) > 1 else 'single'
        session['step3_trustees'] = {'same_as_executor': True, 'trustees': [{}]}
    else:
        # New format: object with executors, executor_type, trustee_data
        session['step2_executors'] = step2_raw.get('executors', [])
        session['step3_executor_type'] = step2_raw.get('executor_type', 'single')
        session['step3_trustees'] = step2_raw.get('trustee_data', {'same_as_executor': True, 'trustees': [{}]})
    step3 = json.loads(will_record.step3_data or '{}')
    session['step3_guardians'] = step3.get('guardians', [])
    session['step3_guardian_allowance'] = step3.get('guardian_allowance', {})
    session['step4_beneficiaries'] = json.loads(will_record.step4_data or '[]')
    session['step5_gifts'] = json.loads(will_record.step5_data or '[]')
    session['step6_residuary'] = json.loads(will_record.step6_data or '{}')
    session['step7_trust'] = json.loads(will_record.step7_data or '{}')
    session['step8_others'] = json.loads(will_record.step8_data or '{}')
    session['completed_steps'] = json.loads(will_record.completed_steps or '[]')
    # Ensure step 10 is marked complete if will has been generated
    if will_record.status in ('generated', 'pending_approval', 'approved') and 10 not in session['completed_steps']:
        session['completed_steps'].append(10)
    # Don't load generated_will_text into session — it makes the cookie too large.
    # preview() and download() now read from DB directly.
    session.pop('generated_will_text', None)
    # Refresh identity registry from DB (preferred) or from saved snapshot
    _refresh_session_person_registry(will_record.client_id)
    if not session.get('person_registry'):
        session['person_registry'] = json.loads(will_record.identities_data or '[]')
    session.modified = True


def upsert_person(client_id, full_name, nric_passport, address=None,
                  date_of_birth=None, nationality=None, gender=None,
                  passport_expiry=None, email=None, phone=None,
                  relationship=None, document_id=None):
    """Add or update a person identity in the registry."""
    if not full_name or not nric_passport:
        return None
    # Try exact match first, then normalized match (strip dashes/spaces)
    existing = Person.query.filter_by(client_id=client_id, nric_passport=nric_passport).first()
    if not existing:
        normalized = nric_passport.replace('-', '').replace(' ', '').upper()
        all_persons = Person.query.filter_by(client_id=client_id).all()
        for p in all_persons:
            p_norm = (p.nric_passport or '').replace('-', '').replace(' ', '').upper()
            if p_norm == normalized:
                existing = p
                break
    if existing:
        existing.full_name = full_name.upper()
        if address:
            existing.address = address
        if date_of_birth:
            existing.date_of_birth = date_of_birth
        if nationality:
            existing.nationality = nationality
        if gender:
            existing.gender = gender
        if passport_expiry:
            existing.passport_expiry = passport_expiry
        if email:
            existing.email = email
        if phone:
            existing.phone = phone
        if relationship is not None:
            existing.relationship = relationship
        if document_id is not None:
            existing.document_id = document_id
        db.session.commit()
        _refresh_session_person_registry(client_id)
        return existing
    else:
        person = Person(
            client_id=client_id,
            full_name=full_name.upper(),
            nric_passport=nric_passport,
            address=address or '',
            date_of_birth=date_of_birth,
            nationality=nationality or 'Malaysian',
            gender=gender,
            passport_expiry=passport_expiry,
            email=email,
            phone=phone,
            relationship=relationship or '',
            document_id=document_id or None,
        )
        db.session.add(person)
        db.session.commit()
        _refresh_session_person_registry(client_id)
        return person


def _refresh_session_person_registry(client_id):
    """Refresh the session person registry from DB.
    Sort: Testator first, then by DOB ascending (oldest first), then by name."""
    persons = Person.query.filter_by(client_id=client_id).all()
    # Sort: Testator first, then by DOB (oldest first, None last), then name
    def _sort_key(p):
        is_testator = 0 if (p.relationship or '').lower() == 'testator' else 1
        dob = p.date_of_birth or ''
        # Normalize DOB to YYYY-MM-DD for sorting (handle DD-MM-YYYY format)
        if dob and len(dob) == 10 and dob[2] == '-' and dob[5] == '-':
            dob = f"{dob[6:10]}-{dob[3:5]}-{dob[0:2]}"
        return (is_testator, dob if dob else '9999-99-99', p.full_name)
    persons.sort(key=_sort_key)
    session['person_registry'] = [
        {'id': p.id, 'full_name': p.full_name, 'nric_passport': p.nric_passport,
         'address': p.address or '', 'date_of_birth': p.date_of_birth or '',
         'nationality': p.nationality or 'Malaysian', 'gender': p.gender or '',
         'passport_expiry': p.passport_expiry or '',
         'email': p.email or '', 'phone': p.phone or '',
         'relationship': p.relationship or '',
         'document_id': p.document_id or ''}
        for p in persons
    ]
    session.modified = True


def _propagate_identity_changes(person_id, new_name, new_nric, old_name=None):
    """When an identity is updated, propagate name/NRIC/relationship/address changes across all step session data."""
    person_data = _get_person_from_registry(person_id) or {}
    new_rel = person_data.get('relationship', '')
    new_addr = person_data.get('address', '')
    new_nationality = person_data.get('nationality', 'Malaysian')
    has_name_change = old_name and old_name.upper() != new_name.upper()

    # Always propagate by person_id (even without name change — catches relationship/address updates)
    # Step 1 (Testator)
    step1 = session.get('step1', {})
    if step1.get('person_id') == person_id:
        step1['full_name'] = new_name
        step1['nric_passport'] = new_nric
        step1['residential_address'] = new_addr
        step1['nationality'] = new_nationality
        session['step1'] = step1

    # Step 2 (Executors)
    for ex in session.get('step2_executors', []):
        if ex.get('person_id') == person_id or (has_name_change and ex.get('full_name', '').upper() == old_name.upper()):
            ex['full_name'] = new_name
            ex['nric_passport'] = new_nric
            ex['address'] = new_addr
            ex['relationship'] = new_rel
            if 'nationality' in ex:
                ex['nationality'] = new_nationality

    # Step 4 (Beneficiaries)
    for ben in session.get('step4_beneficiaries', []):
        if ben.get('person_id') == person_id or (has_name_change and ben.get('full_name', '').upper() == old_name.upper()):
            ben['full_name'] = new_name
            ben['nric_passport_birthcert'] = new_nric
            ben['relationship'] = new_rel
            if 'nationality' in ben:
                ben['nationality'] = new_nationality

    # Step 5 (Gift allocations — match by name)
    for gift in session.get('step5_gifts', []):
        for alloc in gift.get('allocations', []):
            if has_name_change and alloc.get('beneficiary_name', '').upper() == old_name.upper():
                alloc['beneficiary_name'] = new_name

    # Step 6 (Residuary estate)
    res = session.get('step6_residuary', {})
    for mb in res.get('main_beneficiaries', []):
        if mb.get('person_id') == person_id or (has_name_change and mb.get('beneficiary_name', '').upper() == old_name.upper()):
            mb['beneficiary_name'] = new_name

    session.modified = True


def _get_person_from_registry(person_id):
    """Look up a person from session['person_registry'] by ID."""
    for p in session.get('person_registry', []):
        if p['id'] == person_id:
            return p
    return None


def build_will_data():
    """Build WillData model from session data."""
    from models import (
        Testator, Executor, Guardian, GuardianAllowance,
        Beneficiary, Gift, GiftAllocation,
        ResiduaryEstate, ResiduaryBeneficiary,
        TestamentaryTrust, TrustBeneficiary,
        OtherMatters, WillData, Trustee,
    )

    # -- Section A: Testator --------------------------------------------------
    s1 = session.get('step1', {})
    testator = Testator(
        full_name=s1.get('full_name', ''),
        nric_passport=s1.get('nric_passport', ''),
        residential_address=s1.get('residential_address', ''),
        nationality=s1.get('nationality', 'Malaysian'),
        country_of_residence=s1.get('country_of_residence', 'Malaysia'),
        date_of_birth=s1.get('date_of_birth', '01-01-2000'),
        occupation=s1.get('occupation', ''),
        religion=s1.get('religion') or None,
        email=s1.get('email') or None,
        phone=s1.get('phone') or None,
        gender=(s1.get('gender') or 'Male'),
        marital_status=(s1.get('marital_status') or 'Single'),
        has_prior_will=s1.get('has_prior_will', False),
        property_coverage=s1.get('property_coverage', 'Malaysia'),
        contemplation_of_marriage=s1.get('contemplation_of_marriage', False),
        fiance_name=s1.get('fiance_name') or None,
        fiance_nric=s1.get('fiance_nric') or None,
        signing_method=s1.get('signing_method', 'Signature'),
        special_circumstances=s1.get('special_circumstances', []),
        translator_name=s1.get('translator_name') or None,
        translator_nric=s1.get('translator_nric') or None,
        translator_relationship=s1.get('translator_relationship') or None,
        translator_language=s1.get('translator_language') or None,
    )

    # -- Section B: Executors --------------------------------------------------
    # Coerce all session list fields to [] in case parser left them as None
    def _list(key):
        v = session.get(key)
        return v if isinstance(v, list) else []
    executors = [Executor(**e) for e in _list('step2_executors')]

    # -- Section B2: Trustees (separate from executors) -----------------------
    trustee_session = session.get('step3_trustees', {'same_as_executor': True})
    trustee_same_as_executor = trustee_session.get('same_as_executor', True)
    trustees = None
    substitute_trustee = None
    substitute_trustees = None
    if not trustee_same_as_executor:
        trustees_raw = trustee_session.get('trustees', [])
        if trustees_raw:
            trustees = [Trustee(**t) for t in trustees_raw if t.get('full_name')]
            if not trustees:
                trustees = None
        # Support multiple substitute trustees (new format)
        sub_trustees_raw = trustee_session.get('substitute_trustees', [])
        if sub_trustees_raw:
            substitute_trustees = [Trustee(**st) for st in sub_trustees_raw if st.get('full_name')]
            if substitute_trustees:
                substitute_trustee = substitute_trustees[0]  # backward compat
            else:
                substitute_trustees = None
        else:
            # Backward compat: single substitute_trustee
            sub_trustee_raw = trustee_session.get('substitute_trustee', {})
            if sub_trustee_raw and sub_trustee_raw.get('full_name'):
                substitute_trustee = Trustee(**sub_trustee_raw)
                substitute_trustees = [substitute_trustee]

    # -- Section C: Guardians (optional) --------------------------------------
    guardians_data = _list('step3_guardians')
    guardians = [Guardian(**g) for g in guardians_data] if guardians_data else None

    ga_data = session.get('step3_guardian_allowance', {})
    guardian_allowance = (
        GuardianAllowance(**ga_data)
        if ga_data and ga_data.get('payment_mode')
        else None
    )

    # -- Section D: Beneficiaries ---------------------------------------------
    def _normalise_beneficiary(b: dict) -> dict:
        # Drop chat-internal keys & ensure nric_passport_birthcert is set
        bb = {k: v for k, v in b.items() if not k.startswith('_')}
        if not bb.get('nric_passport_birthcert'):
            bb['nric_passport_birthcert'] = (bb.get('nric_passport')
                                             or bb.get('birth_cert_no')
                                             or '')
        # Drop nric_passport since model only knows nric_passport_birthcert
        bb.pop('nric_passport', None)
        bb.pop('relationship_to_testator', None)
        bb.pop('address', None)
        return bb
    beneficiaries = [Beneficiary(**_normalise_beneficiary(b))
                     for b in _list('step4_beneficiaries')
                     if b.get('full_name')]

    # -- Section E: Gifts (optional) ------------------------------------------
    gifts_data = _list('step5_gifts')
    gifts = None
    if gifts_data:
        from models.gift import PropertyDetails, FinancialDetails
        gifts = []
        for gd in gifts_data:
            # 🔥 §10x.129 — chat-saved gifts use {kind, beneficiaries,
            # property_info, substitute_specific}; the WillData model
            # expects {gift_type, allocations, property_details,
            # financial_details}. Normalise here.
            kind = (gd.get('kind') or gd.get('asset_type')
                    or gd.get('gift_type') or 'other').lower()
            if kind in ('bank', 'insurance', 'epf', 'kwsp', 'mutual_fund',
                        'unit_trust', 'shares', 'financial'):
                gift_type = 'financial'
            elif kind == 'property':
                gift_type = 'property'
            else:
                gift_type = 'other'

            # Allocations: prefer pre-built, else build from beneficiaries
            # + substitute_specific so the Phek substitute clause renders.
            allocs_raw = gd.get('allocations') or []
            if not allocs_raw and (gd.get('beneficiaries')
                                   or gd.get('substitute_specific')):
                # Build allocations from main beneficiaries
                subs_list = gd.get('substitute_specific') or []
                # Normalise substitute share to a fraction string
                sub_alloc = [{'beneficiary_name': s.get('name', ''),
                              'share': s.get('share', '1/1')}
                             for s in subs_list if s.get('name')]
                for b in (gd.get('beneficiaries') or []):
                    nm = b.get('name') or ''
                    if not nm:
                        continue
                    allocs_raw.append({
                        'beneficiary_name': nm,
                        'share': b.get('share', '1/1'),
                        'role': 'MB',
                        'substitutes': sub_alloc or None,
                    })
            allocations = [GiftAllocation(**a) for a in allocs_raw]

            prop_details = None
            fin_details = None
            if gift_type == 'property':
                # 🔥 §10x.132 — prefer property_info (chat-saved schema)
                # over property_details (older wizard form schema). When
                # both are present property_info has the up-to-date values
                # the user confirmed in chat.
                pd_info = gd.get('property_info') or {}
                pd_legacy = gd.get('property_details') or {}
                # Merge: info wins when both have a value
                pd = {**pd_legacy, **{k: v for k, v in pd_info.items() if v}}
                # Build a PropertyDetails dict that matches the model's
                # expected field names. The model uses `bandar_pekan` for
                # Mukim (legacy naming) — chat saves it as `mukim`.
                if pd or gd.get('property_address'):
                    raw_addr = (pd.get('property_address') or gd.get('property_address') or '').strip()
                    # Strip noise prefixes from chat-summary text
                    for _pref in ('House at ', 'Shop at ', 'Unit at ', 'Apartment at '):
                        if raw_addr.startswith(_pref):
                            raw_addr = raw_addr[len(_pref):]
                            break
                    # Title type: pre-set field wins; otherwise detect from
                    # title_number prefix (HSD, GRN, GM, PTD…)
                    title_num_raw = (pd.get('title_number') or gd.get('title_number') or '').strip()
                    title_type = (pd.get('title_type') or gd.get('title_type') or '').strip().upper()
                    title_num_clean = title_num_raw
                    if not title_type:
                        import re as _re_tt
                        _tt_m = _re_tt.match(r'^(HSD|HSM|HS\(D\)|HS\(M\)|GRN|GM|GERAN|HAKMILIK|PAJAKAN|PTD|PTM)\s*:?\s*(.+)$',
                                              title_num_raw, _re_tt.IGNORECASE)
                        if _tt_m:
                            title_type = _tt_m.group(1).upper()
                            title_num_clean = _tt_m.group(2).strip()
                    # If we have a title_type but title_number still has its
                    # prefix (e.g. "HSD 251041"), strip it
                    if title_type and title_num_clean.upper().startswith(title_type):
                        title_num_clean = title_num_clean[len(title_type):].strip(' :')
                    pd_norm = {
                        'property_address':   raw_addr,
                        'title_type':         title_type,
                        'title_number':       title_num_clean,
                        'lot_number':         pd.get('lot_number') or gd.get('lot_number') or '',
                        # PropertyDetails uses `bandar_pekan` for Mukim
                        'bandar_pekan':       pd.get('mukim') or gd.get('mukim') or pd.get('bandar_pekan') or '',
                        'daerah':             pd.get('daerah') or gd.get('daerah') or '',
                        'negeri':             pd.get('negeri') or gd.get('negeri') or '',
                    }
                    if any(pd_norm.values()):
                        try:
                            prop_details = PropertyDetails(**pd_norm)
                        except Exception:
                            prop_details = None
            if gift_type == 'financial':
                fd = gd.get('financial_details') or {}
                # Build from chat-saved bank/insurance fields if needed
                if not fd:
                    fd = {
                        'asset_type':       (kind if kind in ('bank', 'insurance', 'epf', 'kwsp', 'mutual_fund', 'unit_trust', 'shares') else 'other'),
                        'institution':      gd.get('institution') or gd.get('bank_name') or gd.get('insurer') or '',
                        'account_number':   gd.get('account_number') or gd.get('policy_number') or '',
                        'country':          gd.get('country') or '',
                    }
                # 🔥 §10x.134 — pass `description` field so the FinancialDetails
                # formatter renders the account-type token ("Saving"/"Current"/
                # "Fixed Deposit"). Phek format: "the monies in my UOB Saving
                # Account No. ..." — the "Saving" word comes from `description`.
                if not fd.get('description'):
                    fd['description'] = (gd.get('account_type')
                                          or gd.get('description')
                                          or '')
                try:
                    fin_details = FinancialDetails(**fd)
                except Exception:
                    fin_details = None
            try:
                gifts.append(Gift(
                    gift_type=gift_type,
                    description=gd.get('description', ''),
                    property_details=prop_details,
                    financial_details=fin_details,
                    allocations=allocations,
                    subject_to_trust=gd.get('subject_to_trust', False),
                    subject_to_guardian_allowance=gd.get('subject_to_guardian_allowance', False),
                    sell_property=gd.get('sell_property', False),
                    substitute_mode=gd.get('substitute_mode', 'equal'),
                    ownership_type=gd.get('ownership_type', 'sole'),
                    testator_share=gd.get('testator_share'),
                    joint_owners=gd.get('joint_owners'),
                    encumbrance_status=gd.get('encumbrance_status', 'clean'),
                    debt_source=gd.get('debt_source'),
                    account_ownership=gd.get('account_ownership', 'individual'),
                ))
            except Exception:
                continue   # skip malformed gift rather than abort whole will

    # -- Section F: Residuary Estate ------------------------------------------
    res_data = session.get('step6_residuary') or {}
    if not isinstance(res_data, dict):
        res_data = {}
    # 🔥 §10x.129 — chat-saved residuary uses {beneficiaries[],
    # substitute_specific[]}; the model expects {main_beneficiaries[],
    # substitute_groups[][]}.
    main_bens_raw = res_data.get('main_beneficiaries') or []
    if not main_bens_raw and res_data.get('beneficiaries'):
        for b in res_data.get('beneficiaries') or []:
            nm = b.get('name') or b.get('full_name') or ''
            if nm:
                main_bens_raw.append({
                    'beneficiary_name': nm,
                    'share':            b.get('share', '100/100'),
                })
    elif not main_bens_raw and res_data.get('residuary_beneficiary_name'):
        main_bens_raw.append({
            'beneficiary_name': res_data['residuary_beneficiary_name'],
            'share':            '100/100',
        })

    def _norm_residuary_ben(mb: dict) -> dict:
        bb = {k: v for k, v in mb.items() if not k.startswith('_')}
        # Pydantic model expects beneficiary_name + share
        if not bb.get('beneficiary_name'):
            bb['beneficiary_name'] = bb.pop('name', '') or bb.get('full_name', '')
        return bb

    main_bens = []
    for mb in main_bens_raw:
        try:
            main_bens.append(ResiduaryBeneficiary(**_norm_residuary_ben(mb)))
        except Exception:
            pass

    sub_groups = []
    for sg in (res_data.get('substitute_groups') or []):
        try:
            sub_groups.append([ResiduaryBeneficiary(**_norm_residuary_ben(sb)) for sb in (sg or [])])
        except Exception:
            pass
    if not sub_groups and res_data.get('substitute_specific'):
        try:
            grp = []
            for sb in res_data['substitute_specific']:
                grp.append(ResiduaryBeneficiary(**_norm_residuary_ben(sb)))
            if grp:
                sub_groups.append(grp)
        except Exception:
            pass
    residuary_estate = ResiduaryEstate(
        main_beneficiaries=main_bens,
        substitute_groups=sub_groups,
        additional_notes=res_data.get('additional_notes') or None,
    )

    # -- Section G: Testamentary Trust (optional) -----------------------------
    trust_data = session.get('step7_trust', {})
    testamentary_trust = None
    if trust_data and trust_data.get('beneficiaries'):
        trust_bens = [TrustBeneficiary(**tb) for tb in trust_data.get('beneficiaries', [])]
        balance_bens = [TrustBeneficiary(**bb) for bb in trust_data.get('balance_beneficiaries', [])]
        testamentary_trust = TestamentaryTrust(
            beneficiaries=trust_bens,
            purposes=trust_data.get('purposes', []),
            duration=trust_data.get('duration') or None,
            assets_from_gifts=trust_data.get('assets_from_gifts', []),
            payment_mode=trust_data.get('payment_mode') or None,
            payment_amount=trust_data.get('payment_amount') or None,
            balance_beneficiaries=balance_bens,
        )

    # -- Section H/I: Other Matters (optional) --------------------------------
    om_data = session.get('step8_others', {})
    other_matters = None
    if om_data:
        other_matters = OtherMatters(**om_data)

    return WillData(
        testator=testator,
        executors=executors,
        trustee_same_as_executor=trustee_same_as_executor,
        trustees=trustees,
        substitute_trustee=substitute_trustee,
        substitute_trustees=substitute_trustees,
        guardians=guardians,
        guardian_allowance=guardian_allowance,
        beneficiaries=beneficiaries,
        gifts=gifts,
        residuary_estate=residuary_estate,
        testamentary_trust=testamentary_trust,
        other_matters=other_matters,
        identities=session.get('person_registry', []),
    )


# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Health / smoke-test endpoint (no auth required)
# ---------------------------------------------------------------------------

@app.route('/api/health')
def api_health():
    """Lightweight smoke test — checks DB and returns git hash + model config."""
    try:
        from config import CLAUDE_MODEL, CLAUDE_MODEL_CHEAP
        client_count = db.session.execute(db.text('SELECT COUNT(*) FROM clients')).scalar()
        return jsonify({
            'ok': True,
            'db_clients': client_count,
            'model': CLAUDE_MODEL,
            'model_cheap': CLAUDE_MODEL_CHEAP,
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# Authentication Routes
# ---------------------------------------------------------------------------

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page and handler."""
    if 'user_id' in session:
        return redirect(url_for('index'))

    tenant = get_tenant()
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        remember = request.form.get('remember') == '1'

        if not email or not password:
            return render_template('login.html', error='Please enter email and password.', tenant=tenant)

        user = User.query.filter(db.func.lower(User.email) == email, User.is_active == True).first()
        if not user or not user.check_password(password):
            return render_template('login.html', error='Invalid email or password.', tenant=tenant)

        session['user_id'] = user.id
        session['user_role'] = user.role
        session['user_name'] = user.name
        session['user_email'] = user.email
        session.permanent = remember
        return redirect(url_for('index'))

    # Get quick-login users for display
    quick_users = User.query.filter_by(is_active=True).order_by(User.role, User.name).all()
    return render_template('login.html', tenant=tenant, quick_users=quick_users)


@app.route('/logout')
def logout():
    """Logout and redirect to login page."""
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


# ---------------------------------------------------------------------------
# User Management Routes (Admin only)
# ---------------------------------------------------------------------------

@app.route('/admin/users')
@role_required('admin')
def admin_users():
    """User management page (admin only)."""
    users = User.query.order_by(User.role, User.name).all()
    return render_template('admin/users.html', users=users)


@app.route('/admin/users/add', methods=['POST'])
@role_required('admin')
def admin_user_add():
    """Add a new user."""
    email = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '').strip()
    name = request.form.get('name', '').strip()
    contact = request.form.get('contact', '').strip()
    role = request.form.get('role', 'advisor')

    if not email or not password or not name:
        flash('Name, email, and password are required.', 'error')
        return redirect(url_for('admin_users'))

    if len(password) < 6:
        flash('Password must be at least 6 characters.', 'error')
        return redirect(url_for('admin_users'))

    if role not in ROLE_PERMS:
        flash('Invalid role.', 'error')
        return redirect(url_for('admin_users'))

    existing = User.query.filter(db.func.lower(User.email) == email).first()
    if existing:
        flash('A user with this email already exists.', 'error')
        return redirect(url_for('admin_users'))

    user = User(email=email, name=name, contact=contact, role=role)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()
    flash(f'User "{name}" ({role}) created successfully.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/users/<user_id>/update', methods=['POST'])
@role_required('admin')
def admin_user_update(user_id):
    """Update user details."""
    user = db.session.get(User, user_id)
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('admin_users'))

    name = request.form.get('name', '').strip()
    contact = request.form.get('contact', '').strip()
    role = request.form.get('role', '').strip()
    password = request.form.get('password', '').strip()

    if name:
        user.name = name
    if contact is not None:
        user.contact = contact
    if role and role in ROLE_PERMS:
        user.role = role
    if password and len(password) >= 6:
        user.set_password(password)

    db.session.commit()
    flash(f'User "{user.name}" updated.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/users/<user_id>/toggle', methods=['POST'])
@role_required('admin')
def admin_user_toggle(user_id):
    """Enable/disable a user."""
    user = db.session.get(User, user_id)
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('admin_users'))
    if user.id == session.get('user_id'):
        flash('You cannot disable your own account.', 'error')
        return redirect(url_for('admin_users'))
    user.is_active = not user.is_active
    db.session.commit()
    status = 'enabled' if user.is_active else 'disabled'
    flash(f'User "{user.name}" {status}.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/users/<user_id>/delete', methods=['POST'])
@role_required('admin')
def admin_user_delete(user_id):
    """Delete a user."""
    user = db.session.get(User, user_id)
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('admin_users'))
    if user.id == session.get('user_id'):
        flash('You cannot delete your own account.', 'error')
        return redirect(url_for('admin_users'))
    name = user.name
    db.session.delete(user)
    db.session.commit()
    flash(f'User "{name}" deleted.', 'success')
    return redirect(url_for('admin_users'))


# ---------------------------------------------------------------------------
# Admin: Firm Settings (logo upload)
# ---------------------------------------------------------------------------

def _logo_dir():
    """Return the logos directory path, creating it if needed."""
    d = os.path.join(DATA_DIR, 'logos')
    os.makedirs(d, exist_ok=True)
    return d


def _get_logo_path():
    """Return the absolute path to the current tenant's logo, or None."""
    tenant = get_tenant()
    domain = tenant.get('brand', 'default').lower()
    logo_dir = _logo_dir()
    for ext in ('png', 'jpg', 'jpeg', 'webp'):
        p = os.path.join(logo_dir, f'{domain}_logo.{ext}')
        if os.path.isfile(p):
            return p
    return None


@app.route('/admin/settings')
@role_required('admin')
def admin_settings():
    """Firm settings page — logo upload etc."""
    logo_path = _get_logo_path()
    logo_url = url_for('admin_serve_logo') if logo_path else None
    return render_template('admin/settings.html', logo_url=logo_url)


@app.route('/admin/will-format-preview')
@role_required('admin', 'approver')
def admin_will_format_preview():
    """
    Render the verbatim Alan & Tan PHEK YI TING sample through the current
    PDF/Word generator. Use this to compare against the original sample and
    iterate on the format until it matches exactly.

    Query params:
      ?format=pdf   → PDF (default)
      ?format=docx  → Microsoft Word .docx (downloads, editable in Word)
      ?download=1   → force download even for PDF
    """
    from documents.sample_will_phek_yi_ting import SAMPLE_WILL_TEXT_PHEK_YI_TING
    from documents.empty_template_will import EMPTY_TEMPLATE_WILL_TEXT

    fmt = (request.args.get('format') or 'pdf').lower()
    use_empty = request.args.get('empty') == '1'
    will_text = EMPTY_TEMPLATE_WILL_TEXT if use_empty else SAMPLE_WILL_TEXT_PHEK_YI_TING
    file_label = 'Empty_Template' if use_empty else 'PHEK_YI_TING_Format_Preview'

    logo = _get_logo_path()
    tenant = get_tenant()
    firm_info = None
    if tenant.get('firm_name'):
        firm_info = {
            'firm_name': tenant.get('firm_name', ''),
            'firm_address': tenant.get('firm_address', ''),
            'firm_phone': tenant.get('firm_phone', ''),
            'firm_email': tenant.get('email_from', ''),
        }

    if fmt == 'docx':
        from documents.will_docx import build_will_docx
        try:
            filepath = build_will_docx(
                will_text,
                firm_info=firm_info,
                logo_path=logo,
                is_draft=True,
            )
            with open(filepath, 'rb') as f:
                data = f.read()
        except Exception as e:
            app.logger.error(f'Will format preview (docx) failed: {e}')
            return f'Preview generation (docx) failed: {e}', 500
        return Response(
            data,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{file_label}.docx"',
            },
        )

    # Default: PDF
    from documents.pdf_generator import generate_pdf
    try:
        filepath = generate_pdf(
            will_text,
            file_label,
            logo_path=logo,
            firm_info=firm_info,
        )
        with open(filepath, 'rb') as f:
            pdf_data = f.read()
    except Exception as e:
        app.logger.error(f'Will format preview failed: {e}')
        return f'Preview generation failed: {e}', 500

    download = request.args.get('download') == '1'
    return Response(
        pdf_data,
        mimetype='application/pdf',
        headers={
            'Content-Disposition': (
                f'attachment; filename="{file_label}.pdf"' if download
                else f'inline; filename="{file_label}.pdf"'
            ),
        },
    )


@app.route('/admin/settings/logo')
@login_required
def admin_serve_logo():
    """Serve the tenant's logo image."""
    logo_path = _get_logo_path()
    if not logo_path:
        return '', 404
    return send_file(logo_path)


@app.route('/admin/settings/upload-logo', methods=['POST'])
@role_required('admin')
def admin_upload_logo():
    """Handle firm logo upload."""
    f = request.files.get('logo')
    if not f or not f.filename:
        flash('No file selected.', 'error')
        return redirect(url_for('admin_settings'))

    # Validate extension
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('png', 'jpg', 'jpeg', 'webp'):
        flash('Only PNG, JPG, or WebP images are allowed.', 'error')
        return redirect(url_for('admin_settings'))

    # Validate size (2MB)
    f.seek(0, 2)
    size = f.tell()
    f.seek(0)
    if size > 2 * 1024 * 1024:
        flash('Logo file must be under 2MB.', 'error')
        return redirect(url_for('admin_settings'))

    # Remove any existing logo for this tenant
    tenant = get_tenant()
    domain = tenant.get('brand', 'default').lower()
    logo_dir = _logo_dir()
    for old_ext in ('png', 'jpg', 'jpeg', 'webp'):
        old = os.path.join(logo_dir, f'{domain}_logo.{old_ext}')
        if os.path.isfile(old):
            os.remove(old)

    # Save new logo
    save_path = os.path.join(logo_dir, f'{domain}_logo.{ext}')
    f.save(save_path)
    flash('Firm logo uploaded successfully.', 'success')
    return redirect(url_for('admin_settings'))


@app.route('/admin/settings/remove-logo', methods=['POST'])
@role_required('admin')
def admin_remove_logo():
    """Remove the firm logo."""
    tenant = get_tenant()
    domain = tenant.get('brand', 'default').lower()
    logo_dir = _logo_dir()
    for ext in ('png', 'jpg', 'jpeg', 'webp'):
        p = os.path.join(logo_dir, f'{domain}_logo.{ext}')
        if os.path.isfile(p):
            os.remove(p)
    flash('Logo removed.', 'success')
    return redirect(url_for('admin_settings'))


@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    """User profile page."""
    user = db.session.get(User, session['user_id'])
    if not user:
        return redirect(url_for('logout'))

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        contact = request.form.get('contact', '').strip()
        password = request.form.get('password', '').strip()
        password2 = request.form.get('password2', '').strip()

        if name:
            user.name = name
            session['user_name'] = name
        if contact is not None:
            user.contact = contact
        if password:
            if len(password) < 6:
                flash('Password must be at least 6 characters.', 'error')
                return render_template('profile.html', user=user)
            if password != password2:
                flash('Passwords do not match.', 'error')
                return render_template('profile.html', user=user)
            user.set_password(password)

        db.session.commit()
        flash('Profile updated.', 'success')
        return redirect(url_for('profile'))

    return render_template('profile.html', user=user)


# ---------------------------------------------------------------------------
# Approval Workflow Routes
# ---------------------------------------------------------------------------

@app.route('/wills/<will_id>/submit-for-approval', methods=['POST'])
@login_required
def will_submit_for_approval(will_id):
    """Submit a generated will for approval."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        flash('Will not found.', 'error')
        return redirect(url_for('will_list'))

    if will_record.status not in ('generated', 'rejected'):
        flash('Only generated or rejected wills can be submitted for approval.', 'error')
        return redirect(url_for('preview'))

    will_record.status = 'pending_approval'
    will_record.submitted_by = session['user_id']
    will_record.submitted_at = datetime.utcnow()
    will_record.approval_remarks = None
    db.session.commit()
    flash('Will submitted for approval.', 'success')
    return redirect(url_for('preview'))


@app.route('/approvals')
@role_required('approver')
def approval_list():
    """List wills pending approval."""
    pending = Will.query.filter_by(status='pending_approval').filter(Will.deleted_at.is_(None)).order_by(Will.submitted_at.desc()).all()
    approved = Will.query.filter_by(status='approved').filter(Will.deleted_at.is_(None)).order_by(Will.approved_at.desc()).limit(20).all()
    rejected = Will.query.filter_by(status='rejected').filter(Will.deleted_at.is_(None)).order_by(Will.updated_at.desc()).limit(20).all()
    return render_template('approvals.html', pending=pending, approved=approved, rejected=rejected)


@app.route('/wills/<will_id>/approve', methods=['POST'])
@role_required('approver')
def will_approve(will_id):
    """Approve a will."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        flash('Will not found.', 'error')
        return redirect(url_for('approval_list'))

    if will_record.status not in ('pending_approval', 'generated'):
        flash('This will is not pending approval.', 'error')
        return redirect(url_for('approval_list'))

    remarks = request.form.get('remarks', '').strip()
    will_record.status = 'approved'
    will_record.approved_by = session['user_id']
    will_record.approved_at = datetime.utcnow()
    will_record.approval_remarks = remarks or None
    db.session.commit()
    flash(f'Will "{will_record.title}" approved.', 'success')
    # If approving from preview page, redirect back there
    if request.referrer and '/preview' in request.referrer:
        return redirect(url_for('preview'))
    return redirect(url_for('approval_list'))


@app.route('/wills/<will_id>/reject', methods=['POST'])
@role_required('approver')
def will_reject(will_id):
    """Reject a will."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        flash('Will not found.', 'error')
        return redirect(url_for('approval_list'))

    if will_record.status != 'pending_approval':
        flash('This will is not pending approval.', 'error')
        return redirect(url_for('approval_list'))

    remarks = request.form.get('remarks', '').strip()
    will_record.status = 'rejected'
    will_record.approved_by = session['user_id']
    will_record.approved_at = datetime.utcnow()
    will_record.approval_remarks = remarks or 'No reason provided.'
    db.session.commit()
    flash(f'Will "{will_record.title}" rejected.', 'info')
    return redirect(url_for('approval_list'))


@app.route('/api/will/<will_id>/cost', methods=['GET'])
@login_required
def api_will_cost(will_id):
    """Return total token cost for a will + per-call-site breakdown.

    Response:
      { ok: true,
        will_id: "...",
        total_usd: 0.001234,
        total_usd_fmt: "$0.001234",
        saas_budget_usd: 300,            # annual plan price
        budget_pct: 0.04,                # % of annual budget consumed
        calls: [{ call_site, calls, input_tokens, output_tokens, cost_usd }] }
    """
    will_record = db.session.get(Will, will_id)
    if not will_record:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404

    from ai.cost_tracker import total_for_will, breakdown_for_will
    total = float(total_for_will(will_id))
    calls = breakdown_for_will(will_id)
    saas_budget = 300.0            # $300/yr annual plan
    budget_pct = (total / saas_budget * 100) if saas_budget else 0

    return jsonify({
        'ok': True,
        'will_id': will_id,
        'total_usd': round(total, 6),
        'total_usd_fmt': f'${total:.4f}',
        'saas_budget_usd': saas_budget,
        'budget_pct': round(budget_pct, 4),
        'calls': calls,
    })


@app.route('/api/will/<will_id>/edit-text', methods=['POST'])
@login_required
def api_will_edit_text(will_id):
    """Save edits to the will text and log the change."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404

    if not will_record.generated_will_text:
        return jsonify({'ok': False, 'error': 'No generated will text to edit'}), 400

    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'ok': False, 'error': 'No text provided'}), 400

    new_text = data['text'].strip()
    if not new_text:
        return jsonify({'ok': False, 'error': 'Will text cannot be empty'}), 400

    # Compute diff summary with specific change details
    old_lines = (will_record.generated_will_text or '').splitlines()
    new_lines = new_text.splitlines()
    added = removed = changed = 0
    change_details = []  # Specific descriptions of what changed
    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, old_lines, new_lines).get_opcodes():
        if tag == 'insert':
            added += (j2 - j1)
            for line in new_lines[j1:j2]:
                snippet = line.strip()[:80]
                if snippet:
                    change_details.append(f'+ "{snippet}"')
        elif tag == 'delete':
            removed += (i2 - i1)
            for line in old_lines[i1:i2]:
                snippet = line.strip()[:80]
                if snippet:
                    change_details.append(f'- "{snippet}"')
        elif tag == 'replace':
            changed += max(i2 - i1, j2 - j1)
            # Show first replaced line
            old_snippet = old_lines[i1].strip()[:60] if i1 < len(old_lines) else ''
            new_snippet = new_lines[j1].strip()[:60] if j1 < len(new_lines) else ''
            if old_snippet and new_snippet:
                change_details.append(f'"{old_snippet}" → "{new_snippet}"')
    parts = []
    if changed:
        parts.append(f"{changed} line{'s' if changed != 1 else ''} changed")
    if added:
        parts.append(f"{added} line{'s' if added != 1 else ''} added")
    if removed:
        parts.append(f"{removed} line{'s' if removed != 1 else ''} removed")
    summary = ', '.join(parts) or 'Minor formatting changes'
    # Append first few change details for specificity
    if change_details:
        detail_str = '; '.join(change_details[:3])
        if len(change_details) > 3:
            detail_str += f' (+{len(change_details) - 3} more)'
        summary += f' — {detail_str}'

    # Get editor name
    editor = db.session.get(User, session['user_id'])
    editor_name = editor.name if editor else 'Unknown'

    # Save edited text
    will_record.generated_will_text = new_text
    will_record.text_edited_by = session['user_id']
    will_record.text_edited_at = datetime.utcnow()

    # Create edit log entry
    log_entry = WillEditLog(
        will_id=will_id,
        edited_by=session['user_id'],
        edited_by_name=editor_name,
        edited_at=datetime.utcnow(),
        summary=summary,
        details='\n'.join(change_details) if change_details else None,
    )
    db.session.add(log_entry)

    # Update the current (latest) version's text to match the edit
    latest_ver = WillVersion.query.filter_by(will_id=will_id).order_by(
        WillVersion.version_number.desc()
    ).first()
    if latest_ver:
        latest_ver.will_text = new_text
    db.session.commit()

    # Save change log file in client folder
    try:
        client = db.session.get(Client, will_record.client_id)
        if client:
            log_dir = os.path.join(UPLOAD_DIR, client.folder_name, 'edit_logs')
            os.makedirs(log_dir, exist_ok=True)
            log_file = os.path.join(log_dir, f'edit_log_{will_record.id[:8]}.txt')
            timestamp = will_record.text_edited_at.strftime('%d %b %Y, %I:%M %p')
            with open(log_file, 'a') as f:
                f.write(f"[{timestamp}] {editor_name} — {summary}\n")

            # Also write a unified diff file for this edit
            diff_lines = list(difflib.unified_diff(
                old_lines, new_lines,
                fromfile='before', tofile='after',
                lineterm='',
            ))
            if diff_lines:
                diff_file = os.path.join(log_dir, f'diff_{will_record.text_edited_at.strftime("%Y%m%d_%H%M%S")}.txt')
                with open(diff_file, 'w') as f:
                    f.write(f"Edit by {editor_name} at {timestamp}\n")
                    f.write(f"Summary: {summary}\n")
                    f.write('=' * 60 + '\n')
                    f.write('\n'.join(diff_lines) + '\n')
    except Exception:
        pass  # Don't fail the API call if file write fails

    return jsonify({
        'ok': True,
        'edited_at': will_record.text_edited_at.strftime('%d %b %Y, %I:%M %p'),
        'summary': summary,
        'editor_name': editor_name,
    })


# ---------------------------------------------------------------------------
# AI Redraft (clean up edited will text)
# ---------------------------------------------------------------------------

@app.route('/api/will/<will_id>/redraft', methods=['POST'])
@login_required
def api_will_redraft(will_id):
    """Send edited will text to Claude AI for cleanup: renumber clauses, fix cross-references."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404

    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'ok': False, 'error': 'No text provided'}), 400

    current_text = data['text'].strip()
    if not current_text:
        return jsonify({'ok': False, 'error': 'Will text is empty'}), 400

    # Call Claude to clean up
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model=os.environ.get('CLAUDE_MODEL', 'claude-sonnet-4-20250514'),
            max_tokens=8000,
            system="""You are a Malaysian will drafting assistant. Your task is to clean up an edited will document.
Rules:
- Renumber all clauses sequentially (1, 2, 3...)
- Fix any cross-references to match new numbering (e.g. "clause 3 above" → correct clause number)
- Fix grammar issues caused by clause removal or reordering
- Preserve ALL remaining text exactly as written — do not add, rewrite, or remove any content
- Keep the exact same formatting style (spacing, capitalization, structure)
- Output ONLY the cleaned-up will text, nothing else — no preamble, no explanation""",
            messages=[{
                'role': 'user',
                'content': f"Clean up this edited will document. Renumber clauses, fix cross-references, fix grammar from any removed clauses. Output only the will text:\n\n{current_text}"
            }],
        )
        # §10x.70 — track this previously-unaudited callsite
        try:
            from ai.cost_tracker import log_usage
            log_usage(response, call_site='app.api_will_redraft',
                      will_id=will_id,
                      client_id=getattr(will_record, 'client_id', None))
        except Exception:
            pass
        cleaned_text = response.content[0].text.strip()
    except Exception as e:
        return jsonify({'ok': False, 'error': f'AI redraft failed: {str(e)}'}), 500

    return jsonify({'ok': True, 'text': cleaned_text})


@app.route('/api/will/<will_id>/ai-edit', methods=['POST'])
@login_required
def api_will_ai_edit(will_id):
    """AI-assisted will editing: user describes changes, AI identifies clause and suggests replacement."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404

    data = request.get_json()
    if not data or 'instruction' not in data or 'will_text' not in data:
        return jsonify({'ok': False, 'error': 'Missing instruction or will_text'}), 400

    instruction = data['instruction'].strip()
    will_text = data['will_text'].strip()
    if not instruction or not will_text:
        return jsonify({'ok': False, 'error': 'Instruction and will text are required'}), 400

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model=os.environ.get('CLAUDE_MODEL', 'claude-sonnet-4-20250514'),
            max_tokens=4000,
            system="""You are a Malaysian will editing assistant. The user wants to make changes to their will.

Your task:
1. Read the user's instruction carefully
2. Identify the specific clause(s) in the will text that need to be changed
3. Generate the suggested replacement text for that clause
4. Explain what you changed and why

Rules:
- Use proper Malaysian legal drafting language
- Use "I hereby devise and bequeath" for property gifts
- Use Malay terms for property: Mukim, Daerah, Negeri
- Use fractions (4/10) not percentages (40%)
- For Malaysian NRIC: "MALAYSIA (NRIC No. XXXXXX-XX-XXXX)"
- Include discharge/lien clause for property gifts
- Keep the same formatting style as the rest of the will
- If the instruction is unclear, ask for clarification in the explanation field

Return ONLY a JSON object:
{
    "original_text": "The exact text from the will that should be replaced (copy verbatim)",
    "suggested_text": "The new text to replace it with",
    "explanation": "Brief explanation of what was changed",
    "clause_number": "The clause number affected (e.g., '5' or '5,6' for multiple)"
}

If the instruction cannot be applied (e.g., refers to something not in the will), return:
{
    "original_text": "",
    "suggested_text": "",
    "explanation": "Explanation of why the change cannot be made",
    "clause_number": ""
}""",
            messages=[{
                'role': 'user',
                'content': f"Here is the current will text:\n\n{will_text}\n\n---\n\nUser's instruction: {instruction}"
            }],
        )
        # §10x.70 — track this previously-unaudited callsite
        try:
            from ai.cost_tracker import log_usage
            log_usage(response, call_site='app.api_will_ai_edit',
                      will_id=will_id,
                      client_id=getattr(will_record, 'client_id', None))
        except Exception:
            pass
        result_text = response.content[0].text.strip()
        if result_text.startswith('```'):
            result_text = result_text.split('\n', 1)[1].rsplit('```', 1)[0].strip()

        import json as json_mod
        result = json_mod.loads(result_text)
        return jsonify({'ok': True, 'result': result})
    except json_mod.JSONDecodeError:
        return jsonify({'ok': True, 'result': {
            'original_text': '',
            'suggested_text': result_text,
            'explanation': 'AI returned a text response instead of structured edit. You may need to apply manually.',
            'clause_number': ''
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': f'AI edit failed: {str(e)}'}), 500


# ---------------------------------------------------------------------------
# Email sending
# ---------------------------------------------------------------------------

def send_will_email(to_email, subject, body_html, attachments=None, tenant=None):
    """Send email via Google Workspace SMTP Relay (IP-based auth) with tenant-specific FROM/CC."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    tenant = tenant or get_tenant()
    from_email = tenant.get('email_from')
    if not from_email:
        raise ValueError('No email_from configured for this tenant')
    cc_list = tenant.get('email_cc', [])

    # Use SMTP_USER as envelope sender (required for SMTP AUTH), Reply-To for the logical sender
    sender_addr = SMTP_USER or from_email
    msg = MIMEMultipart()
    msg['From'] = sender_addr
    msg['To'] = to_email
    msg['Subject'] = subject
    if from_email and from_email != sender_addr:
        msg['Reply-To'] = from_email
    if cc_list:
        msg['Cc'] = ', '.join(cc_list)

    msg.attach(MIMEText(body_html, 'html'))

    # Attach files (list of dicts: {'filename': ..., 'data': bytes, 'mime': ...})
    for att in (attachments or []):
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(att['data'])
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{att["filename"]}"')
        msg.attach(part)

    all_recipients = [to_email] + cc_list
    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
        server.ehlo()
        server.starttls()
        server.ehlo()
        if SMTP_USER and SMTP_PASSWORD:
            server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(sender_addr, all_recipients, msg.as_string())

    return True


@app.route('/api/support/report', methods=['POST'])
@login_required
def api_support_report():
    """Send a bug/issue report with optional screenshot to support@lifa.com.my."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    description = request.form.get('description', '').strip()
    if not description:
        return jsonify(ok=False, error='Description is required'), 400

    page = request.form.get('page', 'Unknown')
    browser = request.form.get('browser', 'Unknown')
    user_name = session.get('user_name', 'Unknown')
    user_email = session.get('user_email', 'Unknown')
    user_role = session.get('user_role', 'Unknown')
    tenant = get_tenant()
    domain = os.environ.get('WILLCRAFT_DOMAIN', 'Unknown')

    subject = f"[WillCraft Issue] {description[:80]}"

    body_html = f"""
    <h2 style="color: #dc2626;">Issue Report from WillCraft</h2>
    <table style="border-collapse: collapse; width: 100%; max-width: 600px;">
        <tr><td style="padding: 8px; border: 1px solid #e5e7eb; font-weight: bold; background: #f9fafb; width: 140px;">Reported By</td>
            <td style="padding: 8px; border: 1px solid #e5e7eb;">{user_name} ({user_email})</td></tr>
        <tr><td style="padding: 8px; border: 1px solid #e5e7eb; font-weight: bold; background: #f9fafb;">Role</td>
            <td style="padding: 8px; border: 1px solid #e5e7eb;">{user_role}</td></tr>
        <tr><td style="padding: 8px; border: 1px solid #e5e7eb; font-weight: bold; background: #f9fafb;">Site</td>
            <td style="padding: 8px; border: 1px solid #e5e7eb;">{domain}</td></tr>
        <tr><td style="padding: 8px; border: 1px solid #e5e7eb; font-weight: bold; background: #f9fafb;">Page</td>
            <td style="padding: 8px; border: 1px solid #e5e7eb;">{page}</td></tr>
        <tr><td style="padding: 8px; border: 1px solid #e5e7eb; font-weight: bold; background: #f9fafb;">Browser</td>
            <td style="padding: 8px; border: 1px solid #e5e7eb; font-size: 11px;">{browser}</td></tr>
    </table>
    <h3 style="margin-top: 16px;">Description</h3>
    <div style="background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; padding: 12px; white-space: pre-wrap;">{description}</div>
    <p style="color: #9ca3af; font-size: 11px; margin-top: 16px;">Sent automatically from WillCraft AI</p>
    """

    try:
        sender_addr = SMTP_USER or tenant.get('email_from', '')
        msg = MIMEMultipart()
        msg['From'] = sender_addr
        msg['To'] = 'support@lifa.com.my'
        msg['Subject'] = subject
        if tenant.get('email_from') and tenant['email_from'] != sender_addr:
            msg['Reply-To'] = tenant['email_from']

        msg.attach(MIMEText(body_html, 'html'))

        # Attach screenshot if provided
        screenshot = request.files.get('screenshot')
        if screenshot and screenshot.filename:
            part = MIMEBase('image', 'png')
            part.set_payload(screenshot.read())
            encoders.encode_base64(part)
            safe_name = screenshot.filename.replace('"', '')
            part.add_header('Content-Disposition', f'attachment; filename="{safe_name}"')
            msg.attach(part)

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            if SMTP_USER and SMTP_PASSWORD:
                server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(sender_addr, ['support@lifa.com.my'], msg.as_string())

        return jsonify(ok=True)
    except Exception as e:
        print(f"[Support] Email failed: {e}")
        return jsonify(ok=False, error=f'Failed to send: {str(e)}'), 500


@app.route('/api/will/<will_id>/send-email', methods=['POST'])
@login_required
def api_will_send_email(will_id):
    """Email the approved will (PDF) to a recipient."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404

    if will_record.status != 'approved':
        return jsonify({'ok': False, 'error': 'Only approved wills can be emailed'}), 403

    # Get recipient email — prefer custom to_email from request body, fallback to client email
    data = request.get_json(silent=True) or {}
    to_email = (data.get('to_email') or '').strip()
    client = db.session.get(Client, will_record.client_id)

    if not to_email:
        if client and client.email:
            to_email = client.email
        else:
            return jsonify({'ok': False, 'error': 'Please enter a recipient email address'}), 400

    if '@' not in to_email:
        return jsonify({'ok': False, 'error': 'Please enter a valid email address'}), 400

    # Generate PDF attachment
    will_text = will_record.generated_will_text or ''
    if not will_text:
        return jsonify({'ok': False, 'error': 'No will text to send'}), 400

    testator_name = (client.full_name if client else '') or 'Client'
    safe_name = "".join(c for c in testator_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_') or 'Will'

    try:
        from documents.pdf_generator import generate_pdf
        logo = None
        if will_record.include_logo:
            logo = _get_logo_path()
        # Build firm info for cover page and prepared-by page
        tenant = get_tenant()
        firm_info = None
        if tenant.get('firm_name'):
            firm_info = {
                'firm_name': tenant.get('firm_name', ''),
                'firm_address': tenant.get('firm_address', ''),
                'firm_phone': tenant.get('firm_phone', ''),
                'firm_email': tenant.get('email_from', ''),
            }
        filepath = generate_pdf(will_text, safe_name, logo_path=logo, firm_info=firm_info)
        with open(filepath, 'rb') as f:
            pdf_data = f.read()
    except Exception as e:
        app.logger.error(f'PDF generation failed: {e}')
        return jsonify({'ok': False, 'error': 'Failed to generate PDF'}), 500

    # Build email
    tenant = get_tenant()
    brand = tenant.get('brand', 'WillCraft AI')
    user_name = session.get('user_name', 'Unknown')
    user_role = session.get('user_role', '')

    # CC: merge user-provided CC, tenant CC, and auto-CC approver for admin/advisor
    cc_list = list(tenant.get('email_cc', []))
    user_cc = (data.get('cc') or '').strip()
    if user_cc:
        for addr in user_cc.split(','):
            addr = addr.strip()
            if addr and '@' in addr and addr not in cc_list:
                cc_list.append(addr)
    if user_role in ('admin', 'advisor'):
        approvers = User.query.filter_by(role='approver', is_active=True).all()
        for ap in approvers:
            if ap.email and ap.email not in cc_list:
                cc_list.append(ap.email)

    # Use user-provided subject and body, with defaults
    subject = (data.get('subject') or '').strip() or f"Last Will and Testament — {testator_name}"
    user_body = (data.get('body') or '').strip()

    if user_body:
        import html as html_mod
        body_escaped = html_mod.escape(user_body).replace('\n', '<br>')
        body_html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <p>{body_escaped}</p>
            <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
            <p style="font-size: 12px; color: #718096;">
                This email and its attachment are confidential and intended solely for the addressee.
            </p>
        </div>
        """
    else:
        body_html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <p>Dear {testator_name},</p>
            <p>Please find attached the Last Will and Testament in PDF format.</p>
            <p>Kindly review the document carefully. If you have any questions or require
            any amendments, please do not hesitate to contact us.</p>
            <br>
            <p>Best regards,<br><strong>{user_name}</strong><br>{brand}</p>
            <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
            <p style="font-size: 12px; color: #718096;">
                This email and its attachment are confidential and intended solely for the addressee.
            </p>
        </div>
        """

    from_email = tenant.get('email_from')
    if not from_email:
        user = db.session.get(User, session.get('user_id'))
        from_email = user.email if user else None
    if not from_email:
        return jsonify({'ok': False, 'error': 'No sender email configured. Contact admin.'}), 400

    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.base import MIMEBase
        from email import encoders

        sender_addr = SMTP_USER or from_email
        msg = MIMEMultipart()
        msg['From'] = sender_addr
        msg['To'] = to_email
        msg['Subject'] = subject
        if from_email and from_email != sender_addr:
            msg['Reply-To'] = from_email
        if cc_list:
            msg['Cc'] = ', '.join(cc_list)
        msg.attach(MIMEText(body_html, 'html'))

        part = MIMEBase('application', 'octet-stream')
        part.set_payload(pdf_data)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{safe_name}_Will.pdf"')
        msg.attach(part)

        all_recipients = [to_email] + cc_list
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            if SMTP_USER and SMTP_PASSWORD:
                server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(sender_addr, all_recipients, msg.as_string())
    except Exception as e:
        app.logger.error(f'Email sending failed: {e}')
        return jsonify({'ok': False, 'error': f'Failed to send email: {str(e)}'}), 500

    sender_name = session.get('user_name', 'Unknown')
    app.logger.info(f'Will {will_id} emailed to {to_email} by {sender_name} (cc: {cc_list})')

    return jsonify({
        'ok': True,
        'sent_to': to_email,
        'cc': cc_list,
        'message': f'Will emailed to {to_email}',
    })


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route('/')
@login_required
def index():
    """Landing page."""
    wills_query = Will.query.filter(Will.deleted_at.is_(None))
    # Approvers see all wills; admin/advisor see only their own
    if session.get('user_role') != 'approver':
        wills_query = wills_query.filter_by(created_by=session.get('user_id'))
    saved_wills = wills_query.order_by(Will.updated_at.desc()).all()
    return render_template('index.html', saved_wills=saved_wills)


# -- Save / Load / Delete Wills ------------------------------------------------

@app.route('/api/will/save', methods=['POST'])
@login_required
def api_will_save():
    """AJAX endpoint to save current session to DB."""
    try:
        will_record = save_will_to_db()
        return jsonify({'ok': True, 'will_id': will_record.id})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/will/toggle-logo', methods=['POST'])
@login_required
def api_will_toggle_logo():
    """Toggle include_logo flag on the current will."""
    will_id = session.get('will_id')
    if not will_id:
        return jsonify({'ok': False, 'error': 'No will in session'}), 400
    wr = db.session.get(Will, will_id)
    if not wr:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404
    data = request.get_json(silent=True) or {}
    wr.include_logo = bool(data.get('include_logo', True))
    db.session.commit()
    return jsonify({'ok': True, 'include_logo': wr.include_logo})


@app.route('/api/will/delete-generated', methods=['POST'])
@login_required
def api_will_delete_generated():
    """Delete the generated will text (not the will record itself)."""
    will_id = session.get('will_id')
    if not will_id:
        return jsonify({'ok': False, 'error': 'No will in session'}), 400
    wr = db.session.get(Will, will_id)
    if not wr:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404
    wr.generated_will_text = None
    wr.status = 'draft'
    wr.submitted_by = None
    wr.submitted_at = None
    wr.approved_by = None
    wr.approved_at = None
    wr.approval_remarks = None
    session.pop('generated_will_text', None)
    db.session.commit()
    return jsonify({'ok': True})


@app.route('/api/will/restore-generated', methods=['POST'])
@login_required
def api_will_restore_generated():
    """Restore the generated will text from the latest version."""
    will_id = session.get('will_id')
    if not will_id:
        return jsonify({'ok': False, 'error': 'No will in session'}), 400
    wr = db.session.get(Will, will_id)
    if not wr:
        return jsonify({'ok': False, 'error': 'Will not found'}), 404
    # Find the latest version
    latest = WillVersion.query.filter_by(will_id=will_id).order_by(
        WillVersion.version_number.desc()
    ).first()
    if not latest or not latest.will_text:
        return jsonify({'ok': False, 'error': 'No version found to restore'}), 404
    wr.generated_will_text = latest.will_text
    wr.status = 'generated'
    session['generated_will_text'] = latest.will_text
    db.session.commit()
    return jsonify({'ok': True, 'version': latest.version_number})


@app.route('/api/will/version/<int:version_id>/delete', methods=['POST'])
@login_required
def api_will_delete_version(version_id):
    """Delete a specific version from version history."""
    will_id = session.get('will_id')
    if not will_id:
        return jsonify({'ok': False, 'error': 'No will in session'}), 400
    version = db.session.get(WillVersion, version_id)
    if not version or version.will_id != will_id:
        return jsonify({'ok': False, 'error': 'Version not found'}), 404
    # Don't allow deleting if it's the only version
    total = WillVersion.query.filter_by(will_id=will_id).count()
    if total <= 1:
        return jsonify({'ok': False, 'error': 'Cannot delete the only version'}), 400
    # Check if deleting the latest version
    latest = WillVersion.query.filter_by(will_id=will_id).order_by(
        WillVersion.version_number.desc()
    ).first()
    is_latest = (version.id == latest.id)
    db.session.delete(version)
    db.session.flush()
    # If we deleted the latest, update will's generated_will_text to the new latest
    if is_latest:
        new_latest = WillVersion.query.filter_by(will_id=will_id).order_by(
            WillVersion.version_number.desc()
        ).first()
        if new_latest:
            wr = db.session.get(Will, will_id)
            if wr:
                wr.generated_will_text = new_latest.will_text
    db.session.commit()
    return jsonify({'ok': True})


@app.route('/api/feedback', methods=['POST'])
@login_required
def api_feedback():
    """Send feedback/issue report to support@lifa.com.my."""
    client_name = request.form.get('client_name', '').strip()
    description = request.form.get('description', '').strip()
    if not description:
        return jsonify({'ok': False, 'error': 'Description is required'}), 400

    user = session.get('user', {})
    user_name = user.get('name', 'Unknown')
    user_email = user.get('email', 'Unknown')
    tenant = get_tenant()
    tenant_host = tenant.get('host', 'unknown')

    subject = f"[WillCraft] Issue Report — {client_name or 'No client'}"
    body_html = f"""
    <h3>Issue Report from WillCraft AI</h3>
    <p><strong>Client Name:</strong> {client_name}</p>
    <p><strong>Reported by:</strong> {user_name} ({user_email})</p>
    <p><strong>Site:</strong> {tenant_host}</p>
    <hr>
    <p><strong>Problem Description:</strong></p>
    <p style="white-space: pre-wrap;">{description}</p>
    """

    attachments = []
    screenshot = request.files.get('screenshot')
    if screenshot and screenshot.filename:
        att_data = screenshot.read()
        attachments.append({
            'filename': screenshot.filename,
            'data': att_data,
            'mime': screenshot.content_type or 'image/png',
        })

    try:
        send_will_email(
            to_email='support@lifa.com.my',
            subject=subject,
            body_html=body_html,
            attachments=attachments,
            tenant=tenant,
        )
        return jsonify({'ok': True})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/wills')
@login_required
def will_list():
    """Unified client+wills page: list all clients with their wills grouped."""
    q = request.args.get('q', '').strip()
    user_role = session.get('user_role', '')
    user_id = session.get('user_id', '')
    # 🔥 Multi-tenant isolation — Client list is now filtered by created_by
    # for non-approver users. Approvers see all clients across the firm.
    base_q = Client.query_for_user(user_id, user_role)
    if q:
        all_clients = base_q.filter(
            db.or_(
                Client.full_name.ilike(f'%{q}%'),
                Client.nric_passport.ilike(f'%{q}%'),
            )
        ).order_by(Client.updated_at.desc()).all()
    else:
        all_clients = base_q.order_by(Client.updated_at.desc()).all()

    # Build grouped data: each client with their wills and stats
    client_groups = []
    for c in all_clients:
        wills_query = Will.query.filter_by(client_id=c.id).filter(Will.deleted_at.is_(None))
        # Approvers see all wills; admin/advisor see only their own
        if user_role != 'approver':
            wills_query = wills_query.filter_by(created_by=user_id)
        wills = wills_query.order_by(Will.updated_at.desc()).all()
        if user_role != 'approver' and not wills:
            continue  # Skip clients with no wills for this user
        doc_count = Document.query.filter_by(client_id=c.id).count()
        # Count generated files on disk
        draft_count = 0
        generated_count = 0
        folder_path = os.path.join(UPLOAD_DIR, c.folder_name)
        drafts_dir = os.path.join(folder_path, 'drafts')
        gen_dir = os.path.join(folder_path, 'generated')
        if os.path.isdir(drafts_dir):
            draft_count = len([f for f in os.listdir(drafts_dir) if os.path.isfile(os.path.join(drafts_dir, f))])
        if os.path.isdir(gen_dir):
            generated_count = len([f for f in os.listdir(gen_dir) if os.path.isfile(os.path.join(gen_dir, f))])
        client_groups.append({
            'client': c,
            'wills': wills,
            'doc_count': doc_count,
            'draft_count': draft_count,
            'generated_count': generated_count,
        })
    return render_template('will_list.html', client_groups=client_groups, search_query=q)


@app.route('/wills/<will_id>/load')
@login_required
def will_load(will_id):
    """Load a saved will into the session."""
    will_record = db.session.get(Will, will_id)
    if not will_record:
        flash('Will not found.', 'error')
        return redirect(url_for('index'))
    load_will_to_session(will_record)
    flash(f'Loaded: {will_record.title}', 'info')
    # If ?goto=preview and will has generated text, go directly to preview
    if request.args.get('goto') == 'preview' and (will_record.generated_will_text or will_record.status in ('generated', 'pending_approval', 'approved')):
        return redirect(url_for('preview'))
    return redirect(url_for('wizard_step_identities'))


@app.route('/wills/<will_id>/delete', methods=['POST'])
@login_required
def will_delete(will_id):
    """Soft-delete a saved will (recoverable for 30 days)."""
    will_record = db.session.get(Will, will_id)
    if will_record:
        will_record.deleted_at = datetime.utcnow()
        db.session.commit()
        # Clear session if we deleted the currently loaded will
        if session.get('will_id') == will_id:
            session.pop('will_id', None)
        flash('Will moved to trash. It can be restored within 30 days.', 'info')
    return redirect(url_for('will_list'))


@app.route('/wills/<will_id>/restore', methods=['POST'])
@login_required
def will_restore(will_id):
    """Restore a soft-deleted will."""
    will_record = db.session.get(Will, will_id)
    if will_record and will_record.deleted_at:
        will_record.deleted_at = None
        db.session.commit()
        flash('Will restored successfully.', 'success')
    return redirect(url_for('trash_list'))


@app.route('/wills/<will_id>/permanent-delete', methods=['POST'])
@login_required
def will_permanent_delete(will_id):
    """Permanently delete a will (admin only, from trash)."""
    role = session.get('user_role')
    if role not in ('admin',):
        flash('Access denied.', 'error')
        return redirect(url_for('trash_list'))
    will_record = db.session.get(Will, will_id)
    if will_record:
        db.session.delete(will_record)
        db.session.commit()
        flash('Will permanently deleted.', 'success')
    return redirect(url_for('trash_list'))


@app.route('/trash')
@login_required
def trash_list():
    """Show soft-deleted wills and probate applications (recoverable for 30 days)."""
    from datetime import timedelta
    cutoff = datetime.utcnow() - timedelta(days=30)

    # Auto-purge items older than 30 days
    expired_wills = Will.query.filter(Will.deleted_at.isnot(None), Will.deleted_at < cutoff).all()
    for w in expired_wills:
        db.session.delete(w)
    expired_probates = ProbateApplication.query.filter(
        ProbateApplication.deleted_at.isnot(None), ProbateApplication.deleted_at < cutoff
    ).all()
    for p in expired_probates:
        # Clean up generated form files
        gen_forms = ProbateGeneratedForm.query.filter_by(probate_id=p.id).all()
        for gf in gen_forms:
            if gf.file_path and os.path.exists(gf.file_path):
                try:
                    os.remove(gf.file_path)
                except OSError:
                    pass
        ProbateGeneratedForm.query.filter_by(probate_id=p.id).delete()
        db.session.delete(p)
    if expired_wills or expired_probates:
        db.session.commit()

    # Fetch soft-deleted items with days_left
    now = datetime.utcnow()
    deleted_wills = Will.query.filter(Will.deleted_at.isnot(None)).order_by(Will.deleted_at.desc()).all()
    for w in deleted_wills:
        w.days_left = max(0, 30 - (now - w.deleted_at).days)
    deleted_probates = ProbateApplication.query.filter(
        ProbateApplication.deleted_at.isnot(None)
    ).order_by(ProbateApplication.deleted_at.desc()).all()
    for p in deleted_probates:
        p.days_left = max(0, 30 - (now - p.deleted_at).days)

    return render_template('trash.html',
                           deleted_wills=deleted_wills,
                           deleted_probates=deleted_probates)


@app.route('/clients/<client_id>/delete', methods=['POST'])
@login_required
def client_delete(client_id):
    """Delete a client and ALL associated data (wills, persons, documents, disk files)."""
    import shutil
    client = db.session.get(Client, client_id)
    if not client:
        flash('Client not found.', 'error')
        return redirect(url_for('will_list'))

    # Delete associated documents from disk
    folder_path = os.path.join(UPLOAD_DIR, client.folder_name)
    if os.path.isdir(folder_path):
        shutil.rmtree(folder_path, ignore_errors=True)

    # Delete DB records (cascade: documents, wills, persons)
    Document.query.filter_by(client_id=client_id).delete()
    Will.query.filter_by(client_id=client_id).delete()
    Person.query.filter_by(client_id=client_id).delete()
    db.session.delete(client)
    db.session.commit()

    # Clear loaded-client / loaded-will references — but NEVER the whole
    # session. session.clear() also wipes user_id, which kicks the user
    # back to the login page mid-task. Only pop the keys that point to
    # the now-deleted client.
    if session.get('client_id') == client_id:
        for _k in ('client_id', 'will_id', 'completed_steps',
                   'identities', 'step1', 'step2', 'step3', 'step4',
                   'step5_gifts', 'step6_residuary', 'step7_trust',
                   'step8_others', 'step9_witnesses', 'witnesses'):
            session.pop(_k, None)

    flash(f'Client "{client.full_name}" and all associated data deleted.', 'info')
    return redirect(url_for('will_list'))


# -- Person Registry API -------------------------------------------------------

@app.route('/api/persons', methods=['GET'])
@login_required
def api_persons_list():
    """Return JSON list of persons for the current client."""
    client_id = session.get('client_id')
    if not client_id:
        return jsonify([])
    persons = Person.query.filter_by(client_id=client_id).order_by(Person.full_name).all()
    return jsonify([
        {'id': p.id, 'full_name': p.full_name, 'nric_passport': p.nric_passport,
         'address': p.address or '', 'date_of_birth': p.date_of_birth or '',
         'nationality': p.nationality or 'Malaysian', 'gender': p.gender or '',
         'passport_expiry': p.passport_expiry or '',
         'email': p.email or '', 'phone': p.phone or '',
         'relationship': p.relationship or '',
         'document_id': p.document_id or ''}
        for p in persons
    ])


@app.route('/api/persons', methods=['POST'])
@login_required
def api_persons_create():
    """Create a new person identity."""
    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    data = request.get_json() or {}
    full_name = (data.get('full_name') or '').strip()
    nric_passport = (data.get('nric_passport') or '').strip()
    if not full_name or not nric_passport:
        return jsonify({'ok': False, 'error': 'Name and NRIC/Passport are required'}), 400
    try:
        person = upsert_person(
            client_id, full_name, nric_passport,
            address=(data.get('address') or '').strip(),
            date_of_birth=(data.get('date_of_birth') or '').strip() or None,
            nationality=(data.get('nationality') or 'Malaysian').strip(),
            gender=(data.get('gender') or '').strip() or None,
            passport_expiry=(data.get('passport_expiry') or '').strip() or None,
            email=(data.get('email') or '').strip() or None,
            phone=(data.get('phone') or '').strip() or None,
            relationship=(data.get('relationship') or '').strip() or None,
            document_id=(data.get('document_id') or '').strip() or None,
        )
    except Exception as e:
        app.logger.error(f'Failed to save identity: {e}')
        return jsonify({'ok': False, 'error': f'Failed to save: {str(e)}'}), 500
    return jsonify({'ok': True, 'person': {
        'id': person.id, 'full_name': person.full_name,
        'nric_passport': person.nric_passport, 'address': person.address or '',
        'nationality': person.nationality or 'Malaysian',
        'date_of_birth': person.date_of_birth or '',
        'gender': person.gender or '',
        'email': person.email or '', 'phone': person.phone or '',
        'passport_expiry': person.passport_expiry or '',
        'relationship': person.relationship or '',
        'document_id': person.document_id or '',
    }})


@app.route('/api/persons/<person_id>', methods=['PUT'])
@login_required
def api_persons_update(person_id):
    """Update an existing person identity."""
    person = db.session.get(Person, person_id)
    if not person:
        return jsonify({'ok': False, 'error': 'Person not found'}), 404
    data = request.get_json() or {}
    old_name = person.full_name  # Capture before update
    if data.get('full_name'):
        person.full_name = data['full_name'].strip().upper()
    if data.get('nric_passport'):
        person.nric_passport = data['nric_passport'].strip()
    if 'address' in data:
        person.address = (data['address'] or '').strip()
    if 'nationality' in data:
        person.nationality = (data['nationality'] or 'Malaysian').strip()
    if 'passport_expiry' in data:
        person.passport_expiry = (data['passport_expiry'] or '').strip() or None
    if 'date_of_birth' in data:
        person.date_of_birth = (data['date_of_birth'] or '').strip() or None
    if 'gender' in data:
        person.gender = (data['gender'] or '').strip() or None
    if 'email' in data:
        person.email = (data['email'] or '').strip() or None
    if 'phone' in data:
        person.phone = (data['phone'] or '').strip() or None
    if 'relationship' in data:
        person.relationship = (data['relationship'] or '').strip() or None
    if 'document_id' in data:
        person.document_id = (data['document_id'] or '').strip() or None
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Failed to update identity: {e}')
        return jsonify({'ok': False, 'error': f'Failed to update: {str(e)}'}), 500
    _refresh_session_person_registry(person.client_id)
    _propagate_identity_changes(person.id, person.full_name, person.nric_passport, old_name)
    save_will_to_db()  # Persist propagated changes
    return jsonify({'ok': True, 'person': {
        'id': person.id, 'full_name': person.full_name,
        'nric_passport': person.nric_passport, 'address': person.address or '',
        'nationality': person.nationality or 'Malaysian',
        'date_of_birth': person.date_of_birth or '',
        'gender': person.gender or '',
        'email': person.email or '', 'phone': person.phone or '',
        'passport_expiry': person.passport_expiry or '',
        'relationship': person.relationship or '',
        'document_id': person.document_id or '',
    }})


@app.route('/api/persons/<person_id>', methods=['DELETE'])
@login_required
def api_persons_delete(person_id):
    """Delete a person identity after checking for references."""
    person = db.session.get(Person, person_id)
    if not person:
        return jsonify({'ok': False, 'error': 'Person not found'}), 404

    # Check if this person is referenced anywhere in the will
    refs = _get_person_references(person_id, person.full_name)
    if refs:
        msg = f"Cannot delete {person.full_name}. This person is assigned as:\n"
        msg += "\n".join(f"• {r}" for r in refs)
        msg += "\n\nPlease remove them from these roles first."
        return jsonify({'ok': False, 'error': msg}), 400

    client_id = person.client_id
    db.session.delete(person)
    db.session.commit()
    _refresh_session_person_registry(client_id)
    return jsonify({'ok': True})


@app.route('/api/persons/<person_id>', methods=['PATCH'])
@login_required
def api_persons_patch(person_id):
    """Partial update of a person — for inline error correction."""
    person = db.session.get(Person, person_id)
    if not person:
        return jsonify({'ok': False, 'error': 'Person not found'}), 404

    data = request.get_json()
    if not data:
        return jsonify({'ok': False, 'error': 'No data provided'}), 400

    # Only allow updating specific fields
    allowed_fields = ['full_name', 'nric_passport', 'address', 'date_of_birth',
                       'gender', 'nationality', 'email', 'phone', 'relationship']
    updated = []
    for field in allowed_fields:
        if field in data:
            setattr(person, field, data[field])
            updated.append(field)

    if updated:
        db.session.commit()
        _refresh_session_person_registry(person.client_id)
        return jsonify({'ok': True, 'updated': updated})

    return jsonify({'ok': False, 'error': 'No valid fields to update'}), 400


@app.route('/api/validate/persons', methods=['GET'])
@login_required
def api_validate_persons():
    """Validate all persons in the current will session."""
    from validation.field_validator import validate_person
    registry = session.get('person_registry', [])
    results = {}
    for p in registry:
        errors = validate_person(p)
        if errors:
            results[p.get('id', '')] = {
                'name': p.get('full_name', ''),
                'errors': errors
            }
    return jsonify({'ok': True, 'results': results})


@app.route('/api/validate/gifts', methods=['GET'])
@login_required
def api_validate_gifts():
    """Validate all gift property details in the current will session."""
    from validation.field_validator import validate_property_details
    gifts = session.get('step5_gifts', [])
    results = {}
    for i, g in enumerate(gifts):
        if g.get('gift_type') == 'property':
            prop = g.get('property_details', {})
            errors = validate_property_details(prop)
            if errors:
                results[str(i)] = {
                    'description': prop.get('property_address', f'Gift {i+1}')[:50],
                    'errors': errors
                }
    return jsonify({'ok': True, 'results': results})


def _get_person_references(person_id, full_name):
    """Find all references to a person across will wizard session data."""
    refs = []

    # Step 2 - Testator
    step1 = session.get('step1', {})
    if step1.get('person_id') == person_id:
        refs.append('Testator')

    # Step 3 - Executors
    for ex in session.get('step2_executors', []):
        if ex.get('person_id') == person_id:
            role = ex.get('role', 'Primary')
            refs.append(f'Executor ({role})')

    # Step 3 - Trustees
    trustee_data = session.get('step3_trustees', {})
    for tr in trustee_data.get('trustees', []):
        if tr.get('person_id') == person_id:
            refs.append('Trustee')
    for st in trustee_data.get('substitute_trustees', []):
        if st.get('person_id') == person_id:
            refs.append('Substitute Trustee')

    # Step 4 - Guardians
    for gdn in session.get('step3_guardians', []):
        if gdn.get('person_id') == person_id:
            role = gdn.get('role', 'Primary')
            refs.append(f'Guardian ({role})')

    # Step 5 - Beneficiaries
    for ben in session.get('step4_beneficiaries', []):
        if ben.get('person_id') == person_id:
            refs.append('Beneficiary')

    # Step 6 - Gift allocations (matched by name)
    for gi, gift in enumerate(session.get('step5_gifts', [])):
        gift_num = gi + 1
        for alloc in gift.get('allocations', []):
            if alloc.get('beneficiary_name') == full_name:
                refs.append(f'Beneficiary of Gift No. {gift_num}')
            for sub in alloc.get('substitutes', []):
                if sub.get('beneficiary_name') == full_name:
                    refs.append(f'Substitute Beneficiary of Gift No. {gift_num}')

    # Step 7 - Residuary Estate
    res = session.get('step6_residuary', {})
    for mb in res.get('main_beneficiaries', []):
        if mb.get('person_id') == person_id:
            refs.append('Residuary Beneficiary')
    for sg in res.get('substitute_groups', []):
        for sb in (sg if isinstance(sg, list) else []):
            if sb.get('person_id') == person_id:
                refs.append('Residuary Substitute Beneficiary')

    # Step 8 - Trust
    trust = session.get('step7_trust', {})
    if trust.get('trustee_person_id') == person_id:
        refs.append('Trust Trustee')

    return refs


# -- Upload & Document API ----------------------------------------------------

@app.route('/api/upload', methods=['POST'])
@login_required
def api_upload():
    """Generic file upload endpoint."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    category = request.form.get('category', 'general')
    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    try:
        from uploads import save_uploaded_file
        folder_name = get_client_folder_name(client_id)
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, category, folder_name=folder_name)
    except ValueError as e:
        return jsonify({'ok': False, 'error': str(e)}), 400
    doc = Document(
        client_id=client_id,
        will_id=session.get('will_id'),
        filename=saved_name,
        original_filename=file.filename,
        file_path=rel_path,
        file_type=file.content_type,
        file_size=file_size,
        category=category,
    )
    db.session.add(doc)
    db.session.commit()
    return jsonify({'ok': True, 'document_id': doc.id, 'filename': saved_name})


@app.route('/api/documents')
@login_required
def api_documents_list():
    """List documents for current client."""
    client_id = session.get('client_id')
    if not client_id:
        return jsonify([])
    docs = Document.query.filter_by(client_id=client_id).order_by(Document.created_at.desc()).all()
    return jsonify([
        {'id': d.id, 'filename': d.original_filename, 'category': d.category,
         'file_size': d.file_size, 'created_at': d.created_at.isoformat()}
        for d in docs
    ])


@app.route('/api/documents/<doc_id>')
@login_required
def api_document_view(doc_id):
    """View/download a specific document."""
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    from config import UPLOAD_DIR
    abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
    if not os.path.exists(abs_path):
        return jsonify({'error': 'File not found on disk'}), 404
    return send_file(abs_path, download_name=doc.original_filename)


@app.route('/api/documents/<doc_id>', methods=['DELETE'])
@login_required
def api_document_delete(doc_id):
    """Delete a specific document."""
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({'ok': False, 'error': 'Document not found'}), 404
    # Remove file from disk
    from config import UPLOAD_DIR
    abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
    if os.path.exists(abs_path):
        os.remove(abs_path)
    # Clear document_id from any linked persons
    linked_persons = Person.query.filter_by(document_id=doc_id).all()
    for p in linked_persons:
        p.document_id = None
    # Clear references from probate applications
    for pa in ProbateApplication.query.filter_by(death_cert_document_id=doc_id).all():
        pa.death_cert_document_id = None
    for pa in ProbateApplication.query.filter_by(will_document_id=doc_id).all():
        pa.will_document_id = None
    db.session.delete(doc)
    db.session.commit()
    if linked_persons:
        _refresh_session_person_registry(linked_persons[0].client_id)
    return jsonify({'ok': True})


@app.route('/api/documents/<doc_id>/translate', methods=['POST'])
@login_required
def api_document_translate(doc_id):
    """Translate a document image from Bahasa Malaysia to English."""
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({'ok': False, 'error': 'Document not found'}), 404
    from config import UPLOAD_DIR
    abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
    if not os.path.isfile(abs_path):
        return jsonify({'ok': False, 'error': 'File not found on disk'}), 404
    try:
        from ai.ocr import translate_document
        translation = translate_document(abs_path)
        return jsonify({'ok': True, 'translation': translation})
    except Exception as e:
        app.logger.error(f'Document translate error: {e}')
        return jsonify({'ok': False, 'error': 'Translation failed. Please try again.'}), 500


# -- OCR Extraction API -------------------------------------------------------

@app.route('/api/ocr/nric', methods=['POST'])
@login_required
def api_ocr_nric():
    """Upload NRIC/passport image, extract data via Claude Vision."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400
    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    try:
        from uploads import save_uploaded_file
        folder_name = get_client_folder_name(client_id)
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, 'nric', folder_name=folder_name)
    except ValueError as e:
        return jsonify({'ok': False, 'error': str(e)}), 400
    abs_path = os.path.join(UPLOAD_DIR, rel_path)
    extracted = None
    ocr_warning = None

    # 🔥 §10x.81 + §10x.82 — Skip the expensive vision call when this is the
    # BACK of an IC we already scanned. Front IC has all key fields
    # (name + NRIC + DOB). Back has only address. The Haiku vision pass
    # on a back-of-IC costs ~$0.014 and adds NO new field beyond what's
    # already known. We do a free Tesseract pre-check and look for ANY
    # of three signals that this IC is for an already-verified Person:
    #
    #   (a) NRIC visible on the image → match Person.nric_passport
    #   (b) Address visible on the image → match Person.address
    #   (c) Full name visible on the image → match Person.full_name
    #
    # ANY signal matching → skip vision. The user uploaded a back-IC
    # whose front we already verified — no need to spend $0.014 to
    # learn nothing new.
    skipped_vision = False
    skip_reason = ''
    matched_person = None
    try:
        from ai.ocr_preprocessor import ocr_extract
        import re as _re_nric
        _quick_text = (ocr_extract(abs_path) or '').upper()
        _quick_lines = _quick_text.replace('\n', ' ')
        existing = Person.query.filter_by(client_id=client_id).all()
        # (a) NRIC match
        _m_nric = _re_nric.search(r'(\d{6})[-\s]?(\d{2})[-\s]?(\d{4})', _quick_text)
        if _m_nric:
            _nric_norm = (_m_nric.group(1) + _m_nric.group(2) + _m_nric.group(3))
            for _p in existing:
                _p_norm = (_p.nric_passport or '').replace('-', '').replace(' ', '')
                if _p_norm == _nric_norm:
                    matched_person, skip_reason = _p, 'nric_match'
                    break
        # (b) Address match — cheap fuzzy: any 3+ shared address tokens
        if not matched_person and len(_quick_text) > 20:
            for _p in existing:
                if not _p.address:
                    continue
                _p_addr_up = _p.address.upper()
                # Pull street/postcode tokens from the Person address
                _p_tokens = set(t for t in _re_nric.split(r'[\s,/\-]+', _p_addr_up)
                                 if len(t) >= 4 and not t.isdigit())
                _shared = sum(1 for t in _p_tokens if t in _quick_lines)
                # Also count postcode hits
                _p_postcode = _re_nric.search(r'\b\d{5}\b', _p_addr_up)
                if _p_postcode and _p_postcode.group(0) in _quick_text:
                    _shared += 1
                if _shared >= 3:
                    matched_person, skip_reason = _p, 'address_match'
                    break
        # (c) Name match — Person's surname + first name both present
        if not matched_person:
            for _p in existing:
                if not _p.full_name:
                    continue
                _name_tokens = [t for t in _p.full_name.upper().split()
                                if len(t) >= 3]
                if len(_name_tokens) >= 2 and all(t in _quick_text for t in _name_tokens[:2]):
                    matched_person, skip_reason = _p, 'name_match'
                    break

        if matched_person:
            extracted = {
                'doc_type': 'nric',
                'full_name': matched_person.full_name or '',
                'nric_number': matched_person.nric_passport or '',
                'date_of_birth': matched_person.date_of_birth or '',
                'address': matched_person.address or '',
                'gender': matched_person.gender or '',
                'nationality': matched_person.nationality or 'Malaysian',
                'passport_expiry': matched_person.passport_expiry or '',
                '_already_known': True,
                '_matched_person_id': matched_person.id,
                '_skip_reason': skip_reason,
                '_savings_usd': 0.014,
            }
            skipped_vision = True
    except Exception:
        pass  # quick-OCR optimisation is best-effort — fall through to vision

    if not skipped_vision:
        try:
            from ai.ocr import extract_nric_data
            extracted = extract_nric_data(abs_path)
        except Exception as e:
            app.logger.error(f'OCR NRIC error: {e}')
            ocr_warning = 'Image unclear — could not scan automatically. File saved. Please fill in the details manually.'

    if extracted:
        # Address fallback: if OCR couldn't read address clearly (empty),
        # check if the same NRIC already has an address on file and use that.
        if not extracted.get('address') and extracted.get('nric_number'):
            nric_norm = extracted['nric_number'].replace('-', '').replace(' ', '').upper()
            existing_persons = Person.query.filter_by(client_id=client_id).all()
            for p in existing_persons:
                p_norm = (p.nric_passport or '').replace('-', '').replace(' ', '').upper()
                if p_norm == nric_norm and p.address:
                    extracted['address'] = p.address
                    extracted['_address_from_existing'] = True
                    break
        # Remove internal confidence field (not needed in response)
        extracted.pop('confidence', None)

    # 🔥 §10x.82 — When the scan was skipped (already-verified person),
    # save the doc as `duplicate` instead of `nric` so the IC walker
    # doesn't add it to the pending list — the user already confirmed
    # this Person via the front IC.
    doc_category = 'duplicate' if skipped_vision else 'nric'
    doc = Document(
        client_id=client_id, will_id=session.get('will_id'),
        filename=saved_name, original_filename=file.filename,
        file_path=rel_path, file_type=file.content_type,
        file_size=file_size, category=doc_category,
        extracted_data=json.dumps(extracted) if extracted else None,
    )
    db.session.add(doc)
    db.session.commit()
    result = {'ok': True, 'document_id': doc.id}
    if extracted:
        result['extracted'] = extracted
    if skipped_vision and matched_person:
        result['already_known'] = True
        result['matched_person'] = {
            'id': matched_person.id,
            'name': matched_person.full_name,
            'relationship': matched_person.relationship,
        }
        result['skip_reason'] = skip_reason
        result['savings_usd'] = 0.014
        result['notice'] = (
            f"This is the back of {matched_person.full_name}'s IC — "
            f"already verified. Skipped scan (saved ~$0.014)."
        )
    if ocr_warning:
        result['warning'] = ocr_warning
    return jsonify(result)


@app.route('/api/ocr/nric/<document_id>', methods=['POST'])
@login_required
def api_ocr_nric_scan(document_id):
    """Run OCR on an already-uploaded NRIC/passport document."""
    doc = db.session.get(Document, document_id)
    if not doc:
        return jsonify({'ok': False, 'error': 'Document not found'}), 404
    abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
    if not os.path.isfile(abs_path):
        return jsonify({'ok': False, 'error': 'File not found on disk'}), 404

    client_id = doc.client_id or session.get('client_id')
    extracted = None
    ocr_warning = None
    try:
        from ai.ocr import extract_nric_data
        extracted = extract_nric_data(abs_path)
    except Exception as e:
        app.logger.error(f'OCR NRIC scan error: {e}')
        ocr_warning = 'Image unclear — could not scan automatically. Please fill in the details manually.'

    if extracted:
        if not extracted.get('address') and extracted.get('nric_number'):
            nric_norm = extracted['nric_number'].replace('-', '').replace(' ', '').upper()
            existing_persons = Person.query.filter_by(client_id=client_id).all()
            for p in existing_persons:
                p_norm = (p.nric_passport or '').replace('-', '').replace(' ', '').upper()
                if p_norm == nric_norm and p.address:
                    extracted['address'] = p.address
                    extracted['_address_from_existing'] = True
                    break
        extracted.pop('confidence', None)
        doc.extracted_data = json.dumps(extracted)
        db.session.commit()

    result = {'ok': True, 'document_id': doc.id}
    if extracted:
        result['extracted'] = extracted
    if ocr_warning:
        result['warning'] = ocr_warning
    return jsonify(result)


@app.route('/api/ocr/property', methods=['POST'])
@login_required
def api_ocr_property():
    """Upload cukai tanah/cukai pintu, extract property data."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400
    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    try:
        from uploads import save_uploaded_file
        folder_name = get_client_folder_name(client_id)
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, 'property', folder_name=folder_name)
    except ValueError as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    # ── Deduplication: skip if identical filename+size already exists for client ──
    existing_dup = Document.query.filter_by(
        client_id=client_id,
        original_filename=file.filename,
        file_size=file_size,
    ).filter(Document.category != 'deleted').first()
    if existing_dup:
        app.logger.info('Skipping duplicate upload: %s (size %s) for client %s',
                        file.filename, file_size, client_id)
        return jsonify({
            'ok': True,
            'document_id': existing_dup.id,
            'document_url': f'/api/documents/{existing_dup.id}',
            'duplicate': True,
            'extracted': json.loads(existing_dup.extracted_data) if existing_dup.extracted_data else None,
        })

    abs_path = os.path.join(UPLOAD_DIR, rel_path)
    extracted = None
    ocr_warning = None
    try:
        from ai.property_extractor import extract_property_data
        doc_type = request.form.get('doc_type', 'general')
        extracted = extract_property_data(abs_path, doc_type=doc_type)
    except Exception as e:
        app.logger.error(f'OCR property error: {e}')
        ocr_warning = 'Image unclear — could not scan automatically. File saved. Please fill in the details manually.'
    doc = Document(
        client_id=client_id, will_id=session.get('will_id'),
        filename=saved_name, original_filename=file.filename,
        file_path=rel_path, file_type=file.content_type,
        file_size=file_size, category='property',
        extracted_data=json.dumps(extracted) if extracted else None,
    )
    db.session.add(doc)
    db.session.commit()
    result = {'ok': True, 'document_id': doc.id, 'document_url': f'/api/documents/{doc.id}'}
    if extracted:
        result['extracted'] = extracted
    if ocr_warning:
        result['warning'] = ocr_warning
    return jsonify(result)


@app.route('/api/ocr/asset', methods=['POST'])
@login_required
def api_ocr_asset():
    """Upload bank/investment statement, extract asset data."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400
    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    try:
        from uploads import save_uploaded_file
        folder_name = get_client_folder_name(client_id)
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, 'financial', folder_name=folder_name)
    except ValueError as e:
        return jsonify({'ok': False, 'error': str(e)}), 400
    abs_path = os.path.join(UPLOAD_DIR, rel_path)
    extracted = None
    ocr_warning = None
    try:
        from ai.asset_extractor import extract_asset_data
        extracted = extract_asset_data(abs_path)
    except Exception as e:
        app.logger.error(f'OCR asset error: {e}')
        ocr_warning = 'Image unclear — could not scan automatically. File saved. Please fill in the details manually.'
    doc = Document(
        client_id=client_id, will_id=session.get('will_id'),
        filename=saved_name, original_filename=file.filename,
        file_path=rel_path, file_type=file.content_type,
        file_size=file_size, category='financial',
        extracted_data=json.dumps(extracted) if extracted else None,
    )
    db.session.add(doc)
    db.session.commit()
    result = {'ok': True, 'document_id': doc.id, 'document_url': f'/api/documents/{doc.id}'}
    if extracted:
        result['extracted'] = extracted
    if ocr_warning:
        result['warning'] = ocr_warning
    return jsonify(result)


@app.route('/client/documents')
@login_required
def client_documents():
    """Client document browser page (legacy — redirects to new client files page)."""
    client_id = session.get('client_id')
    if client_id:
        return redirect(url_for('client_files', client_id=client_id))
    documents = []
    return render_template('client_documents.html', documents=documents)


@app.route('/clients')
@login_required
def clients_list():
    """Redirect to unified /wills page (backward compatibility)."""
    q = request.args.get('q', '')
    if q:
        return redirect(url_for('will_list', q=q))
    return redirect(url_for('will_list'))


@app.route('/clients/<client_id>/files')
@login_required
def client_files(client_id):
    """Browse all files for a specific client: documents, drafts, generated wills."""
    client = db.session.get(Client, client_id)
    if not client:
        flash('Client not found.', 'error')
        return redirect(url_for('will_list'))

    # Uploaded documents from DB
    documents = Document.query.filter_by(client_id=client_id).order_by(Document.created_at.desc()).all()

    # Group docs by category
    doc_groups = {}
    for doc in documents:
        cat = doc.category or 'general'
        if cat not in doc_groups:
            doc_groups[cat] = []
        doc_groups[cat].append(doc)

    # Scan client folder for drafts and generated wills
    drafts = []
    generated = []
    folder_path = os.path.join(UPLOAD_DIR, client.folder_name)
    drafts_dir = os.path.join(folder_path, 'drafts')
    gen_dir = os.path.join(folder_path, 'generated')

    if os.path.isdir(drafts_dir):
        for fname in sorted(os.listdir(drafts_dir), reverse=True):
            fpath = os.path.join(drafts_dir, fname)
            if os.path.isfile(fpath):
                drafts.append({
                    'filename': fname,
                    'size': os.path.getsize(fpath),
                    'modified': os.path.getmtime(fpath),
                    'rel_path': os.path.join(client.folder_name, 'drafts', fname),
                })

    if os.path.isdir(gen_dir):
        for fname in sorted(os.listdir(gen_dir), reverse=True):
            fpath = os.path.join(gen_dir, fname)
            if os.path.isfile(fpath):
                generated.append({
                    'filename': fname,
                    'size': os.path.getsize(fpath),
                    'modified': os.path.getmtime(fpath),
                    'rel_path': os.path.join(client.folder_name, 'generated', fname),
                })

    # Wills in DB for this client
    wills = Will.query.filter_by(client_id=client_id).filter(Will.deleted_at.is_(None)).order_by(Will.updated_at.desc()).all()

    total_docs = sum(len(docs) for docs in doc_groups.values())
    return render_template('client_files.html',
                           client=client, doc_groups=doc_groups,
                           drafts=drafts, generated=generated, wills=wills,
                           total_docs=total_docs)


@app.route('/clients/<client_id>/files/download/<path:rel_path>')
@login_required
def client_file_download(client_id, rel_path):
    """Download a file from the client's folder."""
    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'error': 'Client not found'}), 404
    # Security: ensure path starts with client's folder name
    if not rel_path.startswith(client.folder_name):
        return jsonify({'error': 'Access denied'}), 403
    abs_path = os.path.join(UPLOAD_DIR, rel_path)
    if not os.path.exists(abs_path):
        return jsonify({'error': 'File not found'}), 404
    return send_file(abs_path, download_name=os.path.basename(rel_path))


# -- Upload Existing Will ----------------------------------------------------

@app.route('/upload-will')
@login_required
def upload_will():
    """Page to upload an existing will for parsing."""
    return render_template('upload_will.html')


@app.route('/api/parse-will', methods=['POST'])
@login_required
def api_parse_will():
    """Upload and parse an existing will document, populate session."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ('pdf', 'docx'):
        return jsonify({'ok': False, 'error': 'Only PDF and DOCX files are supported'}), 400

    client_id = session.get('client_id')
    if not client_id:
        client_id = ensure_client()
    try:
        from uploads import save_uploaded_file
        folder_name = get_client_folder_name(client_id)
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, 'wills', folder_name=folder_name)
    except ValueError as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    abs_path = os.path.join(UPLOAD_DIR, rel_path)

    # Save document record
    doc = Document(
        client_id=client_id, will_id=session.get('will_id'),
        filename=saved_name, original_filename=file.filename,
        file_path=rel_path, file_type=file.content_type,
        file_size=file_size, category='wills',
    )
    db.session.add(doc)
    db.session.commit()

    try:
        from ai.will_parser import parse_will_document
        parsed = parse_will_document(abs_path)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Parsing failed: {e}'}), 500

    if parsed.get('error'):
        return jsonify({'ok': False, 'error': parsed['error']}), 500

    # The parser returns objects like {step2_executors: {executors: [...]}} —
    # unwrap the nested arrays so the session matches what the wizard expects.
    def _unwrap(obj, key):
        if isinstance(obj, list):
            return obj
        if isinstance(obj, dict):
            return obj.get(key, []) or []
        return []

    from services.person_registry import ensure_person, normalise_dob as _normalise_dob

    def _ensure_person(name, nric, address='', relationship='', dob='',
                       nationality='Malaysian'):
        return ensure_person(client_id, name, nric=nric, address=address,
                             relationship=relationship, dob=dob, nationality=nationality)

    # ── Step 1: Testator ─────────────────────────────────────────────
    if 'step1_testator' in parsed:
        s1 = dict(parsed['step1_testator'] or {})
        s1['date_of_birth'] = _normalise_dob(s1.get('date_of_birth', ''))
        # Create / update the testator's Person record
        testator_pid = _ensure_person(
            s1.get('full_name'), s1.get('nric_passport'),
            address=s1.get('residential_address', ''),
            relationship='Testator',
            dob=s1.get('date_of_birth', ''),
            nationality=s1.get('nationality', 'Malaysian'),
        )
        if testator_pid:
            s1['person_id'] = testator_pid
        session['step1'] = s1

    # ── Step 2: Executors ────────────────────────────────────────────
    if 'step2_executors' in parsed:
        raw_execs = _unwrap(parsed['step2_executors'], 'executors')
        execs = []
        for ex in raw_execs:
            pid = _ensure_person(
                ex.get('full_name'), ex.get('nric_passport'),
                address=ex.get('address', ''),
                relationship=ex.get('relationship', ''),
                nationality=ex.get('nationality', 'Malaysian'),
            )
            ex_clean = dict(ex)
            if pid:
                ex_clean['person_id'] = pid
            execs.append(ex_clean)
        session['step2_executors'] = execs
        session['step3_executor_type'] = 'joint' if len(execs) > 1 else 'single'
        session['step3_trustees'] = {'same_as_executor': True, 'trustees': [{}]}

    # ── Step 3: Guardians ────────────────────────────────────────────
    if 'step3_guardians' in parsed:
        raw_g = _unwrap(parsed['step3_guardians'], 'guardians')
        guardians = []
        for g in raw_g:
            pid = _ensure_person(g.get('full_name'), g.get('nric_passport'),
                                  address=g.get('address', ''),
                                  relationship=g.get('relationship', ''))
            gd = dict(g)
            if pid:
                gd['person_id'] = pid
            guardians.append(gd)
        session['step3_guardians'] = guardians

    # ── Step 4: Beneficiaries ────────────────────────────────────────
    if 'step4_beneficiaries' in parsed:
        raw_b = _unwrap(parsed['step4_beneficiaries'], 'beneficiaries')
        bens = []
        for b in raw_b:
            pid = _ensure_person(
                b.get('full_name'),
                b.get('nric_passport') or b.get('nric_passport_birthcert'),
                relationship=b.get('relationship', ''),
            )
            bd = dict(b)
            if pid:
                bd['person_id'] = pid
            # Wizard expects nric_passport_birthcert
            if 'nric_passport' in bd and 'nric_passport_birthcert' not in bd:
                bd['nric_passport_birthcert'] = bd.pop('nric_passport')
            bens.append(bd)
        session['step4_beneficiaries'] = bens

    # ── Step 5: Gifts ────────────────────────────────────────────────
    if 'step5_gifts' in parsed:
        session['step5_gifts'] = _unwrap(parsed['step5_gifts'], 'gifts')
    # ── Step 6/7/8: Residuary / Trust / Others ───────────────────────
    if 'step6_residuary' in parsed:
        s6 = parsed['step6_residuary']
        session['step6_residuary'] = s6 if isinstance(s6, dict) else {}
    if 'step7_trust' in parsed:
        s7 = parsed['step7_trust']
        session['step7_trust'] = s7 if isinstance(s7, dict) else {}
    if 'step8_other_matters' in parsed:
        s8 = parsed['step8_other_matters']
        session['step8_others'] = s8 if isinstance(s8, dict) else {}

    # Commit Person creates + refresh registry so wizard pickers see them
    db.session.commit()
    _refresh_session_person_registry(client_id)

    # Mark all steps with imported data as complete so the wizard shows them
    # in green and the user can jump straight to step 10 (Review & Generate).
    completed = set(session.get('completed_steps', []))
    for n, has in [
        (1, 'step1_testator' in parsed),
        (2, 'step1_testator' in parsed),
        (3, 'step2_executors' in parsed),
        (4, 'step3_guardians' in parsed),
        (5, 'step4_beneficiaries' in parsed),
        (6, 'step5_gifts' in parsed),
        (7, 'step6_residuary' in parsed),
        (8, 'step7_trust' in parsed),
        (9, 'step8_other_matters' in parsed),
    ]:
        if has:
            completed.add(n)
    session['completed_steps'] = sorted(completed)

    session['will_title'] = f"Imported: {file.filename}"
    session.modified = True

    # Auto-save to DB
    save_will_to_db()

    return jsonify({'ok': True})


# -- Client Chat (per-client AI inbox) --------------------------------------

def _get_or_create_chat_session(client_id, user_id):
    """One ChatSession per client (the most recent). Create if missing."""
    cs = (ChatSession.query
          .filter_by(client_id=client_id)
          .order_by(ChatSession.created_at.desc())
          .first())
    if cs:
        return cs
    cs = ChatSession(client_id=client_id, created_by=user_id)
    db.session.add(cs)
    db.session.flush()
    return cs


def _get_or_create_active_will(client_id, user_id):
    """Most recently updated non-deleted draft Will for this client, or new one."""
    will = (Will.query
            .filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc())
            .first())
    if will:
        return will
    will = Will(client_id=client_id, created_by=user_id)
    db.session.add(will)
    db.session.flush()
    return will


def _normalise_gifts(gifts: list) -> list:
    """Ensure every gift in step5_data has a `kind` field that the snapshot
    JS can use to display a meaningful label.

    Two historic save paths produced gifts without `kind`:
      1. Old `_try_save_property_gift` before the `kind` field was added —
         these have `document_id` + optional `property_address`/`title_number`.
      2. Wizard-format gifts using `gift_type` instead of `kind`.

    We normalise in place so the JS only needs to check one field.
    """
    if not isinstance(gifts, list):
        return gifts
    result = []
    for g in gifts:
        if not isinstance(g, dict):
            result.append(g)
            continue
        g = dict(g)   # copy so we don't mutate the original
        # --- already has kind ---
        if g.get('kind'):
            result.append(g)
            continue
        # --- wizard format: gift_type → kind ---
        gt = (g.get('gift_type') or '').strip().lower()
        if gt == 'property':
            g['kind'] = 'property'
            # Flatten property_details to top-level so JS checks propAddr etc.
            pd = g.get('property_details') or {}
            for f in ('property_address', 'title_number', 'lot_number',
                      'mukim', 'daerah', 'negeri'):
                if not g.get(f) and pd.get(f):
                    g[f] = pd[f]
            # Flatten allocations → beneficiaries
            if not g.get('beneficiaries') and g.get('allocations'):
                g['beneficiaries'] = [
                    {'name': a.get('beneficiary_name', ''), 'share': a.get('share', '')}
                    for a in g['allocations'] if a.get('beneficiary_name')
                ]
        elif gt in ('financial', 'bank'):
            g['kind'] = 'bank'
            fd = g.get('financial_details') or {}
            if not g.get('bank_name') and fd.get('institution'):
                g['bank_name'] = fd['institution']
            if not g.get('account_number') and fd.get('account_number'):
                g['account_number'] = fd['account_number']
            if not g.get('beneficiaries') and g.get('allocations'):
                g['beneficiaries'] = [
                    {'name': a.get('beneficiary_name', ''), 'share': a.get('share', '')}
                    for a in g['allocations'] if a.get('beneficiary_name')
                ]
        elif gt == 'vehicle':
            g['kind'] = 'vehicle'
        elif gt in ('other', ''):
            g['kind'] = 'other'
        else:
            g['kind'] = gt   # pass through unknown types
        # --- old chat-format: has document_id but no kind → property gift ---
        if not g.get('kind') and g.get('document_id'):
            g['kind'] = 'property'
        result.append(g)
    return result


def _will_data_snapshot(will_record):
    """Parse a Will's step JSON columns into a single dict for the chat planner & UI.
    Also includes the live Person registry from DB so the chat's right-pane
    sees identities as they're added in chat (they don't live in stepN_data
    until the wizard saves)."""
    if not will_record:
        return {}
    def _j(s, default):
        try:
            return json.loads(s) if s else default
        except (json.JSONDecodeError, TypeError):
            return default
    # Live identities from Person table — these are the single source of truth
    # while the chat is running, before the wizard's save_will_to_db serializes
    # them into identities_data.
    persons = (Person.query.filter_by(client_id=will_record.client_id)
               .order_by(Person.full_name.asc()).all())
    identities = [{
        'id': p.id, 'full_name': p.full_name,
        'nric_passport': p.nric_passport, 'relationship': p.relationship or '',
        'date_of_birth': p.date_of_birth or '',
        'document_id': p.document_id or '',
        # 🔥 §10x.123 — include address so the Step 2 testator card can
        # surface family-member addresses as one-click shortcuts.
        'address': p.address or '',
        'gender': getattr(p, 'gender', '') or '',
    } for p in persons]
    # `completed_steps` carries chat-flow markers like 'assets_confirmed'
    # that the planner uses to gate ordering — collect ALL gifts before
    # asking who inherits, just like we collect ALL identities before
    # asking who's the executor.
    completed = _j(will_record.completed_steps, [])
    if not isinstance(completed, list):
        completed = []
    return {
        'will_id': will_record.id,
        'client_id': will_record.client_id,
        'title': will_record.title,
        'status': will_record.status,
        'step1': _j(will_record.step1_data, {}),
        'step2': _j(will_record.step2_data, {}),
        'step3': _j(will_record.step3_data, {}),
        'step4': _j(will_record.step4_data, []),
        'step5': _normalise_gifts(_j(will_record.step5_data, [])),
        'step6': _j(will_record.step6_data, {}),
        'step7': _j(will_record.step7_data, {}),
        'step8': _j(will_record.step8_data, {}),
        'identities': identities,
        'completed_steps': completed,
        'current_stage_num': _current_stage_num(will_record.client_id, will_record),
    }


def _serialise_chat_message(m):
    """Turn a ChatMessage row into a JSON-friendly dict for the UI."""
    def _j(s, default):
        try:
            return json.loads(s) if s else default
        except (json.JSONDecodeError, TypeError):
            return default

    attachment_ids = _j(m.attachments_json, [])
    attachments = []
    if attachment_ids:
        docs = Document.query.filter(Document.id.in_(attachment_ids)).all()
        doc_by_id = {d.id: d for d in docs}
        for did in attachment_ids:
            d = doc_by_id.get(did)
            if not d:
                continue
            try:
                _dex = json.loads(d.extracted_data) if d.extracted_data else {}
            except (json.JSONDecodeError, TypeError):
                _dex = {}
            _manual_review = bool(_dex.get('_manual_review'))
            _name_mismatch = (
                _dex.get('_name_match') is False   # explicitly False, not None
            )
            _ic_mismatch = (
                _dex.get('_ic_match') is False
            )
            _suspected = bool(
                _dex.get('_wrong_upload_suspected')
                or _dex.get('_likely_irrelevant')
                or d.category in ('death_certificate', 'unrelated')
                or _name_mismatch
                or _ic_mismatch
            )
            _wrong_reason = (
                _dex.get('_wrong_reason')
                or _dex.get('_irrelevant_reason')
                or ('Death certificate — not a will asset' if d.category == 'death_certificate' else '')
                or ('Unrelated document' if d.category == 'unrelated' else '')
                or ('Owner name does not match testator' if _name_mismatch else '')
                or ('IC number does not match testator' if _ic_mismatch else '')
            ).strip()[:200]
            attachments.append({
                'id': d.id,
                'filename': d.original_filename,
                'category': d.category,
                'size': d.file_size,
                'purpose': (_dex.get('purpose') or '').strip()[:120],
                'address': (
                    _dex.get('address') or _dex.get('property_address') or ''
                ).strip()[:80],
                'lot_number':   (_dex.get('lot_number') or '').strip()[:60],
                'owner_name':   (_dex.get('owner_name') or '').strip()[:100],
                'name_match':   _dex.get('_name_match'),   # True/False/None
                'ic_match':     _dex.get('_ic_match'),     # True/False/None
                'manual_review': _manual_review,
                'suspected_wrong': _suspected,
                'wrong_reason': _wrong_reason,
            })
    # 🔥 §10x.77 — Strip MACHINE-LANGUAGE markers from user-facing content.
    # The DB content can contain internal markers used by backend logic
    # (cache hashes, dedup tags, planner pivots). They MUST never reach
    # the user UI. EXCEPTION: <!--quickreplies:[...]--> — chat.js parses
    # this to render the action buttons, so it stays.
    _content = m.content or ''
    if _content:
        # Remove every HTML comment EXCEPT quickreplies marker
        def _scrub(text):
            out = []
            i = 0
            while i < len(text):
                start = text.find('<!--', i)
                if start < 0:
                    out.append(text[i:]); break
                out.append(text[i:start])
                end = text.find('-->', start)
                if end < 0:
                    out.append(text[start:]); break
                comment = text[start:end + 3]
                if 'quickreplies:' in comment:
                    out.append(comment)   # keep — chat.js needs it
                # else: drop silently
                i = end + 3
            return ''.join(out)
        _content = _scrub(_content)
        # Tidy any double-blank lines the scrub left behind
        import re as _re_md
        _content = _re_md.sub(r'\n{3,}', '\n\n', _content).rstrip() + '\n'
    return {
        'id': m.id,
        'role': m.role,
        'content': _content,
        'attachments': attachments,
        'clarifying_questions': _j(m.clarifying_questions_json, []),
        'proposed_patch': _j(m.proposed_patch_json, None),
        'advice': _j(m.advice_json, []),
        'applied_at': m.applied_at.isoformat() if m.applied_at else None,
        'rejected_at': m.rejected_at.isoformat() if m.rejected_at else None,
        'created_at': m.created_at.isoformat() if m.created_at else None,
    }


@app.route('/chat/<client_id>')
@login_required
def chat_page(client_id):
    """Render the per-client AI chat page."""
    from services.inbound_address import address_for_client
    client = db.session.get(Client, client_id)
    if not client:
        flash('Client not found.', 'error')
        return redirect(url_for('clients_list'))
    # 🔥 §10x.91 — Pin session to THIS client + their active draft will so
    # the 'Open Wizard' button in the chat header navigates to the right
    # wizard. Without this, session.will_id was either stale (from the
    # last client's wizard) or None, and _refresh_wizard_session_from_db
    # returned early → wizard rendered empty form fields even when the
    # chat had populated step1-step5_data via the Step 1 walkthrough +
    # §10x.42 reconcile.
    session['client_id'] = client.id
    active_will = (Will.query
                   .filter_by(client_id=client.id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc())
                   .first())
    if active_will:
        session['will_id'] = active_will.id
    # MX record is on will.alantanjb.com directly — no subdomain needed
    host = request.host.split(':')[0] if request else 'localhost'
    inbox_address = address_for_client(client, host)
    inbox_enabled = bool(os.environ.get('POSTMARK_INBOUND_USER') and os.environ.get('POSTMARK_INBOUND_PASS'))
    return render_template('chat.html', client=client,
                           inbox_address=inbox_address,
                           inbox_enabled=inbox_enabled)


@app.route('/api/chat/<client_id>/history')
@login_required
def api_chat_history(client_id):
    """Return all messages in the client's chat session + current will snapshot.

    Side effect: if any inbound user message has attachments still in
    `chat_inbox` (vision-classify never finished — usually because the
    background thread died on a redeploy), kick off processing again.
    The thread is daemon-safe and idempotent (already-classified docs are
    skipped). Polling every 5s on the client side means the user will see
    the spinner clear within seconds of the next vision call returning.
    """
    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'ok': False, 'error': 'Client not found'}), 404

    cs = (ChatSession.query
          .filter_by(client_id=client_id)
          .order_by(ChatSession.created_at.desc())
          .first())
    messages = []
    if cs:
        for m in cs.messages:
            messages.append(_serialise_chat_message(m))

    # ── Watchdog: resume stuck inbound processing ────────────────────
    # 🔥 BURN-IN §10x.9 — DO NOT RE-FIRE WHILE ANOTHER THREAD IS WORKING
    # The chat polls /history every 5 s. Without these guards we re-spawn
    # the processor 12 times per minute → 12 duplicate intake/summary
    # cards. Three layers of defence:
    #   (1) skip if processor lock is held for this user_msg (in-flight)
    #   (2) skip if no docs are still chat_inbox (nothing to do)
    #   (3) skip if intake card was already posted (work is done)
    try:
        import threading as _t
        _now = datetime.utcnow()
        for _m in (cs.messages if cs else []):
            if _m.role != 'user':
                continue
            _age = (_now - _m.created_at).total_seconds() if _m.created_at else 0
            if _age < 60:
                continue
            try:
                _ids = json.loads(_m.attachments_json or '[]')
            except Exception:
                _ids = []
            if not _ids:
                continue
            _stuck = (Document.query
                      .filter(Document.id.in_(_ids))
                      .filter(Document.category == 'chat_inbox')
                      .count())
            if _stuck == 0:
                # 🔥 §10x.75 — Race-condition recovery for the §10x.53
                # "Analysis complete — ▶️ Start verify identities" card.
                # The post-classification gate inside the inbound processor
                # only runs while the processor is alive. When all docs
                # transition out of chat_inbox (e.g. last image hits
                # §10x.26 retry terminal state AFTER the processor's
                # final check), no one re-runs the gate and the user is
                # stuck with no action button. The watchdog now closes
                # this loop: if all docs are out of chat_inbox AND no
                # "Analysis complete" card exists yet AND there are
                # pending ICs to verify, post the card here.
                try:
                    _ready_exists = (ChatMessage.query
                                     .filter_by(session_id=cs.id, role='assistant')
                                     .filter(ChatMessage.created_at >= _m.created_at)
                                     .filter(ChatMessage.content.ilike('%Analysis complete%'))
                                     .first())
                    if not _ready_exists:
                        from services.identity_walker import get_pending_ic_documents
                        if get_pending_ic_documents(client_id):
                            _qr = json.dumps([
                                {'label': '▶️ Start — verify identities',
                                 'value': 'inbox start'}
                            ])
                            db.session.add(ChatMessage(
                                session_id=cs.id, role='assistant',
                                content=(
                                    "✅ **Analysis complete.** All exhibits "
                                    "classified — ready to verify identities."
                                    f"\n\n<!--quickreplies:{_qr}-->"
                                ),
                                attachments_json='[]',
                            ))
                            db.session.commit()
                except Exception:
                    try:
                        db.session.rollback()
                    except Exception:
                        pass
                continue   # all classified — nothing to resume
            # (1) lock held = another thread already working
            with _PROCESSING_LOCK:
                if _m.id in _PROCESSING_INFLIGHT:
                    continue
            # 🔥 §10x.29 — DO NOT block re-firing on "intake card exists".
            # The earlier check (3) blocked the watchdog whenever an intake
            # card had been posted, which left stuck `chat_inbox` docs
            # un-retried forever. The §10x.26 retry-counter promotes docs
            # to `needs_review` after 3 attempts, so re-firing is bounded.
            # Card-duplication is prevented INSIDE the processor (at the
            # `_intake_already_posted` check before posting), not here.
            _t.Thread(
                target=_process_inbound_message_async,
                args=(app, _m.id),
                daemon=True,
            ).start()
    except Exception:
        pass   # watchdog is best-effort; never block history load

    active_will = (Will.query
                   .filter_by(client_id=client_id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc())
                   .first())
    snapshot = _will_data_snapshot(active_will)
    return jsonify({'ok': True, 'messages': messages, 'will': snapshot})


@app.route('/api/chat/<client_id>/message', methods=['POST'])
@login_required
def api_chat_message(client_id):
    """Receive a chat message (text + optional file uploads) and return the assistant reply."""
    try:
        return _api_chat_message_impl(client_id)
    except Exception as _top_err:
        import traceback as _tb
        app.logger.error(f'api_chat_message unhandled: {_tb.format_exc()}')
        try:
            db.session.rollback()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': f'Server error: {str(_top_err)[:300]}'}), 500


def _api_chat_message_impl(client_id):
    """Inner implementation — wrapped by api_chat_message for JSON error handling."""
    from uploads import save_uploaded_file
    from ai.file_classifier import classify_file
    from ai.ocr import extract_nric_data
    from ai.chat_planner import plan_turn
    from ai.voice_transcription import is_audio, transcribe

    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'ok': False, 'error': 'Client not found'}), 404

    user_text = (request.form.get('text') or '').strip()
    # §10x.77 — `intent` is the machine token from a quickreply click
    # (e.g. "inbox start"). `text` is the user-visible label
    # (e.g. "▶️ Start — verify identities"). The planner needs to see
    # the intent so its existing pattern matches still fire; the user
    # bubble shows the label so machine language never reaches the UI.
    user_intent = (request.form.get('intent') or '').strip()
    # If the planner downstream uses pattern matching, blend the intent
    # into a single string it can read alongside (or instead of) the
    # display text. We preserve `user_text` as the bubble content and
    # use `user_intent` as the planner's pivot signal.
    files = request.files.getlist('files') if 'files' in request.files else []

    if not user_text and not files:
        return jsonify({'ok': False, 'error': 'Empty message'}), 400

    user_id = session.get('user_id')
    cs = _get_or_create_chat_session(client_id, user_id)

    # Establish cost-tracking context for every Anthropic call made in this
    # request: will_id, client_id, user_id auto-attach to ApiCallLog rows.
    _will_for_ctx = (Will.query
                     .filter_by(client_id=client_id, status='draft')
                     .filter(Will.deleted_at.is_(None))
                     .order_by(Will.updated_at.desc())
                     .first())
    _cost_ctx_ids = {
        'client_id': client_id,
        'will_id': _will_for_ctx.id if _will_for_ctx else None,
        'user_id': user_id,
    }
    try:
        from ai.cost_tracker import track_context as _track_ctx
        _cost_tracker_cm = _track_ctx(**_cost_ctx_ids)
        _cost_tracker_cm.__enter__()
    except Exception:
        _cost_tracker_cm = None

    # 1. Persist the user message first so attachments can FK to it.
    # §10x.77 — content stays as the FRIENDLY user_text (label); from
    # this point on, swap user_text → intent for planner branching, so
    # existing pattern matches like 'inbox start' / 'yes' / 'skip'
    # continue to fire even though the bubble shows the friendly label.
    user_msg = ChatMessage(
        session_id=cs.id, role='user', content=user_text,
        attachments_json='[]',
    )
    db.session.add(user_msg)
    db.session.flush()
    if user_intent:
        user_text = user_intent

    folder_name = client.folder_name
    attachment_ids = []
    artifacts = []
    file_errors = []
    voice_transcripts = []  # collected to merge into user_text

    for f in files:
        if not f or not f.filename:
            continue
        # Audio files go to documents/voice/, get transcribed, and bypass the
        # vision classifier (Whisper handles them, file_classifier would just
        # error trying to send them as images).
        is_voice = is_audio(f.content_type or '') or is_audio(f.filename)
        category_initial = 'voice' if is_voice else 'chat_inbox'
        try:
            saved_name, rel_path, file_size = save_uploaded_file(
                f, client_id, category=category_initial, folder_name=folder_name)
        except ValueError as e:
            file_errors.append(f"{f.filename}: {e}")
            continue

        doc = Document(
            client_id=client_id,
            chat_message_id=user_msg.id,
            filename=saved_name, original_filename=f.filename,
            file_path=rel_path, file_type=f.content_type,
            file_size=file_size, category=category_initial,
        )
        db.session.add(doc)
        db.session.flush()
        attachment_ids.append(doc.id)

        abs_path = os.path.join(UPLOAD_DIR, rel_path)

        if is_voice:
            transcript = transcribe(abs_path)
            if transcript:
                voice_transcripts.append(transcript)
                doc.description = (transcript[:500] or None)
                try:
                    doc.extracted_data = json.dumps({'transcript': transcript})
                except (TypeError, ValueError):
                    pass
            else:
                doc.description = '(transcription failed or unavailable)'
            artifacts.append({
                'document_id': doc.id, 'kind': 'voice',
                'confidence': 'high' if transcript else 'low',
                'extracted': {'transcript': transcript} if transcript else None,
                'original_filename': f.filename,
            })
            continue  # don't run vision classifier on audio

        # 2. Classify (vision)
        classification = classify_file(abs_path)
        kind = classification.get('kind', 'other')

        # 3. Extract per kind
        extracted = None
        try:
            if kind == 'nric':
                extracted = extract_nric_data(abs_path)
            elif kind in ('property_title', 'property_spa', 'property_tax',
                           'property_transfer'):
                # Full property OCR for all four — Borang 14A/16A contains
                # lot_number and title_number which the extractor can read.
                # The KIND tells the downstream walker whether this counts
                # as registered ownership evidence (only property_title does).
                from ai.property_extractor import extract_property_data
                extracted = extract_property_data(abs_path, doc_type='general')
            elif kind in ('utility_bill', 'bank_letter'):
                # Light-touch: no per-field extraction yet — the classifier
                # has already given us `purpose` + `property_hint` which is
                # enough to cluster under a property and show in chat.
                extracted = {}
            elif kind == 'bank_statement':
                from ai.ocr import extract_asset_document
                extracted = extract_asset_document(abs_path, asset_type='bank')
            elif kind == 'vehicle':
                from ai.ocr import extract_asset_document
                extracted = extract_asset_document(abs_path, asset_type='vehicle')
            elif kind in ('insurance', 'epf_kwsp'):
                from ai.ocr import extract_asset_document
                extracted = extract_asset_document(abs_path, asset_type='other')
        except Exception as e:
            extracted = {'error': str(e)}

        # 4. Update Document with classification result
        doc.category = kind if kind != 'other' else 'chat_inbox'
        doc.description = classification.get('reason', '')[:500] if classification.get('reason') else None
        # Persist per-image `purpose`, `property_hint`, `custom_type`,
        # and `person_name` so the chat planner can surface and group.
        purpose      = (classification.get('purpose')      or '').strip()
        prop_hint    = (classification.get('property_hint') or '').strip()
        custom_type  = (classification.get('custom_type')  or '').strip()
        person_name  = (classification.get('person_name')  or '').strip()
        will_relevant = classification.get('will_relevant', True)
        if extracted is None:
            extracted = {}
        if purpose:
            extracted['purpose'] = purpose[:300]
        if prop_hint:
            extracted['property_hint'] = prop_hint[:300]
        if custom_type:
            extracted['custom_type'] = custom_type[:200]
        if person_name:
            extracted['_doc_person_name'] = person_name[:200]

        # ── Wrongly-uploaded detection ────────────────────────────────────
        # death_certificate and unrelated docs are almost always wrong uploads.
        # For any doc with a named person, also compare against testator name.
        if kind in ('death_certificate', 'unrelated'):
            extracted['_wrong_upload_suspected'] = True
            extracted['_wrong_reason'] = (
                'Death certificate detected — likely uploaded by mistake.' if kind == 'death_certificate'
                else 'Document classified as unrelated to asset ownership.'
            )
        elif person_name:
            # Compare against testator name to catch docs about third parties
            try:
                _testator_name = ''
                if _will_for_ctx:
                    _s1 = json.loads(_will_for_ctx.step1_data or '{}')
                    _testator_name = (_s1.get('full_name') or '').strip().upper()
                if not _testator_name:
                    _testator_name = (client.full_name or '').strip().upper()
                if _testator_name and person_name.upper() != _testator_name:
                    # Simple fuzzy check: if NO word from doc person_name appears
                    # in testator name, flag it (avoids false positive on aliases)
                    _doc_words = set(person_name.upper().split())
                    _testator_words = set(_testator_name.split())
                    if not _doc_words.intersection(_testator_words):
                        extracted['_wrong_upload_suspected'] = True
                        extracted['_wrong_reason'] = (
                            f'Document appears to be about "{person_name}", '
                            f'not the testator. Possibly uploaded by mistake.'
                        )
            except Exception:
                pass
        # ── Message context ──────────────────────────────────────────────
        # Store the text the client sent WITH this file. This is the single
        # most reliable context clue — the client typed "my house at Lot
        # 127082, give to Joshua" in the same WhatsApp/email message as
        # they attached the geran. We use it later to back-fill missing
        # OCR fields and to surface intent on the property card.
        # For direct chat uploads we also look at the previous user message
        # (back-to-back: text THEN image is common in WhatsApp-style flows).
        try:
            wa_ctx = _extract_whatsapp_context_for_file(
                user_text or '', doc.original_filename or ''
            )
            if wa_ctx:
                extracted['_message_context'] = wa_ctx[:800]
                extracted['_context_source'] = 'whatsapp_preceding'
            else:
                msg_context_parts = []
                if user_text:
                    msg_context_parts.append(user_text)
                prev_msgs = (ChatMessage.query
                             .filter_by(session_id=cs.id, role='user')
                             .filter(ChatMessage.id != user_msg.id)
                             .order_by(ChatMessage.created_at.desc())
                             .limit(3).all())
                for pm in prev_msgs:
                    txt = (pm.content or '').strip()
                    if txt:
                        msg_context_parts.append(txt)
                if msg_context_parts:
                    extracted['_message_context'] = '\n'.join(msg_context_parts)[:800]
            # WhatsApp timestamp for this image (CLAUDE.md §10i)
            wa_ts = _extract_whatsapp_timestamp_for_file(
                user_text or '', doc.original_filename or ''
            )
            if wa_ts:
                extracted['_msg_timestamp'] = wa_ts
        except Exception:
            pass
        # For low-confidence 'other' docs: cross-check recent chat messages
        # to see if the client mentioned this image (by filename keyword or
        # asset reference). If no mention found, flag as likely irrelevant.
        if kind == 'other' and not will_relevant:
            try:
                recent = _gather_recent_chat_text(client_id)
                fname_stem = doc.original_filename.rsplit('.', 1)[0].lower()
                # Rough check: any asset keywords in context?
                asset_keywords = ('property', 'house', 'land', 'lot', 'geran',
                                  'bank', 'account', 'insurance', 'car', 'vehicle',
                                  'lot', 'mukim', 'title', fname_stem)
                if not any(kw in (recent or '').lower() for kw in asset_keywords):
                    extracted['_likely_irrelevant'] = True
                    extracted['_irrelevant_reason'] = (
                        'Classified as unrecognised document type and no matching '
                        'asset was mentioned in the chat.'
                    )
            except Exception:
                pass
        if extracted:
            try:
                doc.extracted_data = json.dumps(extracted)
            except (TypeError, ValueError):
                doc.extracted_data = None

        # Dedupe: if this IC's extracted name already exists on another
        # nric Document for this client, mark the new one as duplicate
        # so it never enters the walk-through pool.
        is_duplicate = False
        if kind == 'nric' and extracted and not extracted.get('error'):
            is_duplicate = _dedupe_ic_against_existing(client_id, doc, extracted)

        if is_duplicate:
            continue  # don't emit as an artifact — it's a duplicate

        artifacts.append({
            'document_id': doc.id,
            'kind': kind,
            'confidence': classification.get('confidence', 'low'),
            'extracted': extracted,
            'original_filename': f.filename,
        })

    # Merge voice transcripts into the user-visible message text so the planner
    # treats spoken instructions the same as typed ones.
    if voice_transcripts:
        joined = '\n\n'.join(voice_transcripts)
        if user_text:
            user_text = f"{user_text}\n\n_(voice)_ {joined}"
        else:
            user_text = f"_(voice)_ {joined}"
        user_msg.content = user_text

    user_msg.attachments_json = json.dumps(attachment_ids)
    db.session.commit()  # persist user_msg before any deductions / planner

    # 5. Directed flow: try to assign the next pending IC if user replied
    #    with a relationship keyword OR confirmation. If not, see if they
    #    asked to delete the focused doc. Then refresh the pending list so
    #    the planner asks about the NEXT one.
    # Q&A digression: if user asked a side-quest question, answer it as a
    # SEPARATE assistant message. By default we SHORT-CIRCUIT and return
    # only the Q&A — the answer already includes a "↩ Resume <step>" quick
    # reply so the writer can come back. The planner only also runs if the
    # turn ALSO carried files OR the text starts with an action token (yes
    # / skip / confirm / a beneficiary share like "Joshua 50%") — in those
    # cases the user mixed a question with an action and both need handling.
    qa_msg = None
    if user_text:
        try:
            from ai.legal_qa import is_question, answer_question
            looks_like_question = is_question(user_text)
        except Exception as e:
            try:
                import logging
                logging.getLogger(__name__).warning(
                    "is_question failed: %s", e, exc_info=True)
            except Exception:
                pass
            looks_like_question = False
        if looks_like_question:
            current_will = (Will.query.filter_by(client_id=client_id, status='draft')
                            .filter(Will.deleted_at.is_(None))
                            .order_by(Will.updated_at.desc()).first())
            try:
                stage = _current_stage_label(client_id, current_will)
            except Exception:
                stage = ''
            try:
                ans = answer_question(user_text, stage, client_id=client_id,
                                      user_id=session.get('user_id'))
            except Exception as e:
                try:
                    import logging
                    logging.getLogger(__name__).warning(
                        "answer_question crashed: %s", e, exc_info=True)
                except Exception:
                    pass
                # Never let the Q&A path fail silently — emit a fallback so
                # the user always sees SOMETHING acknowledging their question.
                ans = (f"**Answer:** I couldn't reach the legal-Q&A engine just "
                       f"now — please retry in a moment.\n\n_(error: {str(e)[:100]})_")
            if not ans:
                ans = ("**Answer:** I couldn't generate an answer for that — "
                       "try rephrasing, or check the [Legal Library](/library).")
            qa_msg = ChatMessage(session_id=cs.id, role='assistant', content=ans,
                                 attachments_json='[]')
            db.session.add(qa_msg)
            db.session.commit()

            # Decide: should the planner ALSO run after the Q&A?
            #
            # Default: NO — short-circuit. The Q&A answer already includes a
            # "↩ Resume" quick-reply, and re-asking the current step on top
            # of the Q&A buries the answer in noise (this is the bug the
            # writer reported: "no response" — actually the Q&A was there
            # but a giant asset card was rendered after it).
            #
            # Run the planner ONLY when the turn also looks like an action:
            #   - files attached this turn
            #   - text starts with a confirmation/skip/delete token
            #   - text contains an explicit share assignment ("Joshua 50%")
            _kw = user_text.strip().lower()
            _action_starts = ('yes', 'skip', 'delete', 'no', 'confirm',
                              'remove', 'inventory ', 'unlink ',
                              'guardian ', 'trust ', 'trust yes', 'trust skip',
                              'others ', 'residuary skip', 'property fill',
                              'property ', 'address:', 'daerah:', 'negeri:',
                              'mukim:', 'lot:', 'title:',
                              'change ', 'confirm defaults')
            has_share = bool(re.search(r'\b\d{1,3}\s*%|\bequal\b|\b\d+/\d+\b', _kw))
            also_action = (
                bool(files) or bool(artifacts)
                or any(_kw == t.strip() or _kw.startswith(t) for t in _action_starts)
                or has_share
            )
            if not also_action:
                return jsonify({
                    'ok': True,
                    'user_message': _serialise_chat_message(user_msg),
                    'assistant_message': _serialise_chat_message(qa_msg),
                })

    just_assigned = _try_assign_pending_identity(client_id, user_text)
    if not just_assigned:
        just_assigned = _try_skip_pending_identity(client_id, user_text)
    just_deleted = None if just_assigned else _try_delete_pending_identity(client_id, user_text)
    # If user replied "confirm assets" / "i have more to upload" at the
    # asset-inventory gate, stamp the marker (or clear it) so the planner
    # advances past the gate to the per-asset assignment step.
    just_assets_gate = None
    just_inventory = None
    if not just_assigned and not just_deleted:
        # Walk-through actions take priority — the writer is reviewing
        # one specific asset card, so 'inventory confirm' / 'inventory
        # skip' / 'inventory unlink' must hit the per-asset handler
        # before the gate fallback.
        just_inventory = (_try_handle_inbox_action(client_id, user_text)
                          or _try_handle_restart_inbox(client_id, user_text)
                          or _try_handle_restart_gifts(client_id, user_text)
                          or _try_handle_unlink_action(client_id, user_text)
                          # §10x.127 — inventory-completeness gate (must
                          # match the `assets_check ` prefix BEFORE the
                          # general inventory handler that owns
                          # `inventory `)
                          or _try_handle_assets_check(client_id, user_text)
                          or _try_handle_inventory_action(client_id, user_text)
                          # §10hg — H3 placeholder confirm/skip when no pending image
                          or _try_handle_h3_user_match(client_id, user_text)
                          or _try_handle_orphan_claim(client_id, user_text)
                          or _try_handle_doc_assign(client_id, user_text)
                          or _try_handle_h3_property_action(client_id, user_text)
                          # §10hg — conflict resolve replies
                          or _try_handle_message_conflict(client_id, user_text)
                          or _try_handle_property_fill(client_id, user_text)
                          or _try_handle_ownership(client_id, user_text)
                          or _try_handle_encumbrance(client_id, user_text)
                          # §10x.127 — free-text description while in
                          # 'describing' mode. MUST run last so quickreply
                          # paths above get first chance.
                          or _try_handle_assets_describe(client_id, user_text))
        if not just_inventory:
            just_assets_gate = _try_handle_assets_gate(client_id, user_text)
    # If past Step 1, attempt executor save then beneficiaries save then
    # gift-delete then gift-save. Delete BEFORE save so a "delete" reply
    # at a Step-6 property card removes the doc instead of trying to parse
    # it as beneficiary names (which always fails).
    just_executor = None
    just_benef = None
    just_gift_deleted = None
    just_gift = None
    just_guardian = None
    just_trust = None
    just_others = None
    just_residuary_skip = None
    just_testator = None
    if not just_assigned and not just_deleted:
        from services.identity_walker import get_pending_ic_documents as _gpid
        if not _gpid(client_id):  # Step 1 done
            # 🔥 §7 — Step 2 (Testator confirm) runs BEFORE Step 3 (Executor)
            # and BEFORE Step 6 (Specific Gifts). Catch user's confirm
            # click and save Will.step1_data with testator info from Person.
            # 🔥 §10x.122 — testator address save MUST run before
            # _try_confirm_testator (so 'address: 123 Main St' replies
            # save the address) AND before _try_handle_property_fill
            # (so they aren't mistakenly treated as property addresses).
            just_testator = _try_save_testator_address(client_id, user_text)
            if not just_testator:
                just_testator = _try_confirm_testator(client_id, user_text)
            if not just_testator:
                just_executor = _try_save_executor(client_id, user_text)
            if not just_executor:
                just_guardian = _try_handle_guardian_action(client_id, user_text)
                if not just_guardian:
                    just_trust = _try_handle_trust_action(client_id, user_text)
                    if not just_trust:
                        just_others = _try_handle_others_action(client_id, user_text)
                        if not just_others:
                            # 🔥 §10x.116 — Layer 3 substitute residuary FIRST
                            # (so 'residuary substitute ...' isn't accidentally
                            # captured by main). Then main, then skip.
                            just_residuary_skip = _try_save_residuary_substitute(client_id, user_text)
                            if not just_residuary_skip:
                                # 🔥 §10x.114 — Layer 2 main residuary beneficiary
                                just_residuary_skip = _try_save_residuary_main(client_id, user_text)
                            if not just_residuary_skip:
                                just_residuary_skip = _try_handle_residuary_skip(client_id, user_text)
                            if not just_residuary_skip:
                                # 🔥 §10x.115 — explicit confirmation of auto-
                                # populated step4_data via 'beneficiaries confirm'.
                                # Run BEFORE the legacy save handler.
                                just_benef = _try_handle_beneficiaries_confirm(client_id, user_text)
                                if not just_benef:
                                    just_benef = _try_save_beneficiaries(client_id, user_text)
                                if not just_benef:
                                    # §10x.18 mismatch handler runs FIRST so 'mismatch ...'
                                    # quickreplies don't fall to other handlers.
                                    just_mm = _try_handle_mismatch(client_id, user_text)
                                    # §10x.21 role-match handler runs next so its
                                    # 'role_match confirm <id>' quickreplies route correctly.
                                    just_role = (None if just_mm
                                                  else _try_handle_role_match(client_id, user_text))
                                    just_gift_deleted = (None if just_role
                                                          else _try_delete_pending_gift(client_id, user_text))
                                    if not just_role and not just_gift_deleted:
                                        # §10x.23 layered handlers run FIRST so 'bank_l1/l2/l3'
                                        # and 'insurance_l1/l2/l3' route correctly.
                                        just_gift = (_try_save_bank_layered_gift(client_id, user_text)
                                                     or _try_save_insurance_layered_gift(client_id, user_text)
                                                     or _try_save_bank_h3_gift(client_id, user_text)
                                                     or _try_save_insurance_h3_gift(client_id, user_text))
                                        if not just_gift:
                                            just_gift = _try_save_property_gift(client_id, user_text)
                                        if not just_gift:
                                            # Bank-account question handler (FUCK-13).
                                            # Property handler runs first because the bank
                                            # question only fires when no pending props remain.
                                            just_gift = _try_save_bank_gift(client_id, user_text)
    from services.identity_walker import get_pending_ic_documents
    from services.gift_walker import get_pending_gift_documents
    pending_ics = get_pending_ic_documents(client_id)
    pending_gifts = get_pending_gift_documents(client_id)
    recent_text = _gather_recent_chat_text(client_id)

    # Persist any address/field enrichment found by reverse-lookup so that
    # matched addresses survive page refreshes and appear in the wizard.
    # Runs async-safe (best-effort, never raises).
    _persist_property_enrichment(client_id, recent_text)

    # 6. Plan the assistant turn against the current Will state
    active_will = (Will.query
                   .filter_by(client_id=client_id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc())
                   .first())
    will_snapshot = _will_data_snapshot(active_will)
    # Treat any save as "just_assigned" so the planner acknowledges + advances
    just = (just_assigned or just_testator or just_executor or just_benef
            or just_gift_deleted or just_gift or just_assets_gate
            or just_inventory or just_guardian or just_trust
            or just_others or just_residuary_skip)
    will_snapshot['pending_gifts'] = pending_gifts
    will_snapshot['layer2_pending_props'] = _get_layer2_pending_props(client_id)
    # If a property_fill action produced a reply_override (e.g. the "how to
    # type missing fields" prompt), inject it into the plan instead of running
    # the normal planner — it's a simple instructional message, not a full turn.
    # Any inventory action that sets reply_override wants to replace the
    # planner's normal turn (address gate, ownership gate, encumbrance gate,
    # gifts restart, etc.). The presence of reply_override is the selector —
    # only gate/fill results set it; normal confirm/skip results do not.
    _fill_override = (isinstance(just_inventory, dict)
                      and bool(just_inventory.get('reply_override')))
    try:
        plan = plan_turn(user_text, artifacts, will_snapshot,
                         pending_ics=pending_ics, recent_text=recent_text,
                         just_assigned=just, just_deleted=just_deleted)
    except Exception as _plan_err:
        import traceback as _tb
        _err_detail = _tb.format_exc()
        app.logger.error(f'plan_turn crashed: {_err_detail}')
        # If we have a fill override (e.g. restart gifts reply), still send it
        if _fill_override:
            plan = {'reply': just_inventory['reply_override'], 'proposed_patch': None,
                    'clarifying_questions': [], 'advice': [], 'focus_attachments': []}
        else:
            return jsonify({'ok': False, 'error': f'Planner error: {_err_detail[:400]}'}), 500
    if _fill_override:
        plan['reply'] = just_inventory['reply_override']
        # Propagate focus_attachments from the inventory action if provided
        # (e.g. inbox reset attaches all doc IDs so thumbnails render in the reply)
        if just_inventory.get('focus_attachments'):
            plan['focus_attachments'] = just_inventory['focus_attachments']

    if file_errors:
        plan['reply'] = (plan.get('reply') or '') + (
            "\n\n**Some files were rejected:**\n- " + "\n- ".join(file_errors)
        )

    # 🔥 §10x.80 — Post the ack ("✅ Saved X as Y") as a SEPARATE chat
    # bubble BEFORE the next walkthrough card. Mixing them in one bubble
    # was confusing — the user couldn't tell the previous action's
    # confirmation apart from the next prompt.
    ack_text = (plan.get('ack_reply') or '').strip()
    if ack_text:
        ack_msg = ChatMessage(
            session_id=cs.id,
            role='assistant',
            content=ack_text,
            attachments_json='[]',
            target_will_id=active_will.id if active_will else None,
        )
        db.session.add(ack_msg)
        db.session.flush()   # ensure created_at strictly precedes the next card

    asst_msg = ChatMessage(
        session_id=cs.id,
        role='assistant',
        content=plan.get('reply', ''),
        attachments_json=json.dumps(plan.get('focus_attachments') or []),
        clarifying_questions_json=json.dumps(plan.get('clarifying_questions', [])),
        proposed_patch_json=json.dumps(plan['proposed_patch']) if plan.get('proposed_patch') else None,
        advice_json=json.dumps(plan.get('advice', [])),
        target_will_id=active_will.id if active_will else None,
    )
    db.session.add(asst_msg)
    db.session.commit()

    # If a side-quest Q&A reply was produced earlier in this turn, surface it
    # alongside the planner's reply so the user sees BOTH (the answer + the
    # re-asked step) immediately, without waiting on a history refresh.
    extra = {}
    if qa_msg is not None:
        extra['qa_message'] = _serialise_chat_message(qa_msg)
    try:
        if _cost_tracker_cm is not None:
            _cost_tracker_cm.__exit__(None, None, None)
    except Exception:
        pass
    return jsonify({
        'ok': True,
        'user_message': _serialise_chat_message(user_msg),
        'assistant_message': _serialise_chat_message(asst_msg),
        **extra,
    })


@app.route('/api/chat/<client_id>/apply/<message_id>', methods=['POST'])
@login_required
def api_chat_apply(client_id, message_id):
    """Apply the proposed_patch from an assistant message to the active Will."""
    from services.person_registry import ensure_person

    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'ok': False, 'error': 'Client not found'}), 404

    m = db.session.get(ChatMessage, message_id)
    if not m or m.role != 'assistant':
        return jsonify({'ok': False, 'error': 'Message not found'}), 404
    if m.applied_at:
        return jsonify({'ok': False, 'error': 'Already applied'}), 400
    if m.rejected_at:
        return jsonify({'ok': False, 'error': 'Already rejected'}), 400
    if not m.proposed_patch_json:
        return jsonify({'ok': False, 'error': 'No patch on this message'}), 400

    try:
        patch = json.loads(m.proposed_patch_json)
    except json.JSONDecodeError:
        return jsonify({'ok': False, 'error': 'Patch JSON is malformed'}), 500

    user_id = session.get('user_id')
    will = _get_or_create_active_will(client_id, user_id)

    # 1. Apply person actions first so we have person_ids to embed in step JSON
    person_id_by_role = {}
    for p in patch.pop('_persons', []):
        pid = ensure_person(
            client_id,
            p.get('name', ''),
            nric=p.get('nric', ''),
            address=p.get('address', ''),
            relationship=p.get('role', ''),
            dob=p.get('dob', ''),
            nationality=p.get('nationality', 'Malaysian'),
            document_id=p.get('document_id'),
        )
        if pid:
            person_id_by_role[p.get('role', '')] = pid

    # 2. Merge each step patch into the corresponding step JSON column
    step_attrs = {
        'step1': 'step1_data', 'step2': 'step2_data', 'step3': 'step3_data',
        'step4': 'step4_data', 'step5': 'step5_data', 'step6': 'step6_data',
        'step7': 'step7_data', 'step8': 'step8_data',
    }
    for step_key, attr in step_attrs.items():
        if step_key not in patch:
            continue
        new_partial = patch[step_key]
        try:
            current = json.loads(getattr(will, attr) or '{}')
        except json.JSONDecodeError:
            current = {}
        if isinstance(current, list) or isinstance(new_partial, list):
            # List-shaped steps: replace wholesale only if patch is a list
            merged = new_partial if isinstance(new_partial, list) else current
        else:
            merged = dict(current)
            merged.update(new_partial)
        # Embed Testator person_id if we just created/updated one
        if step_key == 'step1' and person_id_by_role.get('Testator'):
            merged['person_id'] = person_id_by_role['Testator']
        setattr(will, attr, json.dumps(merged))

    # 3. Sync Client header from step1 if testator name/nric were just set
    if 'step1' in patch:
        s1 = json.loads(will.step1_data or '{}')
        if s1.get('full_name'):
            client.full_name = s1['full_name']
        if s1.get('nric_passport'):
            client.nric_passport = s1['nric_passport']
        if s1.get('email'):
            client.email = s1['email']
        if s1.get('phone'):
            client.phone = s1['phone']
        will.title = f"Will of {s1.get('full_name', 'Unknown')}"

    m.applied_at = datetime.utcnow()
    m.applied_by = user_id
    m.target_will_id = will.id
    db.session.commit()

    return jsonify({
        'ok': True,
        'message': _serialise_chat_message(m),
        'will': _will_data_snapshot(will),
    })


@app.route('/api/chat/<client_id>/replan/<message_id>', methods=['POST'])
@login_required
def api_chat_replan(client_id, message_id):
    """Re-run the planner over an existing user message's attachments using
    already-classified Document data (no re-call to Claude vision). Deletes
    any prior assistant replies that followed this message and replaces them
    with a fresh one. Useful after planner-code upgrades."""
    from ai.chat_planner import plan_turn
    user_msg = db.session.get(ChatMessage, message_id)
    if not user_msg or user_msg.role != 'user':
        return jsonify({'ok': False, 'error': 'Not a user message'}), 404
    cs = db.session.get(ChatSession, user_msg.session_id)
    if not cs or cs.client_id != client_id:
        return jsonify({'ok': False, 'error': 'Wrong client'}), 403

    try:
        doc_ids = json.loads(user_msg.attachments_json or '[]')
    except (json.JSONDecodeError, TypeError):
        doc_ids = []
    docs = Document.query.filter(Document.id.in_(doc_ids)).all() if doc_ids else []

    KNOWN_KINDS = {'nric', 'property_title', 'property_spa', 'property_tax',
                   'property_transfer', 'utility_bill', 'bank_letter',
                   'bank_statement', 'insurance', 'epf_kwsp', 'vehicle', 'will', 'voice'}
    artifacts = []
    for doc in docs:
        try:
            extracted = json.loads(doc.extracted_data) if doc.extracted_data else None
        except (json.JSONDecodeError, TypeError):
            extracted = None
        kind = doc.category if doc.category in KNOWN_KINDS else 'other'
        artifacts.append({
            'document_id': doc.id, 'kind': kind,
            'confidence': 'high', 'extracted': extracted,
            'original_filename': doc.original_filename,
        })

    # Delete any assistant replies that came AFTER this user message
    n_deleted = (ChatMessage.query
                 .filter(ChatMessage.session_id == cs.id,
                         ChatMessage.role == 'assistant',
                         ChatMessage.created_at > user_msg.created_at)
                 .delete(synchronize_session=False))

    client = db.session.get(Client, cs.client_id)
    active_will = (Will.query.filter_by(client_id=client.id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc()).first())
    from services.identity_walker import get_pending_ic_documents
    pending_ics = get_pending_ic_documents(client.id)
    recent_text = _gather_recent_chat_text(client.id)
    # On replan we want the planner to skip the "intake" stage — pass
    # empty artifacts so it goes straight to walk-through / next step.
    # 🔥 §10x.103 — wrap with track_context so log_usage in nested
    # pipeline calls (asset_pipeline, web_property_clues, geo_resolver,
    # property_locale_verifier) auto-attaches client_id/will_id. Without
    # this every Claude call inside plan_turn fired a §10x.68 contract
    # violation and cost rows had will_id=None.
    from ai.cost_tracker import track_context as _track_ctx_replan
    with _track_ctx_replan(client_id=client.id,
                            will_id=active_will.id if active_will else None,
                            user_id=session.get('user_id')):
        plan = plan_turn(user_msg.content or '', [], _will_data_snapshot(active_will),
                         pending_ics=pending_ics, recent_text=recent_text)

    asst_msg = ChatMessage(
        session_id=cs.id, role='assistant',
        content=plan.get('reply', ''),
        attachments_json=json.dumps(plan.get('focus_attachments') or []),
        clarifying_questions_json=json.dumps(plan.get('clarifying_questions', [])),
        proposed_patch_json=json.dumps(plan['proposed_patch']) if plan.get('proposed_patch') else None,
        advice_json=json.dumps(plan.get('advice', [])),
        target_will_id=active_will.id if active_will else None,
    )
    db.session.add(asst_msg)
    db.session.commit()
    return jsonify({'ok': True, 'message_id': asst_msg.id, 'replaced_replies': n_deleted})


@app.route('/api/legal-library/upload', methods=['POST'])
@login_required
def api_legal_library_upload():
    """Accept a PDF upload (multipart/form-data, field 'file' + optional 'slug')
    and save under data/legal_acts/<slug>.pdf for the legal_qa retrieval."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    f = request.files['file']
    if not f or not f.filename:
        return jsonify({'ok': False, 'error': 'Empty file'}), 400
    if not f.filename.lower().endswith('.pdf'):
        return jsonify({'ok': False, 'error': 'Only PDF accepted'}), 400
    slug = (request.form.get('slug') or f.filename[:-4]).lower()
    slug = re.sub(r'[^a-z0-9_]+', '_', slug).strip('_')
    if not slug:
        return jsonify({'ok': False, 'error': 'Invalid slug'}), 400
    folder = os.path.join(DATA_DIR, 'legal_acts')
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, f"{slug}.pdf")
    f.save(path)
    size_kb = os.path.getsize(path) // 1024
    return jsonify({'ok': True, 'slug': slug, 'size_kb': size_kb,
                    'path': f'data/legal_acts/{slug}.pdf'})


@app.route('/api/legal-library/list')
@login_required
def api_legal_library_list():
    """List currently-loaded acts."""
    from services.legal_library import list_available_acts
    return jsonify({'ok': True, 'acts': list_available_acts()})


@app.route('/library')
@login_required
def legal_library_page():
    """Admin UI to upload + view loaded Acts. Login required (any role)."""
    from services.legal_library import list_available_acts
    return render_template('legal_library.html', acts=list_available_acts())


@app.route('/library/download/<slug>')
@login_required
def legal_library_download(slug):
    """🔥 §10x.148 — serve a library PDF inline (preview) so users can
    open the gold standard guide / any Act / book directly from the
    library page."""
    safe_slug = re.sub(r'[^a-z0-9_]+', '', (slug or '').lower())
    if not safe_slug:
        return jsonify({'ok': False, 'error': 'Invalid slug'}), 400
    path = os.path.join(DATA_DIR, 'legal_acts', f"{safe_slug}.pdf")
    if not os.path.isfile(path):
        return jsonify({'ok': False, 'error': 'Not found'}), 404
    return send_file(path, mimetype='application/pdf',
                      as_attachment=False,
                      download_name=f"{safe_slug}.pdf")


@app.route('/api/legal-library/delete/<slug>', methods=['POST'])
@login_required
def api_legal_library_delete(slug):
    """Delete a single Act PDF from the library."""
    safe_slug = re.sub(r'[^a-z0-9_]+', '', slug.lower())
    if not safe_slug:
        return jsonify({'ok': False, 'error': 'Invalid slug'}), 400
    path = os.path.join(DATA_DIR, 'legal_acts', f"{safe_slug}.pdf")
    if not os.path.isfile(path):
        return jsonify({'ok': False, 'error': 'Not found'}), 404
    try:
        os.remove(path)
        # Library content changed → cached answers may be stale. Purge.
        try:
            from ai.legal_qa import cache_clear_all
            cache_clear_all()
        except Exception:
            pass
        return jsonify({'ok': True})
    except OSError as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/legal-library/cache/stats')
@login_required
def api_legal_library_cache_stats():
    """Return how many cached answers we have, total hits served, and the
    top 10 most-asked questions. Helps quantify token savings."""
    try:
        from database import LegalQACache
        from sqlalchemy import func as _func
        total = db.session.query(_func.count(LegalQACache.id)).scalar() or 0
        total_hits = db.session.query(_func.coalesce(_func.sum(LegalQACache.hits), 0)).scalar() or 0
        top = (LegalQACache.query
               .order_by(LegalQACache.hits.desc(), LegalQACache.last_used_at.desc())
               .limit(10).all())
        return jsonify({
            'ok': True,
            'cached_questions': int(total),
            'total_cache_hits': int(total_hits),
            'estimated_tokens_saved': int(total_hits) * 1500,  # ~1.5k tokens per LLM call
            'top_questions': [{
                'question': r.question_text[:200],
                'hits': r.hits or 0,
                'mode': r.mode,
                'cited_act': r.cited_act,
                'last_used_at': r.last_used_at.isoformat() if r.last_used_at else None,
            } for r in top],
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/legal-library/cache/clear', methods=['POST'])
@login_required
def api_legal_library_cache_clear():
    """Admin: purge the legal Q&A cache (RAM + DB). Use after a big
    library update so stale answers don't keep getting served."""
    try:
        from ai.legal_qa import cache_clear_all
        n = cache_clear_all()
        return jsonify({'ok': True, 'deleted': int(n)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/chat/<client_id>/backfill-extractions', methods=['POST'])
@login_required
def api_chat_backfill_extractions(client_id):
    """Re-run the appropriate extractor on every Document for this client
    where extracted_data is empty (or only an error). Useful after fixing
    extractor bugs without re-uploading."""
    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'ok': False, 'error': 'Client not found'}), 404
    docs = Document.query.filter_by(client_id=client_id).all()
    done = []
    for doc in docs:
        if doc.category not in ('property_title', 'property_tax', 'bank_statement',
                                'vehicle', 'insurance', 'epf_kwsp', 'nric'):
            continue
        try:
            existing = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            existing = {}
        # Skip if already has substantive data (no error + at least one truthy field)
        if existing and not existing.get('error') and any(
            v for k, v in existing.items() if k not in ('error', 'raw') and v
        ):
            continue
        abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
        if not os.path.isfile(abs_path):
            continue
        try:
            if doc.category == 'nric':
                from ai.ocr import extract_nric_data
                ext = extract_nric_data(abs_path)
            elif doc.category in ('property_title', 'property_tax'):
                from ai.property_extractor import extract_property_data
                ext = extract_property_data(abs_path, doc_type='general')
            elif doc.category == 'bank_statement':
                from ai.ocr import extract_asset_document
                ext = extract_asset_document(abs_path, asset_type='bank')
            elif doc.category == 'vehicle':
                from ai.ocr import extract_asset_document
                ext = extract_asset_document(abs_path, asset_type='vehicle')
            else:
                from ai.ocr import extract_asset_document
                ext = extract_asset_document(abs_path, asset_type='other')
            doc.extracted_data = json.dumps(ext)
            done.append({'id': doc.id, 'category': doc.category,
                         'filename': doc.original_filename})
            db.session.commit()
        except Exception as e:
            done.append({'id': doc.id, 'category': doc.category, 'error': str(e)})
    return jsonify({'ok': True, 'updated': len(done), 'details': done[:30]})


@app.route('/api/chat/<client_id>/clear', methods=['POST'])
@login_required
def api_chat_clear(client_id):
    """Delete conversational messages in this client's chat.

    CRITICAL: Source messages — those that contain the original WhatsApp /
    email body with addresses and beneficiary hints typed by the client —
    are NEVER deleted. These are identified by:
      1. Having document attachments (attachments_json is a non-empty list)
      2. Containing '(forwarded via email' in the body (Postmark inbound)

    Only bot replies and pure conversational user messages are deleted.
    Documents are kept and remain linked to their source messages.
    """
    client = db.session.get(Client, client_id)
    if not client:
        return jsonify({'ok': False, 'error': 'Client not found'}), 404
    cs = (ChatSession.query
          .filter_by(client_id=client_id)
          .order_by(ChatSession.created_at.desc())
          .first())
    if not cs:
        return jsonify({'ok': True, 'deleted': 0})

    all_msgs = ChatMessage.query.filter_by(session_id=cs.id).all()
    source_msg_ids: set = set()   # preserve — original WhatsApp/email source
    clearable_ids: list = []

    for m in all_msgs:
        # Keep messages that have document attachments (original upload messages)
        try:
            atts = json.loads(m.attachments_json or '[]')
        except (json.JSONDecodeError, TypeError):
            atts = []
        is_source = bool(atts and isinstance(atts, list) and len(atts) > 0)
        # Also keep messages that are clearly the forwarded WhatsApp email body
        if not is_source and '(forwarded via email' in (m.content or ''):
            is_source = True
        if is_source:
            source_msg_ids.add(m.id)
        else:
            clearable_ids.append(m.id)

    n = 0
    if clearable_ids:
        # Only unlink docs from messages we're actually deleting
        Document.query.filter(Document.chat_message_id.in_(clearable_ids)).update(
            {Document.chat_message_id: None}, synchronize_session=False)
        n = ChatMessage.query.filter(ChatMessage.id.in_(clearable_ids)).delete(
            synchronize_session=False)

    db.session.commit()
    return jsonify({'ok': True, 'deleted': n, 'preserved': len(source_msg_ids)})


@app.route('/api/chat/<client_id>/message/<message_id>', methods=['DELETE'])
@login_required
def api_chat_message_delete(client_id, message_id):
    """Delete a single chat message (and its assistant reply if it's a user
    message). Useful for cleaning up junk without nuking the whole thread."""
    m = db.session.get(ChatMessage, message_id)
    if not m:
        return jsonify({'ok': False, 'error': 'Message not found'}), 404
    cs = db.session.get(ChatSession, m.session_id)
    if not cs or cs.client_id != client_id:
        return jsonify({'ok': False, 'error': 'Wrong client'}), 403
    deleted_ids = [m.id]
    # If this is a user message, also delete the immediately following
    # assistant reply (so they remove as a pair the way they were created).
    if m.role == 'user':
        nxt = (ChatMessage.query
               .filter(ChatMessage.session_id == cs.id,
                       ChatMessage.created_at > m.created_at,
                       ChatMessage.role == 'assistant')
               .order_by(ChatMessage.created_at.asc()).first())
        if nxt:
            deleted_ids.append(nxt.id)
    Document.query.filter(Document.chat_message_id.in_(deleted_ids)).update(
        {Document.chat_message_id: None}, synchronize_session=False)
    ChatMessage.query.filter(ChatMessage.id.in_(deleted_ids)).delete(synchronize_session=False)
    db.session.commit()
    return jsonify({'ok': True, 'deleted': len(deleted_ids)})


@app.route('/api/chat/<client_id>/reject/<message_id>', methods=['POST'])
@login_required
def api_chat_reject(client_id, message_id):
    """Mark a proposed patch as rejected (no changes applied)."""
    m = db.session.get(ChatMessage, message_id)
    if not m or m.role != 'assistant':
        return jsonify({'ok': False, 'error': 'Message not found'}), 404
    if m.applied_at or m.rejected_at:
        return jsonify({'ok': False, 'error': 'Already resolved'}), 400
    m.rejected_at = datetime.utcnow()
    db.session.commit()
    return jsonify({'ok': True, 'message': _serialise_chat_message(m)})


# -- Directed chat helpers --------------------------------------------------

def _gather_recent_chat_text(client_id: str, max_chars: int = 20000) -> str:
    """Concat recent chat content for this client — user messages (all) plus
    recent assistant AI-summary messages (address/lot lookups).

    AI summaries contain the AI-extracted address↔PTD mappings that were
    parsed from the WhatsApp forward. Including them lets _enrich_from_chat_text
    link a lot number on a geran to the street address mentioned in the summary.

    max_chars is intentionally large (20 000) because a 15-property email body
    can easily exceed 8 000 chars. Truncating or skipping the email body means
    lot numbers for later properties (e.g. Property 9 of 15) are never found.
    Individual messages that exceed the remaining budget are TRUNCATED (not
    skipped) so we always get at least the start of every message.
    """
    cs = (ChatSession.query.filter_by(client_id=client_id)
          .order_by(ChatSession.created_at.desc()).first())
    if not cs:
        return ''
    msgs = (ChatMessage.query.filter_by(session_id=cs.id)
            .order_by(ChatMessage.created_at.asc()).all())
    out = []
    total = 0
    for m in msgs:
        if m.role == 'user':
            c = m.content or ''
        elif m.role == 'assistant':
            # Only include assistant messages that look like AI property summaries
            # (they contain lot/PTD numbers + street addresses from the WhatsApp text).
            # Skip short acks, gate prompts, etc.
            c = m.content or ''
            if len(c) < 200:
                continue
            # Strip quick-reply markers and markdown decorations to plain text
            c = re.sub(r'<!--quickreplies:.*?-->', '', c)
            c = re.sub(r'#{1,4}\s+', '', c)
        else:
            continue
        if not c:
            continue
        remaining = max_chars - total
        if remaining <= 0:
            break
        # Truncate oversized messages rather than skipping them entirely —
        # a 15-property email body should never be dropped just because it's long.
        if len(c) > remaining:
            c = c[:remaining]
        out.append(c)
        total += len(c)

    # ── Fallback: _raw_forward_text from Will step6_data ────────────────────
    # If chat messages were cleared (api_chat_clear) the WhatsApp body is
    # preserved in the Will record under step6_data['_raw_forward_text'].
    # Include it so enrichment can still work even after a chat clear.
    if total < max_chars // 2:  # only pull in if we don't have much already
        try:
            _will = (Will.query.filter_by(client_id=client_id, status='draft')
                     .filter(Will.deleted_at.is_(None))
                     .order_by(Will.updated_at.desc()).first())
            if _will and _will.step6_data:
                _s6 = json.loads(_will.step6_data)
                raw_fwd = (_s6.get('_raw_forward_text') or '').strip()
                if raw_fwd and raw_fwd not in '\n\n'.join(out):
                    remaining = max_chars - total
                    if remaining > 0:
                        out.insert(0, raw_fwd[:remaining])  # prepend — it's the source
        except Exception:
            pass

    return '\n\n'.join(out)


def _get_layer2_pending_props(client_id: str) -> list:
    """Properties that completed Layer 1 (inventoried) but not Layer 2 (beneficiary assigned)."""
    try:
        docs = Document.query.filter(
            Document.client_id == client_id,
            Document.category == 'property_title',
        ).all()
        result = []
        for d in docs:
            try:
                ex = json.loads(d.extracted_data) if d.extracted_data else {}
            except (json.JSONDecodeError, TypeError):
                ex = {}
            if (ex.get('_inventoried')
                    and not ex.get('_substitute_assigned')
                    and not ex.get('_skipped')
                    # 🔥 §10x.127 — skipped-not-in-will / user-removed docs
                    # MUST NOT surface as Layer 2 pending. They were
                    # explicitly excluded from the will by the user.
                    and not ex.get('_skipped_not_in_will')
                    and not ex.get('_user_removed')
                    and not ex.get('_orphan_group_skipped')):
                # Skip properties with no NLC identifiers (no lot/title) — these
                # are ghost entries (e.g. address-only photos with no land title data)
                # that cannot form a valid probate gift.
                has_nlc = (ex.get('title_number') or '').strip() or (ex.get('lot_number') or '').strip()
                if not has_nlc:
                    continue
                result.append({'document_id': d.id, 'extracted': ex})
        # Sort by confidence: high-confidence properties first.
        # This ensures they are inventoried first AND claim their addresses
        # first in the enrichment pass.
        try:
            from services.gift_walker import _score_property_confidence
            result.sort(key=lambda p: _score_property_confidence(p.get('extracted') or {}),
                        reverse=True)
        except Exception:
            pass  # sort failure is non-critical
        # ╔════════════════════════════════════════════════════════════════╗
        # ║ 🔥 BURN-IN §10hg — H3 PLACEHOLDERS NEED LAYER-2 TOO 🔥          ║
        # ║ H3 gifts (no Document) are saved into step5_data with empty     ║
        # ║ beneficiaries. They MUST surface in the layer2 queue so the     ║
        # ║ user is asked who inherits them — otherwise residuary fires     ║
        # ║ before beneficiaries are assigned.                              ║
        # ╚════════════════════════════════════════════════════════════════╝
        try:
            will = (Will.query.filter_by(client_id=client_id, status='draft')
                    .filter(Will.deleted_at.is_(None))
                    .order_by(Will.updated_at.desc()).first())
            if will and will.step5_data:
                try:
                    s5 = json.loads(will.step5_data)
                except Exception:
                    s5 = []
                if isinstance(s5, list):
                    seen_doc_ids = {p.get('document_id') for p in result}
                    for idx, g in enumerate(s5):
                        if not isinstance(g, dict):
                            continue
                        if not (g.get('kind') == 'property' or g.get('asset_type') == 'property'):
                            continue
                        if g.get('beneficiaries'):
                            continue   # already has beneficiaries
                        # 🔥 §10x.127 — skipped / removed gifts MUST NOT
                        # come back through the s5 synthesis loop.
                        if (g.get('skipped')
                                or g.get('_ai_summary_skipped')
                                or g.get('_user_removed')
                                or g.get('_skipped_not_in_will')):
                            continue
                        # Image-bound gifts have a real document_id already in
                        # `seen_doc_ids` from the Document loop above. Skip
                        # those — they're handled normally.
                        gid = g.get('document_id')
                        if gid and gid in seen_doc_ids:
                            continue
                        # Also skip if the bound Document itself is
                        # _skipped_not_in_will / _user_removed.
                        if gid:
                            try:
                                _gd = db.session.get(Document, gid)
                                if _gd:
                                    try:
                                        _gex = json.loads(_gd.extracted_data or '{}')
                                    except Exception:
                                        _gex = {}
                                    if (_gex.get('_skipped_not_in_will')
                                            or _gex.get('_user_removed')
                                            or _gex.get('_orphan_group_skipped')):
                                        continue
                            except Exception:
                                pass
                        # Synthetic entry — backed by step5 gift, not a Document.
                        pi = g.get('property_info') or g.get('property_details') or {}
                        synth_ex = {
                            'property_address': pi.get('property_address') or g.get('property_address') or '',
                            'title_number':     pi.get('title_number') or g.get('title_number') or '',
                            'lot_number':       pi.get('lot_number') or g.get('lot_number') or '',
                            'mukim':            pi.get('mukim') or '',
                            'daerah':           pi.get('daerah') or '',
                            'negeri':           pi.get('negeri') or '',
                            '_inventoried':     True,
                            '_h3_step5_idx':    idx,
                            '_ai_summary_idx':  g.get('_ai_summary_idx'),
                            # Carry main-beneficiary state for the Phase-A→B
                            # transition. Stored on the gift itself (not a Doc).
                            '_main_beneficiary_set': bool(g.get('_main_beneficiary_set')),
                            '_main_beneficiaries': g.get('_main_beneficiaries') or [],
                        }
                        result.append({
                            'document_id': gid or f'_h3_synth_{idx}',
                            'extracted':   synth_ex,
                            '_h3_step5_idx': idx,
                        })
        except Exception:
            pass
        return result
    except Exception:
        return []


def _persist_property_enrichment(client_id: str, recent_text: str) -> None:
    """For every pending property document that is missing a street address
    (or other NLC fields), run the chat-text enrichment and persist any
    newly filled fields back to extracted_data in the DB.

    Addresses that have already been matched and persisted to one property are
    excluded from the candidate pool for subsequent properties — preventing
    two properties from being assigned the same street address.

    Runs best-effort; never raises (non-critical path).
    """
    try:
        from ai.chat_planner import (
            _enrich_property_from_siblings,
            _enrich_from_chat_text,
        )
        from services.gift_walker import get_pending_gift_documents

        # 🔥 §10x.52 — vision-enrich sparse property docs once per process.
        # 🔥 §10x.65 — wrap each call in track_context so client_id is
        # logged (otherwise these calls bypass the per-client cap).
        try:
            from ai.file_classifier import vision_extract_property_fields
            from ai.cost_tracker import track_context, is_over_ceiling
            sparse_docs = Document.query.filter(
                Document.client_id == client_id,
                Document.category.in_([
                    'property_title', 'property_spa', 'property_tax',
                    'property_transfer', 'loan_agreement',
                ]),
            ).all()
            # 🔥 §10x.65 — STOP IMMEDIATELY if the client is over their
            # cost ceiling. Don't call vision at all.
            if is_over_ceiling(client_id):
                try:
                    current_app.logger.warning(
                        f'§10x.65 cost ceiling hit for {client_id} — '
                        'skipping vision enrichment.'
                    )
                except Exception:
                    pass
                sparse_docs = []
            for d in sparse_docs:
                try:
                    ex = json.loads(d.extracted_data or '{}') or {}
                except Exception:
                    ex = {}
                if ex.get('_vision_enriched'):
                    continue   # already done
                # Sparse if no street address OR (no lot AND no title)
                has_addr = bool((ex.get('property_address') or '').strip())
                has_lot = bool((ex.get('lot_number') or '').strip())
                has_title = bool((ex.get('title_number') or '').strip())
                if has_addr and (has_lot or has_title):
                    ex['_vision_enriched'] = 'skipped_sufficient'
                    d.extracted_data = json.dumps(ex)
                    continue
                if not d.file_path:
                    continue
                # 🔥 §10x.65 — wrap in track_context so the Claude call
                # gets logged with the correct client_id (cap enforced).
                with track_context(client_id=client_id, will_id=getattr(d, 'will_id', None)):
                    vf = vision_extract_property_fields(d.file_path) or {}
                changed = False
                for k in ('property_address', 'lot_number', 'title_number',
                          'mukim', 'daerah', 'negeri', 'owner_name',
                          'building_name', 'postcode'):
                    if not (ex.get(k) or '').strip() and vf.get(k):
                        ex[k] = vf[k]
                        changed = True
                ex['_vision_enriched'] = True
                d.extracted_data = json.dumps(ex)
                if changed:
                    try:
                        current_app.logger.info(
                            f'§10x.52 vision-enriched doc {d.id[:8]}: '
                            f'addr={ex.get("property_address","")[:40]!r} '
                            f'lot={ex.get("lot_number","")!r} '
                            f'title={ex.get("title_number","")!r}'
                        )
                    except Exception:
                        pass
            db.session.commit()
        except Exception:
            db.session.rollback()

        pend = get_pending_gift_documents(client_id)
        props = pend.get('property') or []

        _ENRICH_FIELDS = ('property_address', 'negeri', 'daerah',
                          'mukim', 'ownership_type', 'ownership_share',
                          '_beneficiary_hint')

        # Sort by confidence so high-confidence properties are enriched first.
        # This ensures they claim their real addresses before low-confidence
        # properties can accidentally steal them.
        try:
            from services.gift_walker import _score_property_confidence
            props = sorted(props,
                           key=lambda p: _score_property_confidence(p.get('extracted') or {}),
                           reverse=True)
        except Exception:
            pass  # non-critical

        # Build the set of addresses ALREADY claimed by other properties
        # (both those with pre-existing addresses AND those we'll match below).
        # This exclusion set grows as we process each property in sequence.
        claimed_addresses: set = set()
        for p in props:
            addr = ((p.get('extracted') or {}).get('property_address') or '').strip().lower()
            if addr:
                claimed_addresses.add(addr)

        # ╔════════════════════════════════════════════════════════════════╗
        # ║ 🔥 BURN-IN §10h — AI SUMMARY IS THE CANONICAL ADDRESS SOURCE 🔥 ║
        # ║ Per §10ha, title docs DO NOT contain street addresses. The      ║
        # ║ canonical address for any property is in the AI Summary text.   ║
        # ║ When a doc's title/lot uniquely matches an AI Summary entry,    ║
        # ║ use THAT entry's address — not whatever the chat-text regex     ║
        # ║ happens to fish out (which historically hallucinates building   ║
        # ║ names like "Condominium Example" from prior AI noise).          ║
        # ╚════════════════════════════════════════════════════════════════╝
        try:
            from ai.chat_planner import _extract_ai_summary_properties
            _ai_summary = _extract_ai_summary_properties(client_id) or []
        except Exception:
            _ai_summary = []

        def _digits_only(s: str) -> str:
            return re.sub(r'\D', '', s or '')

        def _match_doc_to_ai_summary(ex: dict) -> dict:
            """Return the matching AI-Summary entry (or None) for this doc.
            Match priority: title digits → lot digits → (mukim+daerah, single match)."""
            if not _ai_summary:
                return None
            d_title = _digits_only(ex.get('title_number') or '')
            d_lot   = _digits_only(ex.get('lot_number') or '')
            d_mukim = (ex.get('mukim') or '').strip().lower()
            d_daerah = (ex.get('daerah') or '').strip().lower()
            # Title digits — strongest match
            if d_title and len(d_title) >= 4:
                for ap in _ai_summary:
                    a_title = _digits_only(ap.get('title') or '')
                    if a_title and a_title == d_title:
                        return ap
            # Lot digits — strong
            if d_lot and len(d_lot) >= 3:
                for ap in _ai_summary:
                    a_lot = _digits_only(ap.get('lot') or '')
                    if a_lot and a_lot == d_lot:
                        return ap
            # Mukim+daerah — only match if EXACTLY ONE AI prop is in same mukim
            # AND that AI prop has no claimant doc yet (i.e. its address isn't
            # already in claimed_addresses). Otherwise too ambiguous.
            if d_mukim:
                cands = []
                for ap in _ai_summary:
                    a_mukim = (ap.get('mukim') or '').strip().lower()
                    if a_mukim and a_mukim != d_mukim:
                        continue
                    a_daerah = (ap.get('daerah') or '').strip().lower()
                    if d_daerah and a_daerah and a_daerah != d_daerah:
                        continue
                    a_addr = (ap.get('address') or '').strip().lower()
                    if a_addr and a_addr in claimed_addresses:
                        continue   # already taken
                    cands.append(ap)
                if len(cands) == 1:
                    return cands[0]
            return None

        changed = False
        for p in props:
            doc_id = p.get('document_id')
            if not doc_id:
                continue
            ex_orig = p.get('extracted') or {}

            # ── §10h pre-enrichment: AI Summary address takes precedence ───
            # If the doc matches a unique AI Summary entry by title/lot/mukim,
            # set the address from the summary BEFORE chat-text scan touches it.
            # This kills the "hallucinated Condominium Example" class of bugs.
            from ai.chat_planner import _NLC_ADDR_RE as _nlc_re
            ai_match = _match_doc_to_ai_summary(ex_orig)
            if ai_match:
                a_addr = (ai_match.get('address') or '').strip()
                if a_addr:
                    existing_addr_pre = (ex_orig.get('property_address') or '').strip()
                    # Overwrite when missing, NLC-format, or different from canonical.
                    needs_overwrite = (
                        not existing_addr_pre
                        or _nlc_re.match(existing_addr_pre)
                        or existing_addr_pre.lower() != a_addr.lower()
                    )
                    if needs_overwrite:
                        doc = db.session.get(Document, doc_id)
                        if doc:
                            try:
                                stored = json.loads(doc.extracted_data) if doc.extracted_data else {}
                            except (json.JSONDecodeError, TypeError):
                                stored = {}
                            stored['property_address'] = a_addr
                            stored.setdefault('_enriched_from', [])
                            tag = 'ai_summary.canonical_address'
                            if tag not in stored['_enriched_from']:
                                stored['_enriched_from'].append(tag)
                            # Strip the old enriched-from-chat-text tag if present —
                            # it's no longer truthful after this overwrite.
                            stored['_enriched_from'] = [
                                t for t in stored['_enriched_from']
                                if t != 'chat_text.property_address'
                            ]
                            doc.extracted_data = json.dumps(stored)
                            ex_orig = stored
                            changed = True
                            claimed_addresses.add(a_addr.lower())
                            continue   # skip the chat-text enrichment for this doc

            # ── §10h: scrub pre-stored hallucinated address (e.g. "Condominium
            # Example") so we don't carry it forward AND don't claim it.
            existing_addr_check = (ex_orig.get('property_address') or '').strip()
            _NOISE_BUILDING_PRE = ('example', 'sample', 'unknown', 'placeholder',
                                    'untitled', '<address>', 'n/a', 'tbd')
            if existing_addr_check and any(tok in existing_addr_check.lower()
                                           for tok in _NOISE_BUILDING_PRE):
                _doc_pre = db.session.get(Document, doc_id)
                if _doc_pre:
                    try:
                        _stored_pre = json.loads(_doc_pre.extracted_data) if _doc_pre.extracted_data else {}
                    except (json.JSONDecodeError, TypeError):
                        _stored_pre = {}
                    _stored_pre['property_address'] = ''
                    _stored_pre.setdefault('_enriched_from', [])
                    if 'cleared.hallucinated_address' not in _stored_pre['_enriched_from']:
                        _stored_pre['_enriched_from'].append('cleared.hallucinated_address')
                    _doc_pre.extracted_data = json.dumps(_stored_pre)
                    ex_orig = _stored_pre
                    changed = True

            # Already has a REAL (non-NLC) address — add to claimed set and skip.
            # NLC-format addresses (e.g. "H.S.(D) 251041 P.T.D …") are treated
            # the same as missing — we try to find a real street address to replace them.
            existing_addr = (ex_orig.get('property_address') or '').strip()
            if existing_addr and not _nlc_re.match(existing_addr):
                claimed_addresses.add(existing_addr.lower())
                continue

            # Build a modified recent_text with claimed addresses removed
            # so the reverse-lookup can't pick them for this property.
            filtered_text = recent_text
            for claimed in claimed_addresses:
                if claimed and len(claimed) > 8:
                    # Replace matched address with a placeholder so regex skips it
                    filtered_text = re.sub(
                        re.escape(claimed), '___CLAIMED___', filtered_text, flags=re.IGNORECASE
                    )

            # Run enrichment (sibling cross-ref + chat-text scan)
            enriched = _enrich_property_from_siblings(p)
            enriched = _enrich_from_chat_text(enriched, filtered_text)

            # ── §10h Address whitelist: only accept addresses that exist in
            # the AI Summary. Anything else is either OCR noise, a hallucinated
            # building name, or a self-referential echo from a previous chat
            # turn (e.g. "Condominium Example" copied out of an assistant card
            # that was itself populated from an earlier hallucinated extract).
            new_addr = (enriched.get('property_address') or '').strip()
            if new_addr and _ai_summary:
                # Token-based fuzzy match: every AI Summary address contributes
                # a set of distinctive tokens; the enriched address must hit at
                # least one such set with ≥ 2 distinctive tokens (or 1 unit-like
                # token, e.g. "c-30-08").
                _NOISE_BUILDING = ('example', 'sample', 'unknown', 'placeholder',
                                   'untitled', '<address>', 'n/a', 'tbd')
                _new_lc = new_addr.lower()
                if any(tok in _new_lc for tok in _NOISE_BUILDING):
                    enriched['property_address'] = ''   # reject hallucination
                else:
                    def _toks(s: str):
                        return set(t for t in re.split(r'[^a-z0-9]+', (s or '').lower())
                                   if len(t) >= 3)
                    new_toks = _toks(new_addr)
                    matched_any = False
                    for ap in _ai_summary:
                        a_addr = (ap.get('address') or '').strip()
                        if not a_addr:
                            continue
                        a_toks = _toks(a_addr)
                        if not a_toks:
                            continue
                        overlap = new_toks & a_toks
                        # accept if ≥ 2 token overlap
                        if len(overlap) >= 2:
                            matched_any = True
                            break
                    if not matched_any:
                        enriched['property_address'] = ''   # not in AI Summary

            # Check which fields actually changed or improved.
            # For property_address: also count as "newly filled" if the old
            # value was NLC-format and the new value is a real street address.
            newly_filled = {}
            for f in _ENRICH_FIELDS:
                new_val = enriched.get(f)
                old_val = ex_orig.get(f)
                if not new_val:
                    continue
                if not old_val:
                    newly_filled[f] = new_val
                elif f == 'property_address' and _nlc_re.match(old_val) and not _nlc_re.match(new_val):
                    newly_filled[f] = new_val  # NLC → real street address upgrade
            if not newly_filled:
                continue

            # Mark this address as claimed so subsequent properties can't reuse it
            new_addr = (enriched.get('property_address') or '').strip()
            if new_addr:
                claimed_addresses.add(new_addr.lower())

            # Persist back to DB — reload from DB to avoid overwriting concurrent writes
            doc = db.session.get(Document, doc_id)
            if not doc:
                continue
            try:
                stored = json.loads(doc.extracted_data) if doc.extracted_data else {}
            except (json.JSONDecodeError, TypeError):
                stored = {}

            for f, v in newly_filled.items():
                stored[f] = v
            # Track enrichment source for debugging
            stored.setdefault('_enriched_from', [])
            for tag in (enriched.get('_enriched_from') or []):
                if tag not in stored['_enriched_from']:
                    stored['_enriched_from'].append(tag)

            doc.extracted_data = json.dumps(stored)
            changed = True

        if changed:
            try:
                db.session.commit()
            except Exception:
                db.session.rollback()

        # ── AI address matching pass ───────────────────────────────────────
        # For properties still missing an address after the regex scan,
        # call Claude Haiku to match them to the addresses in the WhatsApp
        # text using NLC context clues (mukim/daerah/ownership).
        # Re-fetch props from DB so we use the just-persisted addresses.
        try:
            from ai.chat_planner import ai_match_property_addresses, _NLC_ADDR_RE
            # Reload to pick up any addresses persisted in the regex pass above
            pend2 = get_pending_gift_documents(client_id)
            props2 = pend2.get('property') or []
            # Include properties with NLC-format addresses (e.g. "H.S.(D) 251041 P.T.D …")
            # alongside those with no address at all — both need real street address matching.
            def _needs_real_addr(p):
                addr = (p.get('extracted') or {}).get('property_address', '').strip()
                if not addr:
                    return True
                return bool(_NLC_ADDR_RE.match(addr))
            unmatched = [p for p in props2 if _needs_real_addr(p)]
            # ╔══════════════════════════════════════════════════════════════╗
            # ║  🔥 BURN-IN — HIGH CONFIDENCE FIRST + ONE-CLAIM-ONLY 🔥        ║
            # ║  Sort props by confidence DESC, process matches in the same   ║
            # ║  order, claim each address greedily. Once an address is       ║
            # ║  claimed by a high-confidence match, no later doc can reuse   ║
            # ║  it — even if the AI suggests the same address for a second   ║
            # ║  doc. NO DUPLICATE address assignments. See CLAUDE.md §10g.   ║
            # ╚══════════════════════════════════════════════════════════════╝
            from services.gift_walker import _score_property_confidence
            unmatched.sort(
                key=lambda p: _score_property_confidence(p.get('extracted') or {}),
                reverse=True,
            )
            if unmatched:
                ai_matches = ai_match_property_addresses(
                    unmatched, recent_text, claimed_addresses
                )
                # ╔══════════════════════════════════════════════════════════╗
                # ║  🔥 BURN-IN — WEB-SEARCH VALIDATES EVERY MATCH 🔥          ║
                # ║  CLAUDE.md §10hf — we have the address from the AI       ║
                # ║  matcher → we MUST web-search it for property-type      ║
                # ║  clues, then verify the doc is compatible with those    ║
                # ║  clues. Incompatible matches are downgraded to 'low'    ║
                # ║  + _address_needs_confirm so the user is asked before   ║
                # ║  the binding goes into step5_data.                      ║
                # ╚══════════════════════════════════════════════════════════╝
                try:
                    from ai.chat_planner import validate_matches_with_web_clues
                    ai_matches = validate_matches_with_web_clues(
                        unmatched, ai_matches
                    )
                except Exception:
                    pass  # non-critical — fall through with un-validated matches
                if ai_matches:
                    ai_changed = False
                    # Build (doc_id, match_dict, conf_rank) tuples and process
                    # in CONFIDENCE DESC order so high-conf claims first.
                    _conf_rank = {'high': 3, 'medium': 2, 'low': 1}
                    pairs = []
                    for p in unmatched:
                        did = p.get('document_id')
                        mv = ai_matches.get(did)
                        if not mv:
                            continue
                        if isinstance(mv, dict):
                            conf = (mv.get('confidence') or 'high').lower()
                        else:
                            conf = 'high'
                        # Bonus: combine doc-confidence + match-confidence so
                        # a high-conf doc with high match goes first
                        doc_score = _score_property_confidence(p.get('extracted') or {})
                        rank = doc_score * 10 + _conf_rank.get(conf, 0)
                        pairs.append((rank, p, mv))
                    pairs.sort(key=lambda t: t[0], reverse=True)
                    for _rank, p, match_val in pairs:
                        doc_id2 = p.get('document_id')
                        # Support both plain string and {address, confidence} dict
                        if isinstance(match_val, dict):
                            matched_addr = (match_val.get('address') or '').strip()
                            confidence = match_val.get('confidence', 'high')
                        else:
                            matched_addr = match_val.strip()
                            confidence = 'high'
                        if not matched_addr or len(matched_addr) < 8:
                            continue
                        # ── ONE-CLAIM-ONLY: skip if already taken ──────────
                        if matched_addr.lower() in claimed_addresses:
                            continue
                        doc2 = db.session.get(Document, doc_id2)
                        if not doc2:
                            continue
                        try:
                            stored2 = json.loads(doc2.extracted_data) if doc2.extracted_data else {}
                        except (json.JSONDecodeError, TypeError):
                            stored2 = {}
                        # For high confidence: overwrite any NLC address. For low confidence:
                        # only overwrite if currently NLC or empty; store as pending confirmation.
                        cur_addr = stored2.get('property_address', '').strip()
                        if cur_addr and not _NLC_ADDR_RE.match(cur_addr):
                            continue  # real address already set — don't touch it
                        stored2['property_address'] = matched_addr[:200]
                        stored2['_address_confidence'] = confidence
                        stored2.setdefault('_enriched_from', [])
                        if 'ai_address_match' not in stored2['_enriched_from']:
                            stored2['_enriched_from'].append('ai_address_match')
                        # ── Web-clue validation flags (CLAUDE.md §10hf) ────
                        # When validate_matches_with_web_clues attached web
                        # evidence to this match, persist it so the
                        # walkthrough card can show "🔗 sources: …" and the
                        # type/mukim hints, and so a future re-render can
                        # tell whether this was clue-validated or not.
                        if isinstance(match_val, dict):
                            cs = match_val.get('_clue_status')
                            if cs:
                                stored2['_clue_status'] = cs
                                if 'web_clues' not in stored2['_enriched_from']:
                                    stored2['_enriched_from'].append('web_clues')
                            for k_src, k_dst in (
                                ('_clue_type', '_web_property_type'),
                                ('_clue_mukim', '_web_mukim'),
                                ('_clue_sources', '_web_sources'),
                                ('_clue_reject_reason', '_web_reject_reason'),
                                ('_resolved_mukim', '_resolved_mukim'),
                                ('_mukim_source', '_mukim_source'),
                            ):
                                v_src = match_val.get(k_src)
                                if v_src:
                                    stored2[k_dst] = v_src
                            # Hint-1 verdict (boolean) — persist with explicit None
                            # support so a card can render "unknown" vs True/False.
                            if '_hint1_mukim_ok' in match_val:
                                stored2['_hint1_mukim_ok'] = match_val.get('_hint1_mukim_ok')
                        # For low-confidence matches: flag for user confirmation
                        if confidence in ('low', 'medium'):
                            stored2['_address_needs_confirm'] = True
                        else:
                            stored2.pop('_address_needs_confirm', None)
                        doc2.extracted_data = json.dumps(stored2)
                        claimed_addresses.add(matched_addr.lower())
                        ai_changed = True
                    if ai_changed:
                        try:
                            db.session.commit()
                        except Exception:
                            db.session.rollback()
        except Exception:
            pass  # AI pass is best-effort

    except Exception:
        pass  # never block the main chat flow


def _extract_whatsapp_context_for_file(body: str, filename: str) -> str:
    """Return text messages adjacent to this filename's attachment reference
    in an exported WhatsApp chat log.

    WhatsApp exports (iOS / Android) embed attachment lines like:
      [02/05/26, 13:52] Ahmad: ‎<attached: PHOTO-2026-05-02-13-52-35.jpg>
      [02/05/26, 13:52] Ahmad: IMG-20260502-WA0001.jpg (file attached)

    Clients send context BEFORE or AFTER images — both patterns are common:
      Pattern A — message then image:
        [time] Client: This is my property lot 127082, give to Sarah.
        [time] Client: <attached: photo.jpg>
      Pattern B — image then message:
        [time] Client: <attached: photo.jpg>
        [time] Client: This is my property, please give to daughter.
      Pattern C — multiple images with one surrounding message:
        [time] Client: All these are for lot 127082.
        [time] Client: <attached: photo1.jpg>
        [time] Client: <attached: photo2.jpg>   ← for this image, message is BEFORE

    We look up to 4 text lines BEFORE and up to 3 text lines AFTER, then
    return whichever side has more content (prefer before if equal).
    Returns '' if the filename isn't found (caller falls back to full body).
    """
    if not filename or not body:
        return ''
    fn_lower = filename.lower()
    body_lower = body.lower()
    idx = body_lower.find(fn_lower)
    if idx == -1:
        return ''

    def _is_attach_line(s: str) -> bool:
        sl = s.lower()
        return '<attached:' in sl or 'file attached' in sl

    # ── Look BEFORE the attachment reference ──────────────────────────────
    before_lines = body[:idx].split('\n')
    before_ctx = []
    # Use a large lookback (20 non-empty lines) so that a multi-property
    # WhatsApp forward list (e.g. "1. No. 18 Jln … PTD 207922 …\n2. …\n3. …")
    # sent as a single message before all images is fully captured.
    for line in reversed(before_lines):
        stripped = line.strip()
        if not stripped:
            continue
        if _is_attach_line(stripped):
            # Skip other attachment lines but don't stop — the text we want
            # may be interspersed with other image sends (Pattern C).
            continue
        before_ctx.insert(0, stripped)
        if len(before_ctx) >= 20:
            break

    # ── Look AFTER the attachment reference ───────────────────────────────
    after_start = idx + len(filename)
    after_lines = body[after_start:].split('\n')
    after_ctx = []
    for line in after_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _is_attach_line(stripped):
            # Another attachment line — stop scanning (the text after this
            # belongs to that next image, not ours)
            break
        after_ctx.append(stripped)
        if len(after_ctx) >= 3:
            break

    # ── Pick the richer side (or combine if both non-empty) ───────────────
    if before_ctx and after_ctx:
        # Both sides have text — combine for maximum context
        combined = before_ctx + after_ctx
        return '\n'.join(combined)
    return '\n'.join(before_ctx or after_ctx)


# ╔════════════════════════════════════════════════════════════════════════╗
# ║  🔥 BURN-IN — WhatsApp timestamp per attachment 🔥                       ║
# ║  CLAUDE.md §10i + §10hb: the property card MUST display the WhatsApp    ║
# ║  timestamp of the image and of the adjacent message that's binding it. ║
# ║  We extract the timestamp at INGEST time (this function runs from the   ║
# ║  /api/inbound-email handler) and persist it to extracted_data. Without ║
# ║  this, a later-rendered card has no way to recover the timing — the    ║
# ║  message body may have been compacted or trimmed by then.              ║
# ╚════════════════════════════════════════════════════════════════════════╝
_WA_TIMESTAMP_RE = re.compile(
    # iOS export:  [02/05/26, 13:52:35]   or  [02/05/26, 1:52:35 PM]
    # Android:     [02/05/2026 13:52]      or  02/05/2026, 13:52 -
    r'\[?(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4})[,\s]+(\d{1,2}:\d{2}(?::\d{2})?(?:\s*[AP]M)?)\]?'
)


def _extract_whatsapp_timestamp_for_file(body: str, filename: str) -> str:
    """Return the WhatsApp timestamp (raw, as it appears in the export)
    of the line that references `filename` as an attachment. Returns ''
    if not found. Format examples: '02/05/26, 13:52:35', '02/05/2026 13:52'.

    Used to populate `_msg_timestamp` on each attachment doc at ingest
    time, so the §10i temporal-proximity matcher can compare image and
    message timestamps even after the message body is rotated.
    """
    if not filename or not body:
        return ''
    fn_lower = filename.lower()
    for line in body.splitlines():
        if fn_lower not in line.lower():
            continue
        m = _WA_TIMESTAMP_RE.search(line)
        if m:
            return f"{m.group(1)} {m.group(2)}".strip()
    return ''


_CONFIRM_TOKENS = ('yes', 'confirm', 'correct', 'ok ', 'okay', 'yep', 'yeah', 'true', 'right')
_SKIP_TOKENS = ('skip', 'later', 'pass')
_DELETE_TOKENS = ('delete', 'remove', 'wrong', 'discard', 'trash', 'irrelevant', 'unrelated')


def _dedupe_ic_against_existing(client_id: str, doc, extracted: dict) -> bool:
    """If `doc` is an IC whose extracted name / NRIC / ADDRESS matches an
    existing Person row OR another nric Document for this client, mark
    `doc` as 'duplicate' and return True. Caller should skip emitting
    this doc as an artifact.

    🔥 §10x.85 — three signals checked, ANY match dedups:
      (a) NRIC equal (digit-strip) — same IC card
      (b) Name equal (uppercase strict) — same person
      (c) Address fuzzy match (≥3 distinctive 4+ char tokens shared,
          + postcode bonus) — back-of-IC of an already-known person
          (the back has only address, NOT name/NRIC, so without (c)
          we miss back-of-IC dedup for ICs uploaded via email)

    Checks against BOTH:
      - Existing Person rows (already-confirmed via Step 1 walkthrough)
      - Other nric Documents (front+back uploaded in same batch)

    Without these checks, a 2nd email carrying the back of an
    already-confirmed person's IC creates a brand-new pending IC card
    asking the user to identify someone they've already verified.
    """
    if not extracted:
        return False
    name = (extracted.get('full_name') or '').strip().upper()
    nric_raw = (extracted.get('nric_number') or '').strip()
    nric_digits = re.sub(r'\D', '', nric_raw)
    address = (extracted.get('address') or '').strip().upper()
    if not (name or nric_digits or address):
        return False

    def _addr_tokens(addr: str) -> set:
        return set(t for t in re.split(r'[\s,/\-]+', addr or '')
                    if len(t) >= 4 and not t.isdigit())

    def _addr_matches(a: str, b: str) -> bool:
        if not a or not b:
            return False
        ta, tb = _addr_tokens(a), _addr_tokens(b)
        shared = len(ta & tb)
        # Postcode bonus
        pa = re.search(r'\b\d{5}\b', a)
        pb = re.search(r'\b\d{5}\b', b)
        if pa and pb and pa.group(0) == pb.group(0):
            shared += 1
        return shared >= 3

    # ── Check against existing Person rows (already confirmed) ──────
    # 🔥 §10x.151 — STRONG MATCH FIRST. NRIC equality + Name equality
    # both uniquely identify a person. Address match alone is WEAK
    # (different family members often share an address — e.g. wife's
    # IC has testator's mailing address) and used to incorrectly bind
    # the late-arrival IC to whichever Person query returned first.
    # Sort persons so strong matches resolve before weak ones.
    try:
        persons = Person.query.filter_by(client_id=client_id).all()
    except Exception:
        persons = []

    def _match_strength(p):
        p_name = (p.full_name or '').strip().upper()
        p_nric_digits = re.sub(r'\D', '', (p.nric_passport or '').strip())
        p_addr = (p.address or '').strip().upper()
        if nric_digits and p_nric_digits and nric_digits == p_nric_digits:
            return 3   # NRIC match — strongest
        if name and p_name and name == p_name:
            return 2   # Name match — strong
        if _addr_matches(address, p_addr):
            return 1   # Address match — weak (only if NRIC + name failed)
        return 0
    persons = sorted(persons, key=_match_strength, reverse=True)

    for p in persons:
        p_name = (p.full_name or '').strip().upper()
        p_nric_digits = re.sub(r'\D', '', (p.nric_passport or '').strip())
        p_addr = (p.address or '').strip().upper()
        # 🔥 §10x.151 — refuse address-only match against a Person who
        # already has a different NRIC. Otherwise wife's IC would link
        # to husband's already-verified Person via shared address.
        addr_only = (
            not (nric_digits and p_nric_digits and nric_digits == p_nric_digits)
            and not (name and p_name and name == p_name)
            and _addr_matches(address, p_addr)
        )
        if addr_only and p_nric_digits and nric_digits and p_nric_digits != nric_digits:
            continue   # different NRICs — addr alone isn't enough
        if (nric_digits and p_nric_digits and nric_digits == p_nric_digits) \
           or (name and p_name and name == p_name) \
           or _addr_matches(address, p_addr):
            # 🔥 §10x.143 — H3 PLACEHOLDER BACKFILL
            # If matching Person was created from text (H3 placeholder per
            # §10x.34) — empty NRIC, no document_id — and the IC carries
            # NRIC/address now, BACKFILL the Person before marking the doc
            # as duplicate. Otherwise the Person stays empty and the will
            # generates "(MALAYSIA NRIC No. )" blanks.
            try:
                if not (p.nric_passport or '').strip() and nric_digits:
                    canonical_nric = (extracted.get('nric_number') or '').strip()
                    if canonical_nric:
                        p.nric_passport = canonical_nric
                if not (p.address or '').strip() and (extracted.get('address') or '').strip():
                    p.address = (extracted.get('address') or '').strip()
                if not p.document_id:
                    p.document_id = doc.id   # Link IC doc to placeholder Person
            except Exception:
                pass
            doc.category = 'duplicate'
            doc.description = f'(duplicate of {p.full_name})'
            try:
                ed = json.loads(doc.extracted_data) if doc.extracted_data else {}
            except Exception:
                ed = {}
            ed['_already_known'] = True
            ed['_matched_person_id'] = p.id
            ed['_skip_reason'] = ('nric_match' if nric_digits and p_nric_digits and nric_digits == p_nric_digits
                                   else 'name_match' if name and p_name and name == p_name
                                   else 'address_match')
            doc.extracted_data = json.dumps(ed)
            return True

    # ── Check against other IC Documents (sibling batch dedup) ────
    # Include category='duplicate' too — those are ICs we've already
    # tagged as duplicates, but they still carry rich extracted info
    # (name + NRIC + address). A new IC with matching NRIC should
    # transitively dedup against them.
    siblings = Document.query.filter(
        Document.client_id == client_id,
        Document.category.in_(('nric', 'duplicate')),
        Document.id != doc.id,
    ).all()
    for sib in siblings:
        try:
            sib_ex = json.loads(sib.extracted_data) if sib.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            sib_ex = {}
        sib_name = (sib_ex.get('full_name') or '').strip().upper()
        sib_nric_digits = re.sub(r'\D', '', (sib_ex.get('nric_number') or '').strip())
        sib_addr = (sib_ex.get('address') or '').strip().upper()
        if (nric_digits and sib_nric_digits and nric_digits == sib_nric_digits) \
           or (name and sib_name and name == sib_name) \
           or _addr_matches(address, sib_addr):
            doc.category = 'duplicate'
            doc.description = f'(duplicate of {sib.original_filename or sib.id[:8]})'
            return True
    return False


def _try_delete_pending_identity(client_id: str, user_text: str):
    """If user replies with a delete keyword, bulk-delete every Document
    that shares the focused IC's extracted name (re-uploads of the same
    person create multiple docs; one 'delete' should remove them all so
    the walk-through doesn't loop on duplicates).
    Returns {'name', 'action': 'deleted', 'count'} or None."""
    if not user_text:
        return None
    from services.identity_walker import get_pending_ic_documents
    text_lower = user_text.lower().strip()
    words = set(re.findall(r'\b[a-z]+\b', text_lower))
    if not any(t in words for t in _DELETE_TOKENS):
        return None
    pending = get_pending_ic_documents(client_id)
    if not pending:
        return None
    target = pending[0]
    ex = target['extracted'] or {}
    name = (ex.get('full_name') or '').strip()
    nric = (ex.get('nric_number') or '').strip()

    # Bulk delete: any nric Document where EITHER name OR nric matches.
    # NRIC catches the case where one upload has unreadable name but
    # the same IC number — those would otherwise re-appear in walk-through.
    target_name = name.upper()
    all_nric = Document.query.filter_by(client_id=client_id, category='nric').all()
    count = 0
    for d in all_nric:
        try:
            exd = json.loads(d.extracted_data) if d.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            exd = {}
        d_name = (exd.get('full_name') or '').strip().upper()
        d_nric = (exd.get('nric_number') or '').strip()
        match = (target_name and d_name and target_name == d_name) or \
                (nric and d_nric and nric == d_nric)
        if match:
            d.category = 'deleted'
            d.description = '(removed by user from chat walk-through)'
            count += 1
    if count == 0:
        # No name or NRIC to match on — just delete this single document
        doc = db.session.get(Document, target['document_id'])
        if doc:
            doc.category = 'deleted'
            doc.description = '(removed by user from chat walk-through)'
            count = 1
    db.session.commit()
    label = name or target.get('original_filename', 'this document')
    return {'name': label, 'action': 'deleted', 'count': count}


def _current_stage_num(client_id: str, will) -> int:
    """🔥 §10x.38 + §10x.88 — Numeric current step. Wizard right-pane
    indicator MUST match what the chat planner is currently asking.
    Computed from the EXACT same gates the chat planner uses, in the
    EXACT same order:

      pending IC                                → 1 (Identity walkthrough)
      testator data missing OR not confirmed    → 2 (Confirm Testator)
      < 1 executor OR not confirmed             → 3 (Executor)
      no beneficiaries OR not confirmed         → 5 (Beneficiaries)
      pending gifts in walkthrough              → 6 (Specific Gifts)
      no residuary                              → 7 (Residuary)
      else                                      → 10 (Generate)

    🔥 §10x.88 — auto-populated steps still require user confirmation.
    Earlier this function only checked DATA PRESENCE, so the §10x.42
    reconcile that auto-adds executor / beneficiaries from message
    context made the planner skip Steps 2/3/5 entirely → user dropped
    straight to Step 6. Now Steps 2/3/5 require completed_steps
    markers ('testator_confirmed' / 'executor_confirmed' /
    'beneficiaries_confirmed') BEFORE advancing past them.

    Earlier this function returned 7 when 'assets_confirmed' was in
    completed_steps EVEN IF s5 had 0 saved gifts (walkthrough mid-flight)
    — a wizard / chat desync that confused users.
    """
    from services.identity_walker import get_pending_ic_documents
    if get_pending_ic_documents(client_id):
        return 1
    if not will:
        return 1
    try:
        s1 = json.loads(will.step1_data) if will.step1_data else {}
        s2 = json.loads(will.step2_data) if will.step2_data else {}
        s4 = json.loads(will.step4_data) if will.step4_data else []
        s5 = json.loads(will.step5_data) if will.step5_data else []
        s6 = json.loads(will.step6_data) if will.step6_data else {}
        completed = json.loads(will.completed_steps or '[]')
    except (json.JSONDecodeError, TypeError):
        s1, s2, s4, s5, s6, completed = {}, {}, [], [], {}, []
    if not isinstance(completed, list):
        completed = []
    # 🔥 §10x.88 — auto-populated steps still require user confirmation.
    # Per-step gate: data exists AND user confirmed (marker in
    # completed_steps). Without the second clause, the §10x.42
    # reconciler that auto-adds executor / beneficiaries from message
    # context made the planner skip Steps 2/3/5 entirely → user landed
    # straight on Step 6.
    if not (s1 or {}).get('full_name') or 'testator_confirmed' not in completed:
        return 2  # Step 2 Testator (data missing OR awaiting confirm)
    n_exec = len((s2 or {}).get('executors') or [])
    if n_exec < 1 or 'executor_confirmed' not in completed:
        return 3  # Step 3 Executors
    if not isinstance(s4, list) or len(s4) == 0 or 'beneficiaries_confirmed' not in completed:
        return 5  # Step 5 Beneficiaries
    # 🔥 §10x.38 — Pending-gift check overrides 'assets_confirmed' flag.
    # The flag means asset-inventory phase ended; walkthrough may still
    # be running. As long as ANY gift is pending OR no gifts saved yet
    # while the walkthrough has assets to walk, we're STILL on Step 6.
    try:
        from services.gift_walker import get_pending_gift_documents
        pg = get_pending_gift_documents(client_id) or {}
        total_pending = sum(len(v) for v in pg.values()
                             if isinstance(v, list))
    except Exception:
        total_pending = 0
    has_saved_gifts = isinstance(s5, list) and len(s5) > 0
    if total_pending > 0:
        return 6   # walkthrough has more cards to show
    if 'assets_confirmed' not in completed and not has_saved_gifts:
        return 6   # asset inventory phase not yet acknowledged
    if not s6 or not (s6.get('beneficiaries') or s6.get('residuary_beneficiary_name')):
        return 7   # Step 7 Residuary
    return 10  # done — Generate


def _current_stage_label(client_id: str, will) -> str:
    """Short human-readable label for the planner's current stage. Used by
    the Q&A nudge to remind the user where to come back to."""
    from services.identity_walker import get_pending_ic_documents
    if get_pending_ic_documents(client_id):
        return "Step 1: Identity walk-through"
    if not will:
        return "Step 1: setup"
    try:
        s2 = json.loads(will.step2_data) if will.step2_data else {}
        s4 = json.loads(will.step4_data) if will.step4_data else []
        completed = json.loads(will.completed_steps or '[]')
    except (json.JSONDecodeError, TypeError):
        s2, s4, completed = {}, [], []
    if not isinstance(completed, list):
        completed = []
    n_exec = len((s2 or {}).get('executors') or [])
    if 'assets_confirmed' not in completed:
        return "Step 6: Asset inventory review"
    if n_exec < 2:
        return f"Step 3: {'main' if n_exec == 0 else 'substitute'} Executor"
    if not isinstance(s4, list) or len(s4) == 0:
        return "Step 5: Confirm Beneficiaries"
    return "Step 6: Specific Gifts"


def _try_save_beneficiaries(client_id: str, user_text: str):
    """Step 5 (Beneficiaries) handler. Accepts:
      - 'yes' / 'confirm' → save the auto-suggested likely list
      - 'remove X' / 'remove X and Y' → save likely list minus X (by name OR
        by relationship word, e.g. 'remove sister in law')
      - 'only X, Y' / 'just X and Y' → save only X and Y
    Returns {'name', 'role', 'kind'} or None."""
    if not user_text:
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s4 = json.loads(will.step4_data) if will.step4_data else []
    except (json.JSONDecodeError, TypeError):
        s4 = []
    if not isinstance(s4, list):
        s4 = []
    if len(s4) > 0:
        return None
    try:
        s2 = json.loads(will.step2_data) if will.step2_data else {}
        if not isinstance(s2, dict):
            s2 = {}
    except (json.JSONDecodeError, TypeError):
        s2 = {}
    if len(s2.get('executors') or []) == 0:
        return None

    BENEFICIARY_RELS = {
        'spouse', 'wife', 'husband', 'son', 'daughter', 'father', 'mother',
        'brother', 'sister', 'grandson', 'granddaughter', 'beneficiary',
        'stepson', 'stepdaughter', 'adopted son', 'adopted daughter',
    }
    persons = Person.query.filter_by(client_id=client_id).all()
    eligible = [p for p in persons
                if (p.relationship or '').lower() not in ('testator', 'witness')]

    text_lower = user_text.lower().strip()
    words = set(re.findall(r'\b[a-z\-]+\b', text_lower))

    def _match_persons(phrase):
        """Find Persons whose name OR relationship matches the phrase."""
        out = []
        ph = phrase.strip().lower().replace('-', ' ')
        if not ph:
            return out
        for p in eligible:
            nm = (p.full_name or '').lower()
            rel = (p.relationship or '').lower().replace('-', ' ')
            if (nm and (ph in nm or any(part in nm for part in ph.split() if len(part) > 2))) \
               or (rel and ph in rel):
                out.append(p)
        return out

    def _serialize(p_list):
        return [{
            'full_name': p.full_name, 'nric_passport': p.nric_passport or '',
            'address': p.address or '', 'relationship': p.relationship or '',
            'person_id': p.id,
        } for p in p_list]

    final = None
    # 1) ONLY/JUST X — explicit list
    only_m = re.search(r'\b(?:only|just|keep)\s+(.+)', text_lower)
    if only_m:
        names_text = only_m.group(1)
        # Split by commas, 'and'
        parts = re.split(r',|\band\b', names_text)
        keep = []
        seen = set()
        for part in parts:
            for p in _match_persons(part):
                if p.id not in seen:
                    keep.append(p); seen.add(p.id)
        if keep:
            final = keep

    # 2) REMOVE X — drop named/related people from the likely list
    if final is None and 'remove' in words:
        # Extract everything after 'remove'
        rem_m = re.search(r'\bremove\b\s+(.+?)(?:\.|$)', text_lower)
        if rem_m:
            remove_text = rem_m.group(1)
            parts = re.split(r',|\band\b', remove_text)
            exclude_ids = set()
            for part in parts:
                for p in _match_persons(part):
                    exclude_ids.add(p.id)
            # Build likely list minus excluded
            likely = []
            for p in eligible:
                rel = (p.relationship or '').lower()
                if p.id in exclude_ids:
                    continue
                if rel in BENEFICIARY_RELS or not rel or 'in-law' in rel:
                    likely.append(p)
            final = likely

    # 3) Plain confirm
    if final is None and any(t in words for t in _CONFIRM_TOKENS):
        likely = []
        for p in eligible:
            rel = (p.relationship or '').lower()
            if rel in BENEFICIARY_RELS or not rel or 'in-law' in rel:
                likely.append(p)
        final = likely

    if not final:
        return None
    will.step4_data = json.dumps(_serialize(final))
    # 🔥 §10x.88 — mark Step 5 confirmed so the planner can advance to
    # Step 6. Without this, auto-populated beneficiaries (via §10x.42)
    # made the planner skip Step 5 entirely.
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
        if 'beneficiaries_confirmed' not in completed:
            completed.append('beneficiaries_confirmed')
            will.completed_steps = json.dumps(completed)
    except Exception:
        pass
    db.session.commit()
    names = ', '.join(p.full_name for p in final)
    return {'name': names, 'role': f'{len(final)} beneficiaries', 'kind': 'beneficiaries'}


def _try_handle_property_fill(client_id: str, user_text: str):
    """Let the writer manually supply missing property fields directly in chat.

    Accepted formats (all case-insensitive):
      address: No. 22, Jalan Rimbun, Seri Alam
      daerah: Johor Bahru
      negeri: Johor
      mukim: Plentong
      lot: 127082
      title: HS(D) 251041
      lot_number: 127082
      title_number: HS(D) 251041
      property fill      ← bare "property fill" → show a prompt template

    Also handles the trigger button value 'property fill' by returning a
    prompt text that tells the writer how to type the values.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()

    # ── Bare trigger — show instructions ──────────────────────────────
    if t == 'property fill':
        from services.gift_walker import get_pending_gift_documents
        pend = get_pending_gift_documents(client_id)
        props = pend.get('property') or []
        # Find the first non-inventoried property
        target_ex = {}
        for p in props:
            if not (p.get('extracted') or {}).get('_inventoried'):
                target_ex = p.get('extracted') or {}
                break
        _FIELD_LABELS = [
            ('property_address', 'address'),
            ('title_number',     'title'),
            ('lot_number',       'lot'),
            ('mukim',            'mukim'),
            ('daerah',           'daerah'),
            ('negeri',           'negeri'),
        ]
        missing = [(fld, kw) for fld, kw in _FIELD_LABELS
                   if not (target_ex.get(fld) or '').strip()]
        if not missing:
            return {'name': 'all fields complete', 'role': 'property_fill_noop', 'kind': 'property_fill'}
        lines = ['Type the missing field(s) directly — one per message or all at once:']
        for fld, kw in missing:
            lines.append(f"  `{kw}: <value>`   e.g. `{kw}: ...`")
        return {
            'name': 'fill prompt shown',
            'role': 'property_fill_prompt',
            'kind': 'property_fill',
            'reply_override': '\n'.join(lines),
        }

    # ── Field assignment: "daerah: Johor Bahru" ──────────────────────
    _FIELD_MAP = {
        'address':         'property_address',
        'property_address': 'property_address',
        'title':           'title_number',
        'title_number':    'title_number',
        'lot':             'lot_number',
        'lot_number':      'lot_number',
        'mukim':           'mukim',
        'daerah':          'daerah',
        'negeri':          'negeri',
        'state':           'negeri',
        'district':        'daerah',
        'area':            'area',
    }
    # Match "keyword: value" at the start of the message
    m = re.match(r'^([a-z_]+)\s*:\s*(.+)$', t, re.IGNORECASE)
    if not m:
        return None
    kw = m.group(1).strip().lower()
    raw_val = user_text[m.start(2):].strip()  # preserve original casing
    field = _FIELD_MAP.get(kw)
    if not field:
        return None

    # Find focused property doc
    from services.gift_walker import get_pending_gift_documents
    pend = get_pending_gift_documents(client_id)
    props = pend.get('property') or []
    target = next((p for p in props if not (p.get('extracted') or {}).get('_inventoried')), None)
    if not target:
        return None

    doc = db.session.get(Document, target['document_id'])
    if not doc:
        return None
    try:
        ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
    except (json.JSONDecodeError, TypeError):
        ex = {}

    old_val = (ex.get(field) or '').strip()
    ex[field] = raw_val.strip()
    ex.setdefault('_manually_edited', [])
    if isinstance(ex['_manually_edited'], list):
        ex['_manually_edited'].append(f'{field}={raw_val.strip()[:60]}')

    try:
        doc.extracted_data = json.dumps(ex)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    label = raw_val.strip()[:60]
    return {
        'name': f'{field} → {label}',
        'role': f'manual edit{"" if not old_val else f" (was: {old_val[:40]})"}',
        'kind': 'property_fill',
    }


def _try_handle_ownership(client_id: str, user_text: str):
    """Handle 'ownership: sole' / 'ownership: joint 1/2' commands.

    These are typed or tapped from the property card to confirm whether
    the testator owns the property solely or jointly, and if jointly
    what their undivided share is.

    Accepted formats (case-insensitive):
      ownership: sole
      ownership: joint 1/2
      ownership: joint 50%
      ownership: 1/2         ← bare share implies joint
      ownership: 1/3
    """
    if not user_text:
        return None
    t = user_text.strip()
    m = re.match(r'^ownership\s*:\s*(.+)$', t, re.IGNORECASE)
    if not m:
        return None
    val = m.group(1).strip().lower()

    # Determine type and share from the value
    if val == 'sole' or 'sole' in val:
        new_type  = 'sole'
        new_share = ''
    else:
        new_type = 'joint'
        # Extract the fraction/percentage: "joint 1/2", "1/2", "50%", "joint 50%"
        sh_m = re.search(r'(\d+/\d+|\d+\s*%)', val)
        new_share = sh_m.group(1).strip() if sh_m else ''

    # Find the current pending property
    from services.gift_walker import get_pending_gift_documents
    pend = get_pending_gift_documents(client_id)
    props = pend.get('property') or []
    target = next((p for p in props if not (p.get('extracted') or {}).get('_inventoried')), None)
    if not target:
        return None

    doc = db.session.get(Document, target['document_id'])
    if not doc:
        return None
    try:
        ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
    except (json.JSONDecodeError, TypeError):
        ex = {}

    ex['ownership_type']  = new_type
    ex['ownership_share'] = new_share
    ex.setdefault('_manually_edited', [])
    if isinstance(ex['_manually_edited'], list):
        ex['_manually_edited'].append(f'ownership={new_type}{"/" + new_share if new_share else ""}')

    try:
        doc.extracted_data = json.dumps(ex)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    label = 'Sole owner' if new_type == 'sole' else f'Joint owner — {new_share or "(share TBC)"}'
    return {'name': label, 'role': 'ownership_confirmed', 'kind': 'property_fill'}


def _try_handle_encumbrance(client_id: str, user_text: str):
    """Handle 'encumbered: yes / no / charge / caveat' commands.

    The Malaysian will template needs to know whether a property is clean
    or encumbered so the correct clause can be drafted:
      - Clean: straightforward gift clause
      - Charge: add direction to Executor to discharge bank charge
      - Caveat: add direction to apply to withdraw private caveat

    Accepted formats (case-insensitive):
      encumbered: yes          ← generic 'yes' — writer still types details
      encumbered: no
      encumbered: clean
      encumbered: charge       ← bank loan / mortgage
      encumbered: caveat       ← private caveat
    """
    if not user_text:
        return None
    t = user_text.strip()
    m = re.match(r'^encumbered\s*:\s*(.+)$', t, re.IGNORECASE)
    if not m:
        return None
    val = m.group(1).strip().lower()

    if val in ('no', 'clean', 'false', 'none'):
        enc_confirmed    = False
        enc_type         = ''
    elif val == 'charge' or 'charge' in val or 'mortgage' in val or 'loan' in val:
        enc_confirmed    = True
        enc_type         = 'charge'
    elif val == 'caveat' or 'caveat' in val or 'caveatan' in val:
        enc_confirmed    = True
        enc_type         = 'caveat'
    else:
        # Generic 'yes' — mark confirmed but leave type for writer to fill
        enc_confirmed    = True
        enc_type         = 'other'

    from services.gift_walker import get_pending_gift_documents
    pend = get_pending_gift_documents(client_id)
    props = pend.get('property') or []
    target = next((p for p in props if not (p.get('extracted') or {}).get('_inventoried')), None)
    if not target:
        return None

    doc = db.session.get(Document, target['document_id'])
    if not doc:
        return None
    try:
        ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
    except (json.JSONDecodeError, TypeError):
        ex = {}

    ex['encumbrance_confirmed'] = enc_confirmed
    if enc_type:
        ex['encumbrance_type'] = enc_type
    ex.setdefault('_manually_edited', [])
    if isinstance(ex['_manually_edited'], list):
        ex['_manually_edited'].append(f'encumbrance={"clean" if not enc_confirmed else enc_type or "yes"}')

    try:
        doc.extracted_data = json.dumps(ex)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    label = 'Clean (no charge/caveat)' if not enc_confirmed else f'Encumbered — {enc_type or "details TBC"}'
    return {'name': label, 'role': 'encumbrance_confirmed', 'kind': 'property_fill'}


def _try_handle_inbox_action(client_id: str, user_text: str):
    """Handle inbox-review actions before the property walkthrough starts.

      • 'inbox remove <doc_id>' → soft-delete that specific document
        (wrong upload, noise from email forward, etc.)
      • 'inbox start' → no-op that triggers the planner to proceed
        to the property walkthrough (user has confirmed the inbox)
    """
    if not user_text:
        return None
    t = user_text.strip().lower()

    if t == 'inbox start':
        # 🔥 §10x.53 — guard against premature 'inbox start' clicks while
        # vision classification is still in progress. If any Document for
        # this client is still 'chat_inbox' (not yet classified), tell
        # the user to wait — don't advance the walkthrough.
        try:
            in_progress = Document.query.filter_by(
                client_id=client_id, category='chat_inbox'
            ).count()
        except Exception:
            in_progress = 0
        if in_progress > 0:
            return {
                'name': 'still analysing',
                'role': 'inbox_not_ready',
                'kind': 'inbox_not_ready',
                'reply_override': (
                    f"🔍 Still analysing **{in_progress}** exhibit(s) — "
                    "please wait a moment. The chat will post a 'Ready' "
                    "message and the verify button once classification "
                    "is complete."
                ),
            }
        # User clicked "▶️ Start matching" — nothing to save, just let
        # plan_turn proceed to the normal walkthrough.
        return {
            'name': 'inbox confirmed',
            'role': 'inbox_start',
            'kind': 'inbox_start',
            'reply_override': None,  # no override — let plan_turn show the walkthrough
        }

    if t.startswith('inbox remove '):
        doc_id_raw = user_text.strip()[len('inbox remove '):].strip()
        if not doc_id_raw:
            return None
        doc = db.session.get(Document, doc_id_raw)
        if not doc or doc.client_id != client_id:
            return None
        fname = (doc.original_filename or doc.filename or 'image')[:60]
        try:
            ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            ex = {}
        ex['_inventoried'] = True
        ex['_deleted_by_user'] = True
        doc.extracted_data = json.dumps(ex)
        doc.category = 'deleted'
        doc.description = '(removed at inbox review step)'
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'name': fname,
            'role': 'removed from inbox',
            'kind': 'inbox_removed',
            'reply_override': (
                f"🗑 **{fname}** removed.\n\n"
                "The remaining attachments are still shown above. "
                "Tap **▶️ Start matching** when you're ready, or remove more exhibits."
            ),
        }

    return None


def _try_handle_restart_inbox(client_id: str, user_text: str):
    """Handle 'restart' / 'restart inbox' / 'clean' / 'start fresh'.

    Full reset + shows the inbox review card for ALL existing non-deleted
    documents so the writer can remove noise before the walkthrough starts.

    Steps:
      1. Clear gifts + assets_confirmed (same as restart gifts)
      2. Clear _inventoried/_skipped on all docs
      3. Build an inbox-style review card from the existing docs and return
         it as reply_override — user can remove images, add context, then
         tap ▶️ Start analysis to enter the walkthrough
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    _INBOX_TOKENS = {
        'restart', 'restart inbox', 'restart all', 'clean',
        'start fresh', 'start over', 'start clean', 'reset',
        'reset all', 'redo all', 'redo', 'fresh start', 'clean start',
    }
    if t not in _INBOX_TOKENS:
        return None

    try:
        will = (Will.query.filter_by(client_id=client_id, status='draft')
                .filter(Will.deleted_at.is_(None))
                .order_by(Will.updated_at.desc()).first())
        if not will:
            return None

        # 1. Clear gifts + assets_confirmed
        will.step5_data = '[]'
        try:
            completed = json.loads(will.completed_steps or '[]')
            if not isinstance(completed, list):
                completed = []
            completed = [c for c in completed if c != 'assets_confirmed']
            will.completed_steps = json.dumps(completed)
        except (json.JSONDecodeError, TypeError):
            pass
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

        # 2. Clear _inventoried/_skipped from all active docs
        # Document uses category='deleted' as soft-delete (no deleted_at column)
        _SKIP_CATS = ('nric', 'duplicate', 'deleted', 'voice')
        docs = Document.query.filter(
            Document.client_id == client_id,
            ~Document.category.in_(_SKIP_CATS),
        ).all()
        for d in docs:
            try:
                ex = json.loads(d.extracted_data) if d.extracted_data else {}
            except (json.JSONDecodeError, TypeError):
                ex = {}
            changed = False
            for flag in ('_inventoried', '_skipped'):
                if flag in ex:
                    ex.pop(flag)
                    changed = True
            if changed:
                d.extracted_data = json.dumps(ex)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

        # 3. Build inbox card from existing docs
        from ai.chat_planner import _intake_email_card as _iec

        # Deduplicate by (original_filename, file_size) — same name + same size
        # means the same physical file uploaded/forwarded multiple times.
        # Different photos CAN share a WhatsApp timestamp filename but have
        # different file sizes, so filename alone is too aggressive.
        _key_to_doc: dict = {}
        _chat_msg_ids: set = set()
        for d in docs:
            try:
                ex = json.loads(d.extracted_data) if d.extracted_data else {}
            except (json.JSONDecodeError, TypeError):
                ex = {}
            fname = (d.original_filename or d.filename or '').strip()
            fsize = d.file_size or 0
            _dedup_key = (fname, fsize) if fname else d.id
            # Track chat_message_ids so we can retrieve the original email text
            if d.chat_message_id:
                _chat_msg_ids.add(d.chat_message_id)
            existing = _key_to_doc.get(_dedup_key)
            if existing is None or str(d.id) > str(existing['_doc_id']):
                _key_to_doc[_dedup_key] = {
                    '_doc_id': d.id,
                    'document_id': d.id,
                    'kind': d.category or 'other',
                    'confidence': 'high',
                    'extracted': ex,
                    'original_filename': fname,
                }
        artifacts = list(_key_to_doc.values())

        if not artifacts:
            return {
                'name': 'inbox reset',
                'role': 'inbox_restarted',
                'kind': 'inbox_restart',
                'reply_override': (
                    "♻️ **Reset complete.** No documents found. "
                    "Upload or forward your documents to begin."
                ),
            }

        # Retrieve original email/WhatsApp body from ChatMessage records.
        # This is more reliable than _message_context on docs (which is capped
        # at 800 chars and may miss long email bodies).
        _best_ctx = ''
        if _chat_msg_ids:
            _chat_msgs = (ChatMessage.query
                          .filter(ChatMessage.id.in_(_chat_msg_ids),
                                  ChatMessage.role == 'user')
                          .order_by(ChatMessage.created_at.desc())
                          .all())
            # Pick the most recent non-empty message content
            for _cm in _chat_msgs:
                _ct = (_cm.content or '').strip()
                if _ct and len(_ct) > len(_best_ctx):
                    _best_ctx = _ct

        _will_snap = _will_data_snapshot(will) if will else {}
        inbox_card = _iec(artifacts, _best_ctx, current_will_data=_will_snap)
        _focus_ids = [a['document_id'] for a in artifacts]

        # Spawn background thread to generate the AI summary and post it as
        # a follow-up assistant message — avoids blocking the reset request
        # and prevents gunicorn worker timeout from two sequential AI calls.
        if _best_ctx:
            def _post_ai_summary(app_obj, cid, raw_ctx, session_id, focus_ids):
                with app_obj.app_context():
                    try:
                        from ai.chat_planner import _summarise_message, _clean_email_body
                        cleaned = _clean_email_body(raw_ctx)
                        # 🔥 §10x.76 — inject extracted Document fields
                        _df = []
                        try:
                            import json as _jsdf
                            _props = (Document.query
                                      .filter(Document.client_id == cid)
                                      .filter(Document.category.in_([
                                          'property_title','property_spa','property_tax',
                                          'property_transfer','loan_agreement',
                                          'bank_statement','insurance','vehicle','nric',
                                      ])).all())
                            for _d in _props:
                                try:
                                    _ex = _jsdf.loads(_d.extracted_data) if _d.extracted_data else {}
                                except Exception:
                                    _ex = {}
                                _row = {'kind': _d.category}
                                for _k in ('title_number','lot_number','mukim','daerah',
                                           'negeri','property_address','owner_name','title_type',
                                           'bank_name','account_number','currency',
                                           'insurer','policy_number','full_name','nric_number'):
                                    _v = _ex.get(_k)
                                    if _v: _row[_k] = _v
                                if len(_row) > 1:
                                    _df.append(_row)
                        except Exception:
                            _df = []
                        summary = _summarise_message(cleaned, doc_fields=_df) if cleaned else ''
                        if not summary:
                            summary = '_Could not generate summary — review exhibits below._'
                        import json as _json
                        _quick = _json.dumps([
                            {'label': '▶️ Start matching', 'value': 'inbox start'}
                        ])
                        reply = (
                            "### 📨 AI Summary of your message\n\n"
                            + summary
                            + f"\n\n<!--quickreplies:{_quick}-->"
                        )
                        cs = (ChatSession.query
                              .filter_by(id=session_id).first())
                        if not cs:
                            return
                        msg = ChatMessage(
                            session_id=cs.id, role='assistant',
                            content=reply,
                            attachments_json='[]',  # no repeat thumbnails in summary
                        )
                        db.session.add(msg)
                        db.session.commit()
                    except Exception:
                        pass

            cs_for_thread = _get_or_create_chat_session(client_id, user_id=None)
            threading.Thread(
                target=_post_ai_summary,
                args=(app, client_id, _best_ctx, cs_for_thread.id, _focus_ids),
                daemon=True,
            ).start()

        return {
            'name': f'inbox reset ({len(artifacts)} docs)',
            'role': 'inbox_restarted',
            'kind': 'inbox_restart',
            'reply_override': inbox_card,
            # No thumbnails on the intake card — the AI summary card
            # (posted by background thread) carries all the attachments.
            # This prevents the exhibit grid from appearing twice.
            'focus_attachments': [],
        }
    except Exception as _rst_err:
        import traceback as _tb
        app.logger.error(f'_try_handle_restart_inbox error: {_tb.format_exc()}')
        try:
            db.session.rollback()
        except Exception:
            pass
        return {
            'name': 'restart error',
            'role': 'inbox_restart_error',
            'kind': 'inbox_restart',
            'reply_override': (
                f"⚠️ Reset failed: `{str(_rst_err)[:200]}`\n\n"
                "Please try again or contact support."
            ),
        }


def _try_handle_restart_gifts(client_id: str, user_text: str):
    """Handle 'restart gifts' / 'restart inventory' / 'reset gifts'.

    Clears the entire Step-6 walkthrough so the writer can redo it from
    scratch with the latest document grouping improvements:

      1. Removes all gifts from the active Will's step5_data
      2. Clears 'assets_confirmed' from completed_steps so the planner
         re-enters the asset-inventory loop
      3. Clears _inventoried and _skipped from all property/bank/vehicle
         Documents so they re-appear in the walk

    Does NOT touch the Documents themselves (OCR data is preserved) or
    any other Will step (beneficiaries, executors, etc.).
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    _RESTART_TOKENS = {'restart gifts', 'restart inventory', 'reset gifts',
                       'reset inventory', 'redo gifts', 'redo inventory',
                       'restart step 6', 'restart step6'}
    if t not in _RESTART_TOKENS:
        return None

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None

    # 1. Clear all saved gifts
    will.step5_data = '[]'

    # 2. Remove 'assets_confirmed' from completed_steps so planner re-enters loop
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
        completed = [c for c in completed if c != 'assets_confirmed']
        will.completed_steps = json.dumps(completed)
    except (json.JSONDecodeError, TypeError):
        pass

    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    # 3. Clear _inventoried + _skipped from ALL non-deleted, non-nric docs.
    # Use a negative filter so new kinds (loan_agreement, insurance, etc.)
    # are automatically included without updating this list every time.
    _SKIP_CATS = ('nric', 'duplicate', 'deleted', 'voice')
    docs = Document.query.filter(
        Document.client_id == client_id,
        ~Document.category.in_(_SKIP_CATS),
    ).all()
    changed = 0
    for d in docs:
        try:
            ex = json.loads(d.extracted_data) if d.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            ex = {}
        if ex.get('_inventoried') or ex.get('_skipped'):
            ex.pop('_inventoried', None)
            ex.pop('_skipped', None)
            d.extracted_data = json.dumps(ex)
            changed += 1
    if changed:
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

    return {
        'name': f'Step 6 reset ({changed} docs unlocked)',
        'role': 'gifts_restarted',
        'kind': 'gifts_restart',
        'reply_override': (
            f"♻️ **Step 6 reset.** All {changed} documents have been unlocked and "
            f"the gift list has been cleared.\n\n"
            f"I'll now walk you through all your properties, bank accounts, and "
            f"vehicles again — this time with improved grouping. Just reply to each card."
        ),
    }


def _try_handle_unlink_action(client_id: str, user_text: str):
    """Handle the support-doc picker:
      • 'unlink <doc_id>' → remove that doc from its parent property's
        group by clearing its property_hint + setting category='other',
        so it pops out of the cluster and lands in chat_inbox for the
        writer to manually re-attach.
      • 'unlink done' → clear the parent's _unlink_pending flag so the
        normal property card returns next turn.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if not (t.startswith('unlink ') or t == 'unlink done'):
        return None

    if t == 'unlink done':
        # Find the property currently in unlink mode and clear the flag
        from services.gift_walker import get_pending_gift_documents
        pend = get_pending_gift_documents(client_id)
        for p in (pend.get('property') or []):
            if (p.get('extracted') or {}).get('_unlink_pending'):
                doc = db.session.get(Document, p.get('document_id'))
                if not doc:
                    continue
                try:
                    ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
                    ex.pop('_unlink_pending', None)
                    doc.extracted_data = json.dumps(ex)
                    db.session.commit()
                    return {'name': 'support docs', 'role': 'kept all',
                            'kind': 'unlink_done'}
                except Exception:
                    db.session.rollback()
                    return None
        return None

    # 'unlink <doc_id>'
    parts = t.split(maxsplit=1)
    if len(parts) != 2:
        return None
    doc_id = parts[1].strip()
    doc = db.session.get(Document, doc_id)
    if not doc or doc.client_id != client_id:
        return None
    # Move out of the property cluster: clear identifying fields so
    # _property_group_key returns '' and gift_walker treats it as orphan.
    try:
        ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
    except (json.JSONDecodeError, TypeError):
        ex = {}
    ex['_unlinked_from_property'] = True
    for k in ('title_number', 'lot_number', 'mukim', 'property_address',
              'description', 'property_hint'):
        ex.pop(k, None)
    # Drop into chat_inbox so it doesn't pollute another cluster
    doc.category = 'chat_inbox'
    try:
        doc.extracted_data = json.dumps(ex)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {'name': doc.original_filename or 'doc', 'role': 'unlinked',
            'kind': 'unlink_one'}


def _try_handle_inventory_action(client_id: str, user_text: str):
    """Per-asset walk-through actions issued by the inventory card:

      • 'inventory confirm' → the writer approves THIS asset; mark its
        Document as inventoried (extracted_data._inventoried = True) so
        it disappears from the walk and the next un-reviewed asset
        comes up.
      • 'inventory skip'    → leave the doc untouched but mark inventoried
        anyway so the walk can progress. Writer can revisit later.
      • 'inventory unlink'  → switch into "wrong supporting docs" picker
        mode — emit a list of the support docs with individual remove
        buttons. (Picker UI in next slice; for now stamp a flag the
        planner can read.)

    Each call resolves ONE focused asset (the same priority order the
    walk-through uses: property → bank → vehicle).
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    # Accept the four explicit prefixes AND the bare "delete"/"remove"
    # tokens that the 🗑 quick-reply on the property/bank/vehicle review
    # cards emits. Without the bare-token branch the delete button was a
    # no-op (it fell through to the IC-only delete handler, which returned
    # None because the user is past Step 1).
    _is_delete = (t == 'delete' or t == 'remove'
                  or t.startswith('delete ') or t.startswith('remove ')
                  or t.startswith('inventory delete')
                  or t.startswith('inventory remove'))
    # "inventory ownership …" and "inventory encumbered …" are sub-commands
    # issued by the guided confirm gate (ownership prompt / encumbrance prompt).
    _is_ownership_gate   = t.startswith('inventory ownership')
    _is_encumbrance_gate = t.startswith('inventory encumbered')
    # "inventory nlc …" sub-commands for Gate 3 (NLC completeness).
    _is_nlc_gate         = t.startswith('inventory nlc')
    if not (t.startswith('inventory confirm')
            or t.startswith('inventory skip')
            or t.startswith('inventory unlink')
            or _is_delete
            or _is_ownership_gate
            or _is_encumbrance_gate
            or _is_nlc_gate):
        return None

    from services.gift_walker import get_pending_gift_documents
    pend = get_pending_gift_documents(client_id)
    target = None
    target_kind = None
    for kind in ('property', 'bank', 'vehicle'):
        items = pend.get(kind) or []
        # Pick the first NOT-yet-inventoried one (matches what the chat
        # card showed the writer this turn).
        for it in items:
            if not (it.get('extracted') or {}).get('_inventoried'):
                target = it
                target_kind = kind
                break
        if target:
            break
    if not target:
        return None

    doc = db.session.get(Document, target['document_id'])
    if not doc:
        return None

    if _is_delete:
        # 🗑 Remove this property/bank/vehicle. Soft-delete the focused
        # Document AND any auto-grouped support docs so the writer doesn't
        # have to delete each piece individually. Also stamp _inventoried
        # so the walk-through skips past this slot on the next render.
        try:
            ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            ex = {}
        ex['_inventoried'] = True
        ex['_deleted_by_user'] = True
        try:
            doc.extracted_data = json.dumps(ex)
            doc.category = 'deleted'
            doc.description = (doc.description or '') + ' (removed via chat walk-through)'
            removed_count = 1
            # Cascade-delete any support docs the planner had grouped under
            # this asset (SPA / cukai / utility / extra geran pages — all
            # tied by property_hint or filename neighbourhood).
            for s in (target.get('support_docs') or []):
                sd = db.session.get(Document, s.get('document_id'))
                if not sd:
                    continue
                try:
                    sex = json.loads(sd.extracted_data) if sd.extracted_data else {}
                except (json.JSONDecodeError, TypeError):
                    sex = {}
                sex['_inventoried'] = True
                sex['_deleted_by_user'] = True
                sd.extracted_data = json.dumps(sex)
                sd.category = 'deleted'
                sd.description = (sd.description or '') + ' (cascade-removed with parent property)'
                removed_count += 1
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        if target_kind == 'property':
            label = (ex.get('property_address') or ex.get('title_number')
                     or ex.get('property_hint') or 'property')
        elif target_kind == 'bank':
            label = (ex.get('bank_name') or 'bank account')
        else:
            label = (ex.get('description') or ex.get('vehicle_make') or 'vehicle')
        return {'name': label[:80],
                'role': f'removed ({removed_count} doc{"s" if removed_count != 1 else ""})',
                'kind': f'inventory_deleted_{target_kind}'}

    if t.startswith('inventory unlink'):
        # Stamp a transient marker so the planner shows the support-doc
        # picker on its next render. Real picker UI lands next slice.
        try:
            ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            ex = {}
        ex['_unlink_pending'] = True
        try:
            doc.extracted_data = json.dumps(ex)
            db.session.commit()
        except Exception:
            db.session.rollback()
        label = (ex.get('property_address') or ex.get('title_number')
                 or 'this asset')
        return {'name': label[:80], 'role': 'review supporting docs',
                'kind': 'inventory_unlink_pending'}

    # ── Address gate (property confirm only) ──────────────────────────
    # A Malaysian will property clause MUST describe the property — even
    # a vague address like "No. 5, Jalan Maju, Subang" is essential so
    # the Executor knows which property to deal with.
    #
    # If the writer taps "✅ Looks right" but the title has NO address,
    # prompt them to type it rather than silently saving a blank gift.
    # They can still override with "inventory confirm no address" or by
    # typing `address: leave blank` if they intentionally want it empty.
    _is_confirm = t.startswith('inventory confirm')
    _force_no_addr = ('no address' in t or 'leave blank' in t or
                      'leave address' in t or 'skip address' in t)
    if _is_confirm and target_kind == 'property' and not _force_no_addr:
        try:
            _ex_check = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            _ex_check = {}
        _addr = (_ex_check.get('property_address') or '').strip()
        if not _addr:
            # Address missing — ask: type it now, or confirm without.
            # Two quick-reply buttons keep it simple; no upfront card clutter.
            import json as _json
            _qr = [
                {'label': '✏️ Type address now', 'value': 'address: '},
                {'label': '🏚 No street address — confirm anyway', 'value': 'inventory confirm no address'},
            ]
            _qr_marker = f'<!--quickreplies:{_json.dumps(_qr)}-->'
            return {
                'name': 'address missing',
                'role': 'address_required',
                'kind': 'property_fill',
                'reply_override': (
                    "**⚠️ No property address found.**\n\n"
                    "Address is needed for the will clause so the Executor knows which property to deal with.\n\n"
                    "  • Type it: `address: No. 22, Jalan Rimbun, Taman Seri Alam, Johor`\n"
                    "  • Or tap below if this is agricultural / industrial land with no street address."
                    + _qr_marker
                ),
            }

    # ── Guided confirm gates (property only) ────────────────────────────
    # Sequential 3-step flow triggered by tapping Accept:
    #
    #   Step 1 — Ownership type:  2 buttons  (Sole / Joint)
    #   Step 1b — If Joint, share: 3 buttons  (1/2 / 1/3 / type manually)
    #   Step 2 — Encumbrance:     2 buttons  (Clean / Has loan or caveat)
    #
    # Sub-commands "inventory ownership …" and "inventory encumbered …" are
    # emitted by the quick-reply buttons. They save the answer then re-enter
    # gate logic to show the next prompt or finalise.
    _is_confirm = t.startswith('inventory confirm')

    if target_kind == 'property' and not t.startswith('inventory skip'):
        try:
            _gex = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            _gex = {}

        # ── Parse and save ownership answer ──────────────────────────
        if _is_ownership_gate:
            _ow_rest = t[len('inventory ownership'):].strip()
            if 'sole' in _ow_rest:
                _gex['ownership_type']  = 'sole'
                _gex['ownership_share'] = ''
            elif _ow_rest == 'joint':
                # "joint" with no share yet → show share picker (Step 1b)
                # Don't save yet; return the share prompt immediately.
                _qr_share = [
                    {'label': '1/2 share',      'value': 'inventory ownership joint 1/2'},
                    {'label': '1/3 share',      'value': 'inventory ownership joint 1/3'},
                    {'label': '✏️ Other — type', 'value': 'inventory ownership joint '},
                ]
                return {
                    'name': 'joint share',
                    'role': 'ownership_gate',
                    'kind': 'property_fill',
                    'reply_override': (
                        "**🤝 Joint ownership — what is the testator's share?**\n\n"
                        "_(Select the undivided share, or tap Other and type e.g. `inventory ownership joint 2/5`)_"
                        + f'<!--quickreplies:{json.dumps(_qr_share)}-->'
                    ),
                }
            else:
                # "joint 1/2", "joint 1/3", "joint 2/5" etc.
                _sh = re.search(r'(\d+/\d+|\d+\s*%)', _ow_rest)
                _gex['ownership_type']  = 'joint'
                _gex['ownership_share'] = _sh.group(1).strip() if _sh else _ow_rest.replace('joint', '').strip()
            _gex.setdefault('_manually_edited', [])
            _ow_tag = f"ownership={_gex['ownership_type']}"
            if _gex.get('ownership_share'):
                _ow_tag += '/' + _gex['ownership_share']
            if isinstance(_gex['_manually_edited'], list):
                _gex['_manually_edited'].append(_ow_tag)
            try:
                doc.extracted_data = json.dumps(_gex)
                db.session.commit()
            except Exception:
                db.session.rollback()

        # ── Parse and save encumbrance answer ────────────────────────
        if _is_encumbrance_gate:
            _enc_rest = t[len('inventory encumbered'):].strip()
            if _enc_rest in ('clean', 'no', 'none', 'false'):
                _gex['encumbrance_confirmed'] = False
                _gex['encumbrance_type']      = ''
            elif 'charge' in _enc_rest or 'mortgage' in _enc_rest or 'loan' in _enc_rest:
                _gex['encumbrance_confirmed'] = True
                _gex['encumbrance_type']      = 'charge'
            elif 'caveat' in _enc_rest:
                _gex['encumbrance_confirmed'] = True
                _gex['encumbrance_type']      = 'caveat'
            else:
                _gex['encumbrance_confirmed'] = True
                _gex['encumbrance_type']      = 'charge'  # default to charge if ambiguous
            try:
                doc.extracted_data = json.dumps(_gex)
                db.session.commit()
            except Exception:
                db.session.rollback()

        # ── Gate 1: Ownership type not yet set ───────────────────────
        _ow_type = (_gex.get('ownership_type') or '').strip().lower()
        if not _ow_type and (_is_confirm or _is_ownership_gate):
            _num_own = _gex.get('num_owners') or 1
            try:
                _num_own = int(_num_own)
            except (TypeError, ValueError):
                _num_own = 1
            _ocr_shares = (_gex.get('ownership_shares') or '').strip()
            _ocr_hint   = ''
            if _num_own > 1 or _ocr_shares:
                _ocr_hint = (
                    f"\n\n_OCR detected {_num_own} registered owners"
                    + (f" ({_ocr_shares})" if _ocr_shares else '')
                    + " — likely joint ownership._"
                )

            # 🔥 §10x.95 v2 — use the canonical asset_pipeline binding
            # (per §10hf web-search + §10ha geo bridge + §10hc mukim
            # resolver) to find which AI Summary entry this doc belongs
            # to. The pipeline's bind_assets() runs the full Tier A/B/C
            # cascade. v1 used lexical substring on `_gex_addr in _ap_addr`
            # which silently failed for narrative-format AI Summaries
            # that omit lot/title (user typed only addresses) — see
            # §10x.46 R4 ("FIND THE ROOT CAUSE. DON'T JUST PATCH").
            _ai_hint = ''
            _ai_match_value = ''
            _ai_match_label = ''
            try:
                from ai.chat_planner import _extract_ai_summary_properties
                from services.asset_pipeline import (parse_canonical_assets,
                                                       group_documents,
                                                       bind_assets)
                _ai_props = _extract_ai_summary_properties(client_id) or []
                _items   = parse_canonical_assets(client_id)
                _groups  = group_documents(client_id)
                _bindings = bind_assets(_items, _groups)
                # Find the DocGroup containing this doc, then its Binding
                _doc_group = next(
                    (g for g in _groups if doc.id in g.document_ids), None
                )
                _matched_ai = None
                _matched_via = ''
                if _doc_group:
                    _binding = next(
                        (b for b in _bindings if b.group_id == _doc_group.group_id),
                        None,
                    )
                    if _binding and _binding.tier in ('A', 'B', 'C'):
                        if 0 <= _binding.ai_index < len(_ai_props):
                            _matched_ai = _ai_props[_binding.ai_index]
                            _matched_via = _binding.match_via
                # Fallback for completeness: if pipeline didn't bind (Tier D
                # or no group), keep the lexical substring as a last-resort
                # signal so we don't strictly regress vs v1. This only
                # fires when the pipeline already gave up.
                if not _matched_ai:
                    _gex_addr = (_gex.get('property_address') or '').strip().upper()
                    _gex_lot  = re.sub(r'\D', '', (_gex.get('lot_number') or ''))
                    _gex_title = re.sub(r'\D', '', (_gex.get('title_number') or ''))
                    for _ap in _ai_props:
                        _ap_lot   = re.sub(r'\D', '', _ap.get('lot') or '')
                        _ap_title = re.sub(r'\D', '', _ap.get('title') or '')
                        _ap_addr  = (_ap.get('address') or '').upper()
                        if (_gex_lot and _gex_lot == _ap_lot) \
                           or (_gex_title and _gex_title == _ap_title) \
                           or (_gex_addr and _ap_addr and
                               (_gex_addr in _ap_addr or _ap_addr in _gex_addr)):
                            _matched_ai = _ap
                            _matched_via = 'lexical_fallback'
                            break
                if _matched_ai:
                    own_text = ((_matched_ai.get('ownership') or '') + ' '
                                + (_matched_ai.get('beneficiary') or '')).lower()
                    if any(p in own_text for p in (
                        'jointly', 'joint with', 'share with', 'co-owned',
                        '50/50', '1/2', '50 percent')):
                        _ai_match_value = 'inventory ownership joint'
                        _ai_match_label = '🤝 Joint owner (per your message)'
                        _ai_hint = (
                            f"\n\n📨 _from your message:_ \""
                            f"{(_matched_ai.get('ownership') or _matched_ai.get('beneficiary') or '')[:140]}\""
                        )
                    elif 'sole' in own_text or '100%' in own_text \
                         or '100percent' in own_text or '1/1' in own_text:
                        _ai_match_value = 'inventory ownership sole'
                        _ai_match_label = '👤 Sole owner (per your message)'
                        _ai_hint = (
                            f"\n\n📨 _from your message:_ \""
                            f"{(_matched_ai.get('ownership') or _matched_ai.get('beneficiary') or '')[:140]}\""
                        )
            except Exception:
                pass

            if _ai_match_value:
                # AUTO-DEDUCED — confirm button first, alternates after
                _qr_ow = [
                    {'label': f'✅ {_ai_match_label} — Confirm',
                     'value': _ai_match_value},
                    {'label': '👤 Sole owner instead',
                     'value': 'inventory ownership sole' if _ai_match_value.endswith('joint')
                              else 'inventory ownership joint'},
                ]
            else:
                _qr_ow = [
                    {'label': '👤 Sole owner',  'value': 'inventory ownership sole'},
                    {'label': '🤝 Joint owner', 'value': 'inventory ownership joint'},
                ]
            return {
                'name': 'ownership',
                'role': 'ownership_gate',
                'kind': 'property_fill',
                'reply_override': (
                    "**Step 1 of 2 — Ownership**\n\n"
                    "Is the testator the **sole owner**, or is it **jointly owned** with another person?"
                    + _ai_hint
                    + _ocr_hint
                    + f'<!--quickreplies:{json.dumps(_qr_ow)}-->'
                ),
            }

        # ── Gate 2: Encumbrance not yet confirmed ────────────────────
        _enc_confirmed = _gex.get('encumbrance_confirmed')  # None = not yet answered
        if _enc_confirmed is None and (_is_confirm or _is_ownership_gate or _is_encumbrance_gate):
            _enc_ocr      = (_gex.get('encumbrance') or '').strip()
            _enc_type_ocr = (_gex.get('encumbrance_type') or '').strip().lower()
            _enc_hint     = ''
            if _enc_ocr or _enc_type_ocr:
                _enc_icon = '🏦' if _enc_type_ocr == 'charge' else '🚩'
                _enc_hint = f"\n\n_{_enc_icon} OCR detected: {_enc_ocr[:150]}_"
            _qr_enc = [
                {'label': '✅ Clean title',       'value': 'inventory encumbered clean'},
                {'label': '🏦 Has loan or caveat', 'value': 'inventory encumbered charge'},
            ]
            _enc_marker = f'<!--quickreplies:{json.dumps(_qr_enc)}-->'
            return {
                'name': 'encumbrance',
                'role': 'encumbrance_gate',
                'kind': 'property_fill',
                'reply_override': (
                    "**Step 2 of 2 — Encumbrance**\n\n"
                    "Is there a **bank loan or caveat** registered on this property?"
                    + _enc_hint
                    + "\n\n_(The Executor will be directed to settle the loan or withdraw the caveat.)_"
                    + _enc_marker
                ),
            }

        # ── Gate 3: NLC completeness ──────────────────────────────────
        # Required National Land Code fields for a valid will gift clause.
        # Fires after ownership + encumbrance are both answered.
        # Writer can enter them now (fill prompt) or defer to later.
        _NLC_REQUIRED = [
            ('title_number', 'Title/Geran No.',  'title: HS(D) 12345/2005'),
            ('lot_number',   'Lot/PTD No.',       'lot: 12345'),
            ('mukim',        'Mukim',             'mukim: Tebrau'),
            ('daerah',       'Daerah',            'daerah: Johor Bahru'),
            ('negeri',       'Negeri/State',      'negeri: Johor'),
        ]
        if _is_nlc_gate:
            if t == 'inventory nlc skip':
                _gex['_nlc_deferred'] = True
                try:
                    doc.extracted_data = json.dumps(_gex)
                    db.session.commit()
                except Exception:
                    db.session.rollback()
                # Fall through to stamp _inventoried
            elif t == 'inventory nlc fill':
                _missing_fields = [
                    (key, lbl, eg) for key, lbl, eg in _NLC_REQUIRED
                    if not (_gex.get(key) or '').strip()
                ]
                _fill_lines = '\n'.join(
                    f"  • `{eg}`  →  {lbl}" for _, lbl, eg in _missing_fields
                )
                _qr_nlc_fill = [{'label': '⏭ Fill in later', 'value': 'inventory nlc skip'}]
                return {
                    'name': 'nlc fill prompt',
                    'role': 'nlc_gate',
                    'kind': 'property_fill',
                    'reply_override': (
                        "**📋 Type the missing land registry details:**\n\n"
                        + _fill_lines
                        + "\n\n_Reply with each field (e.g. `mukim: Tebrau`). "
                        "Tap Accept on the property card again when done._"
                        + f'<!--quickreplies:{json.dumps(_qr_nlc_fill)}-->'
                    ),
                }
        else:
            # Not an NLC sub-command — check if gate should fire naturally
            # (i.e. ownership + encumbrance answered but NLC fields still missing).
            if not _gex.get('_nlc_deferred'):
                _ow_type_now  = (_gex.get('ownership_type') or '').strip()
                _enc_now      = _gex.get('encumbrance_confirmed')
                if _ow_type_now and _enc_now is not None:
                    _missing_nlc = [
                        (key, lbl, eg) for key, lbl, eg in _NLC_REQUIRED
                        if not (_gex.get(key) or '').strip()
                    ]
                    if _missing_nlc:
                        _missing_labels = ', '.join(lbl for _, lbl, _ in _missing_nlc)
                        _qr_nlc = [
                            {'label': '✏️ Enter now',    'value': 'inventory nlc fill'},
                            {'label': '⏭ Fill in later', 'value': 'inventory nlc skip'},
                        ]
                        return {
                            'name': 'nlc check',
                            'role': 'nlc_gate',
                            'kind': 'property_fill',
                            'reply_override': (
                                f"**📋 Missing land registry details**\n\n"
                                f"The will clause needs: **{_missing_labels}**.\n\n"
                                f"Do you have this information to hand?"
                                + f'<!--quickreplies:{json.dumps(_qr_nlc)}-->'
                            ),
                        }

    # 'inventory confirm' or 'inventory skip' → mark inventoried
    try:
        ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
    except (json.JSONDecodeError, TypeError):
        ex = {}
    ex['_inventoried'] = True
    if t.startswith('inventory skip'):
        ex['_skipped'] = True
    try:
        doc.extracted_data = json.dumps(ex)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    # Friendly ack label
    if target_kind == 'property':
        label = (ex.get('property_address') or ex.get('title_number')
                 or 'property')
    elif target_kind == 'bank':
        label = (ex.get('bank_name') or 'bank account')
    else:
        label = (ex.get('description') or ex.get('vehicle_make') or 'vehicle')

    action = 'skipped' if t.startswith('inventory skip') else 'reviewed'

    # ── Placeholder in step5_data ─────────────────────────────────────────
    # Insert a "pending beneficiary" placeholder into step5_data immediately
    # when a property is accepted so it shows up in the wizard right-pane
    # snapshot. Phase B of _try_save_property_gift will upsert (replace) it
    # with the full gift entry once beneficiary + substitute are assigned.
    # Skip placeholder if the property is explicitly skipped or non-property.
    if target_kind == 'property' and not t.startswith('inventory skip'):
        will_for_ph = (Will.query.filter_by(client_id=client_id, status='draft')
                       .filter(Will.deleted_at.is_(None))
                       .order_by(Will.updated_at.desc()).first())
        if will_for_ph:
            try:
                gifts_ph = json.loads(will_for_ph.step5_data or '[]')
                if not isinstance(gifts_ph, list):
                    gifts_ph = []
            except (json.JSONDecodeError, TypeError):
                gifts_ph = []
            # ╔══════════════════════════════════════════════════════════╗
            # ║  🔥 BURN-IN — NO DUPLICATE GIFTS IN step5_data 🔥          ║
            # ║  Dedup on document_id AND on (lot, address) signature so   ║
            # ║  two different Document rows for the SAME physical         ║
            # ║  property (OCR title drift: 564662 vs 504662) cannot       ║
            # ║  produce two gift entries. See CLAUDE.md §10f.             ║
            # ╚══════════════════════════════════════════════════════════╝
            from services.gift_walker import (_clean_id_value, _looks_like_garbage,
                                              _norm_addr, _is_strata, _title_signature,
                                              _is_genuinely_different_unit)
            existing_ids = {g.get('document_id') for g in gifts_ph}
            new_lot = _clean_id_value(ex.get('lot_number', '') or '')
            if _looks_like_garbage(new_lot):
                new_lot = ''
            new_lot_digits = re.sub(r'\D', '', new_lot)
            new_addr_sig = _norm_addr(ex.get('property_address', '') or '')[:60]
            new_strata = _is_strata(ex)
            new_title_sig = _title_signature(ex)
            def _gift_field(g, key):
                return ((g.get('property_info') or {}).get(key)
                        or (g.get('property_details') or {}).get(key)
                        or g.get(key) or '')
            duplicate = doc.id in existing_ids
            if not duplicate:
                for g in gifts_ph:
                    g_lot = _clean_id_value(_gift_field(g, 'lot_number'))
                    if _looks_like_garbage(g_lot):
                        g_lot = ''
                    g_lot_digits = re.sub(r'\D', '', g_lot)
                    g_addr_sig = _norm_addr(_gift_field(g, 'property_address'))[:60]
                    # 🔥 STRATA EXCEPTION (§10hd): same lot+addr but genuinely
                    # different title signature (not OCR truncation) → different
                    # unit in same building. Not a duplicate.
                    g_ex = {'title_number': _gift_field(g, 'title_number'),
                            'title_type':   _gift_field(g, 'title_type'),
                            'property_description': _gift_field(g, 'property_description'),
                            'document_type': _gift_field(g, 'document_type')}
                    g_strata = _is_strata(g_ex)
                    g_title_sig = _title_signature(g_ex)
                    strata_diff_units = (
                        (new_strata or g_strata)
                        and _is_genuinely_different_unit(new_title_sig, g_title_sig)
                    )
                    if strata_diff_units:
                        continue   # different strata unit — not a duplicate
                    # Match on lot+address (both non-empty and equal)
                    if (new_lot_digits and g_lot_digits and new_lot_digits == g_lot_digits
                        and new_addr_sig and g_addr_sig and new_addr_sig == g_addr_sig):
                        duplicate = True
                        break
                    # Or just identical addresses if both lots are empty/garbage
                    if (not new_lot_digits and not g_lot_digits
                        and new_addr_sig and g_addr_sig and new_addr_sig == g_addr_sig):
                        duplicate = True
                        break
            if not duplicate:
                # 🔥 §10x.100 — stamp _ai_summary_idx on the placeholder
                # so _ai_props_already_handled (Pass 1) sees this slot as
                # taken. Without this tag, the H3 path later re-asks about
                # this same AI Summary property and appends a duplicate
                # gift. See bug table entry on phantom property gifts.
                _ai_idx_for_placeholder = None
                _doc_in_orphan_group = False
                try:
                    from services.asset_pipeline import (parse_canonical_assets,
                                                          group_documents,
                                                          bind_assets)
                    _ai_items = parse_canonical_assets(client_id)
                    _doc_groups = group_documents(client_id)
                    _ai_bindings = bind_assets(_ai_items, _doc_groups)
                    _doc_grp = next(
                        (g for g in _doc_groups if doc.id in g.document_ids), None
                    )
                    if _doc_grp:
                        _b = next(
                            (b for b in _ai_bindings
                             if b.group_id == _doc_grp.group_id),
                            None,
                        )
                        if _b and _b.tier in ('A', 'B', 'C'):
                            _ai_idx_for_placeholder = _b.ai_index
                        else:
                            # 🔥 §10x.105 — orphan group detection.
                            # `_b is None` (no binding references this
                            # group) OR `_b.tier == 'D'` (binding present
                            # but unmatched). Both mean: pipeline can't
                            # tie this doc to any AI Summary slot.
                            # Refuse to save the placeholder; mark the doc
                            # _inventoried so it doesn't reappear in
                            # pending. Avoids the Marina Cove `0664e07a`
                            # phantom in DocGroup `3b89a4a2`.
                            _doc_in_orphan_group = True
                except Exception:
                    pass

                if _doc_in_orphan_group:
                    # Mark inventoried so walker advances past this doc;
                    # skip the step5_data placeholder insert. The doc
                    # remains attached to the chat for evidence but
                    # doesn't pollute the gift list with a phantom entry.
                    try:
                        ex['_inventoried'] = True
                        ex['_orphan_group_skipped'] = True
                        doc.extracted_data = json.dumps(ex)
                        db.session.commit()
                    except Exception:
                        try: db.session.rollback()
                        except Exception: pass

                placeholder = {
                    'document_id':      doc.id,
                    'kind':             'property',
                    'gift_type':        'property',
                    '_pending_beneficiary': True,   # Phase B will replace this
                    # 🔥 §10x.96 — every Document-bound gift MUST carry
                    # _match_via per §10he Step 5 / verifier R12. The
                    # 'inventory confirm' path means the user explicitly
                    # accepted this property card → match_via='user_confirmed'.
                    '_match_via':       'user_confirmed',
                    '_match_tier':      'A',
                    '_match_confidence': 'high',
                    'property_address': ex.get('property_address', ''),
                    'title_number':     ex.get('title_number', ''),
                    'lot_number':       ex.get('lot_number', ''),
                    'property_details': {
                        'property_address': ex.get('property_address', ''),
                        'title_number':     ex.get('title_number', ''),
                        'lot_number':       ex.get('lot_number', ''),
                        'mukim':            ex.get('mukim', ''),
                        'daerah':           ex.get('daerah', ''),
                        'negeri':           ex.get('negeri', ''),
                    },
                    'allocations':   [],
                    'beneficiaries': [],
                }
                # Only set _ai_summary_idx when pipeline gave us a real
                # binding (Tier A/B/C). Tier D or no binding → leave None
                # so it doesn't falsely claim a slot.
                if _ai_idx_for_placeholder is not None:
                    placeholder['_ai_summary_idx'] = _ai_idx_for_placeholder
                # 🔥 §10x.105 — orphan groups (Tier D) skip placeholder
                # insert. Doc was marked _inventoried above; walker
                # advances past it without polluting step5_data.
                if not _doc_in_orphan_group:
                    gifts_ph.append(placeholder)
                will_for_ph.step5_data = json.dumps(gifts_ph)
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()

    # Auto-stamp `assets_confirmed` if this was the LAST un-reviewed asset.
    # Saves the writer from having to type "confirm assets" at the end —
    # walk-through completion IS the confirmation.
    # ╔══════════════════════════════════════════════════════════════════╗
    # ║  🔥 BURN-IN §10hg — DO NOT AUTO-STAMP UNTIL H3 PROPERTIES HANDLED ║
    # ║  AI Summary may name properties with NO image attached. Those H3  ║
    # ║  cards still need to surface for explicit user Confirm/Skip       ║
    # ║  before we lock the inventory. Counting only Document-row pendings║
    # ║  hides H3 candidates and races past them straight to bank/exec.   ║
    # ╚══════════════════════════════════════════════════════════════════╝
    pend_after = get_pending_gift_documents(client_id)
    any_left = False
    for k in ('property', 'bank', 'vehicle'):
        for it in (pend_after.get(k) or []):
            if not (it.get('extracted') or {}).get('_inventoried'):
                any_left = True
                break
        if any_left:
            break
    # H3 check — any AI-Summary property with no matching image group AND
    # not yet represented in step5_data (placeholder or skip).
    if not any_left:
        try:
            from ai.chat_planner import (_extract_ai_summary_properties,
                                          _classify_property_match,
                                          _ai_props_already_handled)
            _ai_props_h3 = _extract_ai_summary_properties(client_id)
            if _ai_props_h3:
                _will_h3 = (Will.query.filter_by(client_id=client_id, status='draft')
                            .filter(Will.deleted_at.is_(None))
                            .order_by(Will.updated_at.desc()).first())
                _wd_h3 = {'step5': []}
                if _will_h3 and _will_h3.step5_data:
                    try:
                        _s5h = json.loads(_will_h3.step5_data)
                        _wd_h3['step5'] = _s5h if isinstance(_s5h, list) else []
                    except Exception:
                        pass
                _handled_h3 = _ai_props_already_handled(client_id, _ai_props_h3, _wd_h3)
                _all_props_h3 = pend_after.get('property') or []
                _claimed = set()
                _matched = []
                for ap in _ai_props_h3:
                    avail = [g for g in _all_props_h3
                             if g.get('document_id') not in _claimed]
                    cls = _classify_property_match(ap, avail)
                    if cls['variant'] in ('h1', 'h2') and cls.get('group'):
                        _matched.append(True)
                        _claimed.add(cls['group'].get('document_id'))
                    else:
                        _matched.append(False)
                _h3_unhandled = any(
                    (not _handled_h3[i]) and (not _matched[i])
                    for i in range(len(_ai_props_h3))
                )
                if _h3_unhandled:
                    any_left = True
        except Exception:
            pass
    if not any_left:
        will = (Will.query.filter_by(client_id=client_id, status='draft')
                .filter(Will.deleted_at.is_(None))
                .order_by(Will.updated_at.desc()).first())
        if will:
            try:
                completed = json.loads(will.completed_steps or '[]')
            except (json.JSONDecodeError, TypeError):
                completed = []
            if not isinstance(completed, list):
                completed = []
            if 'assets_confirmed' not in completed:
                completed.append('assets_confirmed')
                will.completed_steps = json.dumps(completed)
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()

    return {'name': label[:80], 'role': f'{action} & queued for wizard',
            'kind': f'inventory_{action}_{target_kind}'}


_ASSETS_CONFIRM_TOKENS = (
    'confirm assets', 'thats everything', "that's everything",
    'yes thats all', "yes that's all", 'all uploaded',
    'confirm', 'done', 'looks good', 'looks right',
    "i'll skip specific gifts", 'skip specific gifts', 'no specific gifts',
)
_ASSETS_MORE_TOKENS = (
    'i have more to upload', 'more to upload', 'add more',
    'upload more', 'not done', 'wait', 'one more',
)


# ╔════════════════════════════════════════════════════════════════════════╗
# ║  🔥 BURN-IN §10hg — H3 + CONFLICT HANDLERS                              ║
# ║                                                                          ║
# ║  H3 = AI-Summary property with no image evidence. The placeholder card  ║
# ║  emits 'inventory confirm' / 'inventory skip'. The normal inventory     ║
# ║  handler skips H3 because it requires a Document target. This handler   ║
# ║  catches it and records the AI-Summary slot into step5_data with a     ║
# ║  _h3_placeholder flag (or _ai_summary_skipped flag).                    ║
# ╚════════════════════════════════════════════════════════════════════════╝
def _try_handle_h3_user_match(client_id: str, user_text: str):
    """🔥 §10x.51 Path Y — User confirmed a candidate match from the
    candidate-with-confirm card.

    Click value format: `inventory match h3 <ai_idx> <doc_id>`

    Build the gift via the §10x.48 pipeline (build_gift) but with a
    user_confirmed binding: tier='A', match_via='user_confirmed',
    confidence='high'. Persist to step5_data with _layer1_confirmed=True
    so the next walkthrough turn moves to Layer 2 (main beneficiary).
    """
    if not user_text:
        return None
    t = (user_text or '').strip()
    m = re.match(r'^inventory\s+match\s+h3\s+(\d+)\s+(\S+)$', t, re.IGNORECASE)
    if not m:
        return None
    try:
        ai_idx = int(m.group(1))
    except ValueError:
        return None
    doc_id = m.group(2).strip()

    doc = Document.query.filter_by(id=doc_id, client_id=client_id).first()
    if not doc:
        return None
    active_will = (Will.query.filter_by(client_id=client_id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc()).first())
    if not active_will:
        return None

    try:
        from services.asset_pipeline import (parse_canonical_assets,
                                                group_documents,
                                                build_gift, Binding)
    except Exception:
        return None

    items = parse_canonical_assets(client_id)
    target_ai = next((a for a in items if a.ai_index == ai_idx), None)
    if not target_ai:
        return None
    groups = group_documents(client_id)
    target_group = next((g for g in groups if doc_id in g.document_ids), None)
    if not target_group:
        return None
    binding = Binding(
        ai_index=ai_idx, group_id=target_group.group_id,
        tier='A', match_via='user_confirmed', confidence='high',
        evidence='User confirmed candidate match',
    )
    entry = build_gift(target_ai, binding, target_group)
    entry['_layer1_confirmed'] = True
    entry['_user_confirmed_match'] = True
    entry['_ai_summary_idx'] = ai_idx
    entry['beneficiaries'] = entry.get('beneficiaries') or []
    pi = entry.get('property_info') or {}
    entry.setdefault('testator_share', pi.get('testator_share') or '1/1')
    entry.setdefault('address', pi.get('property_address') or '')

    # Append to step5_data; remove any prior placeholder entry for same ai_idx
    try:
        s5 = json.loads(active_will.step5_data) if active_will.step5_data else []
        if not isinstance(s5, list):
            s5 = []
    except Exception:
        s5 = []
    s5 = [g for g in s5 if g.get('_ai_summary_idx') != ai_idx]
    s5.append(entry)
    active_will.step5_data = json.dumps(s5)
    db.session.commit()

    return {
        'kind': 'inventory_user_matched',
        'ai_idx': ai_idx,
        'doc_id': doc_id,
        'group_id': target_group.group_id,
        'name': (target_ai.fields.get('address') or '')[:60] or 'property',
    }


def _try_handle_doc_assign(client_id: str, user_text: str):
    """🔥 §10x.125 + §10x.126 — handle the user's reply to the
    'identify this image' card.

    Quickreply formats:
      `doc_assign <doc_id> <ai_index>`  → bind to AI Summary property
      `doc_assign skip <doc_id>`         → mark inventoried, skip
      `doc_assign remove <doc_id>`       → soft-delete the doc

    Returns dict on save, None to fall through.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if not t.startswith('doc_assign '):
        return None
    parts = t.split()
    if len(parts) < 3:
        return None

    # Variants:
    #   doc_assign <doc_id> <int>   → bind
    #   doc_assign skip <doc_id>    → skip
    #   doc_assign remove <doc_id>  → remove
    action = parts[1]
    if action == 'skip':
        doc_id = parts[2]
        # 🔥 §10x.126 — single Skip click bulk-skips ALL fully-unreadable
        # orphan docs for this client. User indicated they're not relevant
        # ("not in my will") — likely the same intent applies to all
        # similar unreadable images. Saves clicking Skip per orphan.
        all_docs = Document.query.filter_by(
            client_id=client_id, category='property_title',
        ).all()
        n_skipped = 0
        primary_name = ''
        for d in all_docs:
            try:
                ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
            except Exception:
                ex = {}
            if not isinstance(ex, dict):
                ex = {}
            # Skip THIS doc unconditionally
            if d.id == doc_id:
                primary_name = d.original_filename or doc_id
            # Bulk-skip applies only to OTHER fully-unreadable docs
            elif (ex.get('_inventoried') or ex.get('_skipped_not_in_will')):
                continue  # already skipped
            else:
                # Other docs: only bulk-skip if also fully-unreadable
                addr = (ex.get('property_address') or '').strip()
                title = (ex.get('title_number') or '').strip()
                lot = (ex.get('lot_number') or '').strip()
                mukim = (ex.get('mukim') or '').strip()
                if addr or title or lot or mukim:
                    continue   # this doc is identifiable — don't auto-skip
            ex['_inventoried'] = True
            ex['_skipped_not_in_will'] = True
            d.extracted_data = json.dumps(ex)
            n_skipped += 1
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        role = 'not in will'
        if n_skipped > 1:
            role = f'not in will (+ {n_skipped - 1} similar unreadable)'
        return {
            'kind': 'doc_skipped',
            'name': primary_name[:60],
            'role': role,
        }

    if action == 'remove':
        doc_id = parts[2]
        d = db.session.get(Document, doc_id)
        if not d:
            return None
        d.category = 'deleted'
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'doc_removed',
            'name': (d.original_filename or doc_id)[:60],
            'role': 'wrong upload',
        }

    # Default: bind path → 'doc_assign <doc_id> <ai_index>'
    doc_id = parts[1]
    try:
        ai_idx = int(parts[2])
    except (TypeError, ValueError):
        return None

    d = db.session.get(Document, doc_id)
    if not d:
        return None
    try:
        ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
    except Exception:
        ex = {}
    ex['_user_assigned_ai_idx'] = ai_idx
    ex.pop('_inventoried', None)
    ex.pop('_orphan_group_skipped', None)
    d.extracted_data = json.dumps(ex)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {
        'kind': 'doc_assigned',
        'name': (d.original_filename or doc_id)[:60],
        'role': f'AI[{ai_idx}]',
    }


def _try_handle_assets_check(client_id: str, user_text: str):
    """🔥 §10x.127 — handle the inventory-completeness gate.

    Quickreply values from `_step6_assets_complete_gate_card` and
    `_step6_describe_match_card`:

      • `assets_check yes`             → all assets accounted for; bulk-skip
                                          every isolated property image
      • `assets_check no`              → enter 'describing' mode; the next
                                          free-text user message is matched
                                          against isolated docs
      • `assets_check match <doc_id>`  → user picked a specific isolated
                                          image as the missing asset; assign
                                          it (clears describing marker)
      • `assets_check alternates`      → re-render the match card showing
                                          all candidates instead of the top
      • `assets_check noimage`         → none of the images match — show
                                          upload-prompt card
      • `assets_check text_only`       → save as a text-only asset (placeholder
                                          gift entry); clears markers

    Returns dict on success, None if not applicable.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if not t.startswith('assets_check '):
        return None

    parts = t.split()
    if len(parts) < 2:
        return None
    action = parts[1]

    will = Will.query.filter_by(client_id=client_id).first()
    if not will:
        return None

    try:
        completed = json.loads(will.completed_steps or '[]')
    except Exception:
        completed = []
    if not isinstance(completed, list):
        completed = []

    def _stamp(marker: str):
        if marker not in completed:
            completed.append(marker)

    def _clear(marker: str):
        while marker in completed:
            completed.remove(marker)

    # ── YES — bulk-skip all isolated property images ────────────────
    if action == 'yes':
        # Compute the set of doc IDs the asset_pipeline successfully binds
        # to AI Summary entries. ANY doc NOT in a bound group is "isolated"
        # (either fully unreadable or in an orphan group) and gets skipped.
        bound_doc_ids: set = set()
        try:
            from services.asset_pipeline import (parse_canonical_assets,
                                                  group_documents,
                                                  bind_assets)
            items = parse_canonical_assets(client_id)
            groups = group_documents(client_id)
            bindings = bind_assets(items, groups)
            bound_group_ids = {b.group_id for b in bindings
                               if b.tier in ('A', 'B', 'C')}
            for g in groups:
                if g.group_id in bound_group_ids:
                    for did in g.document_ids:
                        bound_doc_ids.add(did)
        except Exception:
            pass

        all_docs = Document.query.filter_by(
            client_id=client_id, category='property_title'
        ).all()
        n_skipped = 0
        for d in all_docs:
            try:
                ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
            except Exception:
                ex = {}
            if not isinstance(ex, dict):
                ex = {}
            if ex.get('_inventoried') or ex.get('_skipped_not_in_will'):
                continue
            addr = (ex.get('property_address') or '').strip()
            title = (ex.get('title_number') or '').strip()
            lot = (ex.get('lot_number') or '').strip()
            mukim = (ex.get('mukim') or '').strip()
            fully_unreadable = not (addr or title or lot or mukim)
            is_orphan = (bound_doc_ids and d.id not in bound_doc_ids)
            if not (fully_unreadable or is_orphan):
                continue   # legitimately bound — leave for normal walkthrough
            ex['_inventoried'] = True
            ex['_skipped_not_in_will'] = True
            ex['_auto_skipped_reason'] = (
                'user confirmed all assets accounted for (§10x.127)')
            d.extracted_data = json.dumps(ex)
            n_skipped += 1
        _stamp('assets_inventory_confirmed')
        _clear('assets_inventory_describing')
        # Drop any pending describe-state marker
        completed = [c for c in completed
                     if not (isinstance(c, str)
                             and c.startswith('assets_describe_pending:'))]
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'assets_inventory_confirmed',
            'name': '',
            'role': f'auto-skipped {n_skipped} extra image(s)',
        }

    # ── NO — enter describing mode ──────────────────────────────────
    if action == 'no':
        _stamp('assets_inventory_describing')
        _clear('assets_inventory_confirmed')
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'assets_inventory_describing',
            'name': '',
            'role': 'awaiting description',
        }

    # ── MATCH — user picked a specific image as the missing asset ──
    if action == 'match' and len(parts) >= 3:
        doc_id = parts[2]
        d = db.session.get(Document, doc_id)
        if not d:
            return None
        try:
            ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
        except Exception:
            ex = {}
        if not isinstance(ex, dict):
            ex = {}

        # Pull the user's most recent description from chat history so
        # we can populate the doc's address field (so the normal
        # walkthrough no longer treats it as fully-unreadable).
        try:
            from ai.chat_planner import _latest_user_description as _lud
            desc = _lud(client_id)
        except Exception:
            desc = ''
        if desc and not (ex.get('property_address') or '').strip():
            ex['property_address'] = desc[:200]
            ex['_address_source'] = 'user_described'

        # Tag this image as user-claimed; clear the skip/inventoried
        # flags so the doc surfaces as a regular pending property.
        ex['_user_claimed_via_describe'] = True
        ex.pop('_inventoried', None)
        ex.pop('_skipped_not_in_will', None)
        ex.pop('_auto_skipped_reason', None)
        d.extracted_data = json.dumps(ex)

        # Bulk-skip remaining isolated docs (user already chose THIS one
        # as the missing asset; the other isolated images are extras).
        all_docs = Document.query.filter_by(
            client_id=client_id, category='property_title'
        ).all()
        for od in all_docs:
            if od.id == d.id:
                continue
            try:
                oex = json.loads(od.extracted_data or '{}') if od.extracted_data else {}
            except Exception:
                oex = {}
            if not isinstance(oex, dict):
                oex = {}
            if oex.get('_inventoried') or oex.get('_skipped_not_in_will'):
                continue
            addr_o = (oex.get('property_address') or '').strip()
            tit_o = (oex.get('title_number') or '').strip()
            lot_o = (oex.get('lot_number') or '').strip()
            muk_o = (oex.get('mukim') or '').strip()
            if addr_o or tit_o or lot_o or muk_o:
                continue
            oex['_inventoried'] = True
            oex['_skipped_not_in_will'] = True
            oex['_auto_skipped_reason'] = (
                'user matched a different image; remaining are extras (§10x.127)')
            od.extracted_data = json.dumps(oex)

        # Clear describing markers
        _stamp('assets_inventory_confirmed')
        _clear('assets_inventory_describing')
        _clear('assets_describe_alternates')
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'assets_inventory_matched',
            'name': (d.original_filename or doc_id)[:60],
            'role': ('matched: ' + (desc[:40] if desc else 'user-described')),
        }

    # ── ALTERNATES — show all candidates next turn ──────────────────
    if action == 'alternates':
        _stamp('assets_describe_alternates')
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'assets_inventory_alternates',
            'name': '',
            'role': 'show alternates',
        }

    # ── NOIMAGE — none of the images match; ask user to upload ──────
    if action == 'noimage':
        _clear('assets_describe_alternates')
        # Stay in describing mode but hint that the user should upload
        # a clearer image. Frontend handles the upload trigger.
        return {
            'kind': 'assets_inventory_noimage',
            'name': '',
            'role': 'awaiting upload',
        }

    # ── TEXT_ONLY — clear markers; user will type asset details ─────
    if action == 'text_only':
        _stamp('assets_inventory_confirmed')
        _clear('assets_inventory_describing')
        completed = [c for c in completed
                     if not (isinstance(c, str)
                             and c.startswith('assets_describe_pending:'))]
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'assets_inventory_text_only',
            'name': '',
            'role': 'text-only asset acknowledged',
        }

    return None


def _try_handle_assets_describe(client_id: str, user_text: str):
    """🔥 §10x.127 — match free-text user input against isolated property
    docs when the user is in 'describing' mode.

    Triggered ONLY when `assets_inventory_describing` is in
    `Will.completed_steps`. Otherwise no-op.

    Tokenises `user_text`, scores each isolated property doc by overlap
    with extracted fields (lot / title / mukim / address / purpose /
    filename), and stamps a marker `assets_describe_pending:<doc_id>`
    so the planner renders the match card on the next turn.

    Returns dict on success, None if not applicable.
    """
    if not user_text:
        return None
    t = user_text.strip()
    # Skip quickreply prefixes — only free-text messages count as descriptions
    low = t.lower()
    qr_prefixes = ('assets_check', 'inventory ', 'doc_assign', 'orphan_',
                   'unlink ', 'gift ', 'guardian ', 'trust ', 'others ',
                   'residuary ', 'beneficiaries ', 'role_match', 'restart ',
                   'conflict ', 'address:', 'dob:', 'gender:', 'marital:',
                   'occupation:', 'daerah:', 'negeri:', 'mukim:', 'lot:',
                   'title:', 'property ', 'change ', 'confirm ', 'h3 ',
                   'bank_l1', 'bank_l2', 'bank_l3', 'bank_h3',
                   'insurance_l1', 'insurance_l2', 'insurance_l3', 'insurance_h3',
                   'banks generic', 'skip', 'delete', 'yes', 'no', 'remove',
                   'substitute ', 'others ', 'open wizard', 'upload-ic',
                   'walk one by one')
    if any(low.startswith(p) for p in qr_prefixes):
        return None
    # Need at least a few characters of substantive text
    if len(t) < 4:
        return None

    will = Will.query.filter_by(client_id=client_id).first()
    if not will:
        return None
    try:
        completed = json.loads(will.completed_steps or '[]')
    except Exception:
        completed = []
    if not isinstance(completed, list):
        completed = []
    if 'assets_inventory_describing' not in completed:
        return None

    # Tokenise the description
    import re as _re
    toks = set(_re.findall(r'[A-Za-z0-9]{3,}', t.upper()))
    # Drop noise words
    NOISE = {'AND', 'THE', 'FOR', 'WITH', 'FROM', 'INTO', 'THIS', 'THAT',
             'GIVE', 'GIVE', 'PROPERTY', 'ASSET', 'ASSETS', 'WILL', 'NOT',
             'YES', 'PLEASE', 'HOUSE', 'UNIT'}
    toks = toks - NOISE
    if not toks:
        return None

    # Score each isolated doc
    docs = Document.query.filter_by(
        client_id=client_id, category='property_title'
    ).all()
    scored = []
    for d in docs:
        try:
            ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
        except Exception:
            ex = {}
        if not isinstance(ex, dict):
            ex = {}
        if (ex.get('_skipped_not_in_will') or ex.get('_user_removed')):
            continue
        # Build a text blob from this doc's fields + filename
        bag = ' '.join(str(ex.get(k) or '') for k in (
            'property_address', 'title_number', 'lot_number', 'mukim',
            'daerah', 'negeri', 'purpose', 'owner_name', 'building_name',
            'property_description'))
        bag += ' ' + (d.original_filename or '')
        bag_toks = set(_re.findall(r'[A-Za-z0-9]{3,}', bag.upper()))
        overlap = toks & bag_toks
        if not overlap:
            continue
        scored.append({
            'document_id': d.id,
            'original_filename': d.original_filename or '',
            'purpose': (ex.get('purpose') or '')[:140],
            '_match_score': len(overlap),
            '_match_tokens': sorted(overlap),
        })

    scored.sort(key=lambda x: -x['_match_score'])

    # The planner reads the most recent user message directly from chat
    # history when 'assets_inventory_describing' is set, so we don't need
    # to persist anything here — the matching is recomputed on every
    # planner turn against the current set of isolated docs.
    return {
        'kind': 'assets_describe_received',
        'name': (scored[0]['original_filename'] if scored else 'no match'),
        'role': (f'{len(scored)} candidate(s)' if scored else 'no candidates'),
    }


def _try_handle_orphan_claim(client_id: str, user_text: str):
    """🔥 §10x.108 — handle the user's reply to the orphan-group
    disambiguation card.

    Quickreply formats:
      • `orphan_claim <group_id> <ai_idx>` — assign these docs to AI[ai_idx]
      • `orphan_remove <group_id>` — soft-delete the docs
      • `orphan_skip <group_id>`   — mark inventoried, no further questions

    Returns dict on success, None if not applicable.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    parts = t.split()
    if not parts or not parts[0].startswith('orphan_'):
        return None

    action = parts[0]
    if len(parts) < 2:
        return None
    group_id = parts[1]

    try:
        from services.asset_pipeline import group_documents
        groups = group_documents(client_id)
        grp = next((g for g in groups if g.group_id == group_id), None)
        if not grp:
            return None
        doc_ids = list(grp.document_ids or [])
    except Exception:
        return None

    if action == 'orphan_claim':
        if len(parts) < 3:
            return None
        try:
            ai_idx = int(parts[2])
        except (TypeError, ValueError):
            return None
        # Stamp _user_assigned_ai_idx on each doc in the group so future
        # bind_assets() runs treat them as belonging to AI[ai_idx].
        n_updated = 0
        for did in doc_ids:
            d = db.session.get(Document, did)
            if not d:
                continue
            try:
                ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
            except Exception:
                ex = {}
            ex['_user_assigned_ai_idx'] = ai_idx
            ex.pop('_inventoried', None)
            ex.pop('_orphan_group_skipped', None)
            d.extracted_data = json.dumps(ex)
            n_updated += 1
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {
            'kind': 'orphan_claimed',
            'role': f'ai_idx={ai_idx}',
            'name': f'{n_updated} doc(s) -> AI[{ai_idx}]',
        }

    if action == 'orphan_remove':
        n = 0
        for did in doc_ids:
            d = db.session.get(Document, did)
            if not d:
                continue
            d.category = 'deleted'
            n += 1
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {'kind': 'orphan_removed', 'role': 'wrong_upload',
                'name': f'{n} doc(s) removed'}

    if action == 'orphan_skip':
        n = 0
        for did in doc_ids:
            d = db.session.get(Document, did)
            if not d:
                continue
            try:
                ex = json.loads(d.extracted_data or '{}') if d.extracted_data else {}
            except Exception:
                ex = {}
            ex['_inventoried'] = True
            ex['_orphan_group_skipped'] = True
            d.extracted_data = json.dumps(ex)
            n += 1
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {'kind': 'orphan_skipped', 'role': 'not_in_will',
                'name': f'{n} doc(s) skipped'}

    return None


def _try_handle_h3_property_action(client_id: str, user_text: str):
    """Handle 'inventory confirm' / 'inventory skip' when there is no
    image-derived pending property — i.e. the card was an H3 placeholder
    rendered from the AI Summary fallback.

    Returns dict on success, None if not applicable.
    """
    if not user_text:
        return None
    t = (user_text or '').strip().lower()
    is_confirm = t.startswith('inventory h3 confirm')
    is_skip    = t.startswith('inventory h3 skip')
    if not (is_confirm or is_skip):
        return None
    # Distinct h3-prefix quick-reply values mean the regular inventory
    # handler doesn't intercept these — no kind-pending guard required.

    # Now find the AI-Summary list and identify the unhandled H3 slot
    try:
        from ai.chat_planner import (_extract_ai_summary_properties,
                                      _ai_props_already_handled,
                                      _classify_property_match)
    except Exception:
        return None

    ai_props = _extract_ai_summary_properties(client_id)
    if not ai_props:
        return None

    active_will = (Will.query.filter_by(client_id=client_id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc()).first())
    if not active_will:
        return None
    snap = {'step5': _normalise_gifts(json.loads(active_will.step5_data) if active_will.step5_data else [])}
    handled = _ai_props_already_handled(client_id, ai_props, snap)

    # Image groups (still pending = not inventoried)
    try:
        from services.gift_walker import get_pending_gift_documents
        pend = get_pending_gift_documents(client_id)
        all_props = pend.get('property') or []
    except Exception:
        all_props = []
    matched = []
    for ap in ai_props:
        cls = _classify_property_match(ap, all_props)
        matched.append(cls['variant'] in ('h1', 'h2'))

    h3_idx = next((i for i, ap in enumerate(ai_props)
                   if not handled[i] and not matched[i]), None)
    if h3_idx is None:
        return None

    # Append placeholder/skip entry to step5_data
    try:
        s5 = json.loads(active_will.step5_data) if active_will.step5_data else []
        if not isinstance(s5, list):
            s5 = []
    except Exception:
        s5 = []

    # 🔥 §10x.48 Stage 4 — build gift via canonical pipeline so address,
    # lot/title, mukim (via §10ha geo bridge), co_owners, and
    # testator_share all derive from one source of truth.
    target_b = None
    doc_groups = []
    try:
        from services.asset_pipeline import (parse_canonical_assets,
                                                group_documents,
                                                bind_assets, build_gift)
        asset_items = parse_canonical_assets(client_id)
        doc_groups = group_documents(client_id)
        bindings = bind_assets(asset_items, doc_groups)
        target_ai = next((a for a in asset_items if a.ai_index == h3_idx), None)
        target_b = next((b for b in bindings if b.ai_index == h3_idx), None)
        if target_ai and target_b:
            group_by_id = {g.group_id: g for g in doc_groups}
            dg = group_by_id.get(target_b.group_id) if target_b.group_id else None
            entry = build_gift(target_ai, target_b, dg)
        else:
            # Fallback to legacy ap dict if pipeline can't resolve
            entry = None
    except Exception:
        entry = None

    if entry is None:
        # Legacy fallback (pipeline import failed) — minimal entry
        ap = ai_props[h3_idx]
        entry = {
            'kind': 'property', 'asset_type': 'property',
            'property_info': {
                'property_address': ap.get('address') or '',
                'lot_number':       ap.get('lot') or '',
                'title_number':     ap.get('title') or '',
                'mukim':            ap.get('mukim') or '',
                'daerah':           ap.get('daerah') or '',
                'co_owners': [], 'testator_share': '1/1',
            },
            'testator_share': '1/1',
            'address': ap.get('address') or '',
            '_ai_summary_idx': h3_idx,
        }
    # Ensure required Layer-1 / Layer-2 / Layer-3 fields are present.
    entry['beneficiaries'] = entry.get('beneficiaries') or []
    entry['_ai_summary_idx'] = h3_idx
    entry['_layer1_confirmed'] = True if is_confirm else entry.get('_layer1_confirmed', False)
    if is_skip:
        entry['_ai_summary_skipped'] = True
        entry['_h3_placeholder'] = False
    entry.setdefault('ownership_intent', ai_props[h3_idx].get('ownership') or '')
    entry.setdefault('beneficiary_intent', ai_props[h3_idx].get('beneficiary') or '')
    # Keep top-level mirrors for downstream code that reads them.
    pi = entry.get('property_info') or {}
    entry.setdefault('testator_share', pi.get('testator_share') or '1/1')
    entry.setdefault('address', pi.get('property_address') or '')

    # 🔥 §10x.100 — defence-in-depth dedup. Even with §10x.100 stamping
    # _ai_summary_idx on inventory-confirm placeholders, an edge case
    # (e.g. pipeline failed mid-flow, doc not yet grouped) could leave
    # a placeholder un-tagged. Before appending, check whether ANY
    # existing gift already covers this AI Summary slot and upsert
    # instead of duplicate-appending.
    _existing_idx = next(
        (i for i, g in enumerate(s5)
         if isinstance(g, dict)
         and g.get('_ai_summary_idx') == h3_idx),
        None,
    )
    # 🔥 §10x.102 — Fix 3: catch "untagged placeholder whose document is
    # in the AI Summary slot's bound DocGroup". Happens when §10x.99
    # leftover non-determinism caused the inventory-confirm save's
    # bind_assets() lookup to return Tier D, leaving _ai_summary_idx=None.
    # On THIS turn the pipeline binding may now succeed → use the bound
    # group's document_ids to find the orphan placeholder and claim it.
    if _existing_idx is None and target_b and target_b.group_id:
        try:
            _bound_dg = next(
                (g for g in doc_groups if g.group_id == target_b.group_id),
                None,
            )
            if _bound_dg:
                _bound_doc_ids = set(_bound_dg.document_ids or [])
                _existing_idx = next(
                    (i for i, g in enumerate(s5)
                     if isinstance(g, dict)
                     and g.get('document_id') in _bound_doc_ids
                     and g.get('_ai_summary_idx') is None
                     and (g.get('kind') == 'property'
                          or g.get('asset_type') == 'property')),
                    None,
                )
                if _existing_idx is not None:
                    # Stamp the now-resolved slot tag onto the orphan.
                    s5[_existing_idx]['_ai_summary_idx'] = h3_idx
        except Exception:
            pass
    if _existing_idx is not None:
        # Upsert: merge new fields onto existing, keep the older entry
        # in place (preserves user's prior layer1_confirmed / beneficiaries).
        existing = s5[_existing_idx]
        for k, v in entry.items():
            # Don't overwrite already-populated user-facing fields.
            if k in ('beneficiaries', 'allocations', 'substitute_specific',
                     'substitute_mode', '_layer1_confirmed') \
               and existing.get(k):
                continue
            existing[k] = v
        s5[_existing_idx] = existing
    else:
        s5.append(entry)
    active_will.step5_data = json.dumps(s5)
    db.session.commit()

    return {
        'name': (ap.get('name') or 'property')[:60],
        'role': 'h3_placeholder' if is_confirm else 'h3_skipped',
        'kind': 'inventory_reviewed_property' if is_confirm else 'inventory_skipped_property',
    }


def _try_handle_message_conflict(client_id: str, user_text: str):
    """Handle the user's reply to a §10hg conflict-clarification card.

    Quick-replies:
      • 'conflict merge X Y' — treat properties X and Y as ONE
      • 'conflict keep X Y'  — treat them as DIFFERENT
      • free text (passes through to next handler)

    On accept, marks the conflict resolved by appending 'conflict_resolved_<X>_<Y>'
    to completed_steps so the planner stops gating.
    """
    if not user_text:
        return None
    t = (user_text or '').strip().lower()
    if not t.startswith('conflict '):
        return None

    parts = t.split()
    if len(parts) < 4:
        return None
    action = parts[1]   # 'merge' or 'keep'
    try:
        a = int(parts[2]); b = int(parts[3])
    except (TypeError, ValueError):
        return None

    active_will = (Will.query.filter_by(client_id=client_id, status='draft')
                   .filter(Will.deleted_at.is_(None))
                   .order_by(Will.updated_at.desc()).first())
    if not active_will:
        return None

    try:
        completed = json.loads(active_will.completed_steps) if active_will.completed_steps else []
        if not isinstance(completed, list):
            completed = []
    except Exception:
        completed = []
    marker = f'conflict_resolved_{action}_{a}_{b}'
    if marker not in completed:
        completed.append(marker)
    active_will.completed_steps = json.dumps(completed)
    db.session.commit()

    return {
        'name': f'conflict #{a}↔{b}',
        'role': f'conflict_{action}',
        'kind': 'inventory_reviewed_conflict',
    }


def _try_handle_assets_gate(client_id: str, user_text: str):
    """Handle the asset-inventory confirmation step. The planner inserts
    a confirmation gate AFTER all uploads but BEFORE per-asset gifts to
    mirror the identity flow ('collect every IC then assign roles').

    - "confirm assets" / "yes that's everything" → stamp 'assets_confirmed'
      in the will's completed_steps so the planner advances to executor →
      beneficiaries → per-asset gift assignment.
    - "i have more to upload" → just acknowledge so the planner re-asks
      the inventory question on the next turn (giving the user a beat to
      attach more files).
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if not t:
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        completed = json.loads(will.completed_steps or '[]')
    except (json.JSONDecodeError, TypeError):
        completed = []
    if not isinstance(completed, list):
        completed = []
    # Skip if already confirmed — the gate's gone, this isn't the right
    # handler for whatever they typed.
    if 'assets_confirmed' in completed:
        return None
    if any(tok in t for tok in _ASSETS_MORE_TOKENS):
        return {'name': 'asset upload', 'role': 'pending more uploads',
                'kind': 'assets_more'}
    if any(tok == t or t.startswith(tok) for tok in _ASSETS_CONFIRM_TOKENS):
        completed.append('assets_confirmed')
        will.completed_steps = json.dumps(completed)
        db.session.commit()
        return {'name': 'asset inventory', 'role': 'confirmed complete',
                'kind': 'assets_confirmed'}
    return None


def _mark_completed(will, key: str) -> bool:
    """Add `key` to will.completed_steps. Returns True if added, False if
    already present or will is None."""
    if not will:
        return False
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
        if key in completed:
            return False
        completed.append(key)
        will.completed_steps = json.dumps(completed)
        db.session.commit()
        return True
    except Exception:
        try:
            db.session.rollback()
        except Exception:
            pass
        return False


def _get_or_create_will(client_id: str):
    """Return the active draft Will for client, creating one if missing."""
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        will = Will(client_id=client_id, status='draft',
                    title='Draft Will', completed_steps='[]')
        db.session.add(will)
        db.session.commit()
    return will


def _try_handle_guardian_action(client_id: str, user_text: str):
    """Handle guardian step quick-replies and free-text names.

    Recognised inputs:
      'guardian skip'              → mark guardians_confirmed, no guardian set
      'guardian skip substitute'   → mark guardians_confirmed (primary already set)
      'guardian <name>'            → save primary guardian
      '<name>' at guardian step    → handled separately in plan context; not here

    Returns {name, role, kind} or None.
    """
    if not user_text:
        return None
    t = user_text.strip().lower()

    if t == 'guardian skip' or t == 'guardian none':
        will = _get_or_create_will(client_id)
        _mark_completed(will, 'guardians_confirmed')
        return {'name': 'guardians', 'role': 'skipped (no minor children / will set via wizard)',
                'kind': 'guardian_skipped'}

    if t == 'guardian skip substitute':
        will = _get_or_create_will(client_id)
        _mark_completed(will, 'guardians_confirmed')
        return {'name': 'guardians', 'role': 'primary set; no substitute',
                'kind': 'guardian_confirmed'}

    if not t.startswith('guardian '):
        return None
    # 'guardian <name>'
    name = user_text.strip()[9:].strip()  # strip 'guardian ' prefix (case-preserved)
    if not name or len(name) < 2:
        return None
    will = _get_or_create_will(client_id)
    try:
        s3 = json.loads(will.step3_data or '{}') if will.step3_data else {}
        if not isinstance(s3, dict):
            s3 = {}
    except (json.JSONDecodeError, TypeError):
        s3 = {}
    guardians = s3.get('guardians') or []
    has_primary = any(not g.get('is_substitute') for g in guardians)
    new_guardian = {'full_name': name.upper(), 'is_substitute': has_primary}
    guardians.append(new_guardian)
    s3['guardians'] = guardians
    try:
        will.step3_data = json.dumps(s3)
        if has_primary:
            # Substitute added → mark confirmed
            _mark_completed(will, 'guardians_confirmed')
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    role = 'substitute guardian' if has_primary else 'primary guardian'
    return {'name': name.upper(), 'role': role, 'kind': 'guardian_saved'}


def _try_handle_trust_action(client_id: str, user_text: str):
    """Handle testamentary trust step quick-replies.

    Recognised:
      'trust yes'                   → set wants_trust=True, ask for trustee
      'trust skip'                  → mark trust_confirmed, trust_skipped=True
      'trust trustee same as executor' → copy executor name as trustee
      'trust trustee <name>'        → save trustee name
      'trust age <n>'               → save distribution age
      'trust age none'              → no age limit
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if not (t.startswith('trust ') or t == 'trust yes' or t == 'trust skip'):
        return None
    will = _get_or_create_will(client_id)
    try:
        s7 = json.loads(will.step7_data or '{}') if will.step7_data else {}
        if not isinstance(s7, dict):
            s7 = {}
    except (json.JSONDecodeError, TypeError):
        s7 = {}

    if t == 'trust skip':
        s7['trust_skipped'] = True
        will.step7_data = json.dumps(s7)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        _mark_completed(will, 'trust_confirmed')
        return {'name': 'trust', 'role': 'skipped', 'kind': 'trust_skipped'}

    if t == 'trust yes':
        s7['wants_trust'] = True
        will.step7_data = json.dumps(s7)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': 'trust', 'role': 'setting up trust — enter trustee name',
                'kind': 'trust_wants'}

    if t == 'trust trustee same as executor':
        s2 = json.loads(will.step2_data or '{}') if will.step2_data else {}
        executors = (s2.get('executors') or []) if isinstance(s2, dict) else []
        exec_name = (executors[0].get('full_name') or '') if executors else ''
        s7['trustee_name'] = exec_name or 'Executor'
        will.step7_data = json.dumps(s7)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': exec_name or 'Executor', 'role': 'trustee (same as executor)',
                'kind': 'trust_trustee_set'}

    if t.startswith('trust trustee '):
        name = user_text.strip()[14:].strip().upper()
        s7['trustee_name'] = name
        will.step7_data = json.dumps(s7)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': name, 'role': 'trustee saved', 'kind': 'trust_trustee_set'}

    if t.startswith('trust age '):
        age_val = user_text.strip()[10:].strip().lower()
        if age_val in ('none', 'no limit', ''):
            s7['distribution_age'] = None
            role = 'no age limit'
        else:
            try:
                s7['distribution_age'] = int(re.sub(r'[^\d]', '', age_val) or '25')
                role = f'distribute at age {s7["distribution_age"]}'
            except ValueError:
                return None
        will.step7_data = json.dumps(s7)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        # If trustee + age both set, mark confirmed
        if s7.get('trustee_name') and 'distribution_age' in s7:
            _mark_completed(will, 'trust_confirmed')
        return {'name': 'trust', 'role': role, 'kind': 'trust_age_set'}

    return None


def _try_handle_others_action(client_id: str, user_text: str):
    """Handle 'other matters' step.

    Recognised:
      'others confirm'   → mark others_confirmed with default values
      'others skip'      → same as confirm with defaults
      'change <clause>'  → prompt is handled by chat_planner; here we just
                           look for 'change <clause>: <new text>' pattern
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    will = _get_or_create_will(client_id)
    try:
        s8 = json.loads(will.step8_data or '{}') if will.step8_data else {}
        if not isinstance(s8, dict):
            s8 = {}
    except (json.JSONDecodeError, TypeError):
        s8 = {}

    if t in ('others confirm', 'others skip', 'confirm defaults — proceed to review',
             'confirm defaults', 'skip — use all defaults'):
        s8['confirmed'] = True
        s8.pop('_pending_change', None)
        s8.pop('_combined_input_pending', None)
        will.step8_data = json.dumps(s8)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        _mark_completed(will, 'others_confirmed')
        return {'name': 'other matters', 'role': 'confirmed with defaults', 'kind': 'others_confirmed'}

    # 🔥 §10x.119 — open combined-input card. User clicked
    # "✏️ Yes I have specific wishes" / "Update my wishes".
    if t in ('others customize', 'others customise'):
        s8['_combined_input_pending'] = True
        s8.pop('_pending_change', None)
        will.step8_data = json.dumps(s8)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': 'other matters', 'role': 'awaiting wishes',
                'kind': 'others_customize_prompt'}

    # 🔥 §10x.117 — handle bare 'change X' clicks (no colon/value yet).
    # Stamp _pending_change on s8 so the planner renders a follow-up
    # "Type your <X>" prompt; the user's free-text reply on the next
    # turn is saved to that key (handled in the free-text branch below).
    _CHANGE_BARE_KEYS = {
        'change funeral': 'funeral_arrangements',
        'change funeral instructions': 'funeral_arrangements',
        'change funeral wishes': 'funeral_arrangements',
        'change organ donation': 'organ_donation',
        'change organ donation preference': 'organ_donation',
        'change pets': 'pets',
        'change digital assets': 'digital_assets',
        'change digital assets instructions': 'digital_assets',
        'change debts': 'debts',
        'change governing law': 'governing_law',
    }
    if t in _CHANGE_BARE_KEYS:
        s8['_pending_change'] = _CHANGE_BARE_KEYS[t]
        will.step8_data = json.dumps(s8)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': _CHANGE_BARE_KEYS[t], 'role': 'awaiting input',
                'kind': 'others_change_prompt'}

    # Pattern: 'change <clause>: <new value>'  e.g. "change funeral: Cremation preferred"
    m = re.match(r'^change\s+(.+?):\s*(.+)$', user_text.strip(), re.IGNORECASE)
    if m:
        clause_key = re.sub(r'\s+', '_', m.group(1).strip().lower())
        new_val = m.group(2).strip()
        s8[clause_key] = new_val
        s8.pop('_pending_change', None)
        will.step8_data = json.dumps(s8)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
        return {'name': clause_key, 'role': f'updated: {new_val[:60]}', 'kind': 'others_updated'}

    # 🔥 §10x.119 — free-text reply when combined-input card is showing.
    # User clicked "Yes I have specific wishes" → planner showed the
    # one-card combined input → user typed labelled lines like
    #   Funeral: Buddhist rites...
    #   Organ: yes — donate
    #   Pets: cat to Esther
    if s8.get('_combined_input_pending') and user_text and len(user_text) > 1:
        low_t = user_text.strip().lower()
        # Don't capture obvious other-step replies
        if not low_t.startswith(('change ', 'others ', 'confirm', 'yes',
                                 'no', 'skip', 'trust ', 'residuary',
                                 'inventory', 'orphan_', 'bank_', 'insurance_',
                                 'gift ', 'substitute ', 'beneficiaries ',
                                 'role_match')):
            # Parse labelled lines. Recognised label keywords (case-
            # insensitive) → s8 key:
            _LABEL_MAP = {
                'funeral': 'funeral_arrangements',
                'funeral arrangements': 'funeral_arrangements',
                'funeral wishes': 'funeral_arrangements',
                'organ': 'organ_donation',
                'organ donation': 'organ_donation',
                'pet': 'pets',
                'pets': 'pets',
                'digital': 'digital_assets',
                'digital assets': 'digital_assets',
            }
            updated = []
            # Split into lines; skip blanks; for each line look for
            # "Label:" prefix and capture the remainder.
            for raw in (user_text or '').splitlines():
                line = raw.strip()
                if not line:
                    continue
                m = re.match(r'^([A-Za-z][A-Za-z ]{0,40}?)\s*[:\-]\s*(.+)$',
                             line)
                if not m:
                    continue
                label = m.group(1).strip().lower()
                value = m.group(2).strip()
                if label in _LABEL_MAP and value:
                    s8[_LABEL_MAP[label]] = value[:500]
                    updated.append(_LABEL_MAP[label])
            if updated:
                # User provided at least one labelled value — save and
                # confirm. Defaults remain for unprovided keys.
                s8.pop('_combined_input_pending', None)
                s8['confirmed'] = True
                will.step8_data = json.dumps(s8)
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()
                    return None
                _mark_completed(will, 'others_confirmed')
                return {'name': 'personal wishes',
                        'role': f'updated {len(updated)} field(s)',
                        'kind': 'others_updated'}
            # If the reply doesn't look labelled, treat the whole thing
            # as funeral wishes (most common single intent) and confirm.
            if len(user_text.strip()) > 3:
                s8['funeral_arrangements'] = user_text.strip()[:500]
                s8.pop('_combined_input_pending', None)
                s8['confirmed'] = True
                will.step8_data = json.dumps(s8)
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()
                    return None
                _mark_completed(will, 'others_confirmed')
                return {'name': 'funeral_arrangements',
                        'role': f'set as: {user_text.strip()[:60]}',
                        'kind': 'others_updated'}

    # 🔥 §10x.117 — free-text reply when _pending_change is set.
    # User clicked 'change funeral' on the previous turn → planner asked
    # 'Type your funeral wishes' → user's reply lands here.
    pending = (s8.get('_pending_change') or '').strip()
    if pending and user_text and len(user_text) > 1:
        # Don't capture obvious other-step replies
        low_t = user_text.strip().lower()
        if not low_t.startswith(('change ', 'others ', 'confirm', 'yes',
                                 'no', 'skip', 'trust ', 'residuary',
                                 'inventory', 'orphan_', 'bank_', 'insurance_',
                                 'gift ', 'substitute ', 'beneficiaries ',
                                 'role_match')):
            s8[pending] = user_text.strip()[:500]
            s8.pop('_pending_change', None)
            will.step8_data = json.dumps(s8)
            try:
                db.session.commit()
            except Exception:
                db.session.rollback()
            return {'name': pending,
                    'role': f'updated: {user_text.strip()[:60]}',
                    'kind': 'others_updated'}

    return None


def _try_save_bank_layered_gift(client_id, user_text):
    """🔥 §10x.23 — handle 3-layer bank flow:
        bank_l1 confirm/skip/remove   → save Layer 1 placeholder
        bank_l2 main 100% <name> | equal children   → save main beneficiary
        bank_l3 sub <action>          → save substitute beneficiary
    """
    if not user_text:
        return None
    t = user_text.strip()
    low = t.lower()
    if not (low.startswith('bank_l1') or low.startswith('bank_l2') or low.startswith('bank_l3')):
        return None

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except Exception:
        gifts = []

    persons = Person.query.filter_by(client_id=client_id).all()
    spouse = next((p for p in persons
                   if (p.relationship or '').lower() in ('spouse', 'wife', 'husband')), None)
    children = [p for p in persons
                if (p.relationship or '').lower() in ('son', 'daughter')]
    spouse_name = spouse.full_name if spouse else ''
    child_names = [c.full_name for c in children]

    from ai.chat_planner import _extract_ai_summary_banks
    ai_banks = _extract_ai_summary_banks(client_id) or []

    # Helper: find which AI bank we're acting on (the first incomplete one
    # for whatever layer the input represents).
    saved_bank_by_acct = {}
    for g in gifts:
        if not isinstance(g, dict): continue
        if g.get('kind') != 'bank': continue
        ak = re.sub(r'\W+', '', g.get('account_number') or '')
        if ak: saved_bank_by_acct[ak] = g

    # Layer 1
    if low.startswith('bank_l1'):
        # Find first AI bank with no saved entry yet
        target = None
        target_idx = -1
        for i, b in enumerate(ai_banks):
            ak = re.sub(r'\W+', '', b.get('account_number') or '')
            if ak in saved_bank_by_acct:
                continue
            target = b; target_idx = i; break
        if not target:
            return None
        action = low[len('bank_l1'):].strip()
        new_gift = {
            'kind': 'bank', 'asset_type': 'bank',
            '_ai_summary_bank_idx': target_idx,
            '_layer1_confirmed': action == 'confirm',
            'bank_name':       target.get('bank_name'),
            'account_number':  target.get('account_number'),
            'country':         target.get('country'),
            'account_type':    target.get('account_type'),
            'gift_type': 'financial',
            'financial_details': {
                'asset_type':     'bank',
                'institution':    target.get('bank_name') or '',
                'account_number': target.get('account_number') or '',
                'country':        target.get('country') or '',
            },
            'allocations':         [],
            'beneficiaries':       [],
            'substitute_mode':     None,
            'substitute_specific': None,
        }
        if action == 'skip':
            new_gift['skipped'] = True
        elif action == 'remove':
            new_gift['_user_rejected'] = True
        gifts.append(new_gift)
        will.step5_data = json.dumps(gifts)
        db.session.commit()
        return {'name': target.get('bank_name'), 'role': f'l1_{action}', 'kind': 'gift_bank_l1'}

    # Layer 2 (main beneficiary)
    if low.startswith('bank_l2'):
        # Find first saved bank with empty beneficiaries
        target_idx = -1
        for gi, g in enumerate(gifts):
            if (isinstance(g, dict) and g.get('kind') == 'bank'
                    and not g.get('skipped') and not g.get('_user_rejected')
                    and not (g.get('beneficiaries') or [])):
                target_idx = gi; break
        if target_idx < 0:
            return None
        rest = t[len('bank_l2'):].strip().lower()
        main_bens = []
        if rest.startswith('main equal children') and len(child_names) >= 2:
            share = '1/' + str(len(child_names))
            main_bens = [{'name': c, 'share': share} for c in child_names]
        elif rest.startswith('main 100% '):
            nm = t[len('bank_l2 main 100% '):].strip()
            main_bens = [{'name': nm, 'share': '1/1'}]
        elif rest == 'skip':
            gifts[target_idx]['skipped'] = True
            will.step5_data = json.dumps(gifts); db.session.commit()
            return {'name': gifts[target_idx].get('bank_name'), 'role': 'l2_skip', 'kind': 'gift_bank_l2'}
        else:
            return None
        gifts[target_idx]['beneficiaries'] = main_bens
        gifts[target_idx]['allocations'] = [
            {'beneficiary_name': b['name'], 'share': b['share'], 'role': 'MB'}
            for b in main_bens
        ]
        will.step5_data = json.dumps(gifts); db.session.commit()
        return {'name': gifts[target_idx].get('bank_name'),
                'role': f'main {main_bens[0]["name"]}', 'kind': 'gift_bank_l2'}

    # Layer 3 (substitute)
    if low.startswith('bank_l3'):
        target_idx = -1
        for gi, g in enumerate(gifts):
            if (isinstance(g, dict) and g.get('kind') == 'bank'
                    and not g.get('skipped') and not g.get('_user_rejected')
                    and (g.get('beneficiaries') or [])
                    and g.get('substitute_specific') is None
                    and g.get('substitute_mode') in (None, '')):
                target_idx = gi; break
        if target_idx < 0:
            return None
        rest = t[len('bank_l3'):].strip().lower()
        sub_bens = None
        sub_mode = 'specific'
        main_bens = gifts[target_idx].get('beneficiaries') or []
        if rest == 'sub none':
            sub_bens = []; sub_mode = 'none'
        elif rest == 'sub equal children' and len(child_names) >= 2:
            share = '1/' + str(len(child_names))
            sub_bens = [{'name': c, 'share': share} for c in child_names]
        elif rest == 'sub survivors' and len(main_bens) >= 2:
            share = '1/' + str(len(main_bens))
            sub_bens = [{'name': m['name'], 'share': share} for m in main_bens]
        elif rest.startswith('sub 100% '):
            nm = t[len('bank_l3 sub 100% '):].strip()
            sub_bens = [{'name': nm, 'share': '1/1'}]
        else:
            return None
        gifts[target_idx]['substitute_specific'] = sub_bens
        gifts[target_idx]['substitute_mode'] = sub_mode
        # Mirror into allocations[*].substitutes
        if sub_bens:
            for alloc in gifts[target_idx].get('allocations') or []:
                alloc['substitutes'] = [
                    {'beneficiary_name': s['name'], 'share': s['share']}
                    for s in sub_bens
                ]
        will.step5_data = json.dumps(gifts); db.session.commit()
        return {'name': gifts[target_idx].get('bank_name'),
                'role': f'sub {sub_mode}', 'kind': 'gift_bank_l3'}

    return None


def _try_save_insurance_layered_gift(client_id, user_text):
    """🔥 §10x.23 — same 3-layer flow for insurance policies."""
    if not user_text:
        return None
    t = user_text.strip()
    low = t.lower()
    if not (low.startswith('insurance_l1') or low.startswith('insurance_l2') or low.startswith('insurance_l3')):
        return None

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except Exception:
        gifts = []

    persons = Person.query.filter_by(client_id=client_id).all()
    spouse = next((p for p in persons
                   if (p.relationship or '').lower() in ('spouse', 'wife', 'husband')), None)
    children = [p for p in persons
                if (p.relationship or '').lower() in ('son', 'daughter')]
    spouse_name = spouse.full_name if spouse else ''
    child_names = [c.full_name for c in children]

    from ai.chat_planner import _extract_ai_summary_insurance
    ai_ins = _extract_ai_summary_insurance(client_id) or []

    saved_ins_by_pol = {}
    for g in gifts:
        if not isinstance(g, dict): continue
        if g.get('kind') != 'insurance': continue
        pn = re.sub(r'\W+', '', g.get('policy_number') or '')
        if pn: saved_ins_by_pol[pn] = g

    if low.startswith('insurance_l1'):
        target = None; target_idx = -1
        for i, ins in enumerate(ai_ins):
            pn = re.sub(r'\W+', '', ins.get('policy_number') or '')
            if pn in saved_ins_by_pol: continue
            target = ins; target_idx = i; break
        if not target:
            return None
        action = low[len('insurance_l1'):].strip()
        new_gift = {
            'kind': 'insurance', 'asset_type': 'insurance',
            '_ai_summary_insurance_idx': target_idx,
            '_layer1_confirmed': action == 'confirm',
            'insurer':       target.get('insurer'),
            'policy_number': target.get('policy_number'),
            'gift_type': 'financial',
            'financial_details': {
                'asset_type':     'insurance',
                'institution':    target.get('insurer') or '',
                'account_number': target.get('policy_number') or '',
            },
            'allocations':         [],
            'beneficiaries':       [],
            'substitute_mode':     None,
            'substitute_specific': None,
        }
        if action == 'skip':
            new_gift['skipped'] = True
        elif action == 'remove':
            new_gift['_user_rejected'] = True
        gifts.append(new_gift)
        will.step5_data = json.dumps(gifts); db.session.commit()
        return {'name': target.get('insurer'), 'role': f'l1_{action}', 'kind': 'gift_insurance_l1'}

    if low.startswith('insurance_l2'):
        target_idx = -1
        for gi, g in enumerate(gifts):
            if (isinstance(g, dict) and g.get('kind') == 'insurance'
                    and not g.get('skipped') and not g.get('_user_rejected')
                    and not (g.get('beneficiaries') or [])):
                target_idx = gi; break
        if target_idx < 0:
            return None
        rest = t[len('insurance_l2'):].strip().lower()
        main_bens = []
        if rest.startswith('main equal children') and len(child_names) >= 2:
            share = '1/' + str(len(child_names))
            main_bens = [{'name': c, 'share': share} for c in child_names]
        elif rest.startswith('main 100% '):
            nm = t[len('insurance_l2 main 100% '):].strip()
            main_bens = [{'name': nm, 'share': '1/1'}]
        elif rest == 'skip':
            gifts[target_idx]['skipped'] = True
            will.step5_data = json.dumps(gifts); db.session.commit()
            return {'name': gifts[target_idx].get('insurer'), 'role': 'l2_skip', 'kind': 'gift_insurance_l2'}
        else:
            return None
        gifts[target_idx]['beneficiaries'] = main_bens
        gifts[target_idx]['allocations'] = [
            {'beneficiary_name': b['name'], 'share': b['share'], 'role': 'MB'}
            for b in main_bens
        ]
        will.step5_data = json.dumps(gifts); db.session.commit()
        return {'name': gifts[target_idx].get('insurer'),
                'role': f'main {main_bens[0]["name"]}', 'kind': 'gift_insurance_l2'}

    if low.startswith('insurance_l3'):
        target_idx = -1
        for gi, g in enumerate(gifts):
            if (isinstance(g, dict) and g.get('kind') == 'insurance'
                    and not g.get('skipped') and not g.get('_user_rejected')
                    and (g.get('beneficiaries') or [])
                    and g.get('substitute_specific') is None
                    and g.get('substitute_mode') in (None, '')):
                target_idx = gi; break
        if target_idx < 0:
            return None
        rest = t[len('insurance_l3'):].strip().lower()
        sub_bens = None; sub_mode = 'specific'
        main_bens = gifts[target_idx].get('beneficiaries') or []
        if rest == 'sub none':
            sub_bens = []; sub_mode = 'none'
        elif rest == 'sub equal children' and len(child_names) >= 2:
            share = '1/' + str(len(child_names))
            sub_bens = [{'name': c, 'share': share} for c in child_names]
        elif rest == 'sub survivors' and len(main_bens) >= 2:
            share = '1/' + str(len(main_bens))
            sub_bens = [{'name': m['name'], 'share': share} for m in main_bens]
        elif rest.startswith('sub 100% '):
            nm = t[len('insurance_l3 sub 100% '):].strip()
            sub_bens = [{'name': nm, 'share': '1/1'}]
        else:
            return None
        gifts[target_idx]['substitute_specific'] = sub_bens
        gifts[target_idx]['substitute_mode'] = sub_mode
        if sub_bens:
            for alloc in gifts[target_idx].get('allocations') or []:
                alloc['substitutes'] = [
                    {'beneficiary_name': s['name'], 'share': s['share']}
                    for s in sub_bens
                ]
        will.step5_data = json.dumps(gifts); db.session.commit()
        return {'name': gifts[target_idx].get('insurer'),
                'role': f'sub {sub_mode}', 'kind': 'gift_insurance_l3'}

    return None


def _try_save_bank_h3_gift(client_id, user_text):
    """Handler for the per-bank H3 card (§10x.12).
    Quick-reply values: 'bank_h3 confirm 100% <NAME>', 'bank_h3 confirm equal children',
    'bank_h3 skip', 'bank_h3 remove'.

    Saves ONE step5_data gift entry per bank account, with §10x.14 default
    substitute (wife→both children, single→other child, multi→survivors).
    """
    if not user_text:
        return None
    t = user_text.strip()
    if not t.lower().startswith('bank_h3'):
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except Exception:
        gifts = []
    # Need to know which AI-summary bank we're on
    from ai.chat_planner import _extract_ai_summary_banks
    ai_banks = _extract_ai_summary_banks(client_id) or []
    saved_acct = set()
    for g in gifts:
        an = (g.get('account_number') or '').strip()
        if an:
            saved_acct.add(re.sub(r'\W+', '', an))
    target = None
    target_idx = -1
    for i, b in enumerate(ai_banks):
        ack = re.sub(r'\W+', '', b.get('account_number') or '')
        if ack and ack in saved_acct:
            continue
        target = b
        target_idx = i
        break
    if not target:
        return None
    # Identities for substitute defaults (§10x.14)
    persons = Person.query.filter_by(client_id=client_id).all()
    spouse = next((p for p in persons
                   if (p.relationship or '').lower() in ('spouse', 'wife', 'husband')), None)
    children = [p for p in persons
                if (p.relationship or '').lower() in ('son', 'daughter')]
    spouse_name = spouse.full_name if spouse else ''
    child_names = [c.full_name for c in children]
    # Parse the action
    rest = t[len('bank_h3'):].strip().lower()
    if rest == 'skip':
        gifts.append({
            'kind': 'bank',
            'asset_type': 'bank',
            '_ai_summary_bank_idx': target_idx,
            'skipped': True,
            'bank_name': target.get('bank_name'),
            'account_number': target.get('account_number'),
            'beneficiaries': [],
        })
        will.step5_data = json.dumps(gifts)
        db.session.commit()
        return {'name': target.get('bank_name'), 'role': 'skipped',
                'kind': 'gift_bank_h3'}
    if rest == 'remove':
        # Mark dropped (counts toward "handled" so walker advances)
        gifts.append({
            'kind': 'bank',
            'asset_type': 'bank',
            '_ai_summary_bank_idx': target_idx,
            '_user_rejected': True,
            'bank_name': target.get('bank_name'),
            'account_number': target.get('account_number'),
            'beneficiaries': [],
        })
        will.step5_data = json.dumps(gifts)
        db.session.commit()
        return {'name': target.get('bank_name'), 'role': 'removed',
                'kind': 'gift_bank_h3'}
    # Otherwise expect 'confirm 100% <name>' or 'confirm equal children'
    main_bens = []
    sub_bens = []
    if rest.startswith('confirm equal children') and len(child_names) >= 2:
        share = '1/' + str(len(child_names))
        main_bens = [{'name': c, 'share': share} for c in child_names]
        # §10x.14: multi-children main → survivors equal (same set)
        sub_bens = list(main_bens)
    elif rest.startswith('confirm 100% '):
        nm = user_text.strip()[len('bank_h3 confirm 100% '):].strip()
        main_bens = [{'name': nm, 'share': '1/1'}]
        # §10x.14 substitute defaults
        if nm.lower() == spouse_name.lower() and child_names:
            share = '1/' + str(len(child_names))
            sub_bens = [{'name': c, 'share': share} for c in child_names]
        elif nm in child_names:
            others = [c for c in child_names if c != nm]
            if others:
                share = '1/' + str(len(others))
                sub_bens = [{'name': c, 'share': share} for c in others]
        else:
            # Other person (e.g. brother) → all children equal
            if child_names:
                share = '1/' + str(len(child_names))
                sub_bens = [{'name': c, 'share': share} for c in child_names]
    else:
        return None   # unknown bank_h3 action
    if not main_bens:
        return None
    allocations = [{'beneficiary_name': b['name'], 'share': b['share'], 'role': 'MB'}
                   for b in main_bens]
    if sub_bens:
        for alloc in allocations:
            alloc['substitutes'] = [
                {'beneficiary_name': s['name'], 'share': s['share']}
                for s in sub_bens
            ]
    gift_entry = {
        # ── canonical (chat-side reads these) ─────────────────
        'kind': 'bank',
        'asset_type': 'bank',
        '_ai_summary_bank_idx': target_idx,
        'bank_name': target.get('bank_name'),
        'account_number': target.get('account_number'),
        'country': target.get('country'),
        'account_type': target.get('account_type'),
        # ── wizard-compatible mirror (templates/wizard/step6_gifts.html
        #    reads gift_type + financial_details) ──────────────
        'gift_type': 'financial',
        'financial_details': {
            'asset_type':     'bank',
            'institution':    target.get('bank_name') or '',
            'account_number': target.get('account_number') or '',
            'country':        target.get('country') or '',
        },
        # ── beneficiaries ─────────────────────────────────────
        'allocations': allocations,
        'beneficiaries': main_bens,
        'substitute_mode': 'specific' if sub_bens else 'none',
        'substitute_specific': sub_bens or None,
    }
    gifts.append(gift_entry)
    will.step5_data = json.dumps(gifts)
    db.session.commit()
    return {'name': target.get('bank_name'),
            'role': f'{main_bens[0]["name"]} {main_bens[0]["share"]}',
            'kind': 'gift_bank_h3'}


def _try_save_insurance_h3_gift(client_id, user_text):
    """Counterpart to _try_save_bank_h3_gift but for insurance policies (§10x.12).
    Same substitute defaults from §10x.14.
    """
    if not user_text:
        return None
    t = user_text.strip()
    if not t.lower().startswith('insurance_h3'):
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except Exception:
        gifts = []
    from ai.chat_planner import _extract_ai_summary_insurance
    ai_ins = _extract_ai_summary_insurance(client_id) or []
    saved_pol = set()
    for g in gifts:
        pn = (g.get('policy_number') or '').strip()
        if pn:
            saved_pol.add(re.sub(r'\W+', '', pn))
    target = None
    target_idx = -1
    for i, ins in enumerate(ai_ins):
        pol = re.sub(r'\W+', '', ins.get('policy_number') or '')
        if pol and pol in saved_pol:
            continue
        target = ins
        target_idx = i
        break
    if not target:
        return None
    persons = Person.query.filter_by(client_id=client_id).all()
    spouse = next((p for p in persons
                   if (p.relationship or '').lower() in ('spouse', 'wife', 'husband')), None)
    children = [p for p in persons
                if (p.relationship or '').lower() in ('son', 'daughter')]
    spouse_name = spouse.full_name if spouse else ''
    child_names = [c.full_name for c in children]
    rest = t[len('insurance_h3'):].strip().lower()
    if rest == 'skip':
        gifts.append({
            'kind': 'insurance',
            'asset_type': 'insurance',
            '_ai_summary_insurance_idx': target_idx,
            'skipped': True,
            'insurer': target.get('insurer'),
            'policy_number': target.get('policy_number'),
            'beneficiaries': [],
        })
        will.step5_data = json.dumps(gifts)
        db.session.commit()
        return {'name': target.get('insurer'), 'role': 'skipped',
                'kind': 'gift_insurance_h3'}
    if rest == 'remove':
        gifts.append({
            'kind': 'insurance',
            'asset_type': 'insurance',
            '_ai_summary_insurance_idx': target_idx,
            '_user_rejected': True,
            'insurer': target.get('insurer'),
            'policy_number': target.get('policy_number'),
            'beneficiaries': [],
        })
        will.step5_data = json.dumps(gifts)
        db.session.commit()
        return {'name': target.get('insurer'), 'role': 'removed',
                'kind': 'gift_insurance_h3'}
    main_bens = []
    sub_bens = []
    if rest.startswith('confirm equal children') and len(child_names) >= 2:
        share = '1/' + str(len(child_names))
        main_bens = [{'name': c, 'share': share} for c in child_names]
        sub_bens = list(main_bens)
    elif rest.startswith('confirm 100% '):
        nm = user_text.strip()[len('insurance_h3 confirm 100% '):].strip()
        main_bens = [{'name': nm, 'share': '1/1'}]
        if nm.lower() == spouse_name.lower() and child_names:
            share = '1/' + str(len(child_names))
            sub_bens = [{'name': c, 'share': share} for c in child_names]
        elif nm in child_names:
            others = [c for c in child_names if c != nm]
            if others:
                share = '1/' + str(len(others))
                sub_bens = [{'name': c, 'share': share} for c in others]
        else:
            if child_names:
                share = '1/' + str(len(child_names))
                sub_bens = [{'name': c, 'share': share} for c in child_names]
    else:
        return None
    if not main_bens:
        return None
    allocations = [{'beneficiary_name': b['name'], 'share': b['share'], 'role': 'MB'}
                   for b in main_bens]
    if sub_bens:
        for alloc in allocations:
            alloc['substitutes'] = [
                {'beneficiary_name': s['name'], 'share': s['share']}
                for s in sub_bens
            ]
    gift_entry = {
        # ── canonical ────────────────────────────────────────
        'kind': 'insurance',
        'asset_type': 'insurance',
        '_ai_summary_insurance_idx': target_idx,
        'insurer': target.get('insurer'),
        'policy_number': target.get('policy_number'),
        # ── wizard-compatible mirror ─────────────────────────
        'gift_type': 'financial',
        'financial_details': {
            'asset_type':     'insurance',
            'institution':    target.get('insurer') or '',
            'account_number': target.get('policy_number') or '',
        },
        # ── beneficiaries ────────────────────────────────────
        'allocations': allocations,
        'beneficiaries': main_bens,
        'substitute_mode': 'specific' if sub_bens else 'none',
        'substitute_specific': sub_bens or None,
    }
    gifts.append(gift_entry)
    will.step5_data = json.dumps(gifts)
    db.session.commit()
    return {'name': target.get('insurer'),
            'role': f'{main_bens[0]["name"]} {main_bens[0]["share"]}',
            'kind': 'gift_insurance_h3'}


def _try_handle_mismatch(client_id: str, user_text: str):
    """🔥 BURN-IN §10x.18 — handle the text-vs-image clarification card.

    Quickreplies:
        'mismatch use_text   <gift_idx> <field>'  — keep text value
        'mismatch use_image  <gift_idx> <field>'  — overwrite gift with image value
        'mismatch type_manually <gift_idx> <field>' — fall through to manual entry
        'mismatch remove_image <gift_idx>'        — unbind document
    """
    if not user_text:
        return None
    t = user_text.strip()
    if not t.lower().startswith('mismatch '):
        return None
    parts = t.split()
    if len(parts) < 3:
        return None
    action = parts[1].lower()
    try:
        gi = int(parts[2])
    except ValueError:
        return None
    field = parts[3] if len(parts) > 3 else ''

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s5 = json.loads(will.step5_data or '[]')
        if not isinstance(s5, list):
            return None
    except Exception:
        return None
    if not (0 <= gi < len(s5)):
        return None
    g = s5[gi]
    doc_id = g.get('document_id')

    if action == 'use_text':
        # Mark resolved — text wins, no change to gift, just unblock the gate
        pass
    elif action == 'use_image':
        # Overwrite the gift's field with the image's value
        if doc_id:
            doc = db.session.get(Document, doc_id)
            try:
                ex = json.loads(doc.extracted_data) if doc and doc.extracted_data else {}
            except Exception:
                ex = {}
            new_v = (ex.get(field) or '').strip()
            if new_v:
                if g.get('kind') == 'property':
                    g.setdefault('property_info', {})[field] = new_v
                    g[field] = new_v
                elif g.get('kind') == 'bank':
                    g[field] = new_v
                    g.setdefault('financial_details', {})['account_number' if field == 'account_number' else field] = new_v
                elif g.get('kind') == 'insurance':
                    g[field] = new_v
                s5[gi] = g
                will.step5_data = json.dumps(s5)
    elif action == 'type_manually':
        # The user will type the correct value in the next message;
        # for now, mark resolved so the gate doesn't keep firing.
        # Future enhancement: capture a follow-up free-text input.
        pass
    elif action == 'remove_image':
        # Drop the document_id from the gift
        g.pop('document_id', None)
        s5[gi] = g
        will.step5_data = json.dumps(s5)
    else:
        return None

    # Mark resolved in completed_steps so the gate doesn't refire.
    try:
        cs = json.loads(will.completed_steps or '[]')
        if not isinstance(cs, list):
            cs = []
    except Exception:
        cs = []
    marker = f'mismatch_resolved_{gi}_{field}' if field else f'mismatch_resolved_{gi}'
    if marker not in cs:
        cs.append(marker)
        will.completed_steps = json.dumps(cs)
    db.session.commit()
    return {'name': f'gift[{gi}]', 'role': action, 'kind': 'mismatch_resolved'}


def _try_handle_role_match(client_id: str, user_text: str):
    """🔥 BURN-IN §10x.21 — handle the role-match clarification card.

    Quickreplies:
        'role_match confirm <person_id>' → promote that Person to executor
        'role_match manual'              → fall through to generic exec card
        'role_match skip'                → record skip
    """
    if not user_text:
        return None
    t = user_text.strip()
    if not t.lower().startswith('role_match'):
        return None
    rest = t[len('role_match'):].strip().lower()

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None

    if rest == 'manual' or rest == 'skip':
        # Mark as 'manual' so the planner falls through to the generic
        # executor card on the next turn (no DB change needed beyond a
        # marker in completed_steps).
        try:
            cs = json.loads(will.completed_steps or '[]')
            if not isinstance(cs, list):
                cs = []
        except Exception:
            cs = []
        marker = 'role_match_skipped' if rest == 'skip' else 'role_match_manual'
        if marker not in cs:
            cs.append(marker)
            will.completed_steps = json.dumps(cs)
            db.session.commit()
        return {'name': 'role_match', 'role': rest, 'kind': 'role_match'}

    # 'role_match confirm <person_id>'
    if not rest.startswith('confirm '):
        return None
    person_id = t.split()[-1].strip()   # last token is the person_id
    p = db.session.get(Person, person_id)
    if not p or p.client_id != client_id:
        return None

    # Promote this Person to executor relationship
    p.relationship = 'executor'
    db.session.commit()

    # Mirror into step2_data.executors
    try:
        s2 = json.loads(will.step2_data) if will.step2_data else {}
        if not isinstance(s2, dict):
            s2 = {}
    except Exception:
        s2 = {}
    execs = s2.get('executors') or []
    # Avoid duplicate
    if not any((e.get('person_id') == p.id) or
               ((e.get('full_name') or '').upper() == (p.full_name or '').upper())
               for e in execs):
        # Pull the phone from the role mention if available
        phone = ''
        try:
            from services.role_matcher import extract_role_mentions
            mentions = extract_role_mentions(client_id) or []
            exec_mentions = [m for m in mentions if m.get('role') == 'executor']
            if exec_mentions:
                phone = exec_mentions[0].get('phone', '')
        except Exception:
            pass
        execs.append({
            'person_id':     p.id,
            'full_name':     p.full_name,
            'nric_passport': p.nric_passport or '',
            'phone':         phone,
            'relationship':  'executor',
            'role':          'Primary',
            'entry_type':    'individual',
        })
    s2['executors'] = execs
    s2.setdefault('executor_type', 'single')
    s2.setdefault('trustee_data', {'same_as_executor': True, 'trustees': [{}]})
    will.step2_data = json.dumps(s2)
    db.session.commit()

    return {'name': p.full_name, 'role': 'executor', 'kind': 'role_match_confirmed'}


def _try_handle_beneficiaries_confirm(client_id: str, user_text: str):
    """🔥 §10x.115 — handle the user's response to the Step 5 main
    beneficiaries confirmation card.

    Quickreplies:
      • `beneficiaries confirm` → stamp `beneficiaries_confirmed` so
        planner advances to Step 6 (Specific Gifts).
      • `beneficiaries edit` → fall through to free-text editing
        (handled by _try_save_beneficiaries with 'remove X' / 'only X')
    """
    if not user_text:
        return None
    t = user_text.strip().lower()
    if t != 'beneficiaries confirm':
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
    except Exception:
        completed = []
    if 'beneficiaries_confirmed' not in completed:
        completed.append('beneficiaries_confirmed')
        will.completed_steps = json.dumps(completed)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
    try:
        s4 = json.loads(will.step4_data or '[]')
    except Exception:
        s4 = []
    n = len(s4) if isinstance(s4, list) else 0
    return {
        'kind': 'beneficiaries_confirmed',
        'name': f'{n} beneficiaries',
        'role': 'confirmed',
    }


def _try_save_residuary_substitute(client_id: str, user_text: str):
    """🔥 §10x.116 — Step 7 Layer 3 (substitute residuary beneficiary).

    Quickreply formats:
      • `residuary substitute survivors`        — surviving main bens equally
      • `residuary substitute equal others`     — surviving non-main family equally
      • `residuary substitute equal children`   — surviving children equally
      • `residuary substitute 100% <name>`      — single named substitute
      • `residuary substitute none`             — no substitute clause
      • free text 'Joshua 50%, Esther 50%' or 'wife 100%'

    Writes to step6_data:
      substitute_mode: 'specific' | 'equal' | 'survivors' | 'none'
      substitute_specific: [{name, share}, ...]
    """
    if not user_text:
        return None
    t = user_text.strip()
    low = t.lower()
    if not low.startswith('residuary substitute'):
        # Also accept free-text NAMES if planner is at Layer 3.
        # Only fire when step6 has main beneficiaries but no substitute yet.
        will_check = (Will.query.filter_by(client_id=client_id, status='draft')
                      .filter(Will.deleted_at.is_(None))
                      .order_by(Will.updated_at.desc()).first())
        if not will_check:
            return None
        try:
            s6_check = json.loads(will_check.step6_data or '{}')
            if not isinstance(s6_check, dict):
                return None
        except Exception:
            return None
        has_main = bool(s6_check.get('beneficiaries') or s6_check.get('residuary_beneficiary_name'))
        has_sub = (s6_check.get('substitute_specific') is not None
                   or s6_check.get('substitute_mode') in ('specific', 'equal',
                                                          'survivors', 'none'))
        if not has_main or has_sub:
            return None
        # We're at Layer 3 — accept free-text. Apply the same negative
        # filters as _try_save_residuary_main.
        _BAD = ('inventory ', 'orphan_', 'bank_', 'insurance_', 'gift ',
                'substitute ', 'trust ', 'others ', 'guardian ',
                'role_match', 'address:', 'lot:', 'title:', 'mukim:',
                'daerah:', 'negeri:', 'property ', 'change ',
                'confirm assets', 'confirm defaults', 'h3 ',
                'unlink ', 'restart ', 'inbox ', 'conflict ',
                'i have more', 'beneficiaries ', 'residuary skip')
        if any(low.startswith(p) for p in _BAD):
            return None
        if low in ('yes', 'no', 'skip', 'delete', 'confirm', 'remove'):
            return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s6 = json.loads(will.step6_data or '{}')
        if not isinstance(s6, dict): s6 = {}
    except Exception:
        s6 = {}
    if not (s6.get('beneficiaries') or s6.get('residuary_beneficiary_name')):
        return None  # Layer 2 not done yet — skip Layer 3
    if (s6.get('substitute_specific') is not None
        or s6.get('substitute_mode') in ('specific', 'equal',
                                          'survivors', 'none')):
        return None  # already saved

    persons = Person.query.filter_by(client_id=client_id).all()
    eligible = [p for p in persons
                if (p.relationship or '').lower() not in ('testator', 'witness')]
    main_names = {(b.get('name') or b.get('full_name') or '').upper()
                  for b in (s6.get('beneficiaries') or [])
                  if isinstance(b, dict)}
    children = [p for p in eligible
                if (p.relationship or '').lower() in ('son', 'daughter')]

    sub_mode = ''
    sub_list: List[Dict[str, str]] = []

    if low == 'residuary substitute none':
        sub_mode = 'none'
    elif low == 'residuary substitute survivors':
        sub_mode = 'survivors'
        # Survivors of the main list, equal shares
        n = max(len(main_names), 1)
        sub_list = [{'name': n_, 'share': f'1/{n}'}
                    for n_ in (b.get('name') or b.get('full_name')
                                for b in (s6.get('beneficiaries') or []))
                    if n_]
    elif low == 'residuary substitute equal others':
        sub_mode = 'equal'
        others = [p for p in eligible
                  if (p.full_name or '').upper() not in main_names
                  and (p.relationship or '').lower() in
                      ('son', 'daughter', 'spouse', 'wife', 'husband',
                       'father', 'mother', 'brother', 'sister')]
        n = max(len(others), 1)
        sub_list = [{'name': p.full_name, 'share': f'1/{n}'}
                    for p in others]
    elif low == 'residuary substitute equal children':
        sub_mode = 'equal'
        n = max(len(children), 1)
        sub_list = [{'name': c.full_name, 'share': f'1/{n}'}
                    for c in children]
    elif low.startswith('residuary substitute 100%'):
        sub_mode = 'specific'
        nm = t[len('residuary substitute 100%'):].strip()
        # Resolve via name OR relationship
        match = None
        for p in eligible:
            if nm.upper() == (p.full_name or '').upper() \
               or nm.lower() in (p.full_name or '').lower():
                match = p; break
        if match:
            sub_list = [{'name': match.full_name, 'share': '1/1'}]
    else:
        # Free-text: 'Joshua 50%, Esther 50%' or 'wife 100%'
        sub_mode = 'specific'
        parts = [p.strip() for p in re.split(r',|\band\b', t) if p.strip()]
        for part in parts:
            share_m = re.search(
                r'\s+(\d{1,3})\s*%|\s+(equal|equally)\b|\s+(\d+/\d+)\s*$',
                part, flags=re.IGNORECASE)
            if share_m:
                name_token = part[:share_m.start()].strip().rstrip(',')
                if share_m.group(1):
                    share = f'{share_m.group(1)}/100'
                elif share_m.group(2):
                    share = 'equal'
                else:
                    share = share_m.group(3) or ''
            else:
                name_token = part.strip()
                share = 'equal'
            # Resolve
            person = None
            tok_low = name_token.lower()
            for p in eligible:
                rel = (p.relationship or '').lower()
                nm = (p.full_name or '').lower()
                if tok_low in nm or tok_low == rel:
                    person = p; break
                if tok_low in ('wife', 'husband', 'spouse') \
                   and rel in ('spouse', 'wife', 'husband'):
                    person = p; break
            if not person:
                return None
            sub_list.append({'name': person.full_name, 'share': share})
        # Normalise equal shares
        if any(s.get('share') in ('equal', '') for s in sub_list):
            n = len(sub_list)
            for s in sub_list:
                s['share'] = f'1/{n}' if n > 1 else '1/1'

    # Persist + stamp residuary_confirmed (Layer 2 + Layer 3 done)
    s6['substitute_mode'] = sub_mode or 'none'
    s6['substitute_specific'] = sub_list if sub_list else None
    will.step6_data = json.dumps(s6)
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
    except Exception:
        completed = []
    if 'residuary_confirmed' not in completed:
        completed.append('residuary_confirmed')
        will.completed_steps = json.dumps(completed)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {
        'kind': 'residuary_substitute_saved',
        'name': ', '.join(s.get('name', '') for s in sub_list) if sub_list
                else 'no substitute',
        'role': sub_mode,
    }


def _try_save_residuary_main(client_id: str, user_text: str):
    """🔥 §10x.114 — Step 7 (Residuary Estate) main-beneficiary saver.

    Parses natural-language replies to the residuary question:
      'wife 100%'                           → spouse, 100%
      'Joshua 50%, Esther 50%'              → 2 beneficiaries
      'LIM BEE YAN equal, JOSHUA equal'     → 2 beneficiaries equal split
      'Joshua, Esther equal'                → 2 beneficiaries equal split
      '<all-listed-equal-shares default>'   → handled when label clicked

    Writes to will.step6_data:
      {
        'beneficiaries': [{'name': '...', 'share': '...'}],
        'residuary_beneficiary_name': '...',  # legacy compat
      }

    Returns dict on success, None if the input doesn't look like a
    residuary main-beneficiary reply (so the chain falls through to
    _try_handle_residuary_skip / generic handlers).
    """
    if not user_text:
        return None
    t = user_text.strip()
    low = t.lower()

    # Don't intercept SKIP — let _try_handle_residuary_skip handle it
    if low == 'residuary skip':
        return None
    # Don't intercept asset-step replies like 'inventory confirm', 'orphan_*',
    # 'bank_l*', 'insurance_l*', 'gift', 'substitute', 'trust *', 'others *'
    _BAD_PREFIXES = ('inventory ', 'orphan_', 'bank_', 'insurance_',
                     'gift ', 'substitute ', 'trust ', 'others ',
                     'guardian ', 'role_match', 'address:', 'lot:',
                     'title:', 'mukim:', 'daerah:', 'negeri:',
                     'property ', 'change ', 'confirm assets',
                     'confirm defaults', 'h3 ', 'unlink ', 'restart ',
                     'inbox ', 'conflict ', 'i have more',
                     'beneficiaries ', 'residuary substitute',
                     'residuary skip')
    if any(low.startswith(p) for p in _BAD_PREFIXES):
        return None
    # Don't intercept "yes/skip/delete/confirm" alone — too generic
    if low in ('yes', 'no', 'skip', 'delete', 'confirm', 'remove'):
        return None

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None

    # Gate: only fire when planner is at Step 7. Pre-conditions:
    #   - step4_data (beneficiaries) populated
    #   - step5_data (gifts) populated OR assets_confirmed
    #   - step6_data residuary NOT yet saved
    try:
        s4 = json.loads(will.step4_data or '[]')
        if not isinstance(s4, list): s4 = []
    except Exception:
        s4 = []
    if len(s4) == 0:
        return None  # not at Step 5+ yet
    try:
        s6 = json.loads(will.step6_data or '{}')
        if not isinstance(s6, dict): s6 = {}
    except Exception:
        s6 = {}
    # Already saved? Don't re-save.
    if s6.get('beneficiaries') or s6.get('residuary_beneficiary_name'):
        return None
    # Need at least one specific gift OR assets_confirmed to be at Step 7.
    try:
        s5 = json.loads(will.step5_data or '[]')
    except Exception:
        s5 = []
    if not isinstance(s5, list):
        s5 = []
    try:
        completed = json.loads(will.completed_steps or '[]')
    except Exception:
        completed = []
    if len(s5) == 0 and 'assets_confirmed' not in completed:
        return None  # not at Step 7

    # Parse name + share from user text. Support:
    #   'wife 100%' / 'spouse 100%'
    #   'Joshua 50%, Esther 50%'
    #   'LIM BEE YAN equal, JOSHUA equal'
    #   'Joshua, Esther equal'
    persons = Person.query.filter_by(client_id=client_id).all()
    eligible = [p for p in persons
                if (p.relationship or '').lower() not in ('testator', 'witness')]

    def _resolve_name(token: str):
        token_low = token.strip().lower().replace('-', ' ')
        if not token_low:
            return None
        for p in eligible:
            nm = (p.full_name or '').lower()
            rel = (p.relationship or '').lower().replace('-', ' ')
            # Exact name match
            if nm and (token_low == nm or token_low in nm):
                return p
            # First-name token match
            if nm and any(token_low == part for part in nm.split() if len(part) > 2):
                return p
            # Relationship match (wife / spouse / son / daughter / etc.)
            if rel and (token_low == rel or token_low in rel):
                return p
            # Spouse synonyms
            if token_low in ('wife', 'husband', 'spouse') and rel in ('spouse', 'wife', 'husband'):
                return p
        return None

    # Split on commas / 'and' — each part is "name [share]"
    parts = [p.strip() for p in re.split(r',|\band\b', t) if p.strip()]
    parsed: List[Dict[str, str]] = []
    for part in parts:
        # Try to extract a share at the end (e.g. "Joshua 50%" or "wife 100%")
        share_m = re.search(
            r'\s+(\d{1,3})\s*%|\s+(equal|equally)\b|\s+(\d+/\d+)\s*$',
            part, flags=re.IGNORECASE)
        if share_m:
            name_token = part[:share_m.start()].strip().rstrip(',')
            if share_m.group(1):
                share = f'{share_m.group(1)}%'
            elif share_m.group(2):
                share = 'equal'
            else:
                share = share_m.group(3) or ''
        else:
            name_token = part.strip()
            share = 'equal'  # default when no share specified

        person = _resolve_name(name_token)
        if not person:
            # If we can't resolve any name, give up — don't half-save
            return None
        parsed.append({
            'name': person.full_name,
            'share': share,
            'person_id': person.id,
            'relationship': person.relationship or '',
        })

    if not parsed:
        return None

    # Normalise shares: if any part says 'equal', distribute equally
    if any(p.get('share') in ('equal', '') for p in parsed):
        n = len(parsed)
        for p in parsed:
            p['share'] = '1/' + str(n) if n > 1 else '1/1'
    else:
        # If user gave percentages, convert to fractions
        for p in parsed:
            sh = p.get('share') or ''
            pct_m = re.match(r'(\d+)\s*%', sh)
            if pct_m:
                p['share'] = f'{int(pct_m.group(1))}/100'

    # Persist
    s6['beneficiaries'] = parsed
    s6['residuary_beneficiary_name'] = parsed[0]['name']  # legacy compat
    s6.pop('skipped', None)
    will.step6_data = json.dumps(s6)
    # 🔥 §10x.116 — DON'T stamp residuary_confirmed here. Layer 3
    # (substitute) hasn't been answered yet. Stamp only when both
    # main + substitute saved (in skip handler or substitute handler).
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    return {
        'kind': 'residuary_main_saved',
        'name': ', '.join(p['name'] for p in parsed),
        'role': f'{len(parsed)} main residuary beneficiar' + ('y' if len(parsed) == 1 else 'ies'),
    }


def _try_handle_residuary_skip(client_id: str, user_text: str):
    """If user taps 'residuary skip', mark residuary_confirmed with no
    beneficiaries so the planner advances past step 7.

    🔥 BURN-IN (CLAUDE.md §10hg + sleep directive #5):
    BLOCK residuary advancement while ANY specific gift in step5_data
    is incomplete (missing address, missing Layer 1 confirmation, or
    missing beneficiaries). Both layers must be done for every gift
    BEFORE the user can move on to residuary."""
    if not user_text:
        return None
    t = user_text.strip().lower()
    if t != 'residuary skip':
        return None
    will = _get_or_create_will(client_id)

    # ── BURN-IN GATE: every specific gift must have both layers done ──
    try:
        s5 = json.loads(will.step5_data) if will.step5_data else []
        if not isinstance(s5, list):
            s5 = (s5.get('gifts') or []) if isinstance(s5, dict) else []
    except Exception:
        s5 = []
    incomplete = []
    for gi, g in enumerate(s5):
        if not isinstance(g, dict):
            continue
        # Property/bank/vehicle gifts must have beneficiaries (Layer 2).
        # Skipped entries (kind=skip, _ai_summary_skipped, or
        # explicit skipped=True) are exempt — user said no.
        if (g.get('kind') == 'skip'
                or g.get('skipped')
                or g.get('_ai_summary_skipped')):
            continue
        bens = g.get('beneficiaries') or []
        addr = ((g.get('property_info') or {}).get('property_address')
                or g.get('property_address') or '').strip()
        is_property = (g.get('kind') == 'property'
                       or g.get('asset_type') == 'property'
                       or addr)
        if is_property:
            # Layer 1 check
            if g.get('_h3_placeholder') and not g.get('_layer1_confirmed'):
                incomplete.append((gi, 'layer1_not_confirmed'))
                continue
            if not addr:
                incomplete.append((gi, 'missing_address'))
                continue
            if not bens:
                incomplete.append((gi, 'missing_beneficiaries'))
                continue
        elif g.get('kind') == 'bank' or g.get('asset_type') == 'bank' or g.get('bank_name'):
            if not bens:
                incomplete.append((gi, 'bank_missing_beneficiaries'))
                continue
    # ── §10x.12 GATE: every AI-Summary bank+insurance must have a gift ──
    try:
        from ai.chat_planner import (_extract_ai_summary_banks,
                                       _extract_ai_summary_insurance)
        ai_banks = _extract_ai_summary_banks(client_id) or []
        ai_ins = _extract_ai_summary_insurance(client_id) or []
        saved_acct = set()
        saved_pol = set()
        for g in s5:
            if not isinstance(g, dict):
                continue
            an = (g.get('account_number') or '').strip()
            if an:
                saved_acct.add(re.sub(r'\W+', '', an))
            pn = (g.get('policy_number') or '').strip()
            if pn:
                saved_pol.add(re.sub(r'\W+', '', pn))
        for b in ai_banks:
            ack = re.sub(r'\W+', '', b.get('account_number') or '')
            if ack and ack not in saved_acct:
                incomplete.append((-1, f'bank_missing:{b.get("bank_name")}/{b.get("account_number")}'))
        for ins in ai_ins:
            pol = re.sub(r'\W+', '', ins.get('policy_number') or '')
            if pol and pol not in saved_pol:
                incomplete.append((-1, f'insurance_missing:{ins.get("insurer")}/{ins.get("policy_number")}'))
    except Exception:
        pass
    if incomplete:
        return {'name': 'residuary', 'role': 'blocked',
                'kind': 'residuary_blocked',
                'reason': f'{len(incomplete)} specific gifts incomplete: '
                          + ', '.join(f'gift[{i}]={r}' for i, r in incomplete[:5])}

    # Write an empty step6 so the planner sees it as 'done'
    try:
        s6 = json.loads(will.step6_data or '{}') if will.step6_data else {}
        if not isinstance(s6, dict):
            s6 = {}
        s6['skipped'] = True
        s6['beneficiaries'] = []
        will.step6_data = json.dumps(s6)
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {'name': 'residuary', 'role': 'skipped', 'kind': 'residuary_skipped'}


def _try_delete_pending_gift(client_id: str, user_text: str):
    """If user types 'delete' / 'remove' / 'wrong' at a Step-6 gift question
    (property/bank/vehicle), soft-delete the focused Document so it stops
    re-appearing in the walk-through. Returns {'name','action','count'} or
    None.

    This was missing — the chat planner was offering a "Delete" button on
    each property card but no handler picked it up, so taps did nothing
    and the same property kept being asked over and over.
    """
    if not user_text:
        return None
    text_lower = user_text.lower().strip()
    words = set(re.findall(r'\b[a-z]+\b', text_lower))
    if not any(t in words for t in _DELETE_TOKENS):
        return None
    from services.gift_walker import get_pending_gift_documents
    pend = get_pending_gift_documents(client_id)
    # Pick the first pending of any kind, in the same priority the planner
    # would ask about (property → bank → vehicle).
    target = None
    target_kind = None
    for kind in ('property', 'bank', 'vehicle'):
        items = pend.get(kind) or []
        if items:
            target = items[0]
            target_kind = kind
            break
    if not target:
        return None
    doc = db.session.get(Document, target['document_id'])
    if not doc:
        return None
    doc.category = 'deleted'
    doc.description = '(removed by user from chat walk-through)'
    # Also soft-delete any support docs that were grouped with this title
    # (e.g. SPA + cukai tanah images for the same lot). Otherwise they
    # could orphan and clutter the doc library.
    if target_kind == 'property':
        for s in (target.get('support_docs') or []):
            sdoc = db.session.get(Document, s.get('document_id'))
            if sdoc:
                sdoc.category = 'deleted'
                sdoc.description = '(removed with parent property)'
    db.session.commit()
    # Friendly label so the planner can acknowledge the delete cleanly
    ex = target.get('extracted', {}) or {}
    if target_kind == 'property':
        label = ex.get('property_address') or ex.get('title_number') or 'this property'
    elif target_kind == 'bank':
        label = (f"{ex.get('bank_name', '')} {ex.get('account_number', '')}"
                 .strip() or 'this bank account')
    else:
        label = (ex.get('registration_number') or ex.get('vehicle_make')
                 or 'this vehicle')
    return {'name': label[:100], 'action': 'deleted', 'count': 1,
            'kind': f'gift_{target_kind}_delete'}


def _pct_to_frac(share_str: str) -> str:
    """Convert a share string to fractional form for the will clause.

    '50%' → '1/2'  |  '33%' → '1/3'  |  '100%' → '1/1'
    '1/2' → '1/2'  (fractions already fine)
    'equal' → 'equal'  (keyword kept as-is)
    """
    s = (share_str or '').strip()
    if not s or s.lower() in ('equal', 'equally', 'all', ''):
        return s
    if '%' in s:
        try:
            from fractions import Fraction
            n = float(s.rstrip('%'))
            if n == 0:
                return s
            f = Fraction(n / 100).limit_denominator(20)
            # Fraction(1,1) str() → "1" not "1/1"; always use numerator/denominator form
            return f'{f.numerator}/{f.denominator}'
        except Exception:
            return s
    return s  # already fraction or other string


def _try_save_property_gift(client_id: str, user_text: str):
    """Step 6 (Property gift) handler — two-phase: main beneficiary then substitute.

    Phase A (main): user says "Joshua 1/2, Esther 1/2" or "Joshua 50%"
      → shares stored as fractions (50% → 1/2).
      → saves to extracted_data, sets _main_beneficiary_set: True.
      → returns 'gift_main' kind so planner shows substitute prompt next turn.

    Phase B (substitute): _main_beneficiary_set detected.
      Substitute mode from wizard design:
        'equal'   → surviving MBs get equal shares
        'prorata' → surviving MBs get pro-rata shares
        'specific'→ named specific substitute(s)
        'none'    → no substitute clause
      → saves full gift to step5_data → returns 'gift' kind.

    Accepts replies like:
      - "Joshua 1/2, Esther 1/2"      → main (fractional)
      - "Joshua 100%"                  → main → converted to 1/1
      - "Wife"                         → resolves via relationship
      - "skip"                         → skip gift (Phase A)
      - "substitute equal"             → Phase B: equal surviving MBs
      - "substitute prorata"           → Phase B: pro-rata surviving MBs
      - "substitute specific <name>"   → Phase B: named substitute
      - "gift substitute skip"         → Phase B: no substitute
    """
    if not user_text:
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    # 🔥 §10x.90 — Property-gift beneficiary candidate pool ALWAYS merges
    # step4_data + identities + Person rows (minus testator). Earlier we
    # only fell back to identities when step4 was EMPTY, but the §10x.42
    # reconciler often adds ONE auto-detected beneficiary (e.g. wife for
    # bank gifts) which kept step4 non-empty AND prevented the merge.
    # Then a property gift to a NAMED beneficiary (e.g. "to son Joshua")
    # couldn't parse because Joshua wasn't in step4 → bank handler
    # mis-claimed the click → 5 bank accounts wrongly saved.
    try:
        s4_saved = json.loads(will.step4_data or '[]')
        if not isinstance(s4_saved, list): s4_saved = []
    except (json.JSONDecodeError, TypeError):
        s4_saved = []
    try:
        _idents = json.loads(will.identities_data or '[]')
        if not isinstance(_idents, list): _idents = []
    except (json.JSONDecodeError, TypeError):
        _idents = []
    try:
        _s1_name = (json.loads(will.step1_data or '{}') or {}).get('full_name', '').upper()
    except (json.JSONDecodeError, TypeError):
        _s1_name = ''
    # Also pull every Person row — handles the case where identity_walker
    # ensured a Person but identities_data wasn't refreshed.
    try:
        _persons = Person.query.filter_by(client_id=client_id).all()
    except Exception:
        _persons = []
    s4 = list(s4_saved)
    seen_names = {(p.get('full_name') or '').upper().strip() for p in s4}
    seen_names.discard('')
    for i in _idents:
        nm = (i.get('full_name') or '').strip()
        nm_up = nm.upper()
        if nm and nm_up != _s1_name and nm_up not in seen_names:
            s4.append({'full_name': nm, 'relationship': i.get('relationship', '')})
            seen_names.add(nm_up)
    for p in _persons:
        nm = (p.full_name or '').strip()
        nm_up = nm.upper()
        rel = (p.relationship or '').lower()
        if nm and nm_up != _s1_name and rel != 'testator' and nm_up not in seen_names:
            s4.append({'full_name': nm, 'relationship': p.relationship or ''})
            seen_names.add(nm_up)
    if not s4:
        return None

    from services.gift_walker import get_pending_gift_documents, parse_beneficiary_shares

    # Priority: if _get_layer2_pending_props has items, we're in Layer 1/Layer 2
    # interleaving mode and the chat card is showing a property from that queue.
    # _try_save_property_gift MUST target the same doc the card displayed, so
    # we check _get_layer2_pending_props FIRST.
    # Fallback to get_pending_gift_documents only when no Layer 2 queue exists
    # (i.e. normal post-assets_confirmed Step 6 flow).
    _using_layer2_fallback = False
    layer2_pend = _get_layer2_pending_props(client_id)
    if layer2_pend:
        target = layer2_pend[0]
        _using_layer2_fallback = True
    else:
        pend = get_pending_gift_documents(client_id)
        pending_props = pend.get('property') or []
        if not pending_props:
            return None
        target = pending_props[0]

    doc_id = target['document_id']

    # ╔════════════════════════════════════════════════════════════════╗
    # ║ 🔥 BURN-IN §10hg — H3 SYNTHETIC TARGET (no Document)             ║
    # ║ When target represents a step5 H3 placeholder (no Document row),  ║
    # ║ phase-A/B state is stored on the step5 gift itself, not on a Doc. ║
    # ╚════════════════════════════════════════════════════════════════╝
    _h3_step5_idx = target.get('_h3_step5_idx')
    if _h3_step5_idx is None:
        _h3_step5_idx = (target.get('extracted') or {}).get('_h3_step5_idx')
    _is_h3 = _h3_step5_idx is not None

    if _is_h3:
        # Load step5 directly — no Document.
        try:
            _gifts_for_h3 = json.loads(will.step5_data or '[]')
            if not isinstance(_gifts_for_h3, list):
                _gifts_for_h3 = []
        except (json.JSONDecodeError, TypeError):
            _gifts_for_h3 = []
        if not (0 <= _h3_step5_idx < len(_gifts_for_h3)):
            return None
        _h3_gift = _gifts_for_h3[_h3_step5_idx] or {}
        # Phase state is stored ON THE GIFT (not Document.extracted_data).
        doc = None
        doc_ex = {
            '_main_beneficiary_set': bool(_h3_gift.get('_main_beneficiary_set')),
            '_main_beneficiaries':   _h3_gift.get('_main_beneficiaries') or [],
        }
    else:
        # Load the document's extracted_data to check phase state
        doc = db.session.get(Document, doc_id)
        if not doc:
            return None
        try:
            doc_ex = json.loads(doc.extracted_data) if doc.extracted_data else {}
        except (json.JSONDecodeError, TypeError):
            doc_ex = {}

    # Load existing gifts list
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except (json.JSONDecodeError, TypeError):
        gifts = []

    txt = user_text.strip().lower()
    ex_t = (target.get('extracted', {}) or {})
    addr = ex_t.get('property_address', '') or ex_t.get('title_number', '') or 'property'

    # ── Helper: parse beneficiary text ───────────────────────────────
    def _parse_beneficiary(text: str):
        """Parse names + shares; convert % → fraction. Returns list or []."""
        known_names = [p.get('full_name', '') for p in s4 if p.get('full_name')]
        result = parse_beneficiary_shares(text, known_names)
        if not result:
            # Fallback: relationship words
            REL_MAP = {
                'wife': 'spouse', 'husband': 'spouse', 'spouse': 'spouse',
                'son': 'son', 'daughter': 'daughter', 'father': 'father',
                'mother': 'mother', 'children': None, 'kids': None,
            }
            words_lower = re.findall(r'\b[a-z\-]+\b', text.lower())
            matched = []
            for w in words_lower:
                rel = REL_MAP.get(w, w)
                if rel is None:
                    for p in s4:
                        if (p.get('relationship') or '').lower() in ('son', 'daughter'):
                            matched.append(p)
                else:
                    for p in s4:
                        if (p.get('relationship') or '').lower() in (rel, w):
                            matched.append(p)
            seen_n: set = set()
            uniq = []
            for p in matched:
                n = p.get('full_name', '').upper()
                if n and n not in seen_n:
                    seen_n.add(n); uniq.append(p)
            if uniq:
                share = '1/1' if len(uniq) == 1 else 'equal'
                result = [{'name': p['full_name'], 'share': share} for p in uniq]
        # Convert percentages → fractions for NLC will clause format
        for entry in result:
            entry['share'] = _pct_to_frac(entry.get('share') or '1/1')
        return result

    # ── Helper: NLC alert after saving ───────────────────────────────
    def _nlc_alert(ex_data: dict) -> str:
        missing = []
        if not (ex_data.get('title_number') or '').strip():
            missing.append('Title/Geran No.')
        if not (ex_data.get('lot_number') or '').strip():
            missing.append('Lot/PTD No.')
        if not (ex_data.get('mukim') or '').strip():
            missing.append('Mukim')
        if not (ex_data.get('daerah') or '').strip():
            missing.append('Daerah')
        if not (ex_data.get('negeri') or '').strip():
            missing.append('Negeri')
        if not missing:
            return ''
        return (
            "🚨 **Probate alert:** missing **" + ', '.join(missing) + "**"
            " — the lawyer needs these to file _Borang 14A_ at the Pejabat Tanah."
        )

    # ═══════════════════════════════════════════════════════════════
    # PHASE B — substitute beneficiary (main already set)
    # ═══════════════════════════════════════════════════════════════
    if doc_ex.get('_main_beneficiary_set'):
        main_bens = doc_ex.get('_main_beneficiaries') or []

        # Determine substitute mode from user input
        _skip_tokens = ('skip', 'next', 'pass', 'no substitute',
                        'none', 'gift substitute skip', 'no sub')
        substitute_mode = None
        substitute_specific = None  # [{name, share}] for 'specific' mode

        if txt in _skip_tokens:
            substitute_mode = 'none'

        elif txt in ('substitute equal', 'gift substitute equal', 'equal shares',
                     'surviving equal', 'equal'):
            substitute_mode = 'equal'

        elif txt in ('substitute prorata', 'gift substitute prorata', 'pro-rata',
                     'prorata', 'pro rata', 'surviving prorata'):
            substitute_mode = 'prorata'

        elif txt.startswith('substitute specific ') or txt.startswith('gift substitute specific '):
            # "substitute specific SARAH BT ALI" → named specific substitute
            raw = re.sub(r'^(?:gift\s+)?substitute\s+specific\s+', '', txt, flags=re.IGNORECASE).strip()
            substitute_mode = 'specific'
            substitute_specific = [{'name': raw.upper(), 'share': '1/1'}]

        else:
            # Try to parse as a named substitute (free-form)
            sub_parsed = _parse_beneficiary(user_text)
            if sub_parsed:
                substitute_mode = 'specific'
                substitute_specific = sub_parsed
            else:
                # Accept raw text as a named substitute (unnamed person)
                raw = user_text.strip()
                # Allow any text 3+ chars that looks like a name
                if (len(raw) >= 3
                        and re.match(r"^[A-Za-z][A-Za-z '\-,.]+$", raw)):
                    substitute_mode = 'specific'
                    substitute_specific = [{'name': raw.upper(), 'share': '1/1'}]
                else:
                    # Can't understand → keep showing prompt
                    return None

        # Build wizard-compatible gift entry
        # Main allocations: [{beneficiary_name, share, role: 'MB'}]
        allocations = [
            {'beneficiary_name': b.get('name', ''), 'share': b.get('share', '1/1'), 'role': 'MB'}
            for b in main_bens
        ]
        gift_entry = {
            'document_id':    doc_id,
            'kind':           'property',
            'gift_type':      'property',
            # 🔥 §10x.96 — preserve _match_via on Phase B upsert per
            # verifier R12 (§10he Step 5). The original placeholder set
            # this when the user clicked 'inventory confirm'; Phase B
            # MUST carry it through or the gift gets flagged as a
            # silent guess.
            '_match_via':         'user_confirmed',
            '_match_tier':        'A',
            '_match_confidence':  'high',
            '_layer1_confirmed':  True,
            'property_details': {
                'property_address': ex_t.get('property_address', ''),
                'title_number':     ex_t.get('title_number', ''),
                'lot_number':       ex_t.get('lot_number', ''),
                'mukim':            ex_t.get('mukim', '') or ex_t.get('bandar_pekan', ''),
                'daerah':           ex_t.get('daerah', ''),
                'negeri':           ex_t.get('negeri', ''),
                'encumbrance_status': (
                    'encumbered' if ex_t.get('encumbrance_confirmed') else 'clean'
                ),
                'undivided_share': True,
                'testator_share':  ex_t.get('ownership_share', ''),
            },
            'allocations':        allocations,
            'substitute_mode':    substitute_mode or 'none',
            # For 'specific' mode: store substitutes inside each MB's allocations
            # (wizard expects alloc.substitutes) — attach to first MB for simplicity
            'substitute_specific': substitute_specific,
            # Keep legacy fields for backward compat with normalise_gifts
            'property_address': ex_t.get('property_address', ''),
            'title_number':     ex_t.get('title_number', ''),
            'beneficiaries':    main_bens,
        }
        # Embed specific substitutes inside each MB allocation (wizard format)
        if substitute_mode == 'specific' and substitute_specific:
            for alloc in allocations:
                alloc['substitutes'] = [
                    {'beneficiary_name': s.get('name',''), 'share': s.get('share','1/1')}
                    for s in substitute_specific
                ]

        # ╔══════════════════════════════════════════════════════════╗
        # ║  🔥 BURN-IN — NO DUPLICATE GIFTS (upsert path) 🔥          ║
        # ║  Match by document_id OR by (lot_digits, addr_sig) so a   ║
        # ║  sibling doc with OCR-drifted title still upserts onto    ║
        # ║  the same gift. See CLAUDE.md §10f.                       ║
        # ╚══════════════════════════════════════════════════════════╝
        from services.gift_walker import (_clean_id_value, _looks_like_garbage,
                                          _norm_addr, _is_strata, _title_signature,
                                          _is_genuinely_different_unit)
        new_lot = _clean_id_value(ex_t.get('lot_number', '') or '')
        if _looks_like_garbage(new_lot):
            new_lot = ''
        new_lot_digits = re.sub(r'\D', '', new_lot)
        new_addr_sig = _norm_addr(ex_t.get('property_address', '') or '')[:60]
        new_strata = _is_strata(ex_t)
        new_title_sig = _title_signature(ex_t)
        def _gift_field2(g, key):
            return ((g.get('property_info') or {}).get(key)
                    or (g.get('property_details') or {}).get(key)
                    or g.get(key) or '')
        _existing_idx = None
        for i, g in enumerate(gifts):
            if g.get('document_id') == doc_id:
                _existing_idx = i
                break
            g_lot = _clean_id_value(_gift_field2(g, 'lot_number'))
            if _looks_like_garbage(g_lot):
                g_lot = ''
            g_lot_digits = re.sub(r'\D', '', g_lot)
            g_addr_sig = _norm_addr(_gift_field2(g, 'property_address'))[:60]
            # 🔥 STRATA EXCEPTION (§10hd): genuinely different unit in same
            # building → don't upsert. Not OCR truncation of same unit.
            g_ex_for_title = {
                'title_number': _gift_field2(g, 'title_number'),
                'title_type': _gift_field2(g, 'title_type'),
                'property_description': _gift_field2(g, 'property_description'),
                'document_type': _gift_field2(g, 'document_type'),
            }
            g_strata = _is_strata(g_ex_for_title)
            g_title_sig = _title_signature(g_ex_for_title)
            if ((new_strata or g_strata)
                and _is_genuinely_different_unit(new_title_sig, g_title_sig)):
                continue   # different strata unit — not the same gift
            if (new_lot_digits and g_lot_digits and new_lot_digits == g_lot_digits
                and new_addr_sig and g_addr_sig and new_addr_sig == g_addr_sig):
                _existing_idx = i
                break
        # H3 path: replace at the synthetic step5 index.
        if _is_h3:
            if 0 <= _h3_step5_idx < len(gifts):
                # Preserve the H3 marker + AI-Summary linkage so verifier
                # can attribute this gift back to its summary entry.
                _existing = gifts[_h3_step5_idx] or {}
                gift_entry['_h3_placeholder']   = True
                gift_entry['_layer1_confirmed'] = True
                if _existing.get('_ai_summary_idx') is not None:
                    gift_entry['_ai_summary_idx'] = _existing.get('_ai_summary_idx')
                # Drop the doc_id (it was synthetic).
                gift_entry['document_id'] = _existing.get('document_id') or None
                gifts[_h3_step5_idx] = gift_entry
            else:
                gifts.append(gift_entry)
        elif _existing_idx is not None:
            gifts[_existing_idx] = gift_entry
        else:
            gifts.append(gift_entry)
        will.step5_data = json.dumps(gifts)

        # 🔥 §10x.93 — Propagate gift beneficiaries into step4_data so the
        # wizard's Step 5 (Beneficiaries) and Step 6 (Specific Gifts)
        # dropdowns are populated. Without this, the wizard's gift form
        # renders '-- Select Beneficiary --' empty for every saved gift
        # whose beneficiary names aren't already in step4 (the §10x.42
        # reconcile only added Lim Bee Yan there for bank gifts).
        try:
            _s4_existing = json.loads(will.step4_data or '[]')
            if not isinstance(_s4_existing, list):
                _s4_existing = []
        except Exception:
            _s4_existing = []
        _s4_names = {(b.get('full_name') or '').upper().strip()
                     for b in _s4_existing if isinstance(b, dict)}
        _s4_names.discard('')
        _added_to_s4 = False
        # Pull from main + substitute beneficiaries — both are real people
        for entry in (main_bens or []) + (substitute_specific or []):
            nm = (entry.get('name') or '').strip()
            if not nm or nm.upper() in _s4_names:
                continue
            # Find their Person row to get NRIC + relationship + address
            _p = next((p for p in (Person.query.filter_by(client_id=client_id).all() or [])
                       if (p.full_name or '').upper().strip() == nm.upper()), None)
            _s4_existing.append({
                'full_name':     nm,
                'nric_passport': (_p.nric_passport if _p else '') or '',
                'address':       (_p.address if _p else '') or '',
                'relationship':  (_p.relationship if _p else '') or '',
                'person_id':     (_p.id if _p else ''),
                '_added_by':     '§10x.93 propagate from gift',
            })
            _s4_names.add(nm.upper())
            _added_to_s4 = True
        if _added_to_s4:
            will.step4_data = json.dumps(_s4_existing)

        # Clear the phase flag in extracted_data (only if real Document)
        if doc is not None:
            doc_ex['_main_beneficiary_set'] = False
            doc_ex['_substitute_assigned']  = True
            doc.extracted_data = json.dumps(doc_ex)

        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None

        main_desc = ', '.join(f"{b.get('name','?')} {b.get('share','')}" for b in main_bens)
        sub_label = {
            'equal':    'equal shares among survivors',
            'prorata':  'pro-rata among survivors',
            'specific': ', '.join(s.get('name','?') for s in (substitute_specific or [])),
            'none':     'no substitute',
        }.get(substitute_mode or 'none', substitute_mode)
        alert = _nlc_alert(ex_t)
        return {
            'name':  f"{addr[:50]} → {main_desc}",
            'role':  sub_label,
            'kind':  'gift',
            'alert': alert,
        }

    # ═══════════════════════════════════════════════════════════════
    # PHASE A — main beneficiary
    # ═══════════════════════════════════════════════════════════════

    # "skip" on Phase A → skip the whole gift (sentinel entry)
    if txt in ('skip', 'next', 'pass'):
        _skip_entry = {
            'document_id': doc_id,
            'kind': 'property',
            'skipped': True,
            'beneficiaries': [],
        }
        # Upsert: replace placeholder if already present
        _si = next((i for i, g in enumerate(gifts) if g.get('document_id') == doc_id), None)
        if _si is not None:
            gifts[_si] = _skip_entry
        else:
            gifts.append(_skip_entry)
        will.step5_data = json.dumps(gifts)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return None
        return {'name': addr[:80], 'role': 'skipped', 'kind': 'gift_skip'}

    parsed = _parse_beneficiary(user_text)
    if not parsed:
        return None

    # Save main beneficiary intent (Phase A)
    if _is_h3:
        # Persist on the step5 gift itself.
        try:
            _gh3 = json.loads(will.step5_data or '[]')
            if not isinstance(_gh3, list):
                _gh3 = []
        except (json.JSONDecodeError, TypeError):
            _gh3 = []
        if 0 <= _h3_step5_idx < len(_gh3):
            _gh3[_h3_step5_idx]['_main_beneficiary_set'] = True
            _gh3[_h3_step5_idx]['_main_beneficiaries']   = parsed
            will.step5_data = json.dumps(_gh3)
    else:
        doc_ex['_main_beneficiary_set'] = True
        doc_ex['_main_beneficiaries']   = parsed
        doc.extracted_data = json.dumps(doc_ex)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    main_desc = ', '.join(f"{b['name']} {b.get('share','')}" for b in parsed)
    return {
        'name': main_desc,
        'role': f'main beneficiary set for {addr[:40]}',
        'kind': 'gift_main',
    }


def _try_save_bank_gift(client_id: str, user_text: str):
    """Step 6 (Bank gift) handler. Counterpart to _try_save_property_gift.

    The bank question (`_step6_bank_question`) is single-shot:
       "Who inherits all your bank accounts?"
       quick replies: <beneficiary name> | "Walk through one by one"

    Behaviour:
      • If pending bank statements exist AND no bank gift saved yet:
          - "walk one by one" → flip a per-bank walkthrough flag
            (not yet implemented; falls through to None for now).
          - Anything that parses as a beneficiary (or a known person
            name) → save ONE gift entry per pending bank doc, all
            assigned to that beneficiary 100%.

    Returns {'name','role','kind':'gift_bank'} on save, else None.

    🔥 BURN-IN: dedup at insert site by (document_id) + by
    (institution, last4) so the same bank doc never produces two gifts.
    """
    if not user_text:
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None

    from services.gift_walker import (get_pending_gift_documents,
                                      parse_beneficiary_shares)
    pend = get_pending_gift_documents(client_id)
    pending_banks = pend.get('bank') or []
    if not pending_banks:
        return None

    # Already a bank gift saved? Then this question shouldn't fire and
    # any text the user types isn't ours to consume.
    try:
        gifts = json.loads(will.step5_data or '[]')
        if not isinstance(gifts, list):
            gifts = []
    except (json.JSONDecodeError, TypeError):
        gifts = []
    has_bank_gift = any(
        isinstance(g, dict) and (
            g.get('kind') == 'bank'
            or g.get('asset_type') == 'bank'
            or g.get('bank_name')
            or (g.get('property_info') or {}).get('account_no')
            or (g.get('property_details') or {}).get('account_no')
            or (g.get('financial_details') or {}).get('account_number')
        )
        for g in gifts
    )
    if has_bank_gift:
        return None

    txt = user_text.strip().lower()

    # "walk one by one" → not implemented yet; bail so the same prompt
    # re-renders and the user can pick a single name.
    if txt in ('walk one by one', 'walk through one by one',
               'walk-one-by-one', 'one by one'):
        return None

    # Candidate beneficiary pool: ALWAYS merge step4 (saved beneficiaries) +
    # identities + Person table, minus testator. The bank question's quick
    # replies only show step4 names but the user can type anyone — e.g.
    # "wife (Lim Lay Cheng)" who may be in identities but not yet in step4.
    try:
        s4 = json.loads(will.step4_data or '[]')
    except (json.JSONDecodeError, TypeError):
        s4 = []
    try:
        _idents = json.loads(will.identities_data or '[]')
    except (json.JSONDecodeError, TypeError):
        _idents = []
    try:
        _s1_name = (json.loads(will.step1_data or '{}') or {}).get('full_name', '').upper()
    except (json.JSONDecodeError, TypeError):
        _s1_name = ''
    pool = list(s4)
    seen_pool = {(p.get('full_name') or '').upper().strip() for p in pool}
    seen_pool.discard('')
    for i in _idents:
        nm = (i.get('full_name') or '').strip()
        nm_up = nm.upper()
        if nm and nm_up != _s1_name and nm_up not in seen_pool:
            pool.append({'full_name': nm, 'relationship': i.get('relationship', '')})
            seen_pool.add(nm_up)
    # Also pull Person rows so chat-only persons (no wizard identities row)
    # are matchable.
    for p in Person.query.filter_by(client_id=client_id).all():
        nm = (p.full_name or '').strip()
        nm_up = nm.upper()
        if nm and nm_up != _s1_name and nm_up not in seen_pool:
            pool.append({'full_name': nm, 'relationship': p.relationship or ''})
            seen_pool.add(nm_up)
    s4 = pool
    if not s4:
        return None
    known_names = [p.get('full_name', '') for p in s4 if p.get('full_name')]

    parsed = parse_beneficiary_shares(user_text, known_names)
    if not parsed:
        # Fallback: relationship words (wife/spouse → spouse identity)
        REL_MAP = {
            'wife': 'spouse', 'husband': 'spouse', 'spouse': 'spouse',
            'son': 'son', 'daughter': 'daughter',
        }
        words_lower = re.findall(r'\b[a-z\-]+\b', user_text.lower())
        matched = []
        for w in words_lower:
            rel = REL_MAP.get(w, w)
            for p in s4:
                if (p.get('relationship') or '').lower() == rel:
                    matched.append(p)
        seen_n: set = set()
        uniq = []
        for p in matched:
            n = (p.get('full_name') or '').upper()
            if n and n not in seen_n:
                seen_n.add(n); uniq.append(p)
        if uniq:
            share = '1/1' if len(uniq) == 1 else 'equal'
            parsed = [{'name': p['full_name'], 'share': share} for p in uniq]
    if not parsed:
        return None
    # Convert percentages → fractions
    for entry in parsed:
        entry['share'] = _pct_to_frac(entry.get('share') or '1/1')

    # Build per-bank gift entries
    saved = 0
    seen_keys = set()
    # Pre-seed seen_keys with already-present gift signatures so we never
    # double-insert across already-saved banks (defense-in-depth).
    for g in gifts:
        if not isinstance(g, dict):
            continue
        ad = ((g.get('financial_details') or {}).get('account_number')
              or g.get('account_number') or '')
        inst = ((g.get('financial_details') or {}).get('institution')
                or g.get('bank_name') or '')
        last4 = re.sub(r'\D', '', str(ad))[-4:] if ad else ''
        if inst or last4:
            seen_keys.add((inst.strip().upper(), last4))

    bank_descriptors = []
    for b in pending_banks:
        ex = b.get('extracted', {}) or {}
        institution = (ex.get('bank_name') or ex.get('institution')
                       or b.get('bank_name') or '').strip()
        account_no = (ex.get('account_number') or ex.get('account_no')
                      or b.get('account_number') or '').strip()
        account_type = (ex.get('account_type') or 'savings').strip().lower()
        last4 = re.sub(r'\D', '', str(account_no))[-4:]
        sig = (institution.upper(), last4)
        if sig in seen_keys:
            continue
        seen_keys.add(sig)
        doc_id = b.get('document_id')
        allocations = [
            {'beneficiary_name': p.get('name', ''), 'share': p.get('share', '1/1'), 'role': 'MB'}
            for p in parsed
        ]
        gift_entry = {
            'document_id':    doc_id,
            'kind':           'bank',
            'gift_type':      'financial',
            'asset_type':     'bank',
            'bank_name':      institution,
            'account_number': account_no,
            'account_type':   account_type,
            'financial_details': {
                'institution':    institution,
                'account_number': account_no,
                'asset_type':     'bank',
                'account_type':   account_type,
            },
            'allocations':       allocations,
            'beneficiaries':     parsed,
            'substitute_mode':   'none',
            'substitute_specific': None,
        }
        gifts.append(gift_entry)
        saved += 1
        bank_descriptors.append(
            f"{institution or 'Bank'} …{last4}" if last4 else (institution or 'Bank')
        )

    if not saved:
        return None

    will.step5_data = json.dumps(gifts)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None

    benef_desc = ', '.join(f"{p.get('name','?')} {p.get('share','1/1')}" for p in parsed)
    return {
        'name': benef_desc,
        'role': f'{saved} bank account{"s" if saved != 1 else ""} → ' + ', '.join(bank_descriptors),
        'kind': 'gift_bank',
    }


def _try_save_testator_address(client_id: str, user_text: str):
    """🔥 §10x.122 + §10x.123 — capture testator's compulsory
    Step 2 fields from labelled replies: `address:`, `dob:`, `gender:`,
    `marital:`, `occupation:`.

    Gates (per-field):
      - user_text starts with one of the recognised prefixes
      - testator Person row exists (Step 1 done)
      - that field on step1_data is empty (don't overwrite)
      - For 'address:' specifically — Step 6 not yet started (no saved
        gifts) so it's unambiguously the testator address, not a
        property address.

    Writes to BOTH:
      - Person row (where relationship='Testator')
      - Will.step1_data

    Returns dict on save, None to fall through.
    """
    if not user_text:
        return None
    t = user_text.strip()
    low = t.lower()

    # Map prefix → (s1_key, person_attr, label)
    _FIELD_MAP = {
        'address:':    ('residential_address', 'address',          'Address'),
        'dob:':        ('date_of_birth',       'date_of_birth',    'DOB'),
        'gender:':     ('gender',              'gender',           'Gender'),
        'marital:':    ('marital_status',      'marital_status',   'Marital status'),
        'occupation:': ('occupation',          'occupation',       'Occupation'),
    }
    matched_prefix = None
    for prefix in _FIELD_MAP:
        if low.startswith(prefix):
            matched_prefix = prefix
            break
    if not matched_prefix:
        return None
    s1_key, person_attr, label = _FIELD_MAP[matched_prefix]
    raw_value = t[len(matched_prefix):].strip()
    if not raw_value:
        return None  # bare prefix is the prefill — wait for typing

    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s1 = json.loads(will.step1_data or '{}') or {}
    except Exception:
        s1 = {}

    # 🔥 §10x.124 — 'address:' disambiguation:
    #   testator address EMPTY → save here (Step 2 stage)
    #   testator address SET   → return None so _try_handle_property_fill
    #                            can claim it (Step 6 stage)
    # This handles the dispatch-chain order — testator handler runs
    # FIRST but yields to property_fill when testator address already
    # populated.
    if matched_prefix == 'address:':
        if (s1.get('residential_address') or '').strip():
            return None  # testator already has address → property_fill takes over

    # Don't overwrite already-saved values silently. If user types a
    # field that's already set, let it pass — they're updating.
    # (Just save and return.)

    # Validate / normalise per field
    if s1_key == 'gender':
        v_low = raw_value.lower()
        if v_low.startswith('m'): raw_value = 'Male'
        elif v_low.startswith('f'): raw_value = 'Female'
        else: return None  # unrecognised
    elif s1_key == 'marital_status':
        v_low = raw_value.lower()
        if 'marri' in v_low: raw_value = 'Married'
        elif 'singl' in v_low: raw_value = 'Single'
        elif 'widow' in v_low: raw_value = 'Widowed'
        elif 'divor' in v_low: raw_value = 'Divorced'
        # else: keep raw value
    elif s1_key == 'date_of_birth':
        # Accept DD-MM-YYYY, DD/MM/YYYY, YYYY-MM-DD — normalise to DD-MM-YYYY
        import re as _re_dob
        m = _re_dob.match(r'^(\d{1,2})[-/](\d{1,2})[-/](\d{4})$', raw_value)
        if m:
            d, mo, y = m.group(1), m.group(2), m.group(3)
            raw_value = f'{int(d):02d}-{int(mo):02d}-{y}'
        else:
            m = _re_dob.match(r'^(\d{4})[-/](\d{1,2})[-/](\d{1,2})$', raw_value)
            if m:
                y, mo, d = m.group(1), m.group(2), m.group(3)
                raw_value = f'{int(d):02d}-{int(mo):02d}-{y}'
        # else keep raw — user can type 'circa 1963' etc.

    # Find testator Person row
    testator = (Person.query
                .filter_by(client_id=client_id)
                .filter(Person.relationship.ilike('testator'))
                .first())
    if not testator:
        return None

    # Update Person AND step1_data
    if hasattr(testator, person_attr):
        setattr(testator, person_attr, raw_value[:500])
    s1[s1_key] = raw_value[:500]
    # Backfill name/nric/person_id if missing
    s1['full_name'] = s1.get('full_name') or testator.full_name or ''
    s1['nric_passport'] = s1.get('nric_passport') or testator.nric_passport or ''
    s1['person_id'] = s1.get('person_id') or testator.id
    will.step1_data = json.dumps(s1)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {
        'kind': 'testator_field_saved',
        'name': testator.full_name or 'testator',
        'role': f'{label}: {raw_value[:60]}',
    }


def _try_confirm_testator(client_id: str, user_text: str):
    """🔥 §7 — Step 2 confirm handler.

    The chat planner shows a 'Confirm Testator' card after Step 1 is done.
    When the user clicks ✓ Confirm, save the Testator Person row's data
    into Will.step1_data so `_is_confirmed('testator')` returns True and
    the planner advances to Step 3 (Executor) per §6 wizard order.

    Returns {'name', 'role', 'kind': 'testator_confirmed'} or None.
    """
    if not user_text:
        return None
    text_lower = ' ' + user_text.lower().strip() + ' '
    if not any((' ' + c + ' ') in text_lower for c in _CONFIRM_TOKENS):
        return None
    # Find active draft Will
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s1 = json.loads(will.step1_data or '{}') or {}
    except (json.JSONDecodeError, TypeError):
        s1 = {}
    # Already confirmed? Skip — let the next handler claim the click.
    if s1.get('full_name') and s1.get('person_id'):
        return None
    # Pull Testator Person row
    testator = Person.query.filter_by(client_id=client_id, relationship='Testator').first()
    if not testator:
        return None
    # 🔥 §10x.124 — refuse plain Confirm only when ADDRESS is missing.
    # DOB/gender/marital are AUTO-DERIVED here:
    #   DOB     ← NRIC YYMMDD prefix
    #   Gender  ← NRIC last digit (odd=Male, even=Female)
    #   Marital ← family identities (Wife/Husband present → Married)
    # Per Phek Yi Ting template (§10x.24), only address is typed-required.
    addr = (testator.address or '').strip() or (s1.get('residential_address') or '').strip()
    if not addr:
        return {
            'kind': 'testator_address_required',
            'name': testator.full_name or 'testator',
            'role': 'blocked — address missing',
        }
    nric = testator.nric_passport or s1.get('nric_passport', '')
    # DOB from NRIC
    dob = (testator.date_of_birth or '').strip() or (s1.get('date_of_birth') or '').strip()
    if not dob and nric:
        import re as _re_dob
        m = _re_dob.match(r'^(\d{2})(\d{2})(\d{2})[-\s]?\d{2}[-\s]?\d{4}', nric.strip())
        if m:
            yy, mm, dd = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if 1 <= mm <= 12 and 1 <= dd <= 31:
                from datetime import datetime as _dt
                century = 1900 if yy > (_dt.utcnow().year % 100) else 2000
                dob = f'{dd:02d}-{mm:02d}-{century + yy}'
                if hasattr(testator, 'date_of_birth'):
                    testator.date_of_birth = dob
    # Gender from NRIC last digit
    gender = (testator.gender or '').strip() or (s1.get('gender') or '').strip()
    if not gender and nric:
        import re as _re_g
        digits = _re_g.sub(r'\D', '', nric)
        if digits:
            last = digits[-1]
            if last.isdigit():
                gender = 'Male' if int(last) % 2 == 1 else 'Female'
                if hasattr(testator, 'gender'):
                    testator.gender = gender
    # Marital from family
    marital = (s1.get('marital_status') or '').strip()
    if not marital:
        family = Person.query.filter_by(client_id=client_id).all()
        for p in family:
            if (p.relationship or '').lower() in ('wife', 'husband', 'spouse'):
                marital = 'Married'
                break
        if not marital:
            marital = 'Single'  # safe default
    s1.update({
        'full_name':           testator.full_name or '',
        'nric_passport':       testator.nric_passport or '',
        'date_of_birth':       dob,
        'gender':              gender,
        'marital_status':      marital,
        'residential_address': testator.address or s1.get('residential_address', ''),
        'nationality':         testator.nationality or 'Malaysian',
        'gender':              testator.gender or '',
        'email':               testator.email or '',
        'phone':               testator.phone or '',
        'person_id':           testator.id,
    })
    will.step1_data = json.dumps(s1)
    # Sync Client header
    client = db.session.get(Client, client_id)
    if client:
        if testator.full_name:
            client.full_name = testator.full_name
        if testator.nric_passport:
            client.nric_passport = testator.nric_passport
        will.title = f"Will of {testator.full_name or 'Unknown'}"
    # 🔥 §10x.88 — mark Step 2 confirmed so _current_stage_num can
    # advance past it. Without this, auto-populated executor (via
    # §10x.42) made the planner skip from Step 1 → Step 6.
    try:
        completed = json.loads(will.completed_steps or '[]')
        if not isinstance(completed, list):
            completed = []
        if 'testator_confirmed' not in completed:
            completed.append('testator_confirmed')
            will.completed_steps = json.dumps(completed)
    except Exception:
        pass
    db.session.commit()
    return {
        'name': testator.full_name,
        'role': 'testator',
        'kind': 'testator_confirmed',
    }


def _try_save_executor(client_id: str, user_text: str):
    """If chat is at the executor stage and user replied with a name OR
    'yes' (to confirm the suggested candidate), persist to the active
    Will's step2_data. Returns {'name', 'role'} or None."""
    if not user_text:
        return None
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return None
    try:
        s2 = json.loads(will.step2_data or '{}')
        if not isinstance(s2, dict):
            s2 = {}
    except (json.JSONDecodeError, TypeError):
        s2 = {}
    executors = s2.get('executors') or []
    if len(executors) >= 2:
        return None
    role = 'main' if len(executors) == 0 else 'substitute'

    text_lower = user_text.lower().strip()
    words = set(re.findall(r'\b[a-z]+\b', text_lower))

    # Skip — only valid for substitute
    if role == 'substitute' and any(t in words for t in _SKIP_TOKENS):
        s2['_substitute_skipped'] = True
        will.step2_data = json.dumps(s2)
        db.session.commit()
        return {'name': '(no substitute)', 'role': 'substitute', 'action': 'skipped'}

    # Find a Person matching either a name in the text OR (if 'yes')
    # the planner's suggested candidate.
    persons = Person.query.filter_by(client_id=client_id).all()
    chosen = None
    for p in persons:
        nm = (p.full_name or '').lower()
        if nm and nm in text_lower:
            chosen = p
            break

    if not chosen and any(t in words for t in _CONFIRM_TOKENS):
        # User said yes — apply the candidate the planner suggested
        from ai.chat_planner import find_executor_candidate
        identities = [{
            'id': p.id, 'full_name': p.full_name,
            'relationship': p.relationship or '',
            'date_of_birth': p.date_of_birth or '',
            'document_id': p.document_id or '',
        } for p in persons]
        recent = _gather_recent_chat_text(client_id)
        cand = find_executor_candidate(identities, executors, role, recent)
        if cand:
            chosen = next((p for p in persons if p.id == cand['person_id']), None)

    if not chosen:
        return None

    # Don't add the same person twice
    for e in executors:
        if e.get('person_id') == chosen.id:
            return None

    executors.append({
        'full_name': chosen.full_name,
        'nric_passport': chosen.nric_passport or '',
        'address': chosen.address or '',
        'relationship': chosen.relationship or '',
        'person_id': chosen.id,
        'is_substitute': (role == 'substitute'),
    })
    s2['executors'] = executors
    s2['executor_type'] = 'joint' if len(executors) > 1 else 'single'
    s2['trustee_data'] = s2.get('trustee_data') or {'same_as_executor': True, 'trustees': [{}]}
    will.step2_data = json.dumps(s2)
    # 🔥 §10x.88 — mark Step 3 confirmed when a substitute is picked OR
    # when only one executor is needed and the user explicitly chose
    # them (not the §10x.42 reconcile auto-add). The marker advances
    # _current_stage_num past Step 3.
    if role == 'substitute' or s2.get('_substitute_skipped'):
        try:
            completed = json.loads(will.completed_steps or '[]')
            if not isinstance(completed, list):
                completed = []
            if 'executor_confirmed' not in completed:
                completed.append('executor_confirmed')
                will.completed_steps = json.dumps(completed)
        except Exception:
            pass
    db.session.commit()
    return {'name': chosen.full_name, 'role': f'{role} executor', 'kind': 'executor'}


def _try_assign_pending_identity(client_id: str, user_text: str):
    """If user_text plausibly assigns the next pending IC's role, create
    the Person and return {'name', 'role'}. Else return None."""
    if not user_text:
        return None
    from services.identity_walker import get_pending_ic_documents, parse_relationship
    from services.person_registry import ensure_person
    from ai.role_deducer import deduce_roles

    pending = get_pending_ic_documents(client_id)
    if not pending:
        return None
    target = pending[0]
    ex = target['extracted'] or {}
    name = (ex.get('full_name') or '').strip()
    if not name:
        return None

    text_lower = ' ' + user_text.lower().strip() + ' '
    if any((' ' + s + ' ') in text_lower for s in _SKIP_TOKENS):
        return None  # user said skip — leave for next turn

    # 🔥 §10x.30 / §10x.21 — Step 1 (Identity) is for FAMILY relations
    # only. Will-roles (Executor / Trustee / Guardian / Witness /
    # Beneficiary) are set in LATER steps. If a deducer (LLM) returns
    # a will-role here, we SILENTLY MAP it back to the family relation
    # via outsider-elimination. Saving "LIM LAY CHENG = Executor" in
    # Step 1 would make her invisible to the family identity registry
    # and corrupt later steps that key off relationship.
    _WILL_ROLES = {'Executor', 'Trustee', 'Guardian', 'Witness',
                    'Beneficiary'}
    rel = parse_relationship(user_text)
    chosen_role = None
    if rel and rel not in _WILL_ROLES:
        chosen_role = rel
    elif rel in _WILL_ROLES:
        # User typed/clicked a will-role; we only accept family here.
        # Fall through to deducer + outsider-elimination below to find
        # the correct family relation.
        pass
    if not chosen_role and any((' ' + c + ' ') in text_lower for c in _CONFIRM_TOKENS):
        # User said yes/confirm — apply the deduced role from EITHER:
        #   (a) ai.role_deducer (name-verbatim, FAMILY only) OR
        #   (b) services.role_matcher outsider-elimination (§10x.21)
        recent = _gather_recent_chat_text(client_id)
        ded = deduce_roles(recent, [name])
        # 🔒 §10x.30 — Step 1 only accepts FAMILY roles. If deducer
        # returned a will-role (e.g. 'Executor'), discard it so the
        # outsider-elimination fallback runs instead.
        ded_role = (ded.get(name) or {}).get('role') or ''
        if ded_role and ded_role not in _WILL_ROLES:
            chosen_role = ded_role
        else:
            # Fallback: outsider-elimination — see if THIS IC was matched
            # as the executor's family-relation by role_matcher
            try:
                from services.role_matcher import (
                    extract_role_mentions, find_unassigned_ic_candidates,
                    match_role_to_candidates,
                )
                mentions = extract_role_mentions(client_id) or []
                cands = find_unassigned_ic_candidates(client_id) or []
                this_doc_id = target.get('document_id') or ''
                this_nric_digits = ''.join(
                    ch for ch in (ex.get('nric_number') or '') if ch.isdigit())
                for m in mentions:
                    ranked = match_role_to_candidates(m, cands, client_id=client_id)
                    for c, conf, reason in ranked:
                        if conf != 'high':
                            continue
                        cnric = ''.join(
                            ch for ch in (c.get('nric') or '') if ch.isdigit())
                        if (c.get('document_id') == this_doc_id
                                or (cnric and this_nric_digits
                                    and cnric == this_nric_digits)):
                            fam = (m.get('family_relation') or 'sister-in-law').strip()
                            # Normalize to canonical Title-Case used in the
                            # Person table (e.g. 'sister-in-law' → 'Sister-in-law')
                            chosen_role = '-'.join(
                                p.capitalize() for p in fam.split('-')
                            ) if '-' in fam else fam.capitalize()
                            break
                    if chosen_role:
                        break
            except Exception:
                pass

    if not chosen_role:
        return None

    # 🔥 §10x.143b — H3 placeholder backfill from existing IC docs.
    # When confirming an H3 placeholder (no IC), search for any uploaded
    # IC Document whose extracted name matches this person and pull its
    # NRIC + address + doc_id. Fixes the case where Joshua's IC was
    # marked 'duplicate' (sibling-dedup) before any Person row existed,
    # so the H3 placeholder confirm landed without NRIC even though
    # the data was right there in the doc.
    h3_nric = (ex.get('nric_number') or '').strip()
    h3_address = (ex.get('address') or '').strip()
    h3_doc_id = target.get('document_id')
    if not h3_nric and not h3_doc_id and target.get('_h3_placeholder'):
        try:
            name_upper = name.strip().upper()
            ic_docs = Document.query.filter(
                Document.client_id == client_id,
                Document.category.in_(('nric', 'duplicate')),
            ).all()
            for d in ic_docs:
                try:
                    d_ex = json.loads(d.extracted_data or '{}')
                except Exception:
                    d_ex = {}
                d_name = (d_ex.get('full_name') or '').strip().upper()
                d_nric = (d_ex.get('nric_number') or '').strip()
                if d_name == name_upper and d_nric:
                    h3_nric = d_nric
                    h3_address = h3_address or (d_ex.get('address') or '').strip()
                    h3_doc_id = h3_doc_id or d.id
                    break
        except Exception:
            pass

    pid = ensure_person(
        client_id, name,
        nric=h3_nric,
        address=h3_address,
        relationship=chosen_role,
        dob=(ex.get('date_of_birth') or ''),
        nationality=ex.get('nationality') or 'Malaysian',
        document_id=h3_doc_id,
    )
    db.session.commit()

    # 🔥 §10x.143c — propagate backfilled NRIC/address into any step2/4/5
    # entries that already reference this Person (by person_id OR name).
    # Without this, the will generates '(MALAYSIA NRIC No. )' blanks
    # because step2/step4 cached the empty values from the H3 placeholder
    # before the IC arrived.
    try:
        if (h3_nric or h3_address) and pid:
            _propagate_person_to_steps(client_id, pid, name, h3_nric, h3_address)
    except Exception:
        import traceback as _tb
        app.logger.warning(f'§10x.143c propagation failed: {_tb.format_exc()}')
    # 🔥 §10x.42 — Mid-flow identity add MUST trigger downstream
    # reconciliation. If this person is named as a beneficiary in the
    # message (e.g. "All my Bank Savings go my wife (Lim Bee Yan)"),
    # auto-add them to step4 (Beneficiaries) and re-evaluate any
    # already-saved gifts that should reference them.
    try:
        _reconcile_downstream_for_new_identity(client_id, name, chosen_role)
    except Exception:
        import traceback as _tb
        app.logger.error(
            f'§10x.42 reconciliation failed for {name}: {_tb.format_exc()}')
    return {'name': name, 'role': chosen_role, 'kind': 'identity'}


def _propagate_person_to_steps(client_id: str, person_id: str,
                                 name: str, nric: str, address: str) -> None:
    """🔥 §10x.143c — When a Person row gets backfilled with NRIC/address
    via §10x.143b, propagate those values into existing step2 (executors)
    / step4 (beneficiaries) / step5 (gift allocations) entries that
    already reference this Person. Without this, the will generates
    blanks because the step data was cached BEFORE the IC arrived.

    Match by person_id first (strict), fall back to name (case-insensitive).
    Only fills EMPTY fields — never overwrites a non-empty NRIC.
    """
    if not (nric or address):
        return
    # 🔥 §10x.143c — match the LATEST active will (any status except deleted),
    # not just 'draft'. Users may upload Lim Bee Yan IC mid-flow AFTER an
    # initial will generation; the propagation must still update step2/step4
    # on the generated will so subsequent re-generations include the NRIC.
    will = (Will.query.filter_by(client_id=client_id)
            .filter(Will.deleted_at.is_(None))
            .filter(Will.status.in_(('draft', 'generated', 'approved')))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return
    name_upper = (name or '').strip().upper()
    changed = False

    def _matches(entry: dict) -> bool:
        if entry.get('person_id') == person_id:
            return True
        return (entry.get('full_name') or '').strip().upper() == name_upper

    def _update(entry: dict) -> bool:
        local_changed = False
        if nric and not (entry.get('nric_passport') or '').strip():
            entry['nric_passport'] = nric
            local_changed = True
        if address and not (entry.get('address') or '').strip():
            entry['address'] = address
            local_changed = True
        return local_changed

    # step2_data (executors)
    try:
        s2_raw = json.loads(will.step2_data or '{}') or {}
        if isinstance(s2_raw, dict):
            execs = s2_raw.get('executors') or []
            for e in execs:
                if _matches(e) and _update(e):
                    changed = True
            if changed:
                will.step2_data = json.dumps(s2_raw)
    except Exception:
        pass
    # step4_data (beneficiaries)
    try:
        s4 = json.loads(will.step4_data or '[]')
        if isinstance(s4, dict):
            s4_list = s4.get('beneficiaries') or []
        else:
            s4_list = s4
        for b in s4_list:
            if _matches(b) and _update(b):
                changed = True
        if changed:
            will.step4_data = json.dumps(s4)
    except Exception:
        pass
    if changed:
        db.session.commit()


def _reconcile_downstream_for_new_identity(client_id: str, name: str,
                                             role: str) -> None:
    """🔥 §10x.42 + §10x.44 — When a new identity is added mid-flow,
    walk back through ALL downstream steps and inject this person where
    the message text names them in any will-role:
        Step 3: Executor / Substitute Executor
        Step 4: Guardian (only relevant if minor children)
        Step 5: Beneficiary
        Step 8: Trustee
    Each step has its own dispatcher; only fires if message contains
    the trigger pattern.
    """
    if not name or not role:
        return
    will = (Will.query.filter_by(client_id=client_id, status='draft')
            .filter(Will.deleted_at.is_(None))
            .order_by(Will.updated_at.desc()).first())
    if not will:
        return
    person = Person.query.filter_by(client_id=client_id, full_name=name).first()
    if not person:
        return
    from ai.chat_planner import _gather_summary_source_text
    text = (_gather_summary_source_text(client_id) or '')
    text_l = text.lower()
    name_l = name.lower()
    role_l = (role or '').lower()

    # ── Helper: does message name this person/role with given will-role keyword?
    def _named_with(will_role_kw: str) -> bool:
        """e.g. will_role_kw='executor' — check if message has
        'my executor X', 'executor: Y', 'my <role> as executor', etc."""
        kw = will_role_kw.lower()
        # Pattern A: "my <role> X" or "X (my <role>)" — name proximity
        if name_l and name_l in text_l:
            ni = text_l.find(name_l)
            ki = text_l.find(kw)
            if ki >= 0 and ni >= 0 and abs(ki - ni) < 120:
                return True
        # Pattern B: "my <kw> my <family-role>" — KOID-style
        # ("My Executor My Sister in law")
        if role_l and re.search(
            rf'\bmy\s+{re.escape(kw)}[^\.\n]{{0,40}}my\s+{re.escape(role_l)}',
            text_l):
            return True
        # Pattern C: "my <family-role> as <kw>" / "<family-role> as my <kw>"
        if role_l and re.search(
            rf'\b(?:my\s+)?{re.escape(role_l)}[^\.\n]{{0,40}}\bas\s+(?:my\s+)?{re.escape(kw)}',
            text_l):
            return True
        return False

    # ── Step 5 (Beneficiaries): "go to my X", "to my X", "for my X"
    is_beneficiary = False
    if name_l and name_l in text_l:
        for trig in ('go to', 'goes to', 'to my', 'for my'):
            ti = text_l.find(trig); ni = text_l.find(name_l)
            if ti >= 0 and ni >= 0 and abs(ti - ni) < 120:
                is_beneficiary = True; break
    if not is_beneficiary and role_l:
        if re.search(rf'(?:go(?:es)?\s+(?:to\s+)?my\s+{re.escape(role_l)}|'
                      rf'to\s+my\s+{re.escape(role_l)}|'
                      rf'for\s+my\s+{re.escape(role_l)})', text_l):
            is_beneficiary = True
    if is_beneficiary:
        _step4_add_beneficiary(will, person, name, role)

    # ── Step 3 (Executor) — auto-add when message names them as executor
    if _named_with('executor'):
        _step2_add_executor(will, person, name, role)

    # ── Step 4 (Guardian) — auto-add when message names them as guardian
    if _named_with('guardian'):
        _step3_add_guardian(will, person, name, role)

    # ── Step 8 (Trustee) — auto-add when message names them as trustee
    if _named_with('trustee'):
        _step7_add_trustee(will, person, name, role)


def _step4_add_beneficiary(will, person, name, role):
    try:
        s4 = json.loads(will.step4_data) if will.step4_data else []
    except (json.JSONDecodeError, TypeError):
        s4 = []
    if not isinstance(s4, list):
        s4 = []
    if any((b.get('full_name') or '').upper() == name.upper()
            for b in s4 if isinstance(b, dict)):
        return
    s4.append({
        'full_name':              name,
        'nric_passport_birthcert': person.nric_passport or '',
        'relationship':           role,
        'person_id':              person.id,
        'nationality':            person.nationality or 'Malaysian',
        '_added_by':              '§10x.42 reconcile (Step 5: Beneficiary)',
    })
    will.step4_data = json.dumps(s4)
    db.session.commit()
    app.logger.info(f'§10x.42/44 added {name} ({role}) to step4 (Beneficiary)')


def _step2_add_executor(will, person, name, role):
    try:
        s2 = json.loads(will.step2_data) if will.step2_data else {}
    except (json.JSONDecodeError, TypeError):
        s2 = {}
    if not isinstance(s2, dict):
        s2 = {}
    execs = s2.get('executors') or []
    if any((e.get('full_name') or '').upper() == name.upper() for e in execs):
        return
    is_first = len(execs) == 0
    execs.append({
        'full_name':     name,
        'nric_passport': person.nric_passport or '',
        'relationship':  role,
        'address':       person.address or '',
        'role':          'Primary' if is_first else 'Substitute',
        'person_id':     person.id,
        'nationality':   person.nationality or 'Malaysian',
        '_added_by':     '§10x.44 reconcile (Step 3: Executor)',
    })
    s2['executors'] = execs
    will.step2_data = json.dumps(s2)
    db.session.commit()
    app.logger.info(f'§10x.44 added {name} ({role}) to step2 (Executor)')


def _step3_add_guardian(will, person, name, role):
    try:
        s3 = json.loads(will.step3_data) if will.step3_data else {}
    except (json.JSONDecodeError, TypeError):
        s3 = {}
    if not isinstance(s3, dict):
        s3 = {}
    guardians = s3.get('guardians') or []
    if any((g.get('full_name') or '').upper() == name.upper() for g in guardians):
        return
    guardians.append({
        'full_name':     name,
        'nric_passport': person.nric_passport or '',
        'relationship':  role,
        'address':       person.address or '',
        'person_id':     person.id,
        'nationality':   person.nationality or 'Malaysian',
        '_added_by':     '§10x.44 reconcile (Step 4: Guardian)',
    })
    s3['guardians'] = guardians
    will.step3_data = json.dumps(s3)
    db.session.commit()
    app.logger.info(f'§10x.44 added {name} ({role}) to step3 (Guardian)')


def _step7_add_trustee(will, person, name, role):
    try:
        s7 = json.loads(will.step7_data) if will.step7_data else {}
    except (json.JSONDecodeError, TypeError):
        s7 = {}
    if not isinstance(s7, dict):
        s7 = {}
    trustees = s7.get('trustees') or []
    if any((t.get('full_name') or '').upper() == name.upper() for t in trustees):
        return
    trustees.append({
        'full_name':     name,
        'nric_passport': person.nric_passport or '',
        'relationship':  role,
        'address':       person.address or '',
        'person_id':     person.id,
        'nationality':   person.nationality or 'Malaysian',
        '_added_by':     '§10x.44 reconcile (Step 8: Trustee)',
    })
    s7['trustees'] = trustees
    will.step7_data = json.dumps(s7)
    db.session.commit()
    app.logger.info(f'§10x.44 added {name} ({role}) to step7 (Trustee)')


def _try_skip_pending_identity(client_id: str, user_text: str):
    """If the user typed/clicked 'skip', mark the pending IC as skipped
    so the walkthrough advances to the next one. Returns a just_assigned-
    style dict or None."""
    if not user_text:
        return None
    text_lower = ' ' + user_text.lower().strip() + ' '
    if not any((' ' + s + ' ') in text_lower for s in _SKIP_TOKENS):
        return None
    from services.identity_walker import (
        get_pending_ic_documents, skip_pending_ic_document
    )
    pending = get_pending_ic_documents(client_id)
    if not pending:
        return None
    result = skip_pending_ic_document(client_id)
    if not result:
        return None
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return None
    return {
        'name': result['name'],
        'role': 'skipped',
        'kind': 'identity_skipped',
        'skip_count': result.get('skip_count', 1),
    }


# -- Inbound email → chat (Postmark webhook) --------------------------------

_REPLY_QUOTE_RE = re.compile(
    r'(?:^On .{0,80}wrote:.*\Z)|(?:^>+.*$)|(?:^-{2,}\s*Original Message\s*-{2,}.*\Z)',
    re.MULTILINE | re.DOTALL,
)
_INBOUND_FILE_EXTS = {'jpg', 'jpeg', 'png', 'gif', 'heic', 'heif', 'webp', 'bmp', 'pdf',
                       # Audio (voice notes forwarded from WhatsApp etc.)
                       'mp3', 'mp4', 'm4a', 'wav', 'webm', 'ogg', 'oga', 'mpga'}


def _strip_reply_quotes(text: str) -> str:
    """Trim 'On X, Y wrote:' tails and quoted '>' lines from email bodies."""
    if not text:
        return ''
    cleaned = _REPLY_QUOTE_RE.sub('', text)
    return '\n'.join(line for line in cleaned.splitlines() if line.strip()).strip()


def _extract_inbox_to(payload: dict):
    """Find the inbox-formatted recipient in a Postmark inbound payload.
    Accepts both legacy '<slug>-<8hex>@…' and new '<name><ic4>@…' formats.
    """
    from services.inbound_address import short_id_from_address, NEW_ADDR_RE

    def _is_inbox_addr(addr: str) -> bool:
        """True if the address looks like one of our inbox addresses."""
        a = addr.strip().lower()
        # Legacy format: ends in '-<8hex>@…'
        if short_id_from_address(a):
            return True
        # New format: local part is 2-5 letters + 4 digits (e.g. koid5008)
        local = a.split('@')[0] if '@' in a else ''
        if NEW_ADDR_RE.match(local):
            return True
        return False

    for t in payload.get('ToFull') or []:
        addr = (t.get('Email') or '').strip()
        if _is_inbox_addr(addr):
            return addr
    raw = payload.get('To') or ''
    for part in raw.split(','):
        part = part.strip()
        if '<' in part and '>' in part:
            part = part[part.index('<') + 1: part.index('>')]
        if _is_inbox_addr(part):
            return part
    return None


@app.route('/api/inbound-email', methods=['POST'])
def api_inbound_email():
    """Wrapper that surfaces the underlying exception in JSON so failures
    are debuggable end-to-end (instead of a generic 500). Postmark retries
    on 5xx, so we still raise the original status."""
    try:
        return _api_inbound_email_impl()
    except Exception as e:
        traceback.print_exc()
        return jsonify({'ok': False, 'error': f'{type(e).__name__}: {e}'}), 500


def _api_inbound_email_impl():
    """Postmark inbound webhook — turns a forwarded email into a chat message.

    Two-phase: this handler does the FAST parts synchronously (auth, parse,
    save user message + raw attachments to disk) and returns 200 to Postmark
    inside ~1 second. The slow per-attachment AI work (vision classify, IC
    extract, voice transcribe, planner reply) runs in a background thread so
    Postmark's 10s webhook timeout doesn't kill emails with many photos.

    Auth: HTTP Basic, credentials in env vars POSTMARK_INBOUND_USER /
    POSTMARK_INBOUND_PASS. If either is unset, the endpoint refuses
    everything (so a half-configured deploy can't be abused).

    Sender allowlist: opt-in via env var INBOUND_ALLOWED_DOMAINS (comma-
    separated). Default = empty = accept any sender.
    """
    from services.inbound_address import find_client_by_address
    from uploads import MAX_FILE_SIZE

    expected_user = os.environ.get('POSTMARK_INBOUND_USER', '')
    expected_pass = os.environ.get('POSTMARK_INBOUND_PASS', '')
    if not expected_user or not expected_pass:
        return jsonify({'ok': False, 'error': 'Inbound webhook not configured'}), 503
    auth = request.authorization
    if not auth or auth.username != expected_user or auth.password != expected_pass:
        return jsonify({'ok': False, 'error': 'Unauthorized'}), 401

    payload = request.get_json(silent=True) or {}

    matched_to = _extract_inbox_to(payload)
    if not matched_to:
        return jsonify({'ok': True, 'ignored': 'no inbox address in To'}), 200

    subject = (payload.get('Subject') or '').strip()
    client = find_client_by_address(matched_to, hint_subject=subject)
    if not client:
        return jsonify({'ok': True, 'ignored': 'unknown client', 'to': matched_to}), 200

    from_email = (
        ((payload.get('FromFull') or {}).get('Email')) or payload.get('From') or ''
    ).strip().lower()
    if '<' in from_email and '>' in from_email:
        from_email = from_email[from_email.index('<') + 1: from_email.index('>')]
    allowed_raw = os.environ.get('INBOUND_ALLOWED_DOMAINS', '').strip()
    if allowed_raw and allowed_raw != '*':
        allowed = [d.strip().lower() for d in allowed_raw.split(',') if d.strip()]
        if not any(from_email.endswith('@' + d) for d in allowed):
            return jsonify({'ok': True, 'ignored': 'sender not allowlisted', 'from': from_email}), 200

    text_body = (payload.get('TextBody') or '').strip()
    if not text_body and payload.get('HtmlBody'):
        text_body = re.sub(r'<[^>]+>', ' ', payload['HtmlBody'])
        text_body = re.sub(r'\s+', ' ', text_body).strip()
    text_body = _strip_reply_quotes(text_body)

    attachments = payload.get('Attachments') or []
    if not text_body and not attachments:
        return jsonify({'ok': True, 'ignored': 'empty email'}), 200

    # Postmark puts the original email Date header in 'Date' field.
    # Include it so the AI can understand message sequencing when
    # the user forwards a WhatsApp chat (the body will contain the
    # WhatsApp timestamps, and the email Date gives arrival time).
    email_date = (payload.get('Date') or '').strip()
    body_with_meta = f"_(forwarded via email from {from_email})_\n\n"
    if subject:
        body_with_meta += f"**Subject:** {subject}\n\n"
    if email_date:
        body_with_meta += f"**Email date:** {email_date}\n\n"
    body_with_meta += text_body

    cs = _get_or_create_chat_session(client.id, user_id=None)
    user_msg = ChatMessage(
        session_id=cs.id, role='user', content=body_with_meta,
        attachments_json='[]',
    )
    db.session.add(user_msg)
    db.session.flush()

    # ── Persist WhatsApp/email body permanently in the Will record ──────────
    # The ChatMessage can theoretically be cleared by api_chat_clear. As a
    # belt-and-suspenders safeguard, also append the raw text body to the
    # Will's step6_data under '_raw_forward_text'. This field is never cleared
    # by the chat-clear operation and can be used by enrichment even if the
    # ChatMessage is gone.
    try:
        _will = (Will.query
                 .filter_by(client_id=client.id, status='draft')
                 .filter(Will.deleted_at.is_(None))
                 .order_by(Will.updated_at.desc())
                 .first())
        # 🔥 §10x.54 — create draft Will on first inbound email so the
        # raw_forward_text gets persisted. Without this, fresh clients
        # lose the email body permanently and the matcher has no
        # AI Summary fallback. Bug observed when KOID's first email
        # arrived: Will didn't exist yet → text_body never saved →
        # _raw_forward_text length 0 → AI Summary parse returned 0
        # AssetItems → walkthrough showed 1 garbage property card.
        if not _will and text_body:
            _will = Will(
                client_id=client.id, status='draft',
                title=f"Will of {client.full_name or 'Unknown'}",
            )
            db.session.add(_will)
            db.session.flush()
        if _will and text_body:
            try:
                _s6 = json.loads(_will.step6_data) if _will.step6_data else {}
            except (json.JSONDecodeError, TypeError):
                _s6 = {}
            if not isinstance(_s6, dict):
                _s6 = {}
            # Append to existing raw context (multiple forwards accumulate)
            existing = _s6.get('_raw_forward_text', '')
            separator = '\n\n---\n\n' if existing else ''
            _s6['_raw_forward_text'] = (existing + separator + body_with_meta)[:50000]
            _will.step6_data = json.dumps(_s6)
    except Exception:
        pass  # non-critical — ChatMessage is the primary store

    # SYNC: write each attachment to disk + create Document row.
    # Skip vision classify / IC extract / Whisper here — the background
    # thread does those. Postmark sees a 200 within ~1s.
    folder_name = client.folder_name
    attachment_ids = []
    file_errors = []

    for att in attachments:
        name = att.get('Name', 'attachment')
        b64 = att.get('Content', '')
        if not b64:
            continue
        ext = name.rsplit('.', 1)[-1].lower() if '.' in name else ''
        if ext not in _INBOUND_FILE_EXTS:
            file_errors.append(f"{name}: unsupported file type")
            continue
        try:
            data = base64.b64decode(b64)
        except Exception:
            file_errors.append(f"{name}: could not decode")
            continue
        if len(data) > MAX_FILE_SIZE:
            file_errors.append(f"{name}: larger than 20MB")
            continue

        ctype = (att.get('ContentType') or '').lower()

        # ── DEDUP: same physical file re-forwarded? ──────────────────────────
        # Per CLAUDE.md §10c: the canonical dedup key is the SHA256 of the
        # bytes. Postmark / WhatsApp routinely rename "the same image" to
        # different filenames, and re-encoding sometimes shifts file_size,
        # so (filename, file_size) is unreliable. Hash dedup catches both
        # cases AND prevents paying for vision-classification of the same
        # bytes twice. Old rows have content_hash=NULL — fall back to the
        # legacy key so we don't double-ingest until backfill catches up.
        import hashlib as _hashlib
        _content_hash = _hashlib.sha256(data).hexdigest()
        existing = (Document.query
                    .filter_by(client_id=client.id, content_hash=_content_hash)
                    .order_by(Document.created_at.asc()).first())
        if not existing:
            existing = (Document.query
                        .filter_by(client_id=client.id,
                                   original_filename=name, file_size=len(data),
                                   content_hash=None)
                        .order_by(Document.created_at.asc()).first())
        if existing:
            # Backfill content_hash on legacy match so future uploads
            # short-circuit on the indexed column.
            if not existing.content_hash:
                try:
                    existing.content_hash = _content_hash
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            attachment_ids.append(existing.id)
            continue

        # Provisional category — the async worker will retag based on classify
        # (or transcribe for audio). Save under chat_inbox/ for now.
        cat = 'chat_inbox'
        safe_basename = uuid.uuid4().hex[:12] + '.' + ext
        folder = os.path.join(UPLOAD_DIR, folder_name, 'documents', cat)
        os.makedirs(folder, exist_ok=True)
        abs_path = os.path.join(folder, safe_basename)
        with open(abs_path, 'wb') as f:
            f.write(data)
        rel_path = os.path.join(folder_name, 'documents', cat, safe_basename)

        doc = Document(
            client_id=client.id, chat_message_id=user_msg.id,
            filename=safe_basename, original_filename=name,
            file_path=rel_path, file_type=ctype,
            file_size=len(data), category=cat,
            content_hash=_content_hash,
        )
        db.session.add(doc)
        db.session.flush()
        attachment_ids.append(doc.id)

    user_msg.attachments_json = json.dumps(attachment_ids)

    # If the email had any rejected attachments, surface that as an immediate
    # assistant note (no AI work needed) so the user sees something right away.
    if file_errors:
        early_note = ChatMessage(
            session_id=cs.id, role='assistant',
            content=("📎 Some attachments couldn't be saved:\n- " + "\n- ".join(file_errors)),
        )
        db.session.add(early_note)

    db.session.commit()

    # Spawn the slow processing in the background. The webhook returns NOW.
    threading.Thread(
        target=_process_inbound_message_async,
        args=(app, user_msg.id),
        daemon=True,
    ).start()

    return jsonify({
        'ok': True,
        'client_id': client.id,
        'message_id': user_msg.id,
        'attachments': len(attachment_ids),
        'queued_for_processing': True,
    })


# ╔══════════════════════════════════════════════════════════════════╗
# ║ 🔥 BURN-IN §10x.9 — INBOUND PROCESSOR LOCK 🔥                      ║
# ║ The watchdog in /api/chat/<client_id>/history fires every 5 s     ║
# ║ while the chat tab is open. Without a lock, the same user_msg     ║
# ║ gets re-classified by N concurrent threads → N duplicate intake   ║
# ║ cards + N duplicate AI Summary cards posted to chat. The user     ║
# ║ saw 12+ duplicate cards before this lock was added.               ║
# ║                                                                    ║
# ║ Lock is in-process (Python set) per gunicorn worker. Combined     ║
# ║ with the idempotency check below (skip card posting if intake     ║
# ║ card already exists in chat), this guarantees:                    ║
# ║   - At most 1 thread per worker per user_msg classifying at once  ║
# ║   - At most 1 intake card + 1 AI Summary card per user_msg ever   ║
# ╚══════════════════════════════════════════════════════════════════╝
import threading as _proc_threading
_PROCESSING_LOCK = _proc_threading.Lock()
_PROCESSING_INFLIGHT: set = set()   # user_msg_ids currently being processed


def _process_inbound_message_async(app_obj, user_msg_id):
    """Background processing — runs after the webhook has returned 200.

    For each Document attached to the user_msg:
      - audio → Whisper transcribe
      - else  → vision classify (nric/property_title/...) + extract if IC
    Then call the planner over (text + voice transcripts + extracted artifacts)
    and save the assistant ChatMessage.
    """
    # ── In-process lock ─────────────────────────────────────────────
    # Refuse if another thread in this gunicorn worker is already
    # processing this user_msg. Caller (watchdog) just exits silently.
    with _PROCESSING_LOCK:
        if user_msg_id in _PROCESSING_INFLIGHT:
            return
        _PROCESSING_INFLIGHT.add(user_msg_id)
    try:
        _process_inbound_message_async_inner(app_obj, user_msg_id)
    finally:
        with _PROCESSING_LOCK:
            _PROCESSING_INFLIGHT.discard(user_msg_id)


def _process_inbound_message_async_inner(app_obj, user_msg_id):
    """Inner body — wrapped by the lock above."""
    with app_obj.app_context():
        from ai.file_classifier import classify_file
        from ai.ocr import extract_nric_data
        from ai.chat_planner import plan_turn
        from ai.voice_transcription import transcribe, is_audio

        try:
            user_msg = db.session.get(ChatMessage, user_msg_id)
            if not user_msg:
                return
            cs = db.session.get(ChatSession, user_msg.session_id)
            client = db.session.get(Client, cs.client_id) if cs else None
            if not client:
                return

            doc_ids = []
            try:
                doc_ids = json.loads(user_msg.attachments_json or '[]')
            except (json.JSONDecodeError, TypeError):
                doc_ids = []

            artifacts = []
            voice_transcripts = []
            docs = (Document.query.filter(Document.id.in_(doc_ids)).all()
                    if doc_ids else [])

            # ── Batch-first grouping ──────────────────────────────────────
            # STEP 1: Analyse ALL images together BEFORE classifying individually.
            # This is the key insight: a customer sending 5 photos in one
            # WhatsApp message almost always means 5 pages of the same document.
            # We ask Claude to:
            #   (a) Infer relationships from overlapping identifiers (lot, title, acct)
            #   (b) Use the WhatsApp text (before/after images) as primary evidence
            #   (c) Return asset GROUPS — {image_indices, asset_kind, identifiers}
            # Each subsequent classify_file() call gets the group verdict as context,
            # so "page 4 of a blurry geran" is classified correctly instead of 'other'.
            from ai.file_classifier import classify_batch

            image_docs = [d for d in docs
                          if not (is_audio((d.file_type or '').lower())
                                  or is_audio(d.original_filename or ''))]
            batch_group_map: dict = {}   # doc index (into docs list) → group dict

            # 🔥 §10x.66 — skip classify_batch when ≥80% of image docs
            # are already classified. The watchdog re-fires on every
            # 5-second chat poll, but classify_batch was at the TOP of
            # the function — paying $0.05-$0.13 every time even when
            # all docs were already done. Today: 76 batch calls = $0.59
            # mostly redundant. Now: only call batch on FRESH chat_inbox
            # docs, or on first run (when most are unclassified).
            _SKIP_CATS = {'nric', 'property_title', 'property_spa',
                           'property_tax', 'property_transfer',
                           'loan_agreement', 'bank_statement', 'bank_letter',
                           'utility_bill', 'insurance', 'epf_kwsp',
                           'vehicle', 'will', 'death_certificate',
                           'unrelated', 'needs_review', 'deleted',
                           'duplicate', 'voice'}
            unclassified = [d for d in image_docs
                            if (d.category or '') not in _SKIP_CATS]
            should_run_batch = (
                len(image_docs) >= 2
                and len(unclassified) >= max(2, int(len(image_docs) * 0.2))
            )

            if should_run_batch:
                try:
                    # Pass ONLY the unclassified images to the batch — saves
                    # token cost when only a few docs are stuck.
                    image_paths = [os.path.join(UPLOAD_DIR, d.file_path)
                                   for d in unclassified]
                    # 🔥 §10x.65 — wrap in track_context so client_id is logged
                    from ai.cost_tracker import track_context as _tc_batch, is_over_ceiling as _over
                    if _over(client.id):
                        try:
                            current_app.logger.warning(
                                f'§10x.65 ceiling hit for {client.id} — skip classify_batch')
                        except Exception:
                            pass
                        batch_result = {'groups': []}
                    else:
                        with _tc_batch(client_id=client.id):
                            # Use the email body as message context for the batch analysis.
                            # WhatsApp text (lot numbers, beneficiary names) lives here.
                            batch_msg_ctx = (user_msg.content or '')[:600]
                            from ai.file_classifier import classify_batch as _classify_batch
                            batch_result = _classify_batch(image_paths, message_context=batch_msg_ctx)
                    # Build doc → group lookup. Indices in batch_result refer
                    # to the `unclassified` list we passed in, not the full
                    # image_docs list — translate via unclassified[idx].
                    for grp in (batch_result.get('groups') or []):
                        for img_idx in (grp.get('image_indices') or []):
                            if 0 <= img_idx < len(unclassified):
                                batch_group_map[id(unclassified[img_idx])] = grp
                except Exception:
                    pass   # batch analysis failed — fall through to individual classify

            # ── Pre-compute slow vision calls in parallel ──────────────────
            # Vision Sonnet extract takes ~20-30 s per image. Sequential
            # processing of 30 images → ~10 minutes. Run them concurrently
            # in a ThreadPoolExecutor. The classify_file / extract_*_data
            # functions are pure API calls (no DB writes) so thread-safe.
            # We then walk the results sequentially in the main thread to
            # apply DB updates and commit.
            from concurrent.futures import ThreadPoolExecutor

            def _classify_one(doc, abs_path, group_ctx, testator_profile):
                """Worker: returns (classification_dict, extracted_dict).
                Each thread pushes its own Flask app context so that
                cost_tracker.log_usage and any other context-bound code
                works correctly (otherwise telemetry silently fails and
                we lose visibility into per-image API spend).
                ALSO uses cost_tracker.track_context() so each call's
                cost is attributed to this client/will/document."""
                from ai.cost_tracker import track_context as _tc
                # 🔥 §10x.3 — track_context only accepts client_id/will_id/user_id.
                # Passing document_id used to silently break every classify worker
                # ("track_context() got an unexpected keyword 'document_id'") which
                # left all docs in chat_inbox forever. Do NOT add document_id back
                # without first widening the cost_tracker signature.
                with app_obj.app_context(), _tc(
                    will_id=doc.will_id if hasattr(doc, 'will_id') else None,
                    client_id=doc.client_id):
                    # 🔥 §10x.59 — cost ceiling guard. If this client has
                    # already burned $1.50 in the last 24h, skip vision/AI
                    # and fall back to a sentinel result. The chat will
                    # surface "manual review needed" and stop further
                    # spend until the user (or admin) clears the cap.
                    try:
                        from ai.cost_tracker import is_over_ceiling
                        if is_over_ceiling(doc.client_id):
                            return (
                                {'kind': 'other', 'confidence': 'low',
                                 'reason': 'cost ceiling hit — manual review',
                                 'manual_review': True,
                                 '_cost_capped': True},
                                None,
                            )
                    except Exception:
                        pass
                    try:
                        classification = classify_file(
                            abs_path, group_context=group_ctx,
                            testator_profile=testator_profile,
                        )
                    except Exception as e:
                        classification = {'kind': 'other', 'confidence': 'low',
                                           'reason': f'classify error: {e}'}
                    kind = classification.get('kind', 'other')
                    if group_ctx and kind == 'other':
                        bk = group_ctx.get('asset_kind', '')
                        if bk and bk != 'other':
                            kind = bk
                            classification['kind'] = kind
                            classification['confidence'] = 'medium'
                    extracted = None
                    try:
                        if kind == 'nric':
                            extracted = extract_nric_data(abs_path)
                        elif kind in ('property_title', 'property_spa', 'property_tax',
                                       'property_transfer'):
                            from ai.property_extractor import extract_property_data
                            extracted = extract_property_data(abs_path, doc_type='general')
                        elif kind in ('utility_bill', 'bank_letter'):
                            extracted = {}
                        elif kind == 'bank_statement':
                            from ai.ocr import extract_asset_document
                            extracted = extract_asset_document(abs_path, asset_type='bank')
                        elif kind == 'vehicle':
                            from ai.ocr import extract_asset_document
                            extracted = extract_asset_document(abs_path, asset_type='vehicle')
                        elif kind in ('insurance', 'epf_kwsp'):
                            from ai.ocr import extract_asset_document
                            extracted = extract_asset_document(abs_path, asset_type='other')
                    except Exception as e:
                        extracted = {'error': str(e)}
                    return classification, extracted

            # Dispatch parallel work for non-voice docs.
            _testator_profile = {
                'name': (client.full_name or '').strip(),
                'ic':   (client.nric_passport or '').strip(),
            } if client else None
            _ai_results = {}   # doc.id → (classification, extracted)
            _vision_jobs = []
            # 🔥 §10x.58 — watchdog throttle. Docs that already have a real
            # category (nric/property_title/etc) OR a terminal needs_review
            # MUST NOT be re-classified. Re-firing every chat poll burned
            # $10 today. Categorization is monotonic per §10x.2.
            _SKIP_CATEGORIES = {
                'nric', 'property_title', 'property_spa', 'property_tax',
                'property_transfer', 'loan_agreement', 'bank_statement',
                'bank_letter', 'utility_bill', 'insurance', 'epf_kwsp',
                'vehicle', 'will', 'death_certificate', 'unrelated',
                'needs_review', 'deleted', 'duplicate', 'voice',
            }
            for _d in docs:
                _ap = os.path.join(UPLOAD_DIR, _d.file_path)
                if not os.path.isfile(_ap):
                    continue
                _ct = (_d.file_type or '').lower()
                if is_audio(_ct) or is_audio(_d.original_filename or ''):
                    continue   # voice handled in main loop
                # 🔥 §10x.58 — skip already-classified docs
                if (_d.category or '') in _SKIP_CATEGORIES:
                    # Reuse the already-persisted classification so the
                    # downstream loop has data; no API call needed.
                    try:
                        _ex = json.loads(_d.extracted_data or '{}') or {}
                    except Exception:
                        _ex = {}
                    _ai_results[_d.id] = (
                        {'kind': _d.category, 'confidence': 'high',
                         'reason': 'cached from prior pass', '_skipped_cost': True},
                        _ex,
                    )
                    continue
                _gc = batch_group_map.get(id(_d))
                _vision_jobs.append((_d, _ap, _gc))
            if _vision_jobs:
                with ThreadPoolExecutor(max_workers=5) as _pool:
                    _futures = {
                        _pool.submit(_classify_one, _d, _ap, _gc, _testator_profile): _d
                        for (_d, _ap, _gc) in _vision_jobs
                    }
                    for _fut in _futures:
                        _d = _futures[_fut]
                        try:
                            _ai_results[_d.id] = _fut.result(timeout=180)
                        except Exception as _e:
                            _ai_results[_d.id] = (
                                {'kind': 'other', 'reason': f'timeout/err: {_e}',
                                 'confidence': 'low'},
                                {'error': str(_e)}
                            )

            for doc in docs:
                abs_path = os.path.join(UPLOAD_DIR, doc.file_path)
                if not os.path.isfile(abs_path):
                    continue
                ctype = (doc.file_type or '').lower()
                is_voice = is_audio(ctype) or is_audio(doc.original_filename or '')

                if is_voice:
                    # Move to voice/ folder for tidiness
                    if doc.category != 'voice':
                        doc.category = 'voice'
                    transcript = transcribe(abs_path)
                    if transcript:
                        voice_transcripts.append(transcript)
                        doc.description = transcript[:500]
                        try:
                            doc.extracted_data = json.dumps({'transcript': transcript})
                        except (TypeError, ValueError):
                            pass
                    else:
                        doc.description = '(transcription failed or unavailable)'
                    artifacts.append({
                        'document_id': doc.id, 'kind': 'voice',
                        'confidence': 'high' if transcript else 'low',
                        'extracted': {'transcript': transcript} if transcript else None,
                        'original_filename': doc.original_filename,
                    })
                    db.session.commit()
                    continue

                # STEP 2: Use the parallel result computed above.
                # All vision Sonnet / Haiku calls already finished; this loop
                # only does DB updates (sequential to avoid session conflicts).
                group_ctx = batch_group_map.get(id(doc))
                if doc.id in _ai_results:
                    classification, extracted = _ai_results[doc.id]
                else:
                    classification = {'kind': 'other', 'confidence': 'low',
                                       'reason': 'no result'}
                    extracted = None
                kind = classification.get('kind', 'other')
                # 🔥 BURN-IN §10x — NEVER DOWNGRADE A REAL CATEGORY TO chat_inbox.
                # The reprocess watchdog re-runs this loop on every chat-history
                # poll. If a previous run successfully classified the doc as
                # 'property_title' but the current run gets a low-confidence
                # 'other' (network blip, rate limit), we must NOT overwrite.
                # Categorization is monotonic: chat_inbox → real, never the
                # other way around.
                if kind != 'other':
                    doc.category = kind
                elif doc.category in (None, '', 'chat_inbox', 'other'):
                    # 🔥 §10x.26 — TERMINAL STATE for vision failures.
                    # Without this guard, every chat-history poll (every 5s)
                    # re-classified the same 5 unreadable docs forever.
                    # After 3 failed attempts, promote the doc to
                    # 'needs_review' so the watchdog stops re-firing.
                    try:
                        prev_attempts = int(
                            (json.loads(doc.extracted_data or '{}') or {})
                            .get('_classify_attempts', 0) or 0)
                    except Exception:
                        prev_attempts = 0
                    new_attempts = prev_attempts + 1
                    # Guard: `extracted` may still be None at this point
                    # (the `if extracted is None: extracted = {}` line is
                    # below us). Ensure it's a dict before assignment.
                    if extracted is None:
                        extracted = {}
                    extracted['_classify_attempts'] = new_attempts
                    is_unreadable = bool(classification.get('manual_review')) \
                                    or 'unreadable' in (classification.get('reason') or '').lower()
                    if new_attempts >= 3 or is_unreadable:
                        doc.category = 'needs_review'
                        extracted['_terminal_reason'] = (
                            'unreadable_after_retry' if new_attempts >= 3
                            else 'vision_marked_unreadable')
                    else:
                        doc.category = 'chat_inbox'
                # else: keep existing real category
                doc.description = (classification.get('reason') or '')[:500] or None
                purpose = (classification.get('purpose') or '').strip()
                prop_hint = (classification.get('property_hint') or '').strip()
                will_relevant = classification.get('will_relevant', True)
                if extracted is None:
                    extracted = {}
                if purpose:
                    extracted['purpose'] = purpose[:300]
                if prop_hint:
                    extracted['property_hint'] = prop_hint[:300]
                # ── OCR-extracted fields ──────────────────────────────────
                for _ocr_field in ('lot_number', 'title_number', 'property_address',
                                   'owner_name', 'ic_number', 'bank_name',
                                   'mukim', 'daerah', 'negeri'):
                    val = (classification.get(_ocr_field) or '').strip()
                    if val and not extracted.get(_ocr_field):
                        extracted[_ocr_field] = val
                # ── Testator match flags ──────────────────────────────────
                if classification.get('name_match') is not None:
                    extracted['_name_match'] = classification['name_match']
                if classification.get('ic_match') is not None:
                    extracted['_ic_match'] = classification['ic_match']
                # ── Manual review flag (OCR unreadable) ──────────────────
                if classification.get('manual_review'):
                    extracted['_manual_review'] = True

                # STEP 3: Cross-fill missing NLC fields from the batch group
                # identifiers. The batch analysis extracted lot/title/mukim
                # from the clearest image in the group — propagate to all pages.
                if group_ctx:
                    grp_idents = group_ctx.get('identifiers') or {}
                    prop_fields = ('lot_number', 'title_number', 'mukim',
                                   'daerah', 'negeri', 'property_address',
                                   'bank_name', 'account_number', 'reg_number')
                    for field in prop_fields:
                        if not (extracted.get(field) or '').strip():
                            val = (grp_idents.get(field) or '').strip()
                            if val:
                                extracted[field] = val
                                extracted.setdefault('_enriched_from', []).append(
                                    f'batch_group.{field}'
                                )
                    # Store beneficiary hint from WhatsApp text (e.g. "give to Sarah")
                    bh = (group_ctx.get('beneficiary_hint') or '').strip()
                    if bh and not extracted.get('_beneficiary_hint'):
                        extracted['_beneficiary_hint'] = bh
                # Store the text context that came WITH this image.
                # For WhatsApp-forwarded emails the body is a chat log:
                #   [time] Client: please add this property, lot 127082 to sarah
                #   [time] Client: <attached: PHOTO-2026-05-02-13-52-35.jpg>
                # We extract the lines immediately before THIS image's attachment
                # reference — that's the most specific context for this image.
                # Fall back to the full email body if no WhatsApp format found.
                try:
                    msg_body = (user_msg.content or '').strip()
                    wa_ctx = _extract_whatsapp_context_for_file(
                        msg_body, doc.original_filename or ''
                    )
                    if wa_ctx:
                        # WhatsApp-specific context found for this image
                        extracted['_message_context'] = wa_ctx[:800]
                        extracted['_context_source'] = 'whatsapp_preceding'
                    elif msg_body:
                        # No WhatsApp format — store the whole email body
                        extracted['_message_context'] = msg_body[:800]
                        extracted['_context_source'] = 'email_body'
                    # WhatsApp timestamp for this image (CLAUDE.md §10i).
                    # Persisted at ingest because the message body may be
                    # rotated / trimmed by the time the property card renders.
                    wa_ts = _extract_whatsapp_timestamp_for_file(
                        msg_body, doc.original_filename or ''
                    )
                    if wa_ts:
                        extracted['_msg_timestamp'] = wa_ts
                except Exception:
                    pass
                if kind == 'other' and not will_relevant:
                    try:
                        recent = _gather_recent_chat_text(client.id)
                        fname_stem = doc.original_filename.rsplit('.', 1)[0].lower()
                        asset_keywords = ('property', 'house', 'land', 'lot', 'geran',
                                          'bank', 'account', 'insurance', 'car', 'vehicle',
                                          'mukim', 'title', fname_stem)
                        if not any(kw in (recent or '').lower() for kw in asset_keywords):
                            extracted['_likely_irrelevant'] = True
                            extracted['_irrelevant_reason'] = (
                                'Classified as unrecognised document type and no matching '
                                'asset was mentioned in the chat.'
                            )
                    except Exception:
                        pass
                if extracted:
                    try:
                        doc.extracted_data = json.dumps(extracted)
                    except (TypeError, ValueError):
                        doc.extracted_data = None

                # Dedupe ICs by extracted name — if a previous Document
                # for this client already has the same name, mark this
                # new one as 'duplicate' and skip emitting as artifact.
                is_dup = False
                if kind == 'nric' and extracted and not extracted.get('error'):
                    is_dup = _dedupe_ic_against_existing(client.id, doc, extracted)

                if not is_dup:
                    artifacts.append({
                        'document_id': doc.id, 'kind': kind,
                        'confidence': classification.get('confidence', 'low'),
                        'extracted': extracted,
                        'original_filename': doc.original_filename,
                    })
                # Commit per-document so partial progress is visible to chat
                # polling — useful when there are many attachments.
                db.session.commit()

            # Reload user_msg in case its content changed (it didn't, but safe)
            user_msg = db.session.get(ChatMessage, user_msg_id)

            # Merge voice transcripts into the message body for the planner
            text = user_msg.content or ''
            if voice_transcripts:
                joined = '\n\n'.join(voice_transcripts)
                user_msg.content = text + f"\n\n_(voice transcript)_\n{joined}"
                text = user_msg.content
                db.session.commit()

            active_will = (Will.query.filter_by(client_id=client.id, status='draft')
                           .filter(Will.deleted_at.is_(None))
                           .order_by(Will.updated_at.desc()).first())
            from services.identity_walker import get_pending_ic_documents
            pending_ics = get_pending_ic_documents(client.id)
            recent_text = _gather_recent_chat_text(client.id)

            # Run address enrichment now (after classification) so property
            # addresses from the WhatsApp text are matched to documents
            # immediately — no need for the user to send a message first.
            _persist_property_enrichment(client.id, recent_text)

            # ── §10x.9 idempotency check ────────────────────────────
            # If a planner reply was already posted for this user_msg's
            # session, do NOT post another one. The watchdog must not
            # produce duplicates. Match BOTH:
            #   - "📋 N exhibits received" (with attachments)
            #   - "📋 Asset inventory"     (text-only forwards, §10x.28)
            # Either string proves the planner already ran for this
            # user_msg and posted its reply.
            _intake_already_posted = False
            try:
                from sqlalchemy import or_ as _or
                _existing = (ChatMessage.query
                             .filter_by(session_id=cs.id, role='assistant')
                             .filter(ChatMessage.created_at >= user_msg.created_at)
                             .filter(_or(
                                 ChatMessage.content.ilike('%exhibits received%'),
                                 ChatMessage.content.ilike('%Asset inventory%'),
                             ))
                             .first())
                if _existing:
                    _intake_already_posted = True
            except Exception:
                pass

            if _intake_already_posted:
                # Re-run only re-classifies stuck docs; nothing to post.
                return

            # ── §10x.9 idempotency check (AI Summary) ──────────────
            _summary_already_posted = False
            try:
                _existing_sum = (ChatMessage.query
                                 .filter_by(session_id=cs.id, role='assistant')
                                 .filter(ChatMessage.created_at >= user_msg.created_at)
                                 .filter(ChatMessage.content.ilike('%AI Summary of your message%'))
                                 .first())
                if _existing_sum:
                    _summary_already_posted = True
            except Exception:
                pass

            # ── 🔥 §7 / §10x.28 — Post AI Summary FIRST (Step 2) ────────
            # Per CLAUDE.md §7: Receive WhatsApp → SUMMARISE → Decipher
            # images → Identity match. Order matters in the chat: the user
            # should read the summary card BEFORE the planner's "asset
            # inventory" reply so they verify what we deduced from the
            # text first.
            if text and not _summary_already_posted:
                try:
                    from ai.chat_planner import _summarise_message, _clean_email_body
                    import json as _json
                    cleaned = _clean_email_body(text)
                    if cleaned:
                        # 🔥 §10x.76 — collect extracted fields from every
                        # property-class doc for this client and inject
                        # into the summary prompt. Without this the AI
                        # Summary asks the user to supply lot/title numbers
                        # that the vision extractor already pulled.
                        _doc_fields = []
                        try:
                            _prop_docs = (Document.query
                                          .filter(Document.client_id == client.id)
                                          .filter(Document.category.in_([
                                              'property_title','property_spa',
                                              'property_tax','property_transfer',
                                              'loan_agreement','bank_statement',
                                              'insurance','vehicle','nric',
                                          ]))
                                          .all())
                            for _d in _prop_docs:
                                try:
                                    _ex = _json.loads(_d.extracted_data) if _d.extracted_data else {}
                                except Exception:
                                    _ex = {}
                                _row = {'kind': _d.category}
                                for _k in ('title_number','lot_number','mukim','daerah',
                                           'negeri','property_address','owner_name','title_type',
                                           'bank_name','account_number','currency','account_type',
                                           'insurer','policy_number','full_name','nric_number'):
                                    _v = (_ex.get(_k) or '').strip() if isinstance(_ex.get(_k), str) else _ex.get(_k)
                                    if _v:
                                        _row[_k] = _v
                                if len(_row) > 1:   # has at least one extracted field
                                    _doc_fields.append(_row)
                        except Exception:
                            _doc_fields = []
                        summary = _summarise_message(cleaned, doc_fields=_doc_fields)
                        if not summary:
                            summary = '_Could not generate summary — review exhibits above._'
                        # 🔥 §10x.53 — omit the verify-identities button while
                        # docs are still being analysed. Post button later as a
                        # follow-up "Ready" message once classification is done.
                        _in_progress = 0
                        try:
                            _in_progress = Document.query.filter_by(
                                client_id=client.id, category='chat_inbox'
                            ).count()
                        except Exception:
                            _in_progress = 0
                        # 🔥 §10x.60 — embed input hash so subsequent
                        # _summarise_message calls can hit DB cache and
                        # skip the $0.05 Haiku roundtrip.
                        import hashlib as _hashlib
                        _hash16 = _hashlib.sha256(cleaned.encode('utf-8')).hexdigest()[:16]
                        _hash_marker = f'<!--_summary_hash:{_hash16}-->'
                        if _in_progress > 0:
                            tail = (
                                f"\n\n🔍 _Analysing {_in_progress} exhibit(s) — "
                                "the verify-identities button will appear once "
                                "classification completes._"
                            )
                            reply = (
                                "### 📨 AI Summary of your message\n\n"
                                + summary
                                + tail
                                + f"\n{_hash_marker}"
                            )
                        else:
                            _quick = _json.dumps([
                                {'label': '▶️ Start — verify identities',
                                 'value': 'inbox start'}
                            ])
                            reply = (
                                "### 📨 AI Summary of your message\n\n"
                                + summary
                                + f"\n\n<!--quickreplies:{_quick}-->"
                                + f"\n{_hash_marker}"
                            )
                        summary_msg = ChatMessage(
                            session_id=cs.id, role='assistant',
                            content=reply,
                            attachments_json='[]',  # no exhibit thumbnails in summary
                        )
                        db.session.add(summary_msg)
                        db.session.flush()
                        # 🔥 §10x.63 — post-insert duplicate detection.
                        # The pre-check race (two processors both seeing
                        # "no existing summary" because neither committed
                        # yet) means the start-of-function _summary_already_posted
                        # check isn't enough. Re-query AFTER our flush:
                        # if any OTHER AI Summary was posted before ours
                        # (lower created_at), we're the duplicate — delete
                        # self before committing.
                        try:
                            _earlier = (ChatMessage.query
                                        .filter(ChatMessage.session_id == cs.id,
                                                 ChatMessage.role == 'assistant',
                                                 ChatMessage.id != summary_msg.id,
                                                 ChatMessage.created_at < summary_msg.created_at,
                                                 ChatMessage.content.ilike('%AI Summary of your message%'))
                                        .first())
                            if _earlier:
                                db.session.delete(summary_msg)
                        except Exception:
                            pass
                        db.session.commit()
                except Exception:
                    pass  # non-critical — intake card is the primary response

            # ── Now post the planner reply (Step 3+: asset inventory or
            # exhibits-received card depending on whether attachments
            # exist). This intentionally runs AFTER the AI Summary so
            # the user reads Step 2 first per §7 ordering.
            plan = plan_turn(text, artifacts, _will_data_snapshot(active_will),
                             pending_ics=pending_ics, recent_text=recent_text)
            asst_msg = ChatMessage(
                session_id=cs.id, role='assistant',
                content=plan.get('reply', ''),
                attachments_json=json.dumps(plan.get('focus_attachments') or []),
                clarifying_questions_json=json.dumps(plan.get('clarifying_questions', [])),
                proposed_patch_json=json.dumps(plan['proposed_patch']) if plan.get('proposed_patch') else None,
                advice_json=json.dumps(plan.get('advice', [])),
                target_will_id=active_will.id if active_will else None,
            )
            db.session.add(asst_msg)
            db.session.flush()
            # 🔥 §10x.63 — post-insert dedup for intake card. The hazard
            # is two concurrent processors both posting the same card
            # for the SAME user_msg ~ms apart. The dedup window must be
            # tight or it eats legitimate follow-up cards.
            #
            # 🔥 §10x.78 — Earlier this dedup deleted ANY new "exhibits
            # received" card whenever the session had ever posted one,
            # which silently lost intake cards for follow-up emails (a
            # 2nd email arrives hours after the first; its 2 new docs
            # get processed and the intake card is created — then this
            # dedup deletes it because the morning's card still matches
            # the LIKE query). Fix: only dedup against cards posted
            # WITHIN 30 seconds AND for the same user_msg's window
            # (created_at >= user_msg.created_at).
            reply_text = plan.get('reply', '')
            if 'exhibits received' in reply_text or 'Asset inventory' in reply_text:
                try:
                    from sqlalchemy import or_ as _or_op
                    from datetime import timedelta as _td
                    _earlier = (ChatMessage.query
                                .filter(ChatMessage.session_id == cs.id,
                                         ChatMessage.role == 'assistant',
                                         ChatMessage.id != asst_msg.id,
                                         # Only dedup within the same user_msg's
                                         # window — earlier cards from older
                                         # forwards are NOT duplicates of this one.
                                         ChatMessage.created_at >= user_msg.created_at,
                                         ChatMessage.created_at < asst_msg.created_at,
                                         ChatMessage.created_at >= asst_msg.created_at - _td(seconds=30))
                                .filter(_or_op(
                                    ChatMessage.content.ilike('%exhibits received%'),
                                    ChatMessage.content.ilike('%Asset inventory%'),
                                ))
                                .first())
                    if _earlier:
                        db.session.delete(asst_msg)
                except Exception:
                    pass
            db.session.commit()

            # 🔥 §10x.53 — post a "Ready to verify" follow-up message so the
            # user gets an actionable button AFTER classification completes,
            # even though the intake card / AI Summary may have been posted
            # earlier with an "🔍 Analysing..." status.
            try:
                _remaining = Document.query.filter_by(
                    client_id=client.id, category='chat_inbox'
                ).count()
                if _remaining == 0:
                    # Don't post duplicate "Ready" — check if one already exists
                    _ready_exists = (ChatMessage.query
                                     .filter_by(session_id=cs.id, role='assistant')
                                     .filter(ChatMessage.created_at >= user_msg.created_at)
                                     .filter(ChatMessage.content.ilike('%Analysis complete%'))
                                     .first())
                    # Only post if there are pending IC docs (otherwise nothing to verify)
                    _has_pending_ic = False
                    try:
                        from services.identity_walker import get_pending_ic_documents
                        _has_pending_ic = bool(get_pending_ic_documents(client.id))
                    except Exception:
                        pass
                    if not _ready_exists and _has_pending_ic:
                        _ready_quick = json.dumps([
                            {'label': '▶️ Start — verify identities',
                             'value': 'inbox start'}
                        ])
                        _ready = ChatMessage(
                            session_id=cs.id, role='assistant',
                            content=(
                                "✅ **Analysis complete.** All exhibits "
                                "classified — ready to verify identities."
                                f"\n\n<!--quickreplies:{_ready_quick}-->"
                            ),
                            attachments_json='[]',
                        )
                        db.session.add(_ready)
                        db.session.commit()
            except Exception:
                pass   # non-critical
        except Exception:
            traceback.print_exc()
            try:
                db.session.rollback()
            except Exception:
                pass


# -- Step 1: Identity Management ---------------------------------------------

@app.route('/wizard/step/1', methods=['GET', 'POST'])
@login_required
def wizard_step_identities():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        client_id = session.get('client_id')
        if client_id:
            _refresh_session_person_registry(client_id)
        return render_template(
            'wizard/step1_identities.html',
            current_step=1,
            completed_steps=get_completed_steps(),
            persons=session.get('person_registry', []),
        )

    # POST -- validate at least 1 identity exists, then proceed
    persons = session.get('person_registry', [])
    if not persons:
        flash('Please add at least one identity before proceeding.', 'error')
        return redirect(url_for('wizard_step_identities'))
    mark_step_complete(1)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 1})
    return redirect(url_for('wizard_step_testator'))


# -- Step 2: Testator Info (simplified - select identity) --------------------

@app.route('/wizard/step/2', methods=['GET', 'POST'])
@login_required
def wizard_step_testator():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step2_testator.html',
            current_step=2,
            completed_steps=get_completed_steps(),
            data=session.get('step1', {}),
            persons=session.get('person_registry', []),
        )

    # POST -- merge selected identity with testator-specific fields
    person_id = request.form.get('testator_person_id', '')
    person = _get_person_from_registry(person_id)

    dob_raw = request.form.get('date_of_birth', '')
    if dob_raw and '-' in dob_raw and len(dob_raw) == 10:
        parts = dob_raw.split('-')
        if len(parts) == 3 and len(parts[0]) == 4:
            dob_raw = f"{parts[2]}-{parts[1]}-{parts[0]}"

    special = request.form.getlist('special_circumstances')

    session['step1'] = {
        'person_id': person_id,
        'full_name': person['full_name'] if person else request.form.get('full_name', '').strip(),
        'nric_passport': person['nric_passport'] if person else request.form.get('nric_passport', '').strip(),
        'residential_address': person['address'] if person else request.form.get('residential_address', '').strip(),
        'nationality': person.get('nationality', 'Malaysian') if person else request.form.get('nationality', 'Malaysian').strip(),
        'country_of_residence': request.form.get('country_of_residence', 'Malaysia').strip(),
        'date_of_birth': dob_raw or (person.get('date_of_birth', '') if person else ''),
        'occupation': request.form.get('occupation', '').strip(),
        'religion': request.form.get('religion', '').strip() or None,
        'email': request.form.get('email', '').strip() or (person.get('email') if person else None),
        'phone': request.form.get('phone', '').strip() or (person.get('phone') if person else None),
        'gender': request.form.get('gender', 'Male'),
        'marital_status': request.form.get('marital_status', 'Single'),
        'has_prior_will': bool(request.form.get('has_prior_will')),
        'property_coverage': request.form.get('property_coverage', 'Malaysia'),
        'contemplation_of_marriage': bool(request.form.get('contemplation_of_marriage')),
        'fiance_name': request.form.get('fiance_name', '').strip() or None,
        'fiance_nric': request.form.get('fiance_nric', '').strip() or None,
        'signing_method': request.form.get('signing_method', 'Signature'),
        'special_circumstances': special,
        'translator_name': request.form.get('translator_name', '').strip() or None,
        'translator_nric': request.form.get('translator_nric', '').strip() or None,
        'translator_relationship': request.form.get('translator_relationship', '').strip() or None,
        'translator_language': request.form.get('translator_language', '').strip() or None,
    }
    session.modified = True
    mark_step_complete(2)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 2})
    return redirect(url_for('wizard_step_executors'))


# -- Step 3: Executors (select from identities) -----------------------------

@app.route('/wizard/step/3', methods=['GET', 'POST'])
@login_required
def wizard_step_executors():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step3_executors.html',
            current_step=3,
            completed_steps=get_completed_steps(),
            data={
                'executors': session.get('step2_executors', [{}]),
                'executor_type': session.get('step3_executor_type', 'single'),
                'trustee_data': session.get('step3_trustees', {'same_as_executor': False, 'trustees': [{}]}),
            },
            persons=session.get('person_registry', []),
            beneficiaries=session.get('step4_beneficiaries', []),
        )

    # POST -- parse executor and trustee data
    executor_type = request.form.get('executor_type', 'single')
    count = int(request.form.get('executor_count', 1))
    executors = []
    for i in range(count):
        exec_entry_type = request.form.get(f'exec_type_{i}', 'individual').strip()
        role = request.form.get(f'exec_role_{i}', 'Primary')
        if executor_type == 'joint':
            role = 'Joint'
        elif executor_type == 'single':
            role = 'Primary'

        if exec_entry_type == 'corporate':
            corp_name = request.form.get(f'exec_corp_name_{i}', '').strip()
            if corp_name:
                executors.append({
                    'is_corporate': True,
                    'corp_name': corp_name,
                    'corp_reg': request.form.get(f'exec_corp_reg_{i}', '').strip(),
                    'corp_address': request.form.get(f'exec_corp_address_{i}', '').strip(),
                    'full_name': corp_name,
                    'nric_passport': request.form.get(f'exec_corp_reg_{i}', '').strip(),
                    'address': request.form.get(f'exec_corp_address_{i}', '').strip(),
                    'relationship': 'Corporate Trustee',
                    'role': role,
                })
        else:
            person_id = request.form.get(f'exec_person_id_{i}', '').strip()
            person = _get_person_from_registry(person_id)
            if not person:
                continue
            executors.append({
                'person_id': person_id,
                'full_name': person['full_name'],
                'nric_passport': person['nric_passport'],
                'address': person['address'],
                'relationship': request.form.get(f'exec_relationship_{i}', '').strip(),
                'role': role,
                'nationality': person.get('nationality', 'Malaysian'),
            })

    # Substitute executor(s) - supports individual persons or corporate trustees
    sub_exec_count = int(request.form.get('sub_executor_count', 1))
    for i in range(sub_exec_count):
        sub_type = request.form.get(f'sub_exec_type_{i}', 'individual').strip()
        if sub_type == 'corporate':
            corp_name = request.form.get(f'sub_exec_corp_name_{i}', '').strip()
            if corp_name:
                executors.append({
                    'is_corporate': True,
                    'corp_name': corp_name,
                    'corp_reg': request.form.get(f'sub_exec_corp_reg_{i}', '').strip(),
                    'corp_address': request.form.get(f'sub_exec_corp_address_{i}', '').strip(),
                    'full_name': corp_name,  # for display compatibility
                    'nric_passport': request.form.get(f'sub_exec_corp_reg_{i}', '').strip(),
                    'address': request.form.get(f'sub_exec_corp_address_{i}', '').strip(),
                    'relationship': 'Corporate Trustee',
                    'role': 'Substitute',
                })
        else:
            sub_exec_pid = request.form.get(f'sub_exec_person_id_{i}', '').strip()
            if sub_exec_pid:
                sub_person = _get_person_from_registry(sub_exec_pid)
                if sub_person:
                    executors.append({
                        'person_id': sub_exec_pid,
                        'full_name': sub_person['full_name'],
                        'nric_passport': sub_person['nric_passport'],
                        'address': sub_person['address'],
                        'relationship': request.form.get(f'sub_exec_relationship_{i}', '').strip(),
                        'role': 'Substitute',
                        'nationality': sub_person.get('nationality', 'Malaysian'),
                    })

    session['step2_executors'] = executors
    session['step3_executor_type'] = executor_type

    # Parse trustees
    trustee_same = bool(request.form.get('trustee_same_as_executor'))
    trustee_data = {'same_as_executor': trustee_same, 'trustees': [], 'substitute_trustee': {}, 'substitute_trustees': []}

    if not trustee_same:
        trustee_count = int(request.form.get('trustee_count', 1))
        for i in range(trustee_count):
            pid = request.form.get(f'trustee_person_id_{i}', '').strip()
            person = _get_person_from_registry(pid)
            if not person:
                continue
            trustee_data['trustees'].append({
                'person_id': pid,
                'full_name': person['full_name'],
                'nric_passport': person['nric_passport'],
                'address': person['address'],
                'relationship': request.form.get(f'trustee_relationship_{i}', '').strip(),
                'nationality': person.get('nationality', 'Malaysian'),
            })

        # Substitute trustee(s) - now supports multiple joint substitutes
        sub_tr_count = int(request.form.get('sub_trustee_count', 1))
        sub_trustees = []
        for i in range(sub_tr_count):
            sub_tr_pid = request.form.get(f'sub_trustee_person_id_{i}', '').strip()
            if sub_tr_pid:
                sub_tr = _get_person_from_registry(sub_tr_pid)
                if sub_tr:
                    sub_trustees.append({
                        'person_id': sub_tr_pid,
                        'full_name': sub_tr['full_name'],
                        'nric_passport': sub_tr['nric_passport'],
                        'address': sub_tr['address'],
                        'relationship': request.form.get(f'sub_trustee_relationship_{i}', '').strip(),
                        'nationality': sub_tr.get('nationality', 'Malaysian'),
                    })
        trustee_data['substitute_trustees'] = sub_trustees
        # Keep backward compat: set substitute_trustee to first one if any
        if sub_trustees:
            trustee_data['substitute_trustee'] = sub_trustees[0]

    session['step3_trustees'] = trustee_data
    session.modified = True
    mark_step_complete(3)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 3})
    return redirect(url_for('wizard_step_guardians'))


# -- Step 4: Guardians (select from identities, optional) -------------------

@app.route('/wizard/step/4', methods=['GET', 'POST'])
@login_required
def wizard_step_guardians():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step4_guardians.html',
            current_step=4,
            completed_steps=get_completed_steps(),
            data={
                'guardians': session.get('step3_guardians', []),
                'guardian_allowance': session.get('step3_guardian_allowance', {}),
                'exclude_spouse_guardian': session.get('step3_exclude_spouse', False),
                'exclude_spouse_guardian_reason': session.get('step3_exclude_spouse_reason', ''),
            },
            persons=session.get('person_registry', []),
        )

    # POST -- parse guardian selections from identities
    count = int(request.form.get('guardian_count', 0))
    guardians = []
    for i in range(count):
        person_id = request.form.get(f'guardian_person_id_{i}', '').strip()
        person = _get_person_from_registry(person_id)
        if not person:
            continue
        guardians.append({
            'person_id': person_id,
            'full_name': person['full_name'],
            'nric_passport': person['nric_passport'],
            'address': person['address'],
            'relationship': request.form.get(f'guardian_relationship_{i}', '').strip(),
            'role': request.form.get(f'guardian_role_{i}', 'Primary'),
            'nationality': person.get('nationality', 'Malaysian'),
        })

    # Guardian allowance
    ga = {}
    payment_mode = request.form.get('allowance_payment_mode', '').strip()
    if payment_mode:
        ga = {
            'payment_mode': payment_mode,
            'other_mode': request.form.get('allowance_other_mode', '').strip() or None,
            'amount': request.form.get('allowance_amount', '').strip() or None,
            'until_age': int(request.form.get('allowance_until_age', 0) or 0) or None,
            'source_of_payment': request.form.get('allowance_source_of_payment', '').strip() or None,
        }

    session['step3_guardians'] = guardians
    session['step3_guardian_allowance'] = ga
    session['step3_exclude_spouse'] = bool(request.form.get('exclude_spouse_guardian'))
    session['step3_exclude_spouse_reason'] = request.form.get('exclude_spouse_guardian_reason', '').strip() or None
    session.modified = True
    mark_step_complete(4)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 4})
    return redirect(url_for('wizard_step_beneficiaries'))


# -- Step 5: Beneficiaries (select from identities) -------------------------

@app.route('/wizard/step/5', methods=['GET', 'POST'])
@login_required
def wizard_step_beneficiaries():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step5_beneficiaries.html',
            current_step=5,
            completed_steps=get_completed_steps(),
            data={'beneficiaries': session.get('step4_beneficiaries', [{}])},
            persons=session.get('person_registry', []),
            executor_type=session.get('step3_executor_type', 'single'),
            executors=session.get('step2_executors', []),
        )

    # POST -- parse beneficiary selections from identities
    count = int(request.form.get('beneficiary_count', 1))
    beneficiaries = []
    for i in range(count):
        person_id = request.form.get(f'ben_person_id_{i}', '').strip()
        person = _get_person_from_registry(person_id)
        if not person:
            continue
        beneficiaries.append({
            'person_id': person_id,
            'full_name': person['full_name'],
            'nric_passport_birthcert': person['nric_passport'],
            'relationship': request.form.get(f'ben_relationship_{i}', '').strip(),
            'nationality': person.get('nationality', 'Malaysian'),
        })

    session['step4_beneficiaries'] = beneficiaries
    session.modified = True
    mark_step_complete(5)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 5})
    return redirect(url_for('wizard_step_gifts'))


# -- Step 6: Gifts (optional) ------------------------------------------------

def _refresh_wizard_session_from_db():
    """🔥 BURN-IN §10x.17 — sync session ← will.step*_data on every wizard GET.

    Chat handlers write directly to the Will's step*_data columns. The
    wizard reads from session storage (cookie-cached for client editing).
    Without this refresh, the wizard would show stale data after the chat
    confirms gifts. Call this at the top of every wizard step GET handler.
    """
    will_id = session.get('will_id')
    if not will_id:
        return
    w = db.session.get(Will, will_id)
    if not w:
        return
    def _j(s, default):
        try:
            return json.loads(s) if s else default
        except (json.JSONDecodeError, TypeError):
            return default
    s2 = _j(w.step2_data, {})
    if isinstance(s2, dict):
        if 'executors' in s2:   # new format
            session['step2_executors']    = s2.get('executors', [])
            session['step3_executor_type'] = s2.get('executor_type', 'single')
            session['step3_trustees']      = s2.get('trustee_data', {'same_as_executor': True, 'trustees': [{}]})
        else:
            session['step2_executors']    = s2.get('executors', [])
    s1 = _j(w.step1_data, {})
    # 🔥 BURN-IN §10x.17 — If step1_data is empty but a Person with
    # relationship='testator' exists, mirror that Person into step1
    # so wizard Step 2 (Testator) auto-populates from the chat's
    # identity walkthrough.
    if not s1.get('full_name'):
        try:
            tp = (Person.query
                  .filter_by(client_id=w.client_id, relationship='testator')
                  .first())
            if tp:
                s1 = {
                    'full_name':           tp.full_name or '',
                    'nationality':         tp.nationality or 'Malaysian',
                    'nric_passport':       tp.nric_passport or '',
                    'date_of_birth':       tp.date_of_birth or '',
                    'residential_address': tp.address or '',
                    'person_id':           tp.id,
                }
                w.step1_data = json.dumps(s1)
                db.session.commit()
        except Exception:
            db.session.rollback()
    session['step1'] = s1
    s3 = _j(w.step3_data, {})
    session['step3_guardians']         = s3.get('guardians', []) if isinstance(s3, dict) else []
    session['step3_guardian_allowance'] = s3.get('guardian_allowance', {}) if isinstance(s3, dict) else {}
    session['step4_beneficiaries']     = _j(w.step4_data, [])
    s5_raw = _j(w.step5_data, [])
    if isinstance(s5_raw, dict):
        s5_raw = s5_raw.get('gifts', []) or []
    if not isinstance(s5_raw, list):
        s5_raw = []
    # 🔥 §10x.128 → §10x.131 — image enrichment moved out of session.
    # Storing per-gift `documents[]` in session inflated the cookie past
    # Flask's 4093-byte limit. The enrichment now lives in
    # `enrich_gifts_with_documents()` which is called by the Step 6 GET
    # handler at render time and passed to the template directly — never
    # written back to session/cookie.
    session['step5_gifts']     = s5_raw
    session['step6_residuary'] = _j(w.step6_data, {})
    session['step7_trust']     = _j(w.step7_data, {})
    session['step8_others']    = _j(w.step8_data, {})
    # 🔥 §10x.121 — derive WIZARD STEP NUMBERS from DB data for the
    # sidebar's `{% if num in completed_steps %}` check. The DB column
    # `Will.completed_steps` stores MARKER STRINGS ('executor_confirmed',
    # 'residuary_confirmed'); the sidebar wants step NUMBERS [1, 2, 3,
    # 5, 7]. Bridge them here so chat-saved progress shows green
    # checkmarks in the wizard sidebar instead of red asterisks.
    db_completed_markers = _j(w.completed_steps, [])
    if not isinstance(db_completed_markers, list):
        db_completed_markers = []
    completed_nums = []
    # Step 1: Identities — at least 1 Person row beyond testator
    try:
        n_persons = (Person.query
                     .filter(Person.client_id == w.client_id,
                             Person.relationship.notilike('testator'))
                     .count())
        if n_persons >= 1:
            completed_nums.append(1)
    except Exception:
        pass
    # Step 2: Testator — REQUIRES name + NRIC + address (per §10x.124).
    # DOB/gender are auto-derived from NRIC; marital from family.
    # Per Phek Yi Ting template, only address is typed-compulsory.
    if all((s1.get(k) or '').strip() for k in (
            'full_name', 'nric_passport', 'residential_address')):
        completed_nums.append(2)
    # Step 3: Executors — ≥ 1 executor saved
    if isinstance(s2, dict) and len(s2.get('executors') or []) >= 1:
        completed_nums.append(3)
    # Step 4: Guardians — optional. Complete when explicitly set OR
    # when there are no minor children (skip-by-default).
    if isinstance(s3, dict):
        if (s3.get('guardians')
            or 'guardians_confirmed' in db_completed_markers
            or 'guardians_skipped' in db_completed_markers):
            completed_nums.append(4)
    # Step 5: Beneficiaries — ≥ 1 entry saved
    s4_list = _j(w.step4_data, [])
    if isinstance(s4_list, list) and len(s4_list) >= 1:
        completed_nums.append(5)
    # Step 6: Specific Gifts — optional. Complete when assets_confirmed
    # marker is set OR step5_data has gifts.
    if 'assets_confirmed' in db_completed_markers or (
            isinstance(s5_raw, list) and len(s5_raw) >= 1):
        completed_nums.append(6)
    # Step 7: Residuary — beneficiaries OR explicitly skipped
    s6_dict = _j(w.step6_data, {})
    if isinstance(s6_dict, dict) and (
            s6_dict.get('beneficiaries')
            or s6_dict.get('residuary_beneficiary_name')
            or s6_dict.get('skipped')):
        completed_nums.append(7)
    # Step 8: Trust — optional
    s7_dict = _j(w.step7_data, {})
    if isinstance(s7_dict, dict) and (
            s7_dict.get('trustee_name')
            or s7_dict.get('trust_skipped')
            or 'trust_confirmed' in db_completed_markers):
        completed_nums.append(8)
    # Step 9: Other Matters — optional
    s8_dict = _j(w.step8_data, {})
    if isinstance(s8_dict, dict) and (
            s8_dict.get('confirmed')
            or 'others_confirmed' in db_completed_markers):
        completed_nums.append(9)
    session['completed_steps'] = completed_nums
    _refresh_session_person_registry(w.client_id)
    session.modified = True


def _enrich_gifts_with_documents(client_id: str, gifts: list) -> list:
    """🔥 §10x.131 — render-time-only enrichment. Build a SHALLOW COPY of
    each gift with `documents[]` resolved from `document_id` AND mirror
    chat-saved `property_info` / financial fields into the legacy
    `property_details` / `financial_details` shape that step6_gifts.html
    reads. Used by the wizard Step 6 GET handler; result is passed to the
    template and never stored in session (which is cookie-backed and
    would exceed 4 KB).
    """
    if not isinstance(gifts, list):
        return gifts
    # One pipeline run per request to find sibling docs cheaply
    sibling_map = {}   # doc_id → set of sibling doc_ids
    try:
        from services.asset_pipeline import group_documents
        for gp in group_documents(client_id) or []:
            ids = list(gp.document_ids)
            for did in ids:
                sibling_map[did] = [s for s in ids if s != did]
    except Exception:
        pass
    out = []
    for g in gifts:
        if not isinstance(g, dict):
            out.append(g)
            continue
        gg = dict(g)   # shallow copy
        # Mirror chat-schema fields → wizard-template fields so the
        # collapsed summary + form pre-fill have the right values.
        kind = (gg.get('kind') or gg.get('asset_type') or '').lower()
        # Property: chat saves property_info; template reads property_details.
        if kind == 'property' or gg.get('gift_type') == 'property':
            pi = gg.get('property_info') or {}
            pd = dict(gg.get('property_details') or {})
            # Set gift_type for template radio button + summary
            gg['gift_type'] = 'property'
            # Merge: property_info overrides empty property_details fields
            for k in ('property_address', 'title_number', 'lot_number',
                      'daerah', 'negeri'):
                v = pi.get(k) or gg.get(k)
                if v and not pd.get(k):
                    pd[k] = v
            # Mukim: model uses 'bandar_pekan'; chat uses 'mukim'
            mukim_val = pi.get('mukim') or pi.get('bandar_pekan') or pd.get('bandar_pekan') or pd.get('mukim')
            if mukim_val:
                pd.setdefault('bandar_pekan', mukim_val)
                pd.setdefault('mukim', mukim_val)
            tt = pi.get('title_type') or pd.get('title_type')
            if tt:
                pd['title_type'] = tt
            # 🔥 §10x.145 — parse postcode/city/state out of the address
            # string when wizard form fields are empty. The chat saves the
            # full address as one string ("Shop No. 03 Jalan Gunung 4,
            # Seri Alam Masai, 81750 Masai, Johor"); the wizard expects
            # SEPARATE postcode/city/state inputs. Without parsing, the
            # postcode/city/state inputs render blank and the user thinks
            # "address incomplete".
            addr_str = (pd.get('property_address') or '')
            if addr_str:
                import re as _re
                # Postcode: 5 digits between word boundaries
                if not pd.get('postcode'):
                    pm = _re.search(r'\b(\d{5})\b', addr_str)
                    if pm:
                        pd['postcode'] = pm.group(1)
                # State: known Malaysian states at end (case-insensitive)
                if not pd.get('state'):
                    _STATES = ('Johor', 'Kedah', 'Kelantan', 'Melaka', 'Malacca',
                                'Negeri Sembilan', 'Pahang', 'Penang', 'Pulau Pinang',
                                'Perak', 'Perlis', 'Sabah', 'Sarawak', 'Selangor',
                                'Terengganu', 'Kuala Lumpur', 'Labuan', 'Putrajaya')
                    for st in _STATES:
                        if _re.search(rf'\b{_re.escape(st)}\b', addr_str, _re.I):
                            pd['state'] = st
                            break
                # City: 2-3 words after the postcode (e.g. "81750 Masai")
                if not pd.get('city') and pd.get('postcode'):
                    cm = _re.search(
                        rf'\b{pd["postcode"]}\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){{0,2}})',
                        addr_str)
                    if cm:
                        cand = cm.group(1).strip()
                        if cand and cand.lower() not in ('johor', 'kedah'):
                            pd['city'] = cand
            gg['property_details'] = pd
        # Financial: chat saves bank_name/account_number/insurer; template
        # reads financial_details.
        if kind in ('bank', 'insurance', 'epf', 'kwsp', 'mutual_fund',
                    'unit_trust', 'shares', 'financial') \
                or gg.get('gift_type') == 'financial':
            fd = dict(gg.get('financial_details') or {})
            fd.setdefault('institution', gg.get('institution') or gg.get('bank_name')
                          or gg.get('insurer') or fd.get('institution', ''))
            fd.setdefault('account_number', gg.get('account_number')
                          or gg.get('policy_number') or fd.get('account_number', ''))
            fd.setdefault('asset_type', gg.get('asset_type') or kind or fd.get('asset_type', ''))
            gg['gift_type'] = 'financial'
            gg['financial_details'] = fd
        if gg.get('documents'):
            out.append(gg); continue
        doc_id = (gg.get('document_id') or '').strip()
        if not doc_id or doc_id.startswith('_h3_synth_'):
            out.append(gg); continue
        docs_out = []
        seen_ids = set()
        for did in [doc_id] + sibling_map.get(doc_id, []):
            if did in seen_ids:
                continue
            seen_ids.add(did)
            d = db.session.get(Document, did)
            if not d:
                continue
            cat = d.category or ''
            doctype = ('title' if cat == 'property_title'
                       else 'spa' if cat == 'property_spa'
                       else 'cukai_harta' if cat == 'property_tax'
                       else 'financial' if cat in ('bank_statement', 'insurance')
                       else 'document')
            docs_out.append({
                'document_id': d.id,
                'filename':    d.original_filename or '',
                'url':         f'/api/documents/{d.id}',
                'doctype':     doctype,
                'category':    cat,
            })
        if docs_out:
            gg['documents'] = docs_out
        out.append(gg)
    return out


@app.route('/wizard/step/6', methods=['GET', 'POST'])
@login_required
def wizard_step_gifts():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()
        will_id = session.get('will_id')
        client_id = ''
        if will_id:
            _w = db.session.get(Will, will_id)
            if _w:
                client_id = _w.client_id
        gifts_raw = session.get('step5_gifts', [])
        # Render-time enrichment with documents[] — NOT written to session
        gifts_enriched = (_enrich_gifts_with_documents(client_id, gifts_raw)
                           if client_id else gifts_raw)
        return render_template(
            'wizard/step6_gifts.html',
            current_step=6,
            completed_steps=get_completed_steps(),
            data={'gifts': gifts_enriched},
            beneficiaries=session.get('step4_beneficiaries', []),
            persons=session.get('person_registry', []),
        )

    # POST -- parse gifts with nested allocations and structured details
    gift_count = int(request.form.get('gift_count', 0))
    gifts = []
    for gi in range(gift_count):
        gift_type = request.form.get(f'gift_type_{gi}', 'other').strip()
        desc = request.form.get(f'gift_desc_{gi}', '').strip()

        # Parse structured property details
        property_details = {}
        if gift_type == 'property':
            # Undivided share and ownership
            undivided = bool(request.form.get(f'gift_prop_undivided_{gi}'))
            testator_share = request.form.get(f'gift_prop_share_{gi}', '').strip() if undivided else ''
            ownership = 'joint' if undivided else request.form.get(f'gift_prop_ownership_{gi}', 'sole').strip()

            # Encumbrance
            encumbrance = request.form.get(f'gift_prop_encumbrance_{gi}', 'clean').strip()
            debt_source = request.form.get(f'gift_prop_debt_source_{gi}', 'residuary').strip() if encumbrance == 'encumbered' else ''

            # Split address fields
            prop_addr = request.form.get(f'gift_prop_address_{gi}', '').strip()
            postcode = request.form.get(f'gift_prop_postcode_{gi}', '').strip()
            city = request.form.get(f'gift_prop_city_{gi}', '').strip()
            state = request.form.get(f'gift_prop_state_{gi}', '').strip()
            country = request.form.get(f'gift_prop_country_{gi}', 'Malaysia').strip()

            # Strip postcode/city/state from address if already in separate fields
            # This prevents duplication when user selects from existing addresses
            clean_addr = prop_addr
            if postcode and city:
                # Remove trailing duplicates like ", 81100 JOHOR BAHRU, JOHOR, 81100 JOHOR BAHRU, JOHOR..."
                import re
                # Remove any repeated postcode+city+state patterns from address
                dup_pattern = re.compile(
                    r',\s*' + re.escape(postcode) + r'\s+' + re.escape(city) + r'(?:\s*,\s*' + re.escape(state) + r')?',
                    re.IGNORECASE
                )
                # Keep only the first occurrence (part of the original address), remove the rest
                matches = list(dup_pattern.finditer(clean_addr))
                if len(matches) > 1:
                    # Remove all but first match
                    for m in reversed(matches[1:]):
                        clean_addr = clean_addr[:m.start()] + clean_addr[m.end():]
                elif len(matches) == 1:
                    # If postcode/city/state are in separate fields, strip them from address
                    clean_addr = clean_addr[:matches[0].start()].rstrip(', ')

            property_details = {
                'property_address': clean_addr or prop_addr,
                'title_type': request.form.get(f'gift_prop_title_type_{gi}', '').strip(),
                'title_number': request.form.get(f'gift_prop_title_number_{gi}', '').strip(),
                'lot_number': request.form.get(f'gift_prop_lot_number_{gi}', '').strip(),
                'bandar_pekan': request.form.get(f'gift_prop_bandar_{gi}', '').strip(),
                'daerah': request.form.get(f'gift_prop_daerah_{gi}', '').strip(),
                'negeri': state or request.form.get(f'gift_prop_negeri_{gi}', '').strip(),
                'state': state,
                'postcode': postcode,
                'city': city,
                'country': country,
                'ownership_type': ownership,
                'undivided_share': undivided,
                'testator_share': testator_share,
                'encumbrance_status': encumbrance,
                'debt_source': debt_source,
            }
            if not property_details['property_address']:
                continue

        # Parse structured financial details
        financial_details = {}
        if gift_type == 'financial':
            account_ownership = request.form.get(f'gift_fin_ownership_{gi}', 'individual').strip()
            financial_details = {
                'institution': request.form.get(f'gift_fin_institution_{gi}', '').strip(),
                'account_number': request.form.get(f'gift_fin_account_{gi}', '').strip(),
                'asset_type': request.form.get(f'gift_fin_type_{gi}', '').strip(),
                'description': request.form.get(f'gift_fin_desc_{gi}', '').strip(),
                'account_ownership': account_ownership,
            }
            if not financial_details['institution'] and not financial_details['description']:
                continue

        # For "other" type, skip if no description
        if gift_type == 'other' and not desc:
            continue

        subject_to_trust = bool(request.form.get(f'gift_trust_{gi}'))
        subject_to_guardian_allowance = bool(request.form.get(f'gift_guardian_allowance_{gi}'))
        sell_property = bool(request.form.get(f'gift_sell_property_{gi}'))
        substitute_mode = request.form.get(f'gift_{gi}_sub_mode', 'equal')

        alloc_count = int(request.form.get(f'gift_{gi}_alloc_count', 0))
        allocations = []
        for ai_idx in range(alloc_count):
            ben_name = request.form.get(f'gift_{gi}_alloc_name_{ai_idx}', '').strip()
            if not ben_name:
                continue
            # Parse per-MB substitutes (only used when substitute_mode == 'specific')
            subs = []
            if substitute_mode == 'specific':
                sub_count = int(request.form.get(f'gift_{gi}_mb_{ai_idx}_sub_count', 0))
                for si in range(sub_count):
                    sub_name = request.form.get(f'gift_{gi}_mb_{ai_idx}_sub_name_{si}', '').strip()
                    sub_share = request.form.get(f'gift_{gi}_mb_{ai_idx}_sub_share_{si}', '').strip()
                    if sub_name:
                        subs.append({'beneficiary_name': sub_name, 'share': sub_share or '100%'})
            allocations.append({
                'beneficiary_name': ben_name,
                'share': request.form.get(f'gift_{gi}_alloc_share_{ai_idx}', '').strip(),
                'role': request.form.get(f'gift_{gi}_alloc_role_{ai_idx}', 'MB'),
                'substitutes': subs,
            })

        # Parse uploaded document references
        gift_docs_json = request.form.get(f'gift_docs_{gi}', '[]')
        try:
            gift_docs = json.loads(gift_docs_json) if gift_docs_json else []
        except (json.JSONDecodeError, TypeError):
            gift_docs = []

        gifts.append({
            'gift_type': gift_type,
            'description': desc,
            'property_details': property_details,
            'financial_details': financial_details,
            'allocations': allocations,
            'subject_to_trust': subject_to_trust,
            'subject_to_guardian_allowance': subject_to_guardian_allowance,
            'sell_property': sell_property,
            'substitute_mode': substitute_mode,
            'documents': gift_docs,
        })

    # Reorder gifts if user changed order via drag-and-drop or sort
    gift_order_str = request.form.get('gift_order', '')
    if gift_order_str:
        order = [int(x) for x in gift_order_str.split(',') if x.strip().isdigit()]
        gift_map = {i: g for i, g in enumerate(gifts)}
        reordered = [gift_map[i] for i in order if i in gift_map]
        # Include any gifts not in order list (safety)
        for i, g in enumerate(gifts):
            if i not in order:
                reordered.append(g)
        gifts = reordered

    session['step5_gifts'] = gifts
    session.modified = True
    mark_step_complete(6)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 6})
    return redirect(url_for('wizard_step_residuary'))


# -- Step 7: Residuary Estate ------------------------------------------------

@app.route('/wizard/step/7', methods=['GET', 'POST'])
@login_required
def wizard_step_residuary():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step7_residuary.html',
            current_step=7,
            completed_steps=get_completed_steps(),
            data=session.get('step6_residuary', {}),
            beneficiaries=session.get('step4_beneficiaries', []),
            persons=session.get('person_registry', []),
            gifts=session.get('step5_gifts', []),
        )

    # POST -- parse main beneficiaries and substitute groups
    main_count = int(request.form.get('main_beneficiary_count', 0))
    main_beneficiaries = []
    for i in range(main_count):
        # Support both person_id (dropdown) and name (text fallback)
        person_id = request.form.get(f'main_ben_person_id_{i}', '').strip()
        name = request.form.get(f'main_ben_name_{i}', '').strip()
        if person_id:
            person = _get_person_from_registry(person_id)
            if person:
                name = person['full_name']
        if not name:
            continue
        entry = {
            'beneficiary_name': name,
            'share': request.form.get(f'main_ben_share_{i}', '').strip(),
            'group': 'main',
        }
        if person_id:
            entry['person_id'] = person_id
        main_beneficiaries.append(entry)

    # Substitute groups
    sub_group_count = int(request.form.get('substitute_group_count', 0))
    substitute_groups = []
    for gi in range(sub_group_count):
        sub_count = int(request.form.get(f'sub_group_{gi}_count', 0))
        group = []
        for si in range(sub_count):
            person_id = request.form.get(f'sub_group_{gi}_person_id_{si}', '').strip()
            name = request.form.get(f'sub_group_{gi}_name_{si}', '').strip()
            if person_id:
                person = _get_person_from_registry(person_id)
                if person:
                    name = person['full_name']
            if not name:
                continue
            entry = {
                'beneficiary_name': name,
                'share': request.form.get(f'sub_group_{gi}_share_{si}', '').strip(),
                'group': f'substitute_{gi + 1}',
            }
            if person_id:
                entry['person_id'] = person_id
            group.append(entry)
        if group:
            substitute_groups.append(group)

    additional_notes = request.form.get('additional_notes', '').strip() or None

    session['step6_residuary'] = {
        'main_beneficiaries': main_beneficiaries,
        'substitute_groups': substitute_groups,
        'additional_notes': additional_notes,
    }
    session.modified = True
    mark_step_complete(7)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 7})
    return redirect(url_for('wizard_step_trust'))


# -- Step 8: Testamentary Trust (optional) ------------------------------------

@app.route('/wizard/step/8', methods=['GET', 'POST'])
@login_required
def wizard_step_trust():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step8_trust.html',
            current_step=8,
            completed_steps=get_completed_steps(),
            data=session.get('step7_trust', {}),
            beneficiaries=session.get('step4_beneficiaries', []),
            gifts=session.get('step5_gifts', []),
            persons=session.get('person_registry', []),
        )

    # POST -- parse trust data
    ben_count = int(request.form.get('trust_beneficiary_count', 0))
    trust_bens = []
    for i in range(ben_count):
        name = request.form.get(f'trust_ben_name_{i}', '').strip()
        if not name:
            continue
        trust_bens.append({
            'beneficiary_name': name,
            'share': request.form.get(f'trust_ben_share_{i}', '').strip(),
            'role': request.form.get(f'trust_ben_role_{i}', 'MB'),
        })

    purposes = request.form.getlist('purposes')
    other_purpose = request.form.get('purposes_other_text', '').strip() or None

    trust_data = {}
    if trust_bens:
        trust_data = {
            'beneficiaries': trust_bens,
            'purposes': purposes,
            'other_purpose': other_purpose,
            'duration': request.form.get('trust_duration', '').strip() or None,
            'assets_from_gifts': request.form.getlist('gift_references'),
            'property_actions': {},
            'property_residents': {},
            'payment_mode': request.form.get('payment_mode', '').strip() or None,
            'payment_amount': request.form.get('payment_amount', '').strip() or None,
            'other_payment_mode': request.form.get('payment_mode_other', '').strip() or None,
            'balance_of_trust': request.form.get('balance_of_trust', '').strip() or None,
            'separate_trustee': bool(request.form.get('separate_trustee')),
            'trustee_person_id': request.form.get('trustee_person_id', '').strip() or None,
            'trustee_relationship': request.form.get('trustee_relationship', '').strip() or None,
        }
        # Parse per-property actions (reside/lease/sell) and resident selections
        gift_refs = trust_data['assets_from_gifts']
        for ref in gift_refs:
            # Extract gift number from "Gift 1", "Gift 2", etc.
            try:
                gift_num = int(ref.replace('Gift ', ''))
            except (ValueError, AttributeError):
                continue
            action = request.form.get(f'prop_action_{gift_num}', '').strip()
            if action:
                trust_data['property_actions'][ref] = action
            resident = request.form.get(f'prop_resident_{gift_num}', '').strip()
            if resident:
                trust_data['property_residents'][ref] = resident

        # Look up trustee identity
        trustee_pid = trust_data.get('trustee_person_id')
        if trustee_pid:
            trustee_person = _get_person_from_registry(trustee_pid)
            if trustee_person:
                trust_data['trustee_name'] = trustee_person['full_name']
                trust_data['trustee_nric'] = trustee_person['nric_passport']
                trust_data['trustee_address'] = trustee_person['address']

    session['step7_trust'] = trust_data
    session.modified = True
    mark_step_complete(8)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 8})
    return redirect(url_for('wizard_step_others'))


# -- Step 9: Other Matters (optional) ----------------------------------------

@app.route('/wizard/step/9', methods=['GET', 'POST'])
@login_required
def wizard_step_others():
    if request.method == 'GET':
        _refresh_wizard_session_from_db()   # §10x.17
        return render_template(
            'wizard/step9_others.html',
            current_step=9,
            completed_steps=get_completed_steps(),
            data=session.get('step8_others', {}),
        )

    # POST -- parse other matters
    om_data = {}

    terms = request.form.get('terms_of_endearment', '').strip()
    if terms:
        om_data['terms_of_endearment'] = terms

    # Commorientes
    om_data['commorientes_enabled'] = bool(request.form.get('commorientes_enabled'))
    if om_data['commorientes_enabled']:
        om_data['commorientes_days'] = int(request.form.get('commorientes_days', 0) or 0) or None

    # Exclusion
    om_data['exclusion_enabled'] = bool(request.form.get('exclusion_enabled'))
    if om_data['exclusion_enabled']:
        om_data['exclusion_name'] = request.form.get('exclusion_name', '').strip() or None
        om_data['exclusion_nric'] = request.form.get('exclusion_nric', '').strip() or None
        om_data['exclusion_relationship'] = request.form.get('exclusion_relationship', '').strip() or None
        om_data['exclusion_reason'] = request.form.get('exclusion_reason', '').strip() or None

    # Unnamed children
    om_data['unnamed_children_enabled'] = bool(request.form.get('unnamed_children_enabled'))
    if om_data['unnamed_children_enabled']:
        om_data['unnamed_children_spouse_name'] = request.form.get('unnamed_children_spouse_name', '').strip() or None
        om_data['unnamed_children_spouse_nric'] = request.form.get('unnamed_children_spouse_nric', '').strip() or None

    # Joint bank account clause
    om_data['joint_account_clause_enabled'] = bool(request.form.get('joint_account_clause_enabled'))

    # Discharge / lien clause for properties (default ON)
    om_data['discharge_clause_enabled'] = bool(request.form.get('discharge_clause_enabled'))
    om_data['discharge_placement'] = request.form.get('discharge_placement', 'per_property')

    # Testator satisfaction clause (default ON)
    om_data['testator_satisfaction_enabled'] = bool(request.form.get('testator_satisfaction_enabled'))

    # Translator / Interpreter attestation
    om_data['translator_enabled'] = bool(request.form.get('translator_enabled'))
    if om_data['translator_enabled']:
        om_data['translator_name'] = request.form.get('translator_name', '').strip()
        om_data['translator_nric'] = request.form.get('translator_nric', '').strip()
        om_data['translator_language'] = request.form.get('translator_language', '').strip()
        om_data['translator_address'] = request.form.get('translator_address', '').strip()

    additional = request.form.get('additional_instructions', '').strip()
    if additional:
        om_data['additional_instructions'] = additional

    session['step8_others'] = om_data
    session.modified = True
    mark_step_complete(9)
    save_will_to_db()
    if request.form.get('_save_draft'):
        return jsonify({'ok': True, 'step': 9})
    return redirect(url_for('wizard_step_review'))


# -- Step 10: Review ---------------------------------------------------------

@app.route('/wizard/step/10', methods=['GET'])
@login_required
def wizard_step_review():
    _refresh_wizard_session_from_db()   # §10x.17 — Review must reflect chat saves
    # Build the will data model from session
    try:
        will_data = build_will_data()
    except Exception as e:
        flash(f'Error building will data: {e}', 'error')
        return redirect(url_for('wizard_step_identities'))

    # Run validation
    from validation.legal_rules import validate_will_data, get_errors, get_warnings
    validation_results = validate_will_data(will_data)
    errors = get_errors(validation_results)
    warnings = get_warnings(validation_results)
    infos = [r for r in validation_results if r.severity == 'INFO']

    # Build summary data dict for template
    summary = {
        'identities': session.get('person_registry', []),
        'testator': session.get('step1', {}),
        'executors': session.get('step2_executors', []),
        'executor_type': session.get('step3_executor_type', 'single'),
        'trustee_data': session.get('step3_trustees', {'same_as_executor': True}),
        'guardians': session.get('step3_guardians', []),
        'guardian_allowance': session.get('step3_guardian_allowance', {}),
        'beneficiaries': session.get('step4_beneficiaries', []),
        'gifts': session.get('step5_gifts', []),
        'residuary': session.get('step6_residuary', {}),
        'trust': session.get('step7_trust', {}),
        'others': session.get('step8_others', {}),
        'other_matters': session.get('step8_others', {}),
    }

    # Check if a firm logo exists for this tenant
    has_logo = _get_logo_path() is not None

    # Get current include_logo setting from will record (default True)
    include_logo = True
    will_id = session.get('will_id')
    if will_id:
        wr = db.session.get(Will, will_id)
        if wr and wr.include_logo is not None:
            include_logo = wr.include_logo

    return render_template(
        'wizard/step10_review.html',
        current_step=10,
        completed_steps=get_completed_steps(),
        summary=summary,
        will_data=summary,
        validation_results=validation_results,
        validation_errors=errors,
        validation_warnings=warnings,
        validation_infos=infos,
        has_errors=len(errors) > 0,
        has_logo=has_logo,
        include_logo=include_logo,
        has_versions=WillVersion.query.filter_by(will_id=session.get('will_id', '')).count() > 0 if session.get('will_id') else False,
    )


# -- Generate Will -----------------------------------------------------------

@app.route('/wizard/generate', methods=['POST'])
@login_required
def wizard_generate():
    try:
        will_data = build_will_data()
    except Exception as e:
        flash(f'Error building will data: {e}', 'error')
        return redirect(url_for('wizard_step_review'))

    # Run validation -- block on errors
    from validation.legal_rules import validate_will_data, get_errors
    validation_results = validate_will_data(will_data)
    errors = get_errors(validation_results)
    if errors:
        for err in errors:
            flash(f'Validation Error: {err.message}', 'error')
        return redirect(url_for('wizard_step_review'))

    # Draft will using AI (or mock)
    try:
        if ANTHROPIC_API_KEY and ANTHROPIC_API_KEY != 'your-api-key-here':
            from ai.drafter import draft_will
            will_text = draft_will(will_data)
        else:
            from ai.drafter import draft_will_mock
            will_text = draft_will_mock(will_data)
    except Exception as e:
        flash(f'Error generating will: {e}', 'error')
        traceback.print_exc()
        return redirect(url_for('wizard_step_review'))

    # Store generated will text in DB (not session — session cookie too large)
    session['generated_will_text'] = will_text  # temporary for save_will_to_db
    session.modified = True
    mark_step_complete(10)
    save_will_to_db()

    # Save include_logo preference
    include_logo = '1' in request.form.getlist('include_logo')
    will_id = session.get('will_id')
    if will_id:
        wr = db.session.get(Will, will_id)
        if wr:
            wr.include_logo = include_logo
            db.session.commit()

    # Save version history
    will_id = session.get('will_id')
    if will_id:
        # Determine version number
        latest_version = WillVersion.query.filter_by(will_id=will_id).order_by(
            WillVersion.version_number.desc()
        ).first()
        next_version = (latest_version.version_number + 1) if latest_version else 1
        user_name = ''
        if session.get('user_id'):
            u = db.session.get(User, session['user_id'])
            user_name = u.name if u else ''
        note = 'Initial generation' if next_version == 1 else f'Re-generated (version {next_version})'
        version = WillVersion(
            will_id=will_id,
            version_number=next_version,
            will_text=will_text,
            generated_by=session.get('user_id'),
            generated_by_name=user_name,
            note=note,
        )
        db.session.add(version)
        db.session.commit()

    # If approver generated the will, auto-approve it (no submission step needed)
    will_id = session.get('will_id')
    if will_id:
        user = db.session.get(User, session.get('user_id'))
        if user and ROLE_PERMS.get(user.role, {}).get('canApprove'):
            wr = db.session.get(Will, will_id)
            if wr:
                wr.status = 'approved'
                wr.approved_by = user.id
                wr.approved_at = datetime.utcnow()
                wr.approval_remarks = 'Auto-approved (generated by approver)'
                db.session.commit()

    # Remove from session to keep cookie small
    session.pop('generated_will_text', None)
    session.modified = True
    flash('Will generated successfully! You can now view, edit, or download it.', 'info')
    return redirect(url_for('preview'))


# -- Preview -----------------------------------------------------------------

@app.route('/preview')
@login_required
def preview():
    # Read will text from DB (not session) to avoid oversized cookies
    will_text = ''
    will_record = None
    versions = []
    viewing_version = None  # which version is being displayed
    if session.get('will_id'):
        will_record = db.session.get(Will, session['will_id'])
        if will_record:
            # Auto-approve when an approver views a generated will (the approver is the
            # lawyer — Kylie at Alan Tan & Associates — so an explicit submit-then-approve
            # round-trip is unnecessary friction). Only fires once: status must be 'generated'.
            if (will_record.status == 'generated'
                and getattr(g, 'perms', {}).get('canApprove')
                and getattr(g, 'user', None)):
                will_record.status = 'approved'
                will_record.approved_by = g.user.id
                will_record.approved_at = datetime.utcnow()
                will_record.approval_remarks = 'Auto-approved on preview by approver'
                db.session.commit()
            # Load version history
            versions = WillVersion.query.filter_by(will_id=will_record.id).order_by(
                WillVersion.version_number.desc()
            ).all()

            # Check if a specific version is requested
            ver_num = request.args.get('version', type=int)
            if ver_num and versions:
                for v in versions:
                    if v.version_number == ver_num:
                        will_text = v.will_text
                        viewing_version = v
                        break

            # Default: show current (latest) will text
            if not will_text:
                will_text = will_record.generated_will_text or ''
    if not will_text:
        # Fallback to session for backward compat
        will_text = session.get('generated_will_text', '')
    if not will_text:
        # If will was previously generated but text was lost, direct to Step 10 to re-generate
        if will_record and will_record.status in ('generated', 'pending_approval', 'approved'):
            flash('The will text needs to be re-generated. Please click "Generate My Will" below.', 'warning')
        else:
            flash('No will has been generated yet. Please complete the wizard first.', 'warning')
        return redirect(url_for('wizard_step_review'))

    testator_name = session.get('step1', {}).get('full_name', 'Unknown')

    # Look up last editor name and edit logs
    editor_name = None
    edit_logs = []
    client_email = None
    if will_record:
        if will_record.text_edited_by:
            editor = db.session.get(User, will_record.text_edited_by)
            if editor:
                editor_name = editor.name
        edit_logs = WillEditLog.query.filter_by(will_id=will_record.id).order_by(WillEditLog.edited_at.desc()).all()

        # Attach edit logs to versions: for each version, find edits made between
        # this version's creation and the next version's creation
        if versions and edit_logs:
            for vi, v in enumerate(versions):  # versions are desc by version_number
                v_start = v.created_at
                v_end = versions[vi - 1].created_at if vi > 0 else None  # next newer version
                v.edit_logs = [
                    el for el in edit_logs
                    if el.edited_at >= v_start and (v_end is None or el.edited_at < v_end)
                ]
                # Update display time to last edit time if edits exist
                if v.edit_logs:
                    v.last_edited_at = v.edit_logs[0].edited_at  # most recent edit (already desc)
                    v.last_edited_by = v.edit_logs[0].edited_by_name
                else:
                    v.last_edited_at = None
                    v.last_edited_by = None

        # Get client email for the send-email button
        if will_record.client_id:
            client = db.session.get(Client, will_record.client_id)
            if client:
                client_email = client.email

    return render_template(
        'preview.html',
        will_text=will_text,
        testator_name=testator_name,
        will_record=will_record,
        editor_name=editor_name,
        edit_logs=edit_logs,
        client_email=client_email,
        versions=versions,
        viewing_version=viewing_version,
        has_logo=_get_logo_path() is not None,
        include_logo=will_record.include_logo if will_record and will_record.include_logo is not None else True,
    )


# -- Download -----------------------------------------------------------------

@app.route('/download/verification-pdf')
@login_required
def download_verification_pdf():
    """Generate and download verification PDF with documents + field data."""
    from documents.verification_pdf import generate_verification_pdf

    client_id = session.get('client_id')
    if not client_id:
        flash('No client loaded.', 'error')
        return redirect(url_for('wizard_step', step=1))

    # Gather persons with documents
    persons = []
    for p in session.get('person_registry', []):
        person_record = db.session.get(Person, p.get('id'))
        if person_record:
            pdata = {k: getattr(person_record, k, '') for k in
                     ['full_name', 'nric_passport', 'nationality', 'date_of_birth',
                      'gender', 'address', 'relationship', 'document_id']}
            pdata['id'] = person_record.id
            persons.append(pdata)

    # Gather gifts with documents
    gifts = session.get('step5_gifts', [])

    # Build documents map: document_id -> file_path
    documents_map = {}
    doc_ids = [p['document_id'] for p in persons if p.get('document_id')]
    for g in gifts:
        for d in g.get('documents', []):
            did = d.get('document_id', '')
            # Extract document_id from URL if not set directly
            if not did and d.get('url'):
                url = d['url']
                # URL format: /api/documents/UUID
                parts = url.rstrip('/').split('/')
                if len(parts) >= 3:
                    did = parts[-1]
            if did:
                doc_ids.append(did)
                # Also store the mapping so verification PDF can find it
                d['document_id'] = did
    for doc_id in doc_ids:
        doc = db.session.get(Document, doc_id)
        if doc and doc.file_path:
            # file_path may be relative — prepend UPLOAD_DIR if not absolute
            fp = doc.file_path
            if not os.path.isabs(fp):
                fp = os.path.join(UPLOAD_DIR, fp)
            if os.path.exists(fp):
                documents_map[doc_id] = fp

    testator_name = session.get('step1', {}).get('full_name', 'Unknown')

    filepath = generate_verification_pdf(persons, gifts, documents_map, testator_name)
    if not filepath or not os.path.exists(filepath):
        flash('Could not generate verification PDF.', 'error')
        return redirect(url_for('wizard_step', step=10))

    return send_file(filepath, as_attachment=True,
                     download_name=f"Verification_{''.join(c for c in testator_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')}.pdf",
                     mimetype='application/pdf')


@app.route('/download/<fmt>')
@login_required
def download(fmt):
    # Read will text from DB (not session) to avoid oversized cookies
    will_text = ''
    if session.get('will_id'):
        will_record = db.session.get(Will, session['will_id'])
        if will_record:
            will_text = will_record.generated_will_text or ''
    if not will_text:
        will_text = session.get('generated_will_text', '')
    if not will_text:
        flash('No will has been generated yet.', 'warning')
        return redirect(url_for('preview'))

    # Check download permissions based on role and approval status
    user_role = session.get('user_role', '')
    will_id = session.get('will_id')
    if user_role in ('admin', 'advisor') and will_id:
        will_record = db.session.get(Will, will_id)
        if will_record and will_record.status != 'approved':
            flash('This will must be approved before it can be downloaded.', 'warning')
            return redirect(url_for('preview'))

    testator_name = session.get('step1', {}).get('full_name', 'Will')
    safe_name = "".join(c for c in testator_name if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_name = safe_name.replace(' ', '_') or 'Will'

    if fmt == 'docx':
        # Locked format: Alan & Tan / WillCraft standard. Single source of truth
        # for ALL will docx generation. Format approved 2026-05-02 — see
        # documents/will_docx.py + documents/sample_will_phek_yi_ting.py.
        from documents.will_docx import build_will_docx
        tenant = get_tenant()
        firm_info = None
        logo = None
        if tenant.get('firm_name'):
            firm_info = {
                'firm_name': tenant.get('firm_name', ''),
                'firm_address': tenant.get('firm_address', ''),
                'firm_phone': tenant.get('firm_phone', ''),
                'firm_email': tenant.get('email_from', ''),
            }
        if will_id:
            wr = db.session.get(Will, will_id)
            if wr and wr.include_logo:
                logo = _get_logo_path()
        # Determine draft status from the will record (default True for safety)
        is_draft = True
        if will_id:
            wr_status = (db.session.get(Will, will_id) or None)
            if wr_status and getattr(wr_status, 'status', '') in ('approved', 'final'):
                is_draft = False
        filepath = build_will_docx(
            will_text,
            firm_info=firm_info,
            logo_path=logo,
            is_draft=is_draft,
        )
    elif fmt == 'pdf':
        from documents.pdf_generator import generate_pdf
        # Check will's include_logo flag
        logo = None
        if will_id:
            wr = db.session.get(Will, will_id)
            if wr and wr.include_logo:
                logo = _get_logo_path()
        else:
            logo = _get_logo_path()
        # Build firm info for cover page and prepared-by page
        tenant = get_tenant()
        firm_info = None
        if tenant.get('firm_name'):
            firm_info = {
                'firm_name': tenant.get('firm_name', ''),
                'firm_address': tenant.get('firm_address', ''),
                'firm_phone': tenant.get('firm_phone', ''),
                'firm_email': tenant.get('email_from', ''),
            }
        filepath = generate_pdf(will_text, safe_name, logo_path=logo, firm_info=firm_info)
    else:
        flash('Unsupported download format.', 'error')
        return redirect(url_for('preview'))

    # Save persistent copy to client folder
    try:
        client_id = session.get('client_id')
        if client_id:
            client = db.session.get(Client, client_id)
            if client:
                from uploads import save_generated_will
                with open(filepath, 'rb') as f:
                    file_bytes = f.read()
                will_record = db.session.get(Will, session.get('will_id'))
                is_draft = will_record.status == 'draft' if will_record else True
                saved_name, rel_path = save_generated_will(
                    client.folder_name, file_bytes, fmt, is_draft=is_draft
                )
                # Update will status
                if will_record:
                    will_record.status = 'generated'
                    db.session.commit()
    except Exception as e:
        app.logger.warning(f'Could not save persistent copy: {e}')

    mime = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
    }.get(fmt, 'application/octet-stream')

    return send_file(
        filepath,
        as_attachment=True,
        download_name=f'{safe_name}_Will.{fmt}',
        mimetype=mime,
    )


# -- Reset Session ------------------------------------------------------------

@app.route('/reset')
@login_required
def reset():
    # Preserve auth keys during session reset
    user_id = session.get('user_id')
    user_role = session.get('user_role')
    user_name = session.get('user_name')
    user_email = session.get('user_email')
    session.clear()
    if user_id:
        session['user_id'] = user_id
        session['user_role'] = user_role
        session['user_name'] = user_name
        session['user_email'] = user_email
    flash('Your session has been reset. You can start a new will.', 'info')
    return redirect(url_for('index'))


@app.route('/wizard/new', methods=['GET', 'POST'])
@login_required
def wizard_new():
    """Start a brand-new will. Testator name is REQUIRED — collected via the
    "+ New Will" modal (POST). GET without a name falls back to /wills with
    a flash so users can't bypass the modal by deep-linking.
    """
    # Preserve auth keys during session reset
    auth_keys = {k: session.get(k) for k in
                 ('user_id', 'user_role', 'user_name', 'user_email')}

    if request.method == 'POST':
        full_name = (request.form.get('full_name') or '').strip()
        nric = (request.form.get('nric_passport') or '').strip()
        if not full_name:
            flash('Please enter the testator\'s full name to start a new will.', 'error')
            return redirect(url_for('will_list'))

        session.clear()
        for k, v in auth_keys.items():
            if v: session[k] = v
        session['step1'] = {
            'full_name': full_name,
            'nric_passport': nric,
            'nationality': 'Malaysian',
        }
        # Create Client + Person immediately so Wills list / Client folder UI
        # show the right name from the very first page load.
        client_id = ensure_client()
        from services.person_registry import ensure_person
        pid = ensure_person(client_id, full_name, nric=nric, relationship='Testator')
        if pid:
            session['step1']['person_id'] = pid
        db.session.commit()
        return redirect(url_for('wizard_step_identities'))

    # GET — was the entry point for "+ New Will" before; now require the modal.
    flash('Please use the "+ New Will" button so we can capture the testator\'s name.', 'warning')
    return redirect(url_for('will_list'))


# ---------------------------------------------------------------------------
# Probate Application Module
# ---------------------------------------------------------------------------

MALAYSIAN_COURTS = [
    'JOHOR BAHRU', 'KUALA LUMPUR', 'SHAH ALAM', 'PUTRAJAYA', 'GEORGE TOWN',
    'IPOH', 'KUANTAN', 'KOTA BHARU', 'KUALA TERENGGANU', 'MELAKA',
    'SEREMBAN', 'ALOR SETAR', 'KANGAR', 'KOTA KINABALU', 'KUCHING',
    'MUAR', 'BATU PAHAT', 'SEGAMAT', 'KLANG', 'TAIPING',
]

MALAYSIAN_STATES = [
    'JOHOR DARUL TAKZIM', 'SELANGOR DARUL EHSAN', 'WILAYAH PERSEKUTUAN KUALA LUMPUR',
    'WILAYAH PERSEKUTUAN PUTRAJAYA', 'PULAU PINANG', 'PERAK DARUL RIDZUAN',
    'PAHANG DARUL MAKMUR', 'KELANTAN DARUL NAIM', 'TERENGGANU DARUL IMAN',
    'MELAKA', 'NEGERI SEMBILAN DARUL KHUSUS', 'KEDAH DARUL AMAN', 'PERLIS INDERA KAYANGAN',
    'SABAH', 'SARAWAK',
]


def _sync_probate_from_will(probate, will_record):
    """Copy testator/executor data from linked will into probate fields if empty.
    Makes probate the single source of truth for all downstream validation/generation.
    Returns (changed: bool, errors: list[str]) — errors explain why fields couldn't be synced."""
    errors = []
    if not will_record:
        errors.append('No linked will record found')
        return False, errors
    changed = False

    # Extract testator from will step1_data
    step1 = json.loads(will_record.step1_data or '{}')
    testator = step1 if step1 else {}
    if not testator:
        errors.append(f'Will ({will_record.id}) has no testator data (step1_data is empty)')

    # Extract executor from will — try step3_data first (executors), then step2_data
    executor = {}
    step3 = json.loads(will_record.step3_data or '[]')
    if isinstance(step3, list) and step3:
        executor = step3[0]
    elif isinstance(step3, dict) and (step3.get('full_name') or step3.get('person_id')):
        executor = step3
    # If step3 didn't yield a valid executor (e.g. it's guardians data), try step2
    if not executor or not (executor.get('full_name') or executor.get('person_id')):
        step2 = json.loads(will_record.step2_data or '{}')
        executors = step2.get('executors', []) if isinstance(step2, dict) else step2
        if isinstance(executors, list) and executors:
            executor = executors[0]
    if not executor or not (executor.get('full_name') or executor.get('person_id')):
        errors.append(f'Will ({will_record.id}) has no executor data (step2_data and step3_data are empty)')

    # Resolve person details from identity registry
    identities = json.loads(will_record.identities_data or '[]')
    id_lookup = {p.get('id', ''): p for p in identities} if identities else {}

    if executor.get('person_id'):
        if executor['person_id'] in id_lookup:
            person = id_lookup[executor['person_id']]
            executor = {**executor, **person}
        else:
            errors.append(f'Executor person_id "{executor["person_id"]}" not found in will identities ({len(identities)} identities)')

    if testator.get('person_id'):
        if testator['person_id'] in id_lookup:
            person = id_lookup[testator['person_id']]
            testator = {**testator, **person}
        else:
            errors.append(f'Testator person_id "{testator["person_id"]}" not found in will identities ({len(identities)} identities)')

    # Sync deceased fields from testator
    if not probate.deceased_name and testator.get('full_name'):
        probate.deceased_name = testator['full_name']
        changed = True
    elif not probate.deceased_name and not testator.get('full_name'):
        errors.append('Cannot sync deceased name: will testator has no full_name')

    if not probate.deceased_nric and testator.get('nric_passport'):
        probate.deceased_nric = testator['nric_passport']
        changed = True
    elif not probate.deceased_nric and not testator.get('nric_passport'):
        errors.append('Cannot sync deceased NRIC: will testator has no nric_passport')

    if not probate.deceased_address and (testator.get('residential_address') or testator.get('address')):
        probate.deceased_address = testator.get('residential_address') or testator.get('address')
        changed = True

    # Sync applicant fields from executor
    if not probate.applicant_name and executor.get('full_name'):
        probate.applicant_name = executor['full_name']
        changed = True
    elif not probate.applicant_name and not executor.get('full_name'):
        errors.append('Cannot sync applicant name: will executor has no full_name')

    if not probate.applicant_nric and executor.get('nric_passport'):
        probate.applicant_nric = executor['nric_passport']
        changed = True
    elif not probate.applicant_nric and not executor.get('nric_passport'):
        errors.append('Cannot sync applicant NRIC: will executor has no nric_passport')

    if not probate.applicant_address and executor.get('address'):
        probate.applicant_address = executor['address']
        changed = True
    if not probate.applicant_relationship and executor.get('relationship'):
        probate.applicant_relationship = executor['relationship']
        changed = True

    if changed:
        db.session.commit()

    # Log errors for debugging
    if errors:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f'Probate {probate.id} sync from will {will_record.id}: {"; ".join(errors)}')

    return changed, errors


def _validate_probate_data(probate, will_record, recommendations):
    """Check for missing required data per form and return warnings.
    Always checks probate fields only — will data should already be synced."""
    warnings = {}  # form_code -> list of missing field descriptions
    rec_codes = {r['form_code'] for r in recommendations if r.get('recommended')}

    # Pre-compute common checks — probate fields are the single source of truth
    has_deceased_name = bool(probate.deceased_name)
    has_deceased_nric = bool(probate.deceased_nric)
    has_deceased_addr = bool(probate.deceased_address)
    has_applicant_name = bool(probate.applicant_name)
    has_applicant_nric = bool(probate.applicant_nric)
    has_applicant_addr = bool(probate.applicant_address)
    has_applicant_rel = bool(probate.applicant_relationship)
    has_court = bool(probate.court_location)
    has_court_state = bool(probate.court_state)
    has_firm_name = bool(probate.firm_name)
    has_firm_addr = bool(probate.firm_address)
    has_firm_ref = bool(probate.firm_reference)
    has_lawyer = bool(probate.lawyer_name)
    has_bar = bool(probate.lawyer_bar_number)
    has_dod = bool(probate.date_of_death)
    has_pod = bool(probate.place_of_death)
    has_death_cert = bool(probate.death_cert_number)

    assets = json.loads(probate.assets_data or '[]')
    props = [a for a in assets if a.get('asset_type') == 'property']
    bens = json.loads(probate.beneficiaries_data or '[]') if probate.beneficiaries_data else []

    # Fields used by all/most forms
    common_fields = []
    if not has_deceased_name: common_fields.append('Deceased Name (Step 2)')
    if not has_deceased_nric: common_fields.append('Deceased NRIC (Step 2)')
    if not has_applicant_name: common_fields.append('Applicant Name (Step 2)')
    if not has_applicant_nric: common_fields.append('Applicant NRIC (Step 2)')
    if not has_court: common_fields.append('Court Location (Step 3)')
    if not has_court_state: common_fields.append('Court State (Step 3)')

    for code in rec_codes:
        missing = list(common_fields)  # start with common missing fields

        # Death details
        if code in ('doc01', 'doc02', 'doc03'):
            if not has_dod: missing.append('Date of Death (Step 1)')
            if not has_pod: missing.append('Place of Death (Step 1)')
        if code == 'doc02':
            if not has_death_cert: missing.append('Death Certificate Number (Step 1)')

        # Firm details
        if code in ('doc01', 'doc08', 'form14a', 'form346'):
            if not has_firm_name: missing.append('Firm Name (Step 3)')
            if not has_firm_addr: missing.append('Firm Address (Step 3)')

        # Lawyer details
        if code in ('form14a', 'form346'):
            if not has_lawyer: missing.append('Lawyer Name (Step 3)')
            if not has_bar: missing.append('Bar Council Number (Step 3)')

        # Applicant details
        if code in ('doc01', 'doc02', 'doc03', 'form346'):
            if not has_applicant_addr: missing.append('Applicant Address (Step 2)')
            if not has_applicant_rel: missing.append('Applicant Relationship (Step 2)')

        # Deceased address
        if code in ('doc01', 'doc02'):
            if not has_deceased_addr: missing.append('Deceased Address (Step 2)')

        # Witnesses
        if code == 'doc04':
            if not probate.witness1_name: missing.append('Witness 1 Name (Step 4)')
            if not probate.witness1_nric: missing.append('Witness 1 NRIC (Step 4)')
            if not probate.witness1_address: missing.append('Witness 1 Address (Step 4)')
        if code == 'doc05':
            if not probate.witness2_name: missing.append('Witness 2 Name (Step 4)')
            if not probate.witness2_nric: missing.append('Witness 2 NRIC (Step 4)')
            if not probate.witness2_address: missing.append('Witness 2 Address (Step 4)')

        # Property details
        if code in ('form14a', 'form346'):
            if props:
                p0 = props[0]
                if not p0.get('title_number'): missing.append('Property Title Number (Step 6)')
                if not p0.get('lot_number'): missing.append('Property Lot Number (Step 6)')
                if not p0.get('mukim'): missing.append('Property Mukim (Step 6)')
            else:
                missing.append('No properties entered (Step 6)')

        # Beneficiaries
        if code == 'doc07':
            if not bens: missing.append('No beneficiaries entered (Step 5)')

        # Assets schedule
        if code == 'doc06':
            if not assets: missing.append('No assets entered (Step 6)')

        if missing:
            warnings[code] = missing
    return warnings


def _get_probate_context(probate_id):
    """Load probate app + will data for template context."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        return None, None, {}

    is_la = probate.application_type == 'la'
    will_record = db.session.get(Will, probate.will_id) if probate.will_id else None

    # Auto-sync will data into probate fields (single source of truth)
    sync_errors = []
    if will_record and not is_la:
        _changed, sync_errors = _sync_probate_from_will(probate, will_record)

    if is_la:
        # LA: deceased/applicant from probate fields
        testator = {
            'full_name': probate.deceased_name or '',
            'nric_passport': probate.deceased_nric or '',
            'residential_address': probate.deceased_address or '',
        }
        executor = {
            'full_name': probate.applicant_name or '',
            'nric_passport': probate.applicant_nric or '',
            'address': probate.applicant_address or '',
            'relationship': probate.applicant_relationship or '',
        }
        will_title = f'LA — {probate.deceased_name or "Unnamed"}'
        client_name = probate.applicant_name or ''
    else:
        # Probate: always read from probate fields (already synced from will above)
        testator = {
            'full_name': probate.deceased_name or '',
            'nric_passport': probate.deceased_nric or '',
            'residential_address': probate.deceased_address or '',
        }
        executor = {
            'full_name': probate.applicant_name or '',
            'nric_passport': probate.applicant_nric or '',
            'address': probate.applicant_address or '',
            'relationship': probate.applicant_relationship or '',
        }
        if will_record:
            will_title = will_record.title
            client_name = will_record.client.full_name if will_record.client else ''
        else:
            will_title = f'Probate — {probate.deceased_name or "Manual Entry"}'
            client_name = probate.applicant_name or ''

    # Determine which steps are complete (green tick) — require ALL essential fields
    # Step 1: Death details only
    step1_ok = all([
        probate.date_of_death,
        probate.place_of_death,
    ])
    # Step 2: Deceased & Executor details
    step2_ok = all([
        probate.deceased_name,
        probate.deceased_nric,
        probate.applicant_name,
        probate.applicant_nric,
    ])
    # Step 3: Court & Firm — case_number and firm_reference are OPTIONAL (assigned later)
    step3_ok = all([
        probate.court_location,
        probate.court_state,
        probate.firm_name,
        probate.firm_address,
        probate.lawyer_name,
        probate.lawyer_bar_number,
    ])
    # Step 4: At least one witness with name, NRIC, address
    step4_ok = all([
        probate.witness1_name,
        probate.witness1_nric,
        probate.witness1_address,
    ])
    # Step 5: At least one beneficiary
    _bens_data = json.loads(probate.beneficiaries_data or '[]') if probate.beneficiaries_data else []
    step5_ok = len(_bens_data) > 0 and all(
        b.get('full_name') or b.get('beneficiary_name') for b in _bens_data
    )
    # Step 6: At least one asset entered
    _assets_data = json.loads(probate.assets_data or '[]') if probate.assets_data else []
    step6_ok = len(_assets_data) > 0

    # Build completed steps set (non-sequential — each step independent)
    completed_steps = set()
    if step1_ok: completed_steps.add(1)
    if step2_ok: completed_steps.add(2)
    if step3_ok: completed_steps.add(3)
    if step4_ok: completed_steps.add(4)
    if step5_ok: completed_steps.add(5)
    if step6_ok: completed_steps.add(6)
    if probate.status in ('generated', 'pending_approval', 'approved', 'rejected'):
        completed_steps.add(7)

    all_steps_complete = all(s in completed_steps for s in range(1, 7))  # Steps 1-6 all done

    no_will = not is_la and not will_record  # Manual will upload (probate without linked will)

    return probate, will_record, {
        'probate': probate,
        'probate_id': probate_id,
        'is_la': is_la,
        'no_will': no_will,
        'will_title': will_title,
        'client_name': client_name,
        'testator': testator,
        'executor': type('Obj', (), executor) if executor else None,
        'completed_steps': completed_steps,
        'all_steps_complete': all_steps_complete,
        'sync_errors': sync_errors,
        # Keep max_completed_step for backward compatibility
        'max_completed_step': max(completed_steps) if completed_steps else 0,
    }


@app.route('/probate')
@login_required
def probate_list():
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    q = request.args.get('q', '').strip()
    base_q = ProbateApplication.query.filter(ProbateApplication.deleted_at.is_(None))
    if q:
        applications = base_q.filter(
            db.or_(
                ProbateApplication.deceased_name.ilike(f'%{q}%'),
                ProbateApplication.deceased_nric.ilike(f'%{q}%'),
                ProbateApplication.applicant_name.ilike(f'%{q}%'),
                ProbateApplication.case_number.ilike(f'%{q}%'),
            )
        ).order_by(ProbateApplication.created_at.desc()).all()
    else:
        applications = base_q.order_by(ProbateApplication.created_at.desc()).all()
    approved_wills = Will.query.filter_by(status='approved').filter(Will.deleted_at.is_(None)).order_by(Will.approved_at.desc()).all()
    return render_template('probate/list.html', applications=applications, approved_wills=approved_wills, search_query=q)


@app.route('/probate/<probate_id>/delete', methods=['POST'])
@login_required
def probate_delete(probate_id):
    """Soft-delete a probate application (recoverable for 30 days)."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('probate_list'))
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        flash('Application not found.', 'error')
        return redirect(url_for('probate_list'))
    deceased = probate.deceased_name or 'Unknown'
    probate.deleted_at = datetime.utcnow()
    db.session.commit()
    flash(f'Probate application for "{deceased}" moved to trash. It can be restored within 30 days.', 'success')
    return redirect(url_for('probate_list'))


@app.route('/probate/<probate_id>/restore', methods=['POST'])
@login_required
def probate_restore(probate_id):
    """Restore a soft-deleted probate application."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('trash_list'))
    probate = db.session.get(ProbateApplication, probate_id)
    if probate and probate.deleted_at:
        probate.deleted_at = None
        db.session.commit()
        flash(f'Probate application restored successfully.', 'success')
    return redirect(url_for('trash_list'))


@app.route('/probate/<probate_id>/permanent-delete', methods=['POST'])
@login_required
def probate_permanent_delete(probate_id):
    """Permanently delete a probate application (admin only, from trash)."""
    role = session.get('user_role')
    if role not in ('admin',):
        flash('Access denied.', 'error')
        return redirect(url_for('trash_list'))
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        flash('Application not found.', 'error')
        return redirect(url_for('trash_list'))
    # Delete generated form files from disk
    gen_forms = ProbateGeneratedForm.query.filter_by(probate_id=probate_id).all()
    for gf in gen_forms:
        if gf.file_path and os.path.exists(gf.file_path):
            try:
                os.remove(gf.file_path)
            except OSError:
                pass
    ProbateGeneratedForm.query.filter_by(probate_id=probate_id).delete()
    deceased = probate.deceased_name or 'Unknown'
    db.session.delete(probate)
    db.session.commit()
    flash(f'Probate application for "{deceased}" permanently deleted.', 'success')
    return redirect(url_for('trash_list'))


@app.route('/probate/new-la')
@app.route('/probate/new-probate')  # External will (no WillCraft will linked)
@login_required
def probate_new_la():
    """Create a new LA or external-will probate application."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    # Determine type from URL
    app_type = 'probate' if request.path.endswith('new-probate') else 'la'
    tenant = get_tenant()
    probate = ProbateApplication(
        application_type=app_type,
        filing_year=str(datetime.now().year),
        created_by=session.get('user_id'),
        firm_name=tenant.get('firm_name', 'Tetuan Alan Tan & Associates'),
        firm_address=tenant.get('firm_address', '24-01 & 24-02, Jalan Kempas Utama 2/4, Taman Kempas Utama, 81300 Johor Bahru, Johor'),
        firm_phone=tenant.get('firm_phone', '07-588 5979'),
        lawyer_name=tenant.get('lawyer_name', 'FAIZUL HANAFI BIN TOKIRAN'),
        lawyer_bar_number=tenant.get('lawyer_bar_number', 'BC/F/167'),
    )
    db.session.add(probate)
    db.session.commit()
    return redirect(f'/probate/{probate.id}/step/1')


@app.route('/probate/new/<will_id>')
@login_required
def probate_new(will_id):
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    will_record = db.session.get(Will, will_id)
    if not will_record:
        flash('Will not found.', 'error')
        return redirect(url_for('wills_list'))
    # Check if probate already exists for this will
    existing = ProbateApplication.query.filter_by(will_id=will_id).filter(ProbateApplication.deleted_at.is_(None)).first()
    if existing:
        return redirect(f'/probate/{existing.id}/step/1')
    # Create new probate application
    tenant = get_tenant()
    probate = ProbateApplication(
        will_id=will_id,
        client_id=will_record.client_id,
        filing_year=str(datetime.now().year),
        created_by=session.get('user_id'),
        firm_name=tenant.get('firm_name', 'Tetuan Alan Tan & Associates'),
        firm_address=tenant.get('firm_address', '24-01 & 24-02, Jalan Kempas Utama 2/4, Taman Kempas Utama, 81300 Johor Bahru, Johor'),
        firm_phone=tenant.get('firm_phone', '07-588 5979'),
        lawyer_name=tenant.get('lawyer_name', 'FAIZUL HANAFI BIN TOKIRAN'),
        lawyer_bar_number=tenant.get('lawyer_bar_number', 'BC/F/167'),
    )
    db.session.add(probate)
    db.session.commit()
    # Auto-populate deceased/applicant from linked will
    _changed, _errors = _sync_probate_from_will(probate, will_record)
    return redirect(f'/probate/{probate.id}/step/1')


def _classify_asset(description):
    """Classify an asset description into: property, bank, vehicle, investment, other."""
    if not description:
        return 'other'
    desc = description.lower()
    # Property keywords
    if any(kw in desc for kw in ['land', 'house', 'apartment', 'flat', 'condo', 'bungalow',
            'terrace', 'semi-d', 'lot', 'title', 'geran', 'hakmilik', 'strata',
            'property', 'rumah', 'tanah', 'building', 'premises', 'unit',
            'jalan', 'lorong', 'taman', 'kampung', 'no.', 'no ', 'address']):
        return 'property'
    # Bank / financial account keywords
    if any(kw in desc for kw in ['bank', 'account', 'savings', 'current account', 'fixed deposit',
            'akaun', 'maybank', 'cimb', 'rhb', 'hong leong', 'public bank',
            'ambank', 'bsn', 'affin', 'hsbc', 'ocbc', 'uob', 'standard chartered',
            'alliance', 'kwsp', 'epf', 'tabung', 'amanah', 'asb', 'asp']):
        return 'bank'
    # Vehicle keywords
    if any(kw in desc for kw in ['car', 'vehicle', 'motor', 'kereta', 'toyota', 'honda',
            'proton', 'perodua', 'mercedes', 'bmw', 'nissan', 'mazda',
            'registration', 'plate number', 'nombor pendaftaran']):
        return 'vehicle'
    # Investment keywords
    if any(kw in desc for kw in ['share', 'stock', 'unit trust', 'bond', 'investment',
            'dividend', 'securities', 'saham', 'bursa', 'insurance', 'policy',
            'insurans', 'takaful', 'mutual fund']):
        return 'investment'
    return 'other'


@app.route('/probate/<probate_id>/save-ocr-data', methods=['POST'])
@login_required
def probate_save_ocr_data(probate_id):
    """Save OCR-extracted data (witnesses, beneficiaries, assets) for later steps."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        return jsonify(ok=False, error='Not found'), 404
    data = request.get_json(force=True)
    # Merge into form_data_json
    existing = json.loads(probate.form_data_json or '{}')
    existing['ocr_extracted'] = data
    probate.form_data_json = json.dumps(existing)

    # Auto-fill deceased fields (overwrite from scan — user explicitly confirmed)
    if data.get('deceased_name'):
        probate.deceased_name = data['deceased_name']
    if data.get('deceased_nric'):
        probate.deceased_nric = data['deceased_nric']
    if data.get('deceased_address'):
        probate.deceased_address = data['deceased_address']

    # Auto-fill executor/applicant fields
    if data.get('applicant_name'):
        probate.applicant_name = data['applicant_name']
    if data.get('applicant_nric'):
        probate.applicant_nric = data['applicant_nric']
    if data.get('applicant_address'):
        probate.applicant_address = data['applicant_address']
    if data.get('applicant_relationship'):
        probate.applicant_relationship = data['applicant_relationship']

    # Auto-fill witness fields
    if data.get('witness1_name'):
        probate.witness1_name = data['witness1_name']
    if data.get('witness1_nric'):
        probate.witness1_nric = data['witness1_nric']
    if data.get('witness1_address'):
        probate.witness1_address = data['witness1_address']
    if data.get('witness2_name'):
        probate.witness2_name = data['witness2_name']
    if data.get('witness2_nric'):
        probate.witness2_nric = data['witness2_nric']
    if data.get('witness2_address'):
        probate.witness2_address = data['witness2_address']

    # Auto-fill beneficiaries
    bens = []
    i = 0
    while data.get(f'beneficiary_{i}_name'):
        bens.append({
            'full_name': data.get(f'beneficiary_{i}_name', ''),
            'nric_passport': data.get(f'beneficiary_{i}_nric', ''),
            'relationship': data.get(f'beneficiary_{i}_relationship', ''),
            'address': data.get(f'beneficiary_{i}_address', ''),
        })
        i += 1
    if bens:
        probate.beneficiaries_data = json.dumps(bens)

    # Auto-fill assets
    assets = []
    i = 0
    while data.get(f'asset_{i}'):
        desc = data[f'asset_{i}']
        # Try to parse JSON if the asset was stored as JSON string
        try:
            asset_obj = json.loads(desc) if desc.startswith('{') else None
        except (json.JSONDecodeError, AttributeError):
            asset_obj = None
        if asset_obj and isinstance(asset_obj, dict):
            # Normalize: OCR uses 'type', DB uses 'asset_type'
            if asset_obj.get('type') and not asset_obj.get('asset_type'):
                asset_obj['asset_type'] = asset_obj.pop('type')
            if not asset_obj.get('asset_type'):
                asset_obj['asset_type'] = _classify_asset(asset_obj.get('description', ''))
            if 'estimated_value' not in asset_obj:
                asset_obj['estimated_value'] = ''
            assets.append(asset_obj)
        else:
            assets.append({'asset_type': _classify_asset(desc), 'description': desc, 'estimated_value': ''})
        i += 1
    if assets:
        probate.assets_data = json.dumps(assets)

    probate.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(ok=True)


@app.route('/probate/<probate_id>/step/1', methods=['GET', 'POST'])
@login_required
def probate_step1(probate_id):
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        probate.death_cert_number = request.form.get('death_cert_number', '').strip()
        probate.date_of_death = request.form.get('date_of_death', '').strip()
        probate.time_of_death = request.form.get('time_of_death', '').strip()
        probate.place_of_death = request.form.get('place_of_death', '').strip()
        probate.estate_value_estimate = request.form.get('estate_value_estimate', '').strip()
        doc_id = request.form.get('death_cert_document_id', '').strip()
        if doc_id:
            probate.death_cert_document_id = doc_id
        # Deceased/applicant fields moved to step 2
        # Will document upload (for LA with external will)
        will_doc_id = request.form.get('will_document_id', '').strip()
        if will_doc_id:
            probate.will_document_id = will_doc_id
        # Store extracted will data for auto-populating later steps
        will_extracted = request.form.get('will_extracted_data', '').strip()
        if will_extracted:
            try:
                ext = json.loads(will_extracted)
                # Pre-populate witnesses from extracted will data
                witnesses = ext.get('witnesses', [])
                if witnesses and not probate.witness1_name:
                    if len(witnesses) >= 1:
                        probate.witness1_name = witnesses[0].get('full_name', '')
                        probate.witness1_nric = witnesses[0].get('nric_number', '')
                        if witnesses[0].get('address'):
                            probate.witness1_address = witnesses[0]['address']
                    if len(witnesses) >= 2:
                        probate.witness2_name = witnesses[1].get('full_name', '')
                        probate.witness2_nric = witnesses[1].get('nric_number', '')
                        if witnesses[1].get('address'):
                            probate.witness2_address = witnesses[1]['address']
                # Pre-populate beneficiaries from extracted will data
                bens = ext.get('beneficiaries', [])
                existing_bens = json.loads(probate.beneficiaries_data or '[]')
                if bens and not existing_bens:
                    ben_list = []
                    for b in bens:
                        ben_list.append({
                            'full_name': b.get('full_name', ''),
                            'nric_passport': b.get('nric_number', ''),
                            'relationship': b.get('relationship', ''),
                            'address': b.get('address', ''),
                        })
                    probate.beneficiaries_data = json.dumps(ben_list)
                # Pre-populate assets from extracted will data
                assets = ext.get('assets', [])
                existing_assets = json.loads(probate.assets_data or '[]')
                if assets and not existing_assets:
                    asset_list = []
                    for a in assets:
                        atype = a.get('type', 'other')
                        asset_list.append({
                            'asset_type': atype if atype in ('property', 'bank', 'vehicle') else 'other',
                            'description': a.get('description', ''),
                        })
                    probate.assets_data = json.dumps(asset_list)
            except (json.JSONDecodeError, Exception):
                pass  # Silently skip if parsing fails
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        return redirect(f'/probate/{probate_id}/step/2')

    ctx['probate_step'] = 1
    # Look up uploaded document objects for display
    if probate.death_cert_document_id:
        ctx['death_cert_doc'] = db.session.get(Document, probate.death_cert_document_id)
    if probate.will_document_id:
        ctx['will_doc'] = db.session.get(Document, probate.will_document_id)
    return render_template('probate/step1_death.html', **ctx)


@app.route('/probate/<probate_id>/step/2', methods=['GET', 'POST'])
@login_required
def probate_step2(probate_id):
    """Step 2: Deceased & Executor details."""
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        probate.deceased_name = request.form.get('deceased_name', '').strip()
        probate.deceased_nric = request.form.get('deceased_nric', '').strip()
        probate.deceased_address = request.form.get('deceased_address', '').strip()
        probate.applicant_name = request.form.get('applicant_name', '').strip()
        probate.applicant_nric = request.form.get('applicant_nric', '').strip()
        probate.applicant_address = request.form.get('applicant_address', '').strip()
        probate.applicant_relationship = request.form.get('applicant_relationship', '').strip()
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        return redirect(f'/probate/{probate_id}/step/3')

    ctx['probate_step'] = 2
    return render_template('probate/step2_executor.html', **ctx)


@app.route('/probate/<probate_id>/step/3', methods=['GET', 'POST'])
@login_required
def probate_step3(probate_id):
    """Step 3: Court & Law Firm details."""
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        probate.court_location = request.form.get('court_location', '').strip()
        probate.court_state = request.form.get('court_state', '').strip()
        probate.case_number = request.form.get('case_number', '').strip()
        probate.filing_year = request.form.get('filing_year', '').strip()
        probate.firm_name = request.form.get('firm_name', '').strip()
        probate.firm_address = request.form.get('firm_address', '').strip()
        probate.firm_phone = request.form.get('firm_phone', '').strip()
        probate.firm_fax = request.form.get('firm_fax', '').strip()
        probate.firm_reference = request.form.get('firm_reference', '').strip()
        probate.lawyer_name = request.form.get('lawyer_name', '').strip()
        probate.lawyer_nric = request.form.get('lawyer_nric', '').strip()
        probate.lawyer_bar_number = request.form.get('lawyer_bar_number', '').strip()
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        return redirect(f'/probate/{probate_id}/step/4')

    ctx['probate_step'] = 3
    ctx['courts'] = MALAYSIAN_COURTS
    ctx['states'] = MALAYSIAN_STATES
    ctx['current_year'] = str(datetime.now().year)
    return render_template('probate/step2_court.html', **ctx)


@app.route('/probate/<probate_id>/step/4', methods=['GET', 'POST'])
@login_required
def probate_step4(probate_id):
    """Step 4: Witnesses."""
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        probate.witness1_name = request.form.get('witness1_name', '').strip()
        probate.witness1_nric = request.form.get('witness1_nric', '').strip()
        probate.witness1_address = request.form.get('witness1_address', '').strip()
        probate.witness2_name = request.form.get('witness2_name', '').strip()
        probate.witness2_nric = request.form.get('witness2_nric', '').strip()
        probate.witness2_address = request.form.get('witness2_address', '').strip()
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        return redirect(f'/probate/{probate_id}/step/5')

    ctx['probate_step'] = 4
    return render_template('probate/step3_witnesses.html', **ctx)


@app.route('/probate/<probate_id>/step/5', methods=['GET', 'POST'])
@login_required
def probate_step5(probate_id):
    """Step 5: Beneficiaries list."""
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        bens_json = request.form.get('beneficiaries_json', '[]')
        try:
            bens = json.loads(bens_json)
        except json.JSONDecodeError:
            bens = []
        probate.beneficiaries_data = json.dumps(bens)
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        if request.headers.get('X-Save-Only'):
            return jsonify(ok=True)
        return redirect(f'/probate/{probate_id}/step/6')

    # Pre-populate from will data if beneficiaries_data is empty
    existing_bens = json.loads(probate.beneficiaries_data or '[]')
    if not existing_bens and will_record:
        will_bens = json.loads(will_record.step4_data or '[]')
        # Build identity lookup from will's person registry for address extraction
        identity_lookup = {}
        will_identities = json.loads(will_record.identities_data or '[]')
        for person in will_identities:
            identity_lookup[person.get('id', '')] = person
        for b in will_bens:
            # Look up address from will's identity registry via person_id
            addr = b.get('address', '')
            if not addr and b.get('person_id'):
                person = identity_lookup.get(b['person_id'], {})
                addr = person.get('address', '')
            existing_bens.append({
                'full_name': b.get('full_name', b.get('beneficiary_name', '')),
                'nric_passport': b.get('nric_passport_birthcert', b.get('nric_passport', '')),
                'relationship': b.get('relationship', ''),
                'address': addr,
            })

    ctx['probate_step'] = 5
    ctx['beneficiaries_json'] = json.dumps(existing_bens)
    return render_template('probate/step4_beneficiaries.html', **ctx)


@app.route('/probate/<probate_id>/step/6', methods=['GET', 'POST'])
@login_required
def probate_step6(probate_id):
    """Step 6: Assets & Liabilities schedule."""
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    if request.method == 'POST':
        assets_json = request.form.get('assets_json', '[]')
        try:
            assets = json.loads(assets_json)
        except json.JSONDecodeError:
            assets = []
        probate.assets_data = json.dumps(assets)
        probate.updated_at = datetime.utcnow()
        db.session.commit()
        if request.headers.get('X-Save-Only'):
            return jsonify(ok=True)
        return redirect(f'/probate/{probate_id}/step/7')

    # Pre-populate from will gifts if assets_data is empty and will exists
    existing_assets = json.loads(probate.assets_data or '[]')
    if not existing_assets and will_record:
        gifts = json.loads(will_record.step5_data or '[]')
        for g in gifts:
            if g.get('gift_type') == 'property':
                details = g.get('property_details', {})
                existing_assets.append({
                    'asset_type': 'property',
                    'description': details.get('address', g.get('description', '')),
                    'title_number': details.get('title_number', ''),
                    'lot_number': details.get('lot_number', ''),
                    'mukim': details.get('mukim', ''),
                    'value': '',
                })
            elif g.get('gift_type') == 'financial':
                fin = g.get('financial_details', {})
                existing_assets.append({
                    'asset_type': 'bank',
                    'bank_name': fin.get('institution', ''),
                    'account_number': fin.get('account_number', ''),
                    'value': '',
                })

    # Build exhibit prefix
    exec_data = ctx.get('executor')
    exec_name = exec_data.full_name if exec_data and hasattr(exec_data, 'full_name') else ''
    exhibit_prefix = ''.join(w[0] for w in exec_name.split() if w) if exec_name else 'APP'

    ctx['probate_step'] = 6
    ctx['assets_json'] = json.dumps(existing_assets)
    ctx['exhibit_prefix'] = exhibit_prefix
    return render_template('probate/step5_assets.html', **ctx)


@app.route('/probate/<probate_id>/step/7', methods=['GET'])
@login_required
def probate_step7(probate_id):
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    from documents.probate_generator import recommend_forms, FORM_FIELDS
    recommendations = recommend_forms(will_record, probate)

    # Build actual values lookup for form field display
    exec_data = ctx.get('executor')
    testator = ctx.get('testator') or {}
    field_values = {
        'Deceased name & NRIC': f"{testator.get('full_name', '')} ({testator.get('nric_passport', '')})" if testator.get('full_name') else '',
        'Deceased name': testator.get('full_name', ''),
        'Deceased address': testator.get('residential_address', ''),
        'Date of death': probate.date_of_death or '',
        'Time of death': probate.time_of_death or '',
        'Place of death': probate.place_of_death or '',
        'Death certificate number': probate.death_cert_number or '',
        'Applicant (Executor) name & NRIC': f"{exec_data.full_name} ({exec_data.nric_passport})" if exec_data and exec_data.full_name else '',
        'Applicant name & NRIC': f"{exec_data.full_name} ({exec_data.nric_passport})" if exec_data and exec_data.full_name else '',
        'Applicant address': (exec_data.address if exec_data and hasattr(exec_data, 'address') else '') or '',
        'Applicant relationship': (exec_data.relationship if exec_data and hasattr(exec_data, 'relationship') else '') or '',
        'Court location': probate.court_location or '',
        'Case number': probate.case_number or 'Not set',
        'Firm name & address': f"{probate.firm_name or ''}, {probate.firm_address or ''}" if probate.firm_name else '',
        'Firm phone & fax': f"Tel: {probate.firm_phone or ''}, Fax: {probate.firm_fax or ''}",
        'Firm reference': probate.firm_reference or '',
        'Witness 1 name & NRIC': f"{probate.witness1_name or ''} ({probate.witness1_nric or ''})" if probate.witness1_name else '',
        'Witness 1 address': probate.witness1_address or '',
        'Witness 2 name & NRIC': f"{probate.witness2_name or ''} ({probate.witness2_nric or ''})" if probate.witness2_name else '',
        'Witness 2 address': probate.witness2_address or '',
        'Estate value': probate.estate_value_estimate or '',
        'Exhibit references': 'Auto-generated',
        'Lawyer name': probate.lawyer_name or '',
        'Lawyer bar council number': probate.lawyer_bar_number or '',
        'Filing year': probate.filing_year or '',
    }
    # Assets summary values
    _assets = json.loads(probate.assets_data or '[]')
    _props = [a for a in _assets if a.get('asset_type') == 'property']
    _banks = [a for a in _assets if a.get('asset_type') == 'bank']
    _vehicles = [a for a in _assets if a.get('asset_type') == 'vehicle']
    _others = [a for a in _assets if a.get('asset_type') == 'other']
    _investments = [a for a in _assets if a.get('asset_type') == 'investment']
    _liabs = [a for a in _assets if a.get('asset_type') == 'liability']
    field_values['Properties (title, lot, mukim, address)'] = f'{len(_props)} properties' if _props else 'N/A'
    field_values['Bank accounts (bank, account no., value)'] = f'{len(_banks)} accounts' if _banks else 'N/A'
    field_values['Vehicles (desc, reg no., engine, chassis)'] = f'{len(_vehicles)} vehicles' if _vehicles else 'N/A'
    field_values['Investment accounts (CDS, unit trust, etc.)'] = f'{len(_investments)} accounts' if _investments else 'N/A'
    field_values['Other assets (description, value)'] = f'{len(_others)} items' if _others else 'N/A'
    field_values['Liabilities (description, value)'] = f'{len(_liabs)} items' if _liabs else 'N/A'
    # Beneficiaries from probate (single source of truth)
    _bens = json.loads(probate.beneficiaries_data or '[]') if probate.beneficiaries_data and probate.beneficiaries_data != '[]' else []
    field_values['Beneficiary names & NRIC'] = ', '.join(b.get('full_name', b.get('beneficiary_name', '')) for b in _bens[:5]) if _bens else ''
    field_values['Beneficiary relationships'] = ', '.join(b.get('relationship', '') for b in _bens[:5]) if _bens else ''
    if _props:
        p0 = _props[0]
        field_values['Property title number'] = p0.get('title_number', '')
        field_values['Property lot number'] = p0.get('lot_number', '')
        field_values['Property mukim'] = p0.get('mukim', '')

    # Merge with template info and field mapping
    templates = ProbateFormTemplate.query.order_by(ProbateFormTemplate.sort_order).all()
    tpl_map = {t.form_code: t for t in templates}
    for rec in recommendations:
        tpl = tpl_map.get(rec['form_code'])
        if tpl:
            rec['form_name'] = tpl.form_name
            rec['form_name_malay'] = tpl.form_name_malay
            rec['description'] = tpl.description
        ff = FORM_FIELDS.get(rec['form_code'])
        if ff:
            rec['fields'] = [(name, source, field_values.get(name, '')) for name, source in ff['fields']]
        else:
            rec['fields'] = []

    # Check for previously generated forms
    generated_forms = ProbateGeneratedForm.query.filter_by(probate_id=probate_id).all()
    gen_list = []
    import re as _re
    for gf in generated_forms:
        tpl = tpl_map.get(gf.form_code)
        # Scan for unfilled placeholders in generated file
        missing_fields = []
        if gf.file_path and os.path.exists(gf.file_path):
            try:
                from docx import Document as _DocxDoc
                _doc = _DocxDoc(gf.file_path)
                _seen = set()
                for _p in _doc.paragraphs:
                    for _m in _re.findall(r'\{\{(\w+)\}\}', _p.text):
                        if _m not in _seen:
                            _seen.add(_m)
                            missing_fields.append(_m.replace('_', ' ').title())
                for _t in _doc.tables:
                    for _r in _t.rows:
                        for _c in _r.cells:
                            for _m in _re.findall(r'\{\{(\w+)\}\}', _c.text):
                                if _m not in _seen:
                                    _seen.add(_m)
                                    missing_fields.append(_m.replace('_', ' ').title())
            except Exception:
                pass
        gen_list.append({
            'form_code': gf.form_code,
            'form_name': tpl.form_name if tpl else gf.form_code,
            'file_path': gf.file_path,
            'missing_fields': missing_fields,
        })

    # Build exhibit prefix from applicant initials
    exec_data = ctx.get('executor')
    exec_name = exec_data.full_name if exec_data and hasattr(exec_data, 'full_name') else ''
    exhibit_prefix = ''.join(w[0] for w in exec_name.split() if w) if exec_name else 'APP'

    # Validation: check for missing required info per form
    validation_warnings = _validate_probate_data(probate, will_record, recommendations)

    # Load invoices & receipts for this probate
    receipt_docs = Document.query.filter(
        Document.description.like(f'probate:{probate_id}|%'),
        Document.category.in_(['probate_invoice', 'probate_receipt'])
    ).order_by(Document.created_at.desc()).all()
    receipts = [{
        'id': d.id,
        'filename': d.original_filename,
        'category': d.category.replace('probate_', ''),
        'description': (d.description or '').split('|', 1)[-1],
        'file_size': d.file_size,
        'created_at': d.created_at.strftime('%d %b %Y, %I:%M %p') if d.created_at else '',
    } for d in receipt_docs]

    # Build filing checklist — check actual data completeness for form generation
    form_data = json.loads(probate.form_data_json or '{}')
    manual_checks = form_data.get('filing_checklist', {})
    assets_list = json.loads(probate.assets_data or '[]')
    has_property = any(a.get('asset_type') == 'property' for a in assets_list)
    property_assets = [a for a in assets_list if a.get('asset_type') == 'property']
    has_property_docs = all(a.get('_doc_id') for a in property_assets) if property_assets else False
    exec_nric = (exec_data.nric_passport if exec_data and hasattr(exec_data, 'nric_passport') else '') or ''
    testator = ctx.get('testator')

    # Death cert info complete? (cert number optional — nice to have)
    death_info_ok = bool(probate.date_of_death and probate.place_of_death)
    death_missing = []
    if not probate.date_of_death:
        death_missing.append('Date of death')
    if not probate.place_of_death:
        death_missing.append('Place of death')
    death_warnings = []
    if not probate.death_cert_number:
        death_warnings.append('Death cert number (optional)')

    # Assets info complete?
    assets_ok = len(assets_list) > 0
    assets_missing = [] if assets_ok else ['No assets entered']

    # Will info complete?
    will_ok = bool(will_record)
    will_missing = [] if will_ok else ['No approved will linked']

    # Beneficiary info from probate (single source of truth)
    beneficiaries = json.loads(probate.beneficiaries_data or '[]') if probate.beneficiaries_data and probate.beneficiaries_data != '[]' else []
    ben_ok = len(beneficiaries) > 0
    ben_missing = [] if ben_ok else ['No beneficiaries entered']

    # Executor info?
    exec_ok = bool(exec_nric and exec_data and exec_data.full_name)
    exec_missing = []
    if not exec_data or not exec_data.full_name:
        exec_missing.append('Executor name')
    if not exec_nric:
        exec_missing.append('Executor NRIC')

    # Property title info? Manual key-in (title_number) is sufficient
    prop_missing = []
    if has_property:
        for i, p in enumerate(property_assets):
            if not p.get('title_number'):
                prop_missing.append(f'Property {i+1}: Title number missing')
            if not p.get('description'):
                prop_missing.append(f'Property {i+1}: Description missing')
    prop_ok = has_property and not prop_missing

    # Build detail data for each checklist item
    _testator = ctx.get('testator') or {}
    death_details = [
        ('Deceased Name', _testator.get('full_name', '')),
        ('NRIC', _testator.get('nric_passport', '')),
        ('Date of Death', probate.date_of_death or ''),
        ('Time of Death', probate.time_of_death or ''),
        ('Place of Death', probate.place_of_death or ''),
        ('Death Cert No.', probate.death_cert_number or ''),
        ('Estate Value', probate.estate_value_estimate or ''),
    ]

    assets_details = []
    for a in assets_list:
        atype = a.get('asset_type', '')
        if atype == 'property':
            assets_details.append(('Property', f"{a.get('description', '')} — Title: {a.get('title_number', 'N/A')}"))
        elif atype == 'bank':
            assets_details.append(('Bank', f"{a.get('bank_name', '')} — Acc: {a.get('account_number', '')} — RM {a.get('value', '')}"))
        elif atype == 'vehicle':
            assets_details.append(('Vehicle', f"{a.get('description', '')} — {a.get('reg_number', '')}"))
        elif atype == 'other':
            assets_details.append(('Other', f"{a.get('description', '')} — RM {a.get('value', '')}"))
        elif atype == 'liability':
            assets_details.append(('Liability', f"{a.get('description', '')} — RM {a.get('value', '')}"))
    if not assets_details:
        assets_details.append(('', 'No assets entered'))

    will_details = []
    if will_record:
        will_details.append(('Will Title', will_record.title or ''))
        will_details.append(('Status', will_record.status or ''))
        will_details.append(('Approved', will_record.approved_at.strftime('%d %b %Y') if will_record.approved_at else 'N/A'))
    else:
        will_details.append(('', 'No approved will linked'))

    ben_details = []
    for b in beneficiaries:
        bname = b.get('full_name', b.get('beneficiary_name', b.get('name', '')))
        bnric = b.get('nric_passport_birthcert', b.get('nric_passport', ''))
        brel = b.get('relationship', '')
        detail = f"{brel} — NRIC: {bnric}" if bnric else brel
        ben_details.append((bname, detail))
    if not ben_details:
        ben_details.append(('', 'No beneficiaries in will'))

    exec_details = [
        ('Name', exec_data.full_name if exec_data and hasattr(exec_data, 'full_name') else ''),
        ('NRIC', exec_nric),
        ('Relationship', exec_data.relationship if exec_data and hasattr(exec_data, 'relationship') else ''),
        ('Address', exec_data.address if exec_data and hasattr(exec_data, 'address') else ''),
    ]

    prop_details = []
    for p in property_assets:
        prop_details.append(('Description', p.get('description', '')))
        prop_details.append(('Title No.', p.get('title_number', '')))
        prop_details.append(('Lot No.', p.get('lot_number', '')))
        prop_details.append(('Mukim', p.get('mukim', '')))
    if not prop_details:
        prop_details.append(('', 'No property in estate'))

    filing_checklist = [
        {'key': 'death_cert', 'label': '<strong>Death Certificate</strong> — date, place, cert number',
         'exhibit': f'{exhibit_prefix}-1',
         'complete': death_info_ok,
         'missing': death_missing,
         'warnings': death_warnings,
         'details': death_details,
         'checked': manual_checks.get('death_cert', death_info_ok)},
        {'key': 'assets_schedule', 'label': '<strong>Schedule of Assets &amp; Liabilities</strong>',
         'exhibit': f'{exhibit_prefix}-2',
         'complete': assets_ok,
         'missing': assets_missing,
         'details': assets_details,
         'checked': manual_checks.get('assets_schedule', assets_ok)},
        {'key': 'original_will', 'label': '<strong>Original Will</strong> (certified true copy)',
         'exhibit': f'{exhibit_prefix}-3',
         'complete': will_ok,
         'missing': will_missing,
         'details': will_details,
         'checked': manual_checks.get('original_will', will_ok)},
        {'key': 'beneficiary_list', 'label': '<strong>Beneficiary List</strong>',
         'exhibit': f'{exhibit_prefix}-4',
         'complete': ben_ok,
         'missing': ben_missing,
         'details': ben_details,
         'checked': manual_checks.get('beneficiary_list', ben_ok)},
        {'key': 'executor_nric', 'label': 'Executor NRIC &amp; details',
         'exhibit': None,
         'complete': exec_ok,
         'missing': exec_missing,
         'details': exec_details,
         'checked': manual_checks.get('executor_nric', exec_ok)},
        {'key': 'property_titles', 'label': '<strong>Property title documents</strong> (Hakmilik)',
         'exhibit': None, 'conditional': True, 'condition_met': has_property,
         'complete': prop_ok,
         'missing': prop_missing,
         'details': prop_details,
         'checked': manual_checks.get('property_titles', prop_ok)},
    ]

    ctx['probate_step'] = 7
    ctx['recommendations'] = recommendations
    ctx['generated_forms'] = gen_list
    ctx['exhibit_prefix'] = exhibit_prefix
    ctx['validation_warnings'] = validation_warnings
    ctx['receipts'] = receipts
    ctx['filing_checklist'] = filing_checklist
    ctx['has_property'] = has_property
    ctx['beneficiaries'] = beneficiaries
    return render_template('probate/step6_review.html', **ctx)


@app.route('/probate/<probate_id>/generate', methods=['POST'])
@login_required
def probate_generate(probate_id):
    probate, will_record, ctx = _get_probate_context(probate_id)
    if not probate:
        flash('Probate application not found.', 'error')
        return redirect(url_for('probate_list'))

    # Block generation if steps are incomplete (approver/admin can override)
    role = session.get('user_role')
    if not ctx.get('all_steps_complete') and role not in ('admin', 'approver'):
        flash('Please complete all steps (1-6) before generating forms.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')

    selected_codes = request.form.getlist('forms')
    if not selected_codes:
        flash('Please select at least one form to generate.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')

    # Build template paths map
    templates = ProbateFormTemplate.query.all()
    tpl_map = {t.form_code: os.path.join(os.path.dirname(__file__), t.file_path) for t in templates}
    tpl_name_map = {t.form_code: t.form_name for t in templates}

    # Output directory
    if probate.client_id:
        client = db.session.get(Client, probate.client_id)
        folder = client.folder_name if client else 'unknown'
    else:
        # LA without client — use probate ID as folder
        folder = f'la_{probate.id[:8]}'
    output_dir = os.path.join(DATA_DIR, 'clients', folder, 'probate')
    os.makedirs(output_dir, exist_ok=True)

    from documents.probate_generator import generate_probate_forms
    results = generate_probate_forms(probate, will_record, selected_codes, tpl_map, output_dir)

    # Delete old generated forms for this probate
    ProbateGeneratedForm.query.filter_by(probate_id=probate_id).delete()

    # Save generated form records
    for r in results:
        gf = ProbateGeneratedForm(
            probate_id=probate_id,
            form_code=r['form_code'],
            form_name=tpl_name_map.get(r['form_code'], r['form_code']),
            file_path=r['file_path'],
        )
        db.session.add(gf)

    probate.status = 'generated'
    probate.selected_forms = json.dumps(selected_codes)
    probate.updated_at = datetime.utcnow()
    db.session.commit()

    flash(f'Successfully generated {len(results)} probate form(s).', 'success')
    return redirect(f'/probate/{probate_id}/step/7')


@app.route('/probate/<probate_id>/submit-approval', methods=['POST'])
@login_required
def probate_submit_approval(probate_id):
    """Submit generated forms for approver review."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        flash('Application not found.', 'error')
        return redirect(url_for('probate_list'))
    if probate.status not in ('generated', 'rejected'):
        flash('Forms must be generated before submitting for approval.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')
    probate.status = 'pending_approval'
    probate.submitted_by = session.get('user_id')
    probate.submitted_at = datetime.utcnow()
    probate.approval_notes = None  # Clear previous rejection notes
    db.session.commit()
    flash('Forms submitted for approval.', 'success')
    return redirect(f'/probate/{probate_id}/step/7')


@app.route('/probate/<probate_id>/approve', methods=['POST'])
@login_required
def probate_approve(probate_id):
    """Approve probate forms (approver only)."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('probate_list'))
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        flash('Application not found.', 'error')
        return redirect(url_for('probate_list'))
    notes = request.form.get('approval_notes', '').strip()
    probate.status = 'approved'
    probate.approved_by = session.get('user_id')
    probate.approved_at = datetime.utcnow()
    probate.approval_notes = notes or None
    db.session.commit()
    flash(f'Probate forms for {probate.deceased_name or "estate"} approved.', 'success')
    return redirect(f'/probate/{probate_id}/step/7')


@app.route('/probate/<probate_id>/reject', methods=['POST'])
@login_required
def probate_reject(probate_id):
    """Request changes on probate forms (approver only)."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        flash('Access denied.', 'error')
        return redirect(url_for('probate_list'))
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        flash('Application not found.', 'error')
        return redirect(url_for('probate_list'))
    notes = request.form.get('approval_notes', '').strip()
    probate.status = 'rejected'
    probate.approved_by = session.get('user_id')
    probate.approved_at = datetime.utcnow()
    probate.approval_notes = notes or 'Changes requested'
    db.session.commit()
    flash(f'Changes requested for {probate.deceased_name or "estate"}.', 'info')
    return redirect(f'/probate/{probate_id}/step/7')


@app.route('/probate/<probate_id>/preview/<form_code>')
@login_required
def probate_preview(probate_id, form_code):
    """Serve generated form as inline PDF for browser preview."""
    gf = ProbateGeneratedForm.query.filter_by(probate_id=probate_id, form_code=form_code).first()
    if not gf or not os.path.exists(gf.file_path):
        flash('Form not found.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')
    from documents.probate_generator import convert_to_pdf
    import shutil
    tmp_dir = tempfile.mkdtemp()
    tmp_copy = os.path.join(tmp_dir, os.path.basename(gf.file_path))
    shutil.copy2(gf.file_path, tmp_copy)
    pdf_path = convert_to_pdf(tmp_copy)
    if pdf_path and os.path.exists(pdf_path):
        resp = send_file(pdf_path, mimetype='application/pdf',
                         download_name=f'{gf.form_name or form_code}.pdf')
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        return resp
    flash('PDF conversion failed.', 'error')
    return redirect(f'/probate/{probate_id}/step/7')


@app.route('/probate/<probate_id>/form-html/<form_code>')
@login_required
def probate_form_html(probate_id, form_code):
    """Convert generated DOCX to editable HTML."""
    gf = ProbateGeneratedForm.query.filter_by(probate_id=probate_id, form_code=form_code).first()
    if not gf or not os.path.exists(gf.file_path):
        return '<p>Form not found</p>', 404
    import shutil, subprocess
    tmp_dir = tempfile.mkdtemp()
    tmp_copy = os.path.join(tmp_dir, os.path.basename(gf.file_path))
    shutil.copy2(gf.file_path, tmp_copy)
    # Convert DOCX to HTML using LibreOffice
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'html', tmp_copy, '--outdir', tmp_dir],
        capture_output=True, text=True, timeout=30
    )
    html_file = os.path.splitext(tmp_copy)[0] + '.html'
    if not os.path.exists(html_file):
        return '<p>Conversion failed</p>', 500
    with open(html_file, 'r', encoding='utf-8', errors='replace') as f:
        html_content = f.read()
    # Wrap in editable page with save functionality
    page_html = f'''<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{
    font-family: 'Times New Roman', serif;
    margin: 0; padding: 20px 40px;
    background: white;
  }}
  /* Highlight unfilled placeholders */
  .missing-field {{
    background: #fef2f2;
    border: 1px dashed #ef4444;
    padding: 1px 4px;
    border-radius: 3px;
    color: #dc2626;
    font-weight: bold;
  }}
  /* Toolbar */
  #toolbar {{
    position: sticky; top: 0; z-index: 100;
    background: #1f2937; color: white;
    padding: 8px 16px;
    margin: -20px -40px 20px -40px;
    display: flex; align-items: center; justify-content: space-between;
    font-family: system-ui, sans-serif;
  }}
  #toolbar button {{
    padding: 4px 12px; border: none; border-radius: 4px;
    font-size: 12px; font-weight: 600; cursor: pointer;
  }}
  #toolbar .save-btn {{ background: #16a34a; color: white; }}
  #toolbar .save-btn:hover {{ background: #15803d; }}
  #toolbar .save-btn:disabled {{ background: #6b7280; cursor: not-allowed; }}
  #toolbar .translate-btn {{ background: #2563eb; color: white; }}
  #toolbar .translate-btn:hover {{ background: #1d4ed8; }}
  #toolbar .translate-btn:disabled {{ background: #6b7280; cursor: not-allowed; }}
  #toolbar .edit-toggle {{ background: #7c3aed; color: white; }}
  #toolbar .edit-toggle:hover {{ background: #6d28d9; }}
  #toolbar .edit-toggle.active {{ background: #f59e0b; }}
  #toolbar .status {{ font-size: 11px; color: #9ca3af; margin-left: 8px; }}
  #toolbar .title {{ font-size: 13px; font-weight: 600; }}
  /* Make content editable look nice */
  #doc-content {{ outline: none; min-height: 80vh; }}
  #doc-content:focus {{ outline: none; }}
  #doc-content table {{ border-collapse: collapse; }}
  #doc-content td, #doc-content th {{ border: 1px solid #ccc; padding: 4px 8px; }}
</style>
</head>
<body>
<div id="toolbar">
  <span class="title">{gf.form_name or form_code}</span>
  <div style="display:flex;align-items:center;gap:8px;">
    <span id="status" class="status"></span>
    <button class="translate-btn" id="translate-btn" onclick="translateToEnglish()" title="Translate Bahasa Melayu to English">Translate to English</button>
    <button class="edit-toggle" id="edit-toggle" onclick="toggleEdit()" title="Toggle editing mode">Edit: ON</button>
    <button class="save-btn" id="save-btn" onclick="saveChanges()">Save</button>
  </div>
</div>
<div id="doc-content" contenteditable="true">
''' + html_content + '''
</div>
<script>
let _dirty = false;
const docEl = document.getElementById('doc-content');

// Highlight {{PLACEHOLDERS}}
function highlightMissing() {
  const walker = document.createTreeWalker(docEl, NodeFilter.SHOW_TEXT, null, false);
  const nodes = [];
  while (walker.nextNode()) nodes.push(walker.currentNode);
  nodes.forEach(node => {
    if (/\\{\\{\\w+\\}\\}/.test(node.textContent)) {
      const span = document.createElement('span');
      span.innerHTML = node.textContent.replace(/\\{\\{(\\w+)\\}\\}/g,
        '<span class="missing-field">{{$1}}</span>');
      node.parentNode.replaceChild(span, node);
    }
  });
}
highlightMissing();

docEl.addEventListener('input', () => {
  _dirty = true;
  document.getElementById('status').textContent = 'Unsaved changes';
  document.getElementById('status').style.color = '#fbbf24';
});

async function saveChanges() {
  const btn = document.getElementById('save-btn');
  const status = document.getElementById('status');
  btn.disabled = true;
  btn.textContent = 'Saving...';
  status.textContent = '';

  // Get the edited HTML content
  const html = docEl.innerHTML;
  try {
    const res = await fetch('/probate/''' + probate_id + '''/form-html/''' + form_code + '''', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ html: html })
    });
    const data = await res.json();
    if (data.ok) {
      status.textContent = 'Saved successfully!';
      status.style.color = '#4ade80';
      _dirty = false;
      // Notify parent to refresh if needed
      if (window.parent && window.parent.onFormSaved) window.parent.onFormSaved();
    } else {
      status.textContent = 'Save failed: ' + (data.error || 'Unknown');
      status.style.color = '#f87171';
    }
  } catch(e) {
    status.textContent = 'Save failed: ' + e.message;
    status.style.color = '#f87171';
  }
  btn.disabled = false;
  btn.textContent = 'Save';
}

// Toggle edit mode
let _editMode = true;
function toggleEdit() {
  _editMode = !_editMode;
  docEl.contentEditable = _editMode ? 'true' : 'false';
  const btn = document.getElementById('edit-toggle');
  btn.textContent = 'Edit: ' + (_editMode ? 'ON' : 'OFF');
  btn.classList.toggle('active', _editMode);
}

// Translate to English using Claude API
async function translateToEnglish() {
  if (!confirm('Translate this document from Bahasa Melayu to English?')) return;
  const btn = document.getElementById('translate-btn');
  const status = document.getElementById('status');
  btn.disabled = true;

  // Spinner + countdown timer
  let elapsed = 0;
  const est = 20; // estimated seconds
  btn.innerHTML = '&#9889; Translating...';
  status.innerHTML = '<span style="display:inline-flex;align-items:center;gap:4px;">' +
    '<svg style="width:14px;height:14px;animation:spin 1s linear infinite;" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2v4m0 12v4m-7.07-3.93l2.83-2.83m8.49-8.49l2.83-2.83M2 12h4m12 0h4m-3.93 7.07l-2.83-2.83M6.34 6.34L3.51 3.51"/></svg>' +
    '<span id="translate-timer">~' + est + 's remaining</span></span>';
  status.style.color = '#60a5fa';

  // Add spinner keyframes if not present
  if (!document.getElementById('spin-style')) {
    const s = document.createElement('style');
    s.id = 'spin-style';
    s.textContent = '@keyframes spin{from{transform:rotate(0deg)}to{transform:rotate(360deg)}}';
    document.head.appendChild(s);
  }

  const timer = setInterval(() => {
    elapsed++;
    const rem = Math.max(0, est - elapsed);
    const el = document.getElementById('translate-timer');
    if (el) {
      if (rem > 0) el.textContent = '~' + rem + 's remaining';
      else el.textContent = 'Almost done...';
    }
  }, 1000);

  try {
    const res = await fetch('/probate/''' + probate_id + '''/translate-form/''' + form_code + '''', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ html: docEl.innerHTML })
    });
    clearInterval(timer);
    const data = await res.json();
    if (data.ok && data.translated_html) {
      docEl.innerHTML = data.translated_html;
      _dirty = true;
      status.innerHTML = '&#10003; Translated in ' + elapsed + 's — Review and Save';
      status.style.color = '#4ade80';
    } else {
      status.textContent = 'Translation failed: ' + (data.error || 'Unknown');
      status.style.color = '#f87171';
    }
  } catch(e) {
    clearInterval(timer);
    status.textContent = 'Translation failed: ' + e.message;
    status.style.color = '#f87171';
  }
  btn.disabled = false;
  btn.textContent = 'Translate to English';
}

window.addEventListener('beforeunload', (e) => {
  if (_dirty) { e.preventDefault(); e.returnValue = ''; }
});
</script>
</body></html>'''
    resp = make_response(page_html)
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return resp


@app.route('/probate/<probate_id>/form-html/<form_code>', methods=['POST'])
@login_required
def probate_form_html_save(probate_id, form_code):
    """Save edited HTML back to DOCX."""
    gf = ProbateGeneratedForm.query.filter_by(probate_id=probate_id, form_code=form_code).first()
    if not gf or not os.path.exists(gf.file_path):
        return jsonify(ok=False, error='Form not found'), 404
    data = request.get_json()
    if not data or 'html' not in data:
        return jsonify(ok=False, error='No HTML content provided'), 400
    import subprocess
    # Write edited HTML to temp file
    tmp_dir = tempfile.mkdtemp()
    html_path = os.path.join(tmp_dir, 'edited.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(f'''<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>{data['html']}</body></html>''')
    # Convert HTML back to DOCX using LibreOffice (requires explicit filter)
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'docx:MS Word 2007 XML', html_path, '--outdir', tmp_dir],
        capture_output=True, text=True, timeout=30
    )
    new_docx = os.path.join(tmp_dir, 'edited.docx')
    if not os.path.exists(new_docx):
        return jsonify(ok=False, error='Conversion failed'), 500
    # Replace the generated file
    import shutil
    shutil.copy2(new_docx, gf.file_path)
    gf.generated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(ok=True)


@app.route('/probate/<probate_id>/translate-form/<form_code>', methods=['POST'])
@login_required
def probate_translate_form(probate_id, form_code):
    """Translate form HTML from Bahasa Melayu to English using Claude API."""
    data = request.get_json()
    if not data or 'html' not in data:
        return jsonify(ok=False, error='No HTML content'), 400

    html_content = data['html']
    try:
        import anthropic
        from config import ANTHROPIC_API_KEY, CLAUDE_MODEL_FAST
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        response = client.messages.create(
            model=CLAUDE_MODEL_FAST,
            max_tokens=8000,
            messages=[{
                "role": "user",
                "content": f"""Translate this Malaysian legal court document from Bahasa Melayu to English.

IMPORTANT RULES:
1. Keep ALL HTML tags, attributes, and structure EXACTLY as-is — only translate the text content
2. Keep ALL names, NRIC numbers, addresses, dates, case numbers, and form numbers unchanged
3. Keep legal citations and section references unchanged (e.g., "Seksyen 12" → "Section 12")
4. Use proper English legal terminology (e.g., "Pemohon" → "Applicant", "Si Mati" → "the Deceased")
5. Keep "PEMOHON" as "APPLICANT", "Probet" as "Probate"
6. Return ONLY the translated HTML, no explanations

Common translations:
- Dalam Mahkamah Tinggi Malaya → In the High Court of Malaya
- Dalam Perkara Mengenai Harta Pusaka → In the Matter of the Estate of
- Saman Pemula → Originating Summons
- Afidavit → Affidavit
- Senarai Benefisiari → List of Beneficiaries
- Perintah → Order
- Sumpah → Oath/Affirmation

HTML to translate:
{html_content}"""
            }]
        )
        # §10x.70 — track this previously-unaudited callsite
        try:
            from ai.cost_tracker import log_usage
            # Resolve client_id via the probate record if available
            _cid = None
            try:
                _pb = db.session.get(Probate, probate_id)
                _cid = getattr(_pb, 'client_id', None) if _pb else None
            except Exception:
                pass
            log_usage(response, call_site='app.probate_translate_form',
                      client_id=_cid)
        except Exception:
            pass
        translated = response.content[0].text.strip()
        # Remove markdown code fences if present
        if translated.startswith('```'):
            translated = translated.split('\n', 1)[1] if '\n' in translated else translated[3:]
        if translated.endswith('```'):
            translated = translated[:-3].rstrip()
        return jsonify(ok=True, translated_html=translated)
    except Exception as e:
        return jsonify(ok=False, error=str(e)), 500


@app.route('/probate/<probate_id>/download/<form_code>')
@login_required
def probate_download(probate_id, form_code):
    fmt = request.args.get('format', 'docx')  # docx or pdf
    gf = ProbateGeneratedForm.query.filter_by(probate_id=probate_id, form_code=form_code).first()
    if not gf or not os.path.exists(gf.file_path):
        flash('Form not found.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')

    # Use proper form name for download filename
    safe_name = (gf.form_name or form_code).replace(' ', '_').replace('/', '_')

    if fmt == 'pdf':
        from documents.probate_generator import convert_to_pdf
        pdf_path = convert_to_pdf(gf.file_path)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(pdf_path, as_attachment=True, download_name=f'{safe_name}.pdf')
        flash('PDF conversion failed. Downloading .docx instead.', 'error')

    return send_file(gf.file_path, as_attachment=True, download_name=f'{safe_name}.docx')


@app.route('/probate/<probate_id>/reupload/<form_code>', methods=['POST'])
@login_required
def probate_reupload(probate_id, form_code):
    """Replace a generated form with an edited DOCX upload."""
    gf = ProbateGeneratedForm.query.filter_by(probate_id=probate_id, form_code=form_code).first()
    if not gf:
        return jsonify(ok=False, error='Form not found'), 404
    if 'file' not in request.files:
        return jsonify(ok=False, error='No file uploaded'), 400
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify(ok=False, error='Only .docx files are accepted'), 400
    # Overwrite the existing generated file
    file.save(gf.file_path)
    gf.generated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(ok=True)


@app.route('/probate/<probate_id>/download-all')
@login_required
def probate_download_all(probate_id):
    fmt = request.args.get('format', 'docx')  # docx or pdf
    forms = ProbateGeneratedForm.query.filter_by(probate_id=probate_id).all()
    if not forms:
        flash('No generated forms found.', 'error')
        return redirect(f'/probate/{probate_id}/step/7')

    from documents.probate_generator import create_zip
    zip_path = os.path.join(tempfile.gettempdir(), f'probate_{probate_id[:8]}.zip')
    form_files = [{'form_code': f.form_code, 'file_path': f.file_path, 'form_name': f.form_name} for f in forms]
    create_zip(form_files, zip_path, as_pdf=(fmt == 'pdf'))
    return send_file(zip_path, as_attachment=True, download_name=f'probate_forms_{probate_id[:8]}.zip')


@app.route('/api/probate/<probate_id>/send-email', methods=['POST'])
@login_required
def api_probate_send_email(probate_id):
    """Email all generated probate forms as a ZIP attachment."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        return jsonify({'ok': False, 'error': 'Probate application not found'}), 404

    forms = ProbateGeneratedForm.query.filter_by(probate_id=probate_id).all()
    if not forms:
        return jsonify({'ok': False, 'error': 'No generated forms to send'}), 400

    data = request.get_json(silent=True) or {}
    to_email = (data.get('to_email') or '').strip()
    if not to_email or '@' not in to_email:
        return jsonify({'ok': False, 'error': 'Please enter a valid email address'}), 400

    fmt = data.get('format', 'pdf')  # pdf or docx

    # Build ZIP attachment
    from documents.probate_generator import create_zip
    zip_path = os.path.join(tempfile.gettempdir(), f'probate_email_{probate_id[:8]}.zip')
    form_files = [{'form_code': f.form_code, 'file_path': f.file_path, 'form_name': f.form_name} for f in forms]
    create_zip(form_files, zip_path, as_pdf=(fmt == 'pdf'))
    with open(zip_path, 'rb') as f:
        zip_data = f.read()

    # Determine sender & CC
    tenant = get_tenant()
    user_role = session.get('user_role', '')
    user_name = session.get('user_name', 'Unknown')
    brand = tenant.get('brand', 'WillCraft AI')

    # CC: merge user-provided CC, tenant CC, and auto-CC approver for admin/advisor
    cc_list = list(tenant.get('email_cc', []))
    user_cc = (data.get('cc') or '').strip()
    if user_cc:
        for addr in user_cc.split(','):
            addr = addr.strip()
            if addr and '@' in addr and addr not in cc_list:
                cc_list.append(addr)
    if user_role in ('admin', 'advisor'):
        approvers = User.query.filter_by(role='approver', is_active=True).all()
        for ap in approvers:
            if ap.email and ap.email not in cc_list:
                cc_list.append(ap.email)

    # Use user-provided subject and body, with defaults
    deceased_name = probate.deceased_name or 'the estate'
    subject = (data.get('subject') or '').strip() or f"Probate Forms — {deceased_name}"
    user_body = (data.get('body') or '').strip()

    if user_body:
        # Convert plain text body to HTML (preserve line breaks)
        import html as html_mod
        body_escaped = html_mod.escape(user_body).replace('\n', '<br>')
        body_html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <p>{body_escaped}</p>
            <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
            <p style="font-size: 12px; color: #718096;">
                This email and its attachments are confidential and intended solely for the addressee.
            </p>
        </div>
        """
    else:
        body_html = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <p>Dear Sir/Madam,</p>
            <p>Please find attached the probate court forms for the estate of <strong>{deceased_name}</strong>.</p>
            <p>The attached ZIP file contains {len(forms)} form(s) in {fmt.upper()} format.</p>
            <br>
            <p>Best regards,<br><strong>{user_name}</strong><br>{brand}</p>
            <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 20px 0;">
            <p style="font-size: 12px; color: #718096;">
                This email and its attachments are confidential and intended solely for the addressee.
            </p>
        </div>
        """

    from_email = tenant.get('email_from')
    if not from_email:
        user = db.session.get(User, session.get('user_id'))
        from_email = user.email if user else None
    if not from_email:
        return jsonify({'ok': False, 'error': 'No sender email configured. Contact admin.'}), 400

    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.base import MIMEBase
        from email import encoders

        sender_addr = SMTP_USER or from_email
        msg = MIMEMultipart()
        msg['From'] = sender_addr
        msg['To'] = to_email
        msg['Subject'] = subject
        if from_email and from_email != sender_addr:
            msg['Reply-To'] = from_email
        if cc_list:
            msg['Cc'] = ', '.join(cc_list)
        msg.attach(MIMEText(body_html, 'html'))

        part = MIMEBase('application', 'octet-stream')
        part.set_payload(zip_data)
        encoders.encode_base64(part)
        ext = 'pdf' if fmt == 'pdf' else 'docx'
        part.add_header('Content-Disposition', f'attachment; filename="probate_forms_{probate_id[:8]}.zip"')
        msg.attach(part)

        all_recipients = [to_email] + cc_list
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            if SMTP_USER and SMTP_PASSWORD:
                server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(sender_addr, all_recipients, msg.as_string())

        app.logger.info(f'Probate {probate_id} forms emailed to {to_email} by {user_name} (cc: {cc_list})')
        return jsonify({
            'ok': True,
            'sent_to': to_email,
            'cc': cc_list,
            'message': f'Forms emailed to {to_email}',
        })
    except Exception as e:
        app.logger.error(f'Probate email failed: {e}')
        return jsonify({'ok': False, 'error': f'Failed to send email: {str(e)}'}), 500


@app.route('/api/ocr/death-cert', methods=['POST'])
@login_required
def api_ocr_death_cert():
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400

    from uploads import save_uploaded_file
    client_id = session.get('client_id', 'temp')
    try:
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, category='death_certificate')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    abs_path = os.path.join(UPLOAD_DIR, rel_path)

    from ai.ocr import extract_death_cert_data
    try:
        extracted = extract_death_cert_data(abs_path)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'OCR failed: {str(e)}'}), 500

    if 'error' in extracted:
        return jsonify({'ok': False, 'error': extracted['error'], 'extracted': extracted})

    # Save document record
    doc = Document(
        client_id=client_id,
        filename=saved_name,
        original_filename=file.filename,
        file_path=rel_path,
        file_type=file.filename.rsplit('.', 1)[-1].lower(),
        file_size=file_size,
        category='death_certificate',
        extracted_data=json.dumps(extracted),
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({'ok': True, 'document_id': doc.id, 'extracted': extracted})


@app.route('/api/ocr/will-document', methods=['POST'])
@login_required
def api_ocr_will_document():
    """Upload a will document (PDF/image) and OCR extract testator, executors, witnesses, beneficiaries, assets."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400

    from uploads import save_uploaded_file
    client_id = session.get('client_id', 'temp')
    try:
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, category='will_document')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    abs_path = os.path.join(UPLOAD_DIR, rel_path)

    from ai.ocr import extract_will_data
    try:
        extracted = extract_will_data(abs_path)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'OCR failed: {str(e)}'}), 500

    if 'error' in extracted:
        return jsonify({'ok': False, 'error': extracted['error'], 'extracted': extracted})

    # Save document record
    doc = Document(
        client_id=client_id,
        filename=saved_name,
        original_filename=file.filename,
        file_path=rel_path,
        file_type=file.filename.rsplit('.', 1)[-1].lower(),
        file_size=file_size,
        category='will_document',
        extracted_data=json.dumps(extracted),
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({'ok': True, 'document_id': doc.id, 'extracted': extracted})


@app.route('/api/ocr/asset-doc', methods=['POST'])
@login_required
def api_ocr_asset_doc():
    """Upload an asset document (title, bank statement, vehicle card, etc.) and OCR it."""
    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    fmt_err = _validate_ocr_file(file)
    if fmt_err:
        return jsonify({'ok': False, 'error': fmt_err}), 400

    asset_type = request.form.get('asset_type', 'other')

    from uploads import save_uploaded_file
    client_id = session.get('client_id', 'temp')
    try:
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, category=f'asset_{asset_type}')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    abs_path = os.path.join(UPLOAD_DIR, rel_path)

    from ai.ocr import extract_asset_document
    try:
        extracted = extract_asset_document(abs_path, asset_type)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'OCR failed: {str(e)}'}), 500

    if 'error' in extracted:
        return jsonify({'ok': False, 'error': extracted['error']})

    doc = Document(
        client_id=client_id,
        filename=saved_name,
        original_filename=file.filename,
        file_path=rel_path,
        file_type=file.filename.rsplit('.', 1)[-1].lower(),
        file_size=file_size,
        category=f'asset_{asset_type}',
        extracted_data=json.dumps(extracted),
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({'ok': True, 'document_id': doc.id, 'extracted': extracted})


@app.route('/api/probate/<probate_id>/checklist', methods=['POST'])
@login_required
def api_probate_checklist(probate_id):
    """Save filing checklist state."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        return jsonify(ok=False, error='Not found'), 404
    data = request.get_json(silent=True) or {}
    form_data = json.loads(probate.form_data_json or '{}')
    form_data['filing_checklist'] = data.get('checklist', {})
    probate.form_data_json = json.dumps(form_data)
    probate.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(ok=True)


@app.route('/api/probate/<probate_id>/upload-receipt', methods=['POST'])
@login_required
def api_probate_upload_receipt(probate_id):
    """Upload an invoice or payment receipt for a probate application."""
    probate = db.session.get(ProbateApplication, probate_id)
    if not probate:
        return jsonify({'ok': False, 'error': 'Probate application not found'}), 404

    if 'file' not in request.files:
        return jsonify({'ok': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'ok': False, 'error': 'No file selected'}), 400

    doc_category = request.form.get('category', 'invoice')  # invoice or receipt
    description = request.form.get('description', '')

    from uploads import save_uploaded_file
    client_id = probate.client_id or 'temp'
    try:
        saved_name, rel_path, file_size = save_uploaded_file(file, client_id, category=f'probate_{doc_category}')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

    doc = Document(
        client_id=client_id,
        filename=saved_name,
        original_filename=file.filename,
        file_path=rel_path,
        file_type=file.filename.rsplit('.', 1)[-1].lower(),
        file_size=file_size,
        category=f'probate_{doc_category}',
        description=f'probate:{probate_id}|{description}',
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({
        'ok': True,
        'document': {
            'id': doc.id,
            'filename': doc.original_filename,
            'category': doc_category,
            'description': description,
            'file_size': doc.file_size,
            'created_at': doc.created_at.isoformat(),
        }
    })


@app.route('/api/probate/<probate_id>/receipts')
@login_required
def api_probate_receipts(probate_id):
    """List invoices and payment receipts for a probate application."""
    docs = Document.query.filter(
        Document.description.like(f'probate:{probate_id}|%'),
        Document.category.in_(['probate_invoice', 'probate_receipt'])
    ).order_by(Document.created_at.desc()).all()

    return jsonify([{
        'id': d.id,
        'filename': d.original_filename,
        'category': d.category.replace('probate_', ''),
        'description': (d.description or '').split('|', 1)[-1],
        'file_size': d.file_size,
        'created_at': d.created_at.isoformat(),
    } for d in docs])


# Admin: Probate Template Management
@app.route('/admin/probate-templates')
@role_required('admin')
def admin_probate_templates():
    templates = ProbateFormTemplate.query.order_by(ProbateFormTemplate.sort_order).all()
    flash_msg = request.args.get('msg')
    return render_template('admin/probate_templates.html', templates=templates, flash_msg=flash_msg)


@app.route('/admin/probate-templates/<form_code>/upload', methods=['POST'])
@role_required('admin')
def admin_probate_template_upload(form_code):
    tpl = ProbateFormTemplate.query.filter_by(form_code=form_code).first()
    if not tpl:
        flash('Template not found.', 'error')
        return redirect(url_for('admin_probate_templates'))

    file = request.files.get('template')
    if not file or not file.filename:
        flash('No file selected.', 'error')
        return redirect(url_for('admin_probate_templates'))

    # Save custom template
    custom_dir = os.path.join(os.path.dirname(__file__), 'probate_templates', 'custom')
    os.makedirs(custom_dir, exist_ok=True)
    ext = file.filename.rsplit('.', 1)[-1].lower()
    custom_path = os.path.join(custom_dir, f'{form_code}.{ext}')
    file.save(custom_path)

    tpl.file_path = f'probate_templates/custom/{form_code}.{ext}'
    tpl.is_default = False
    tpl.updated_at = datetime.utcnow()
    db.session.commit()

    return redirect(url_for('admin_probate_templates', msg=f'Template for {tpl.form_name} updated.'))


@app.route('/admin/probate-templates/<form_code>/reset', methods=['POST'])
@role_required('admin')
def admin_probate_template_reset(form_code):
    tpl = ProbateFormTemplate.query.filter_by(form_code=form_code).first()
    if not tpl:
        flash('Template not found.', 'error')
        return redirect(url_for('admin_probate_templates'))

    # Map form_code back to default file
    default_files = {
        'doc01': 'doc01_saman_pemula.docx', 'doc02': 'doc02_afidavit_probet.docx',
        'doc03': 'doc03_sumpah_pentadbiran.docx', 'doc04': 'doc04_afidavit_saksi_1.docx',
        'doc05': 'doc05_afidavit_saksi_2.docx', 'doc06': 'doc06_jadual_aset.docx',
        'doc07': 'doc07_senarai_benefisiari.docx', 'doc08': 'doc08_notis_peguamcara.docx',
        'form14a': 'form14a_land_transfer.docx', 'form346': 'form346_personal_rep.docx',
    }
    default_file = default_files.get(form_code, f'{form_code}.docx')
    tpl.file_path = f'probate_templates/{default_file}'
    tpl.is_default = True
    tpl.updated_at = datetime.utcnow()
    db.session.commit()

    return redirect(url_for('admin_probate_templates', msg=f'Template for {tpl.form_name} reset to default.'))


@app.route('/probate/template/<form_code>/view')
@login_required
def probate_template_view(form_code):
    """Convert template to PDF and serve inline for browser viewing."""
    tpl = ProbateFormTemplate.query.filter_by(form_code=form_code).first()
    if not tpl:
        flash('Template not found.', 'error')
        return redirect(url_for('probate_list'))
    template_path = os.path.join(os.path.dirname(__file__), tpl.file_path)
    if not os.path.exists(template_path):
        flash('Template file not found on disk.', 'error')
        return redirect(url_for('probate_list'))
    # Detect actual file extension for correct download names
    actual_ext = os.path.splitext(template_path)[1].lstrip('.') or 'docx'
    fmt = request.args.get('format', 'pdf')
    if fmt in ('docx', 'doc'):
        return send_file(template_path, as_attachment=True,
                         download_name=f'{form_code}_template.{actual_ext}')
    # Convert to PDF for in-browser viewing
    from documents.probate_generator import convert_to_pdf
    import shutil
    tmp_dir = tempfile.mkdtemp()
    tmp_copy = os.path.join(tmp_dir, os.path.basename(template_path))
    shutil.copy2(template_path, tmp_copy)
    pdf_path = convert_to_pdf(tmp_copy)
    if pdf_path and os.path.exists(pdf_path):
        resp = send_file(pdf_path, mimetype='application/pdf',
                         download_name=f'{form_code}_template.pdf')
        # Prevent browser caching stale PDFs
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        return resp
    # Fallback: download in original format
    return send_file(template_path, as_attachment=True,
                     download_name=f'{form_code}_template.{actual_ext}')


@app.route('/probate/template/<form_code>/translate', methods=['POST'])
@login_required
def probate_template_translate(form_code):
    """Translate a probate form template from Malay to English using AI."""
    tpl = ProbateFormTemplate.query.filter_by(form_code=form_code).first()
    if not tpl:
        return jsonify(ok=False, error='Template not found'), 404
    template_path = os.path.join(os.path.dirname(__file__), tpl.file_path)
    if not os.path.exists(template_path):
        return jsonify(ok=False, error='File not found'), 404
    # Convert to PDF first, then translate using vision
    from documents.probate_generator import convert_to_pdf
    import shutil
    tmp_dir = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, os.path.basename(template_path))
    shutil.copy2(template_path, tmp_docx)
    pdf_path = convert_to_pdf(tmp_docx)
    if not pdf_path or not os.path.exists(pdf_path):
        return jsonify(ok=False, error='Could not convert template to PDF'), 500
    try:
        from ai.ocr import translate_document
        translation = translate_document(pdf_path)
        return jsonify(ok=True, translation=translation)
    except Exception as e:
        app.logger.error(f'Template translate error: {e}')
        return jsonify(ok=False, error='Translation failed. Please try again.'), 500


@app.route('/probate/template/<form_code>/replace', methods=['POST'])
@login_required
def probate_template_replace(form_code):
    """Upload a replacement template for a form (admin/approver only)."""
    role = session.get('user_role')
    if role not in ('admin', 'approver'):
        return jsonify(ok=False, error='Access denied'), 403
    tpl = ProbateFormTemplate.query.filter_by(form_code=form_code).first()
    if not tpl:
        return jsonify(ok=False, error='Template not found'), 404
    file = request.files.get('template')
    if not file or not file.filename:
        return jsonify(ok=False, error='No file selected'), 400
    custom_dir = os.path.join(os.path.dirname(__file__), 'probate_templates', 'custom')
    os.makedirs(custom_dir, exist_ok=True)
    ext = file.filename.rsplit('.', 1)[-1].lower()
    custom_path = os.path.join(custom_dir, f'{form_code}.{ext}')
    file.save(custom_path)
    tpl.file_path = f'probate_templates/custom/{form_code}.{ext}'
    tpl.is_default = False
    tpl.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(ok=True, message=f'Template for {tpl.form_name} updated.')


@app.route('/admin/debug/property_groups/<client_id>')
@login_required
def admin_debug_property_groups(client_id):
    """Debug endpoint: show all property_title docs and their NLC data for a client."""
    try:
        docs = Document.query.filter(
            Document.client_id == client_id,
            Document.category == 'property_title'
        ).order_by(Document.created_at.asc()).all()
        result = []
        for d in docs:
            try:
                ex = json.loads(d.extracted_data) if d.extracted_data else {}
            except Exception:
                ex = {}
            result.append({
                'id': d.id,
                'filename': d.original_filename,
                'lot': ex.get('lot_number', ''),
                'title': ex.get('title_number', ''),
                'mukim': ex.get('mukim', ''),
                'daerah': ex.get('daerah', ''),
                'negeri': ex.get('negeri', ''),
                'address': ex.get('property_address', ''),
                'inventoried': ex.get('_inventoried', False),
                'skipped': ex.get('_skipped', False),
                'substitute_assigned': ex.get('_substitute_assigned', False),
                'pending_beneficiary': ex.get('_pending_beneficiary', False),
                'enriched_from': ex.get('_enriched_from', []),
                'created_at': d.created_at.isoformat() if d.created_at else '',
            })
        return jsonify({'count': len(result), 'docs': result})
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500


@app.route('/admin/debug/dedupe_property_docs/<client_id>', methods=['POST'])
@login_required
def admin_dedupe_property_docs(client_id):
    """Soft-delete duplicate property_title docs for a client.

    Keeps the EARLIEST document for each (filename, lot, title) combination.
    Marks later duplicates as category='deleted'.
    Safe to call multiple times (idempotent).
    """
    try:
        docs = Document.query.filter(
            Document.client_id == client_id,
            Document.category == 'property_title'
        ).order_by(Document.created_at.asc()).all()

        seen = {}   # key → first doc id
        deleted_ids = []
        kept_ids = []

        for d in docs:
            try:
                ex = json.loads(d.extracted_data) if d.extracted_data else {}
            except Exception:
                ex = {}
            # Dedupe key: filename + lot + title (ignores mukim/address OCR variance)
            key = (
                (d.original_filename or '').strip().lower(),
                (ex.get('lot_number') or '').strip().lower(),
                (ex.get('title_number') or '').strip().lower(),
            )
            if key in seen:
                # Already have an earlier copy — soft-delete this one
                d.category = 'deleted'
                d.description = f'(duplicate of {seen[key][:8]} — auto-deduped)'
                deleted_ids.append(d.id)
            else:
                seen[key] = d.id
                kept_ids.append(d.id)

        db.session.commit()
        return jsonify({
            'kept': len(kept_ids),
            'deleted': len(deleted_ids),
            'kept_ids': kept_ids,
            'deleted_ids': deleted_ids,
        })
    except Exception as e:
        import traceback
        db.session.rollback()
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', '1') == '1'
    app.run(debug=debug, port=port)
