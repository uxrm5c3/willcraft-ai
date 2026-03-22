"""Generate a Word document (.docx) from will text using python-docx.

Professional Malaysian will format:
- Running header: "LAST WILL AND TESTAMENT OF [TESTATOR NAME]" on every page
- Running footer: Page number + Testator/Witness1/Witness2 signature spaces on every page
- Proper signing/attestation page layout at the end
"""

import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _extract_testator_name(will_text: str) -> str:
    """Extract testator name from will text."""
    lines = will_text.strip().split('\n')
    for i, line in enumerate(lines):
        if 'LAST WILL AND TESTAMENT OF' in line.upper():
            after = line.upper().replace('LAST WILL AND TESTAMENT OF', '').strip()
            if after:
                return after
            for j in range(i + 1, min(i + 3, len(lines))):
                candidate = lines[j].strip()
                if candidate:
                    return candidate.upper()
    return "THE TESTATOR"


def _split_signing_page(will_text: str):
    """Split will text into main content and signing page section."""
    markers = ['Signature of the Testator', 'SIGNATURE OF THE TESTATOR']
    for marker in markers:
        idx = will_text.find(marker)
        if idx >= 0:
            return will_text[:idx].rstrip(), will_text[idx:]
    return will_text, ""


def _add_header(doc, testator_name: str):
    """Add running header with testator name to the document."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # First line: "LAST WILL AND TESTAMENT OF"
    para1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para1.paragraph_format.space_after = Pt(0)
    para1.paragraph_format.space_before = Pt(0)
    run1 = para1.add_run("LAST WILL AND TESTAMENT OF")
    run1.bold = True
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(9)

    # Second line: testator name
    para2 = header.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_after = Pt(6)
    para2.paragraph_format.space_before = Pt(0)
    run2 = para2.add_run(testator_name)
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(9)

    # Add bottom border to header
    pPr = para2._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)


def _add_footer(doc):
    """Add running footer with page number and Testator/Witness signature spaces.
    Layout: Page N | ___Testator___ | ___Witness 1___ | ___Witness 2___
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Top border line above footer
    border_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    border_para.text = ""
    border_para.paragraph_format.space_after = Pt(4)
    border_para.paragraph_format.space_before = Pt(0)
    pPr_b = border_para._element.get_or_add_pPr()
    pBdr_b = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="1" w:color="999999"/></w:pBdr>')
    pPr_b.append(pBdr_b)

    # Single-row table: [Page N] [Testator sig] [Witness 1 sig] [Witness 2 sig]
    table = footer.add_table(1, 4, Cm(14))
    table.autofit = False

    # Set table to full width and remove all borders
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    # Full width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    def _set_cell_width(cell, width_pct):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_pct}" w:type="pct"/>')
        tcPr.append(tcW)

    def _make_sig_cell(cell, label):
        """Create a cell with signature line + label below."""
        _set_cell_width(cell, 1250)  # 25% each
        # Signature line
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_before = Pt(0)
        run1 = p1.add_run("______________________")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Label
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.space_before = Pt(1)
        run2 = p2.add_run(label)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(7)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Col 0: Page number (narrower)
    cell_0 = table.cell(0, 0)
    _set_cell_width(cell_0, 1000)  # 20%
    p0 = cell_0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.space_before = Pt(6)
    run_page = p0.add_run("Page ")
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # Page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_page._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_page._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_page._element.append(fldChar2)

    # Col 1-3: Signature spaces
    _make_sig_cell(table.cell(0, 1), "Testator")
    _make_sig_cell(table.cell(0, 2), "Witness 1")
    _make_sig_cell(table.cell(0, 3), "Witness 2")


def _add_signing_page(doc):
    """Add the attestation/signing page with proper Rockwill layout."""
    # Page break before signing page
    para_break = doc.add_paragraph()
    run_break = para_break.add_run()
    run_break.add_break(WD_BREAK.PAGE)

    # Helper to add a label-field row
    def add_sig_row(label_text, with_line=True):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        # Remove borders
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            tbl.insert(0, tblPr)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        # Label cell
        cell0 = table.cell(0, 0)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        p0.paragraph_format.space_before = Pt(2)
        run0 = p0.add_run(label_text)
        run0.font.name = 'Times New Roman'
        run0.font.size = Pt(11)

        # Set column width for label (approx 40%)
        tc0 = cell0._element
        tcPr0 = tc0.get_or_add_tcPr()
        tcW0 = parse_xml(f'<w:tcW {nsdecls("w")} w:w="3600" w:type="dxa"/>')
        tcPr0.append(tcW0)

        # Field cell with underline
        cell1 = table.cell(0, 1)
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
        if with_line:
            # Add bottom border to the cell content
            pPr = p1._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            run1 = p1.add_run(" ")
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(11)

        return table

    # Testator signature
    add_sig_row("Signature of the Testator:")
    doc.add_paragraph('')  # spacer

    # Date
    t = add_sig_row("Date of this Will:")
    # Add (dd/mm/yyyy) hint in the field
    cell = t.cell(0, 1)
    p = cell.paragraphs[0]
    for run in p.runs:
        run.clear()
    hint_run = p.add_run("                                                                    (dd/mm/yyyy)")
    hint_run.font.name = 'Times New Roman'
    hint_run.font.size = Pt(9)
    hint_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')  # spacer

    # Attestation clause
    attest = doc.add_paragraph()
    attest.paragraph_format.space_before = Pt(6)
    attest.paragraph_format.space_after = Pt(12)
    run_attest = attest.add_run(
        "This Last Will and Testament was signed by the Testator in the presence "
        "of us both and attested by us in the presence of both Testator and of each other:"
    )
    run_attest.font.name = 'Times New Roman'
    run_attest.font.size = Pt(11)

    # First Witness
    add_sig_row("Signature of First Witness:")
    add_sig_row("First Witness Full Name:")
    add_sig_row("First Witness Identification:")
    add_sig_row("First Witness Address:")
    add_sig_row("")  # Address line 2
    add_sig_row("")  # Address line 3
    add_sig_row("First Witness Contact Number:")

    doc.add_paragraph('')  # spacer

    # Second Witness
    add_sig_row("Signature of Second Witness:")
    add_sig_row("Second Witness Full Name:")
    add_sig_row("Second Witness Identification:")
    add_sig_row("Second Witness Address:")
    add_sig_row("")  # Address line 2
    add_sig_row("")  # Address line 3
    add_sig_row("Second Witness Contact Number:")

    # End of Document marker
    doc.add_paragraph('')
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_para.paragraph_format.space_before = Pt(18)
    run_end = end_para.add_run("End of Document")
    run_end.bold = True
    run_end.font.name = 'Times New Roman'
    run_end.font.size = Pt(10)


def _add_cover_page(doc, testator_name: str, will_text: str, firm_info: dict = None, logo_path: str = None):
    """Add a cover page with firm logo, testator name, and NRIC."""
    import re

    if not firm_info:
        return

    # Extract NRIC from will text
    nric_match = re.search(r'NRIC\s*No\.?\s*[:\s]*(\d{6}-\d{2}-\d{4})', will_text)
    nric = nric_match.group(1) if nric_match else ''

    # Add firm logo if available
    if logo_path and os.path.isfile(logo_path):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_before = Pt(60)
            logo_para.paragraph_format.space_after = Pt(20)
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(3))
        except Exception:
            pass

    # Firm address
    if firm_info.get('firm_address'):
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_para.paragraph_format.space_after = Pt(60)
        run = addr_para.add_run(firm_info['firm_address'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Thin line separator
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/></w:pBdr>')
    line_para._element.get_or_add_pPr().append(pBdr)

    # Title: "The Last Will & Testament"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(16)
    title_para.paragraph_format.space_after = Pt(0)
    title_para.paragraph_format.line_spacing = 2.0
    run = title_para.add_run('The Last Will & Testament')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # "of"
    of_para = doc.add_paragraph()
    of_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    of_para.paragraph_format.space_after = Pt(0)
    of_para.paragraph_format.line_spacing = 2.0
    run = of_para.add_run('of')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # Testator name (no underline — matches law firm format)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(8)
    run = name_para.add_run(testator_name)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)

    # NRIC
    if nric:
        nric_para = doc.add_paragraph()
        nric_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nric_para.paragraph_format.space_after = Pt(10)
        run = nric_para.add_run(f'(NRIC No. {nric})')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Page break after cover
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_prepared_by_page(doc, firm_info: dict = None):
    """Add 'Prepared By' page at the end of the document."""
    if not firm_info:
        return

    # Page break before
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Spacers
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Decorative line
    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line1.add_run('_' * 30)
    run.font.size = Pt(10)

    # "PREPARED BY:"
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(20)
    run = heading.add_run('PREPARED BY:')
    run.bold = True
    run.underline = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Decorative line
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line2.add_run('_' * 30)
    run.font.size = Pt(10)

    # Firm name
    if firm_info.get('firm_name'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(firm_info['firm_name'].upper())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # "ADVOCATES & SOLICITORS"
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    run = sub.add_run('ADVOCATES & SOLICITORS')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Address
    if firm_info.get('firm_address'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(firm_info['firm_address'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Phone
    if firm_info.get('firm_phone'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"TEL NO: {firm_info['firm_phone']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Email
    if firm_info.get('firm_email'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"EMAIL: {firm_info['firm_email']}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Decorative line
    line3 = doc.add_paragraph()
    line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line3.add_run('_' * 30)
    run.font.size = Pt(10)


def generate_docx(will_text: str, filename_base: str = "Will",
                  firm_info: dict = None, logo_path: str = None) -> str:
    """
    Generate a Word document from the will text.

    Args:
        will_text: The full will text to render.
        filename_base: Base name for the output file (without extension).
        firm_info: Optional dict with firm details for cover and prepared-by pages.
        logo_path: Optional path to firm logo for cover page.

    Returns:
        Absolute path to the generated .docx file.
    """
    doc = Document()
    testator_name = _extract_testator_name(will_text)
    main_content, signing_section = _split_signing_page(will_text)

    # -- Page setup -----------------------------------------------------------
    section = doc.sections[0]
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # -- Cover page (if firm info available) -----------------------------------
    _add_cover_page(doc, testator_name, will_text, firm_info, logo_path)

    # -- Add running header ---------------------------------------------------
    _add_header(doc, testator_name)

    # -- Add running footer ---------------------------------------------------
    _add_footer(doc)

    # -- Default font style ---------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # -- Parse and add main content -------------------------------------------
    # Collapse consecutive blank lines to max 1 to prevent excessive whitespace
    lines = main_content.split('\n')
    prev_blank = False
    for line in lines:
        stripped = line.strip()

        if not stripped:
            if prev_blank:
                continue  # Skip consecutive blank lines
            prev_blank = True
            # Add minimal spacer paragraph instead of full empty line
            spacer = doc.add_paragraph('')
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(6)
            continue
        prev_blank = False

        # Skip title lines (already in header)
        if 'LAST WILL AND TESTAMENT OF' in stripped.upper():
            continue

        # Skip per-page footer content (already in running footer)
        if 'Continued on' in stripped and ('next page' in stripped.lower() or 'Page' in stripped):
            continue
        if stripped.startswith('Page|') or stripped.startswith('Page |'):
            continue
        # Skip lines that are just "Testator    Witness 1    Witness 2" footer patterns
        if ('Testator' in stripped and 'Witness 1' in stripped and 'Witness 2' in stripped):
            continue
        # Skip lines of just underscores (signature lines in footer)
        if stripped.replace('_', '').replace(' ', '').replace('|', '') == '':
            continue

        # Detect "THE REST OF THE PAGE IS INTENTIONALLY LEFT BLANK"
        if 'REST OF THE PAGE IS INTENTIONALLY LEFT BLANK' in stripped.upper():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(24)
            run = para.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            continue

        # Detect title-like lines (all uppercase and short)
        is_heading = (
            stripped.isupper()
            and len(stripped) < 80
            and not stripped.startswith('(')
            and not stripped.startswith('-')
        )

        # Detect numbered clause headings
        is_clause_heading = False
        if len(stripped) > 2 and stripped[0].isdigit():
            dot_pos = stripped.find('.')
            if 0 < dot_pos <= 3:
                rest = stripped[dot_pos + 1:].strip()
                if rest and rest.isupper():
                    is_clause_heading = True

        # Detect section headings
        is_section_heading = stripped in (
            'Revocation', 'Appointment of Executor(s)', 'Appointment of Guardian(s)',
            'Non Residuary Gift(s)', 'Residuary Estate', 'Declaration',
            'Testamentary Trust', 'Guardian Allowance', 'Contemplation of Marriage',
        )

        if is_section_heading:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.keep_with_next = True  # Prevent heading alone at page bottom
            run = para.add_run(stripped)
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        elif is_heading:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.keep_with_next = True
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14 if 'LAST WILL' in stripped else 12)
            run.font.name = 'Times New Roman'
        elif is_clause_heading:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.keep_with_next = True
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            para.paragraph_format.space_before = Pt(8)
        else:
            para = doc.add_paragraph()
            # Preserve leading whitespace / indentation
            if line.startswith('    ') or line.startswith('\t'):
                para.paragraph_format.left_indent = Cm(1.27)
            run = para.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # -- Add signing page -----------------------------------------------------
    if signing_section:
        _add_signing_page(doc)

    # -- Add "Prepared By" page (if firm info available) ----------------------
    _add_prepared_by_page(doc, firm_info)

    # -- Save to temp file ----------------------------------------------------
    tmp_dir = tempfile.mkdtemp()
    filepath = os.path.join(tmp_dir, f'{filename_base}_Will.docx')
    doc.save(filepath)
    return filepath
