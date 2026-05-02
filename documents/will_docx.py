"""Render a will as a Microsoft Word .docx using the Alan Tan & Associates
established format.

Mirrors the same format spec as documents/pdf_generator.py — Times New Roman
11pt body, 1.5 line spacing, justified body, bold+underlined section headings,
hanging indent for numbered clauses, 3-line signature footer (Testator / Witness 1
/ Witness 2), no header/footer on the cover page.

Used by the admin Will Format Preview so the user can:
  - download the sample as Word
  - open in Microsoft Word / Pages
  - track changes / red-line edit the format
  - send feedback back to refine the templates

Public API: build_will_docx(will_text, output_path, firm_info=None) -> str
"""

import os
import re
import tempfile
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ── Format constants — matched to PDF extraction (pdfplumber) of the target ──
# Target PDF font sizes per pdfplumber:
#   Body: 11pt | Section heading: 12pt | Page header: 9pt | Footer: 8pt
#   Cover title/of/name: 16pt | Cover NRIC: 14pt | Cover firm addr: 10pt
#   Prepared By heading: 14pt | "ADVOCATES & SOLICITORS": 12pt | Address: 11pt
FONT_FAMILY = 'Times New Roman'
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADING = Pt(12)
FONT_SIZE_PAGE_HEADER = Pt(9)
FONT_SIZE_PAGE_NUMBER = Pt(8)
FONT_SIZE_SIG_LABEL = Pt(8)
FONT_SIZE_COVER_TITLE = Pt(16)
FONT_SIZE_COVER_OF = Pt(16)
FONT_SIZE_COVER_NAME = Pt(16)
FONT_SIZE_COVER_NRIC = Pt(14)
FONT_SIZE_COVER_FIRM_ADDR = Pt(10)


# ── Helpers ───────────────────────────────────────────────────────────
def _set_font(run, *, size=None, bold=False, italic=False, underline=False):
    run.font.name = FONT_FAMILY
    # Word needs both ascii and east-asia explicit
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_FAMILY)
    rFonts.set(qn('w:hAnsi'), FONT_FAMILY)
    if size is not None:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline


def _set_paragraph_format(paragraph, *, alignment=None, space_before_pt=0,
                           space_after_pt=0, line_spacing=1.5,
                           left_indent_pt=None, first_line_indent_pt=None):
    pf = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing
    if left_indent_pt is not None:
        pf.left_indent = Pt(left_indent_pt)
    if first_line_indent_pt is not None:
        pf.first_line_indent = Pt(first_line_indent_pt)


def _add_page_number_field(paragraph):
    """Insert a Word PAGE field that auto-numbers."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    _set_font(run, size=FONT_SIZE_PAGE_NUMBER)


def _add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph (used under page header)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_section_margins(section, *, top_in=1.0, bottom_in=1.25,
                         left_in=1.0, right_in=1.0,
                         header_in=0.5, footer_in=0.6):
    section.top_margin = Inches(top_in)
    section.bottom_margin = Inches(bottom_in)
    section.left_margin = Inches(left_in)
    section.right_margin = Inches(right_in)
    section.header_distance = Inches(header_in)
    section.footer_distance = Inches(footer_in)


def _add_draft_watermark(section):
    """Add a diagonal "DRAFT" watermark to every page in `section`.
    Uses Word VML inside the section's header XML — visible on every page,
    diagonal, light gray, large. Mirrors the original Alan & Tan sample."""
    header = section.header
    # Inject VML watermark as raw OXML inside a header paragraph
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    r = p.add_run()
    vml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" '
        'adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">'
        '<v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>'
        '<v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/>'
        '<v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>'
        '<v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>'
        '<v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/>'
        '<v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>'
        '<v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/></v:formulas>'
        '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="custom"/>'
        '<o:lock v:ext="edit" text="t" shapetype="t"/></v:shapetype>'
        '<v:shape id="WaterMark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:600pt;height:200pt;'
        'rotation:-30;z-index:-251654144;mso-position-horizontal:center;'
        'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
        'mso-position-vertical-relative:margin" '
        'fillcolor="#dddddd" stroked="f">'
        '<v:fill opacity=".4"/>'
        '<v:textpath style="font-family:&quot;Times New Roman&quot;;font-size:1pt" '
        'string="DRAFT"/>'
        '</v:shape></w:pict>'
    )
    from docx.oxml import parse_xml
    pict_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + vml + '</w:r>'
    )
    elem = parse_xml(pict_xml)
    p._p.append(elem[0])


def _build_body_header(section, testator_name: str, logo_path: Optional[str] = None):
    """Build the running header for body/signing pages:
    [logo (small, centered)] / 'LAST WILL AND TESTAMENT OF' / '{TESTATOR_NAME}'
    + horizontal rule. Logo appears on EVERY page per Alan & Tan format.
    """
    header = section.header
    # Logo — first paragraph (if provided)
    p_logo = header.paragraphs[0]
    p_logo.text = ''
    if logo_path and os.path.isfile(logo_path):
        run = p_logo.add_run()
        try:
            run.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            pass
    _set_paragraph_format(p_logo, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    # "LAST WILL AND TESTAMENT OF" — small bold
    p1 = header.add_paragraph()
    r1 = p1.add_run('LAST WILL AND TESTAMENT OF')
    _set_font(r1, size=FONT_SIZE_PAGE_HEADER, bold=True)
    _set_paragraph_format(p1, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    # Testator name in running header — same 9pt bold as the line above
    # (matches target PDF — both lines are 9pt per pdfplumber extraction)
    p2 = header.add_paragraph()
    r2 = p2.add_run(testator_name.upper())
    _set_font(r2, size=FONT_SIZE_PAGE_HEADER, bold=True)
    _set_paragraph_format(p2, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)

    _add_horizontal_line(p2)


def _build_body_footer(section):
    """Build the running footer: a full-width horizontal rule across the page,
    then 'Page X' on the left + 3 signature BOXES (Testator / Witness 1 /
    Witness 2) — labels sit INSIDE each box at the bottom.
    """
    footer = section.footer
    p0 = footer.paragraphs[0]
    p0.text = ''

    # Full-page horizontal line ABOVE the signing boxes (per target PDF)
    _set_paragraph_format(p0, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing=1.0, space_before_pt=0, space_after_pt=4)
    pPr = p0._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_rule = OxmlElement('w:bottom')
    bottom_rule.set(qn('w:val'), 'single')
    bottom_rule.set(qn('w:sz'), '6')
    bottom_rule.set(qn('w:space'), '4')
    bottom_rule.set(qn('w:color'), '000000')
    pBdr.append(bottom_rule)
    pPr.append(pBdr)

    # Single-row 4-column table: [Page X] [Testator box] [Witness 1 box] [Witness 2 box]
    table = footer.add_table(rows=1, cols=4, width=Inches(6.5))
    table.autofit = False
    widths = [Inches(0.8), Inches(1.9), Inches(1.9), Inches(1.9)]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]

    # Box row height
    from docx.oxml import OxmlElement as _Ox
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trH = _Ox('w:trHeight')
    trH.set(qn('w:val'), '720')   # ~36pt
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

    # Cell 0 — "Page X" plain text (no border)
    cell_pg = table.rows[0].cells[0]
    cell_pg.paragraphs[0].text = ''
    p_page = cell_pg.paragraphs[0]
    r_page = p_page.add_run('Page ')
    _set_font(r_page, size=FONT_SIZE_PAGE_NUMBER)
    _add_page_number_field(p_page)
    _set_paragraph_format(p_page, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    # Anchor "Page X" to the bottom of its cell so it lines up with the labels
    tcPr_pg = cell_pg._element.get_or_add_tcPr()
    vAlign_pg = _Ox('w:vAlign')
    vAlign_pg.set(qn('w:val'), 'bottom')
    tcPr_pg.append(vAlign_pg)

    # Cells 1-3 — Testator / Witness 1 / Witness 2 boxes with label INSIDE
    # at the bottom (no separate label row outside)
    labels = ['Testator', 'Witness 1', 'Witness 2']
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i + 1]
        cell.paragraphs[0].text = ''

        # Full borders on every edge
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = _Ox('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = _Ox(f'w:{edge}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:color'), '000000')
            tcBorders.append(b)
        tcPr.append(tcBorders)

        # Vertically anchor cell content to the bottom — label sits at the
        # bottom of the box; signature space is above it
        vAlign = _Ox('w:vAlign')
        vAlign.set(qn('w:val'), 'bottom')
        tcPr.append(vAlign)

        # Label paragraph — centered, at the bottom of the cell
        p_lbl = cell.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        _set_font(r_lbl, size=FONT_SIZE_SIG_LABEL)
        _set_paragraph_format(p_lbl, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=1.0, space_before_pt=0, space_after_pt=2)


def _add_run_with_emphasis(paragraph, text: str):
    """Add text to a paragraph, applying bold to UPPERCASE NAME (NRIC ...) tokens.
    This is a heuristic to emulate the bold-name styling in the sample without
    needing structured input.
    """
    # Pattern matches: "NAME (MALAYSIA NRIC No. xxx-xx-xxxx)" or similar
    # Bold the name + parenthesised ID together.
    pattern = re.compile(
        r'([A-Z][A-Z\s/\'\.&]+?)\s*(\((?:MALAYSIA\s+NRIC|NRIC|Identification|Passport)\s*(?:No\.?|NO\.?)\s*[\w\-/\s]+?\))'
    )

    pos = 0
    for m in pattern.finditer(text):
        # Text before the match
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            _set_font(r, size=FONT_SIZE_BODY)
        # Bold name + parenthesised ID
        r = paragraph.add_run(text[m.start():m.end()])
        _set_font(r, size=FONT_SIZE_BODY, bold=True)
        pos = m.end()
    # Tail
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        _set_font(r, size=FONT_SIZE_BODY)


# Section headings used in Alan & Tan format
_SECTION_HEADINGS = {
    'Revocation', 'Appointment of Executor(s)', 'Appointment of Guardian(s)',
    'Non-Residuary Gift(s)', 'Non Residuary Gift(s)',
    'Residuary Estate', 'Declaration',
    'Testamentary Trust', 'Guardian Allowance', 'Contemplation of Marriage',
}


def _split_signing(will_text: str):
    markers = ['Signature of the Testator', 'SIGNATURE OF THE TESTATOR']
    for m in markers:
        idx = will_text.find(m)
        if idx >= 0:
            return will_text[:idx].rstrip(), will_text[idx:]
    return will_text, ''


def _extract_testator_name(will_text: str) -> str:
    lines = will_text.strip().split('\n')
    for i, line in enumerate(lines):
        if 'LAST WILL AND TESTAMENT OF' in line.upper():
            after = line.upper().replace('LAST WILL AND TESTAMENT OF', '').strip()
            if after:
                return after
            for j in range(i + 1, min(i + 3, len(lines))):
                cand = lines[j].strip()
                if cand:
                    return cand.upper()
    return 'THE TESTATOR'


def build_will_docx(will_text: str, output_path: Optional[str] = None,
                     firm_info: Optional[dict] = None,
                     logo_path: Optional[str] = None,
                     is_draft: bool = True) -> str:
    """Render the will text as .docx and return the path.

    Args:
        will_text: Full will body + signature text
        firm_info: dict with firm_name, firm_address, firm_phone, firm_email
        logo_path: optional path to PNG/JPG logo embedded on cover page
        is_draft: when True, applies a DRAFT watermark to every page
    """
    doc = Document()

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.line_spacing = 1.5

    testator_name = _extract_testator_name(will_text)
    main_text, signing_text = _split_signing(will_text)

    # ── COVER PAGE — section 0: no header text/footer, optional watermark ──
    cover_section = doc.sections[0]
    _set_section_margins(cover_section, top_in=1.2, bottom_in=1.2)
    cover_section.different_first_page_header_footer = False
    cover_section.header.paragraphs[0].text = ''
    cover_section.footer.paragraphs[0].text = ''
    if is_draft:
        _add_draft_watermark(cover_section)

    # Cover layout per Alan & Tan reference — logo + address VERTICALLY
    # CENTERED on the page; title block sits in the lower half.

    # Top spacer to push logo+address to vertical mid-page
    for _ in range(8):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=1.0)

    # (1) Logo — centered, mid-page
    if logo_path and os.path.isfile(logo_path):
        p = doc.add_paragraph()
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=1.0, space_before_pt=4, space_after_pt=4)
        run = p.add_run()
        try:
            run.add_picture(logo_path, height=Inches(1.4))
        except Exception:
            pass

    # (2) Firm address below logo
    if firm_info and firm_info.get('firm_address'):
        p = doc.add_paragraph()
        r = p.add_run(firm_info['firm_address'])
        _set_font(r, size=FONT_SIZE_COVER_FIRM_ADDR)
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=1.0, space_before_pt=4, space_after_pt=12)

    # (3) Space between address and title block
    for _ in range(3):
        p = doc.add_paragraph()
        _set_paragraph_format(p, line_spacing=1.0)

    # (4) Title block
    for line, fs in [
        ('The Last Will & Testament', FONT_SIZE_COVER_TITLE),
        ('of', FONT_SIZE_COVER_OF),
        (testator_name, FONT_SIZE_COVER_NAME),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        _set_font(r, size=fs, bold=True)
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=1.4, space_before_pt=4, space_after_pt=4)

    # NRIC line — extract from will text (14pt per target PDF)
    nric_match = re.search(r'NRIC\s*No\.?\s*[:\s]*(\d{6}-\d{2}-\d{4})', will_text)
    if nric_match:
        p = doc.add_paragraph()
        r = p.add_run(f'(NRIC No. {nric_match.group(1)})')
        _set_font(r, size=FONT_SIZE_COVER_NRIC, bold=True)
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              line_spacing=1.4, space_before_pt=4)

    # ── BODY PAGES — new section with header + footer ──
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_margins(body_section)
    body_section.different_first_page_header_footer = False
    # Critical: unlink from previous section so cover stays blank
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False

    _build_body_header(body_section, testator_name, logo_path=logo_path)
    _build_body_footer(body_section)
    if is_draft:
        _add_draft_watermark(body_section)

    # ── Body content ──
    # Skip the first occurrence of "LAST WILL AND TESTAMENT OF" + testator
    # name lines (already in the running header).
    lines = main_text.split('\n')
    skip_count = 0
    for i, ln in enumerate(lines):
        if 'LAST WILL AND TESTAMENT' in ln.upper():
            skip_count = i + 1
            # Also skip a following testator-name line and any blank
            j = skip_count
            while j < len(lines) and (
                lines[j].strip() == '' or lines[j].strip().upper() == testator_name.upper()
            ):
                j += 1
                skip_count = j
            break

    body_lines = lines[skip_count:]

    # Per target PDF: testator name on page 2 — 14pt bold centered, with
    # ~2-3 lines of empty space ABOVE (between running header and the name)
    # and ~2-3 lines BELOW (between name and "This Will is made by me…").

    # 2-3 blank lines ABOVE the testator name
    for _ in range(3):
        p_blank = doc.add_paragraph()
        _set_paragraph_format(p_blank, line_spacing=1.0, space_before_pt=0, space_after_pt=0)

    p_name = doc.add_paragraph()
    r_name = p_name.add_run(testator_name.upper())
    _set_font(r_name, size=Pt(14), bold=True)
    _set_paragraph_format(p_name, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before_pt=0, space_after_pt=0, line_spacing=1.2)

    # 2-3 blank lines BELOW the testator name
    for _ in range(3):
        p_blank = doc.add_paragraph()
        _set_paragraph_format(p_blank, line_spacing=1.0, space_before_pt=0, space_after_pt=0)

    i = 0
    blank_streak = 0
    while i < len(body_lines):
        line = body_lines[i]
        stripped = line.strip()
        if not stripped:
            blank_streak += 1
            # Allow only one blank-line spacer
            if blank_streak == 1:
                p = doc.add_paragraph()
                _set_paragraph_format(p, line_spacing=1.0, space_before_pt=0, space_after_pt=0)
            i += 1
            continue
        blank_streak = 0

        # Section heading?
        if stripped in _SECTION_HEADINGS or stripped.rstrip(':') in _SECTION_HEADINGS:
            p = doc.add_paragraph()
            r = p.add_run(stripped.rstrip(':'))
            _set_font(r, size=FONT_SIZE_HEADING, bold=True, underline=True)
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  space_before_pt=14, space_after_pt=8, line_spacing=1.5)
            i += 1
            continue

        # Numbered clause "1. ", "2. ", etc.
        clause_match = re.match(r'^(\d{1,2}\.)\s+(.+)$', stripped)
        if clause_match:
            num = clause_match.group(1)
            body_text = clause_match.group(2)
            # Collect continuation lines (until blank or next clause/heading)
            j = i + 1
            while j < len(body_lines):
                next_stripped = body_lines[j].strip()
                if not next_stripped:
                    break
                if next_stripped in _SECTION_HEADINGS:
                    break
                if re.match(r'^\d{1,2}\.\s', next_stripped):
                    break
                if re.match(r'^\([a-z]\)\s', next_stripped):
                    break
                body_text += ' ' + next_stripped
                j += 1
            i = j

            p = doc.add_paragraph()
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                  space_before_pt=0, space_after_pt=12,
                                  left_indent_pt=28, first_line_indent_pt=-28,
                                  line_spacing=1.5)
            # Number prefix
            r_num = p.add_run(f'{num}\t')
            _set_font(r_num, size=FONT_SIZE_BODY)
            # Body with bold-name emphasis
            _add_run_with_emphasis(p, body_text)
            continue

        # Sub-clause (a), (b) ...
        sub_match = re.match(r'^(\([a-z]\))\s+(.+)$', stripped)
        if sub_match:
            letter = sub_match.group(1)
            body_text = sub_match.group(2)
            j = i + 1
            while j < len(body_lines):
                next_stripped = body_lines[j].strip()
                if not next_stripped:
                    break
                if next_stripped in _SECTION_HEADINGS:
                    break
                if re.match(r'^\d{1,2}\.\s', next_stripped):
                    break
                if re.match(r'^\([a-z]\)\s', next_stripped):
                    break
                body_text += ' ' + next_stripped
                j += 1
            i = j
            p = doc.add_paragraph()
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                  space_before_pt=0, space_after_pt=10,
                                  left_indent_pt=64, first_line_indent_pt=-32,
                                  line_spacing=1.5)
            r_num = p.add_run(f'{letter}\t')
            _set_font(r_num, size=FONT_SIZE_BODY)
            _add_run_with_emphasis(p, body_text)
            continue

        # Blank-page marker line
        if 'remaining page is intentionally left blank' in stripped.lower():
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            _set_font(r, size=FONT_SIZE_BODY)
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  space_before_pt=12, space_after_pt=6, line_spacing=1.5)
            i += 1
            continue

        # Plain paragraph (e.g. opening statement, charge note, etc.)
        # Collect continuation lines
        body_text = stripped
        j = i + 1
        while j < len(body_lines):
            next_stripped = body_lines[j].strip()
            if not next_stripped:
                break
            if next_stripped in _SECTION_HEADINGS:
                break
            if re.match(r'^\d{1,2}\.\s', next_stripped):
                break
            if re.match(r'^\([a-z]\)\s', next_stripped):
                break
            body_text += ' ' + next_stripped
            j += 1
        i = j

        p = doc.add_paragraph()
        _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_before_pt=0, space_after_pt=12, line_spacing=1.5)
        _add_run_with_emphasis(p, body_text)

    # ── SIGNATURE PAGE (page 5) — boxed signatures, separate section so we
    # can suppress the body-page footer (no Testator/Witness 1/Witness 2 sig
    # lines at the bottom of the signing page) per Alan & Tan format.
    if signing_text.strip():
        from docx.enum.text import WD_BREAK
        from docx.oxml.ns import qn as _qn

        # New section for signing page → independent footer (just "Page X")
        sign_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_section_margins(sign_section)
        sign_section.header.is_linked_to_previous = False
        sign_section.footer.is_linked_to_previous = False

        # Logo + LAST WILL AND TESTAMENT OF header same as body
        _build_body_header(sign_section, testator_name, logo_path=logo_path)

        # Footer = just "Page X" (no sig boxes)
        sf = sign_section.footer
        for p in list(sf.paragraphs):
            p.text = ''
        for tbl in list(sf.tables):
            tbl._element.getparent().remove(tbl._element)
        p_pg = sf.paragraphs[0]
        r_pg = p_pg.add_run('Page ')
        _set_font(r_pg, size=FONT_SIZE_PAGE_NUMBER)
        _add_page_number_field(p_pg)
        _set_paragraph_format(p_pg, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)

        if is_draft:
            _add_draft_watermark(sign_section)

        # 2-3 blank lines between running header and "Signature of the Testator:"
        for _ in range(3):
            p_blank = doc.add_paragraph()
            _set_paragraph_format(p_blank, line_spacing=1.0,
                                  space_before_pt=0, space_after_pt=0)

        # Page 5 signing layout — 2-column table so ALL fillable lines start
        # at the SAME x position regardless of label length (matches target).
        # Label column ~2.4", line column ~4.0" → lines all aligned.

        def _build_signing_table():
            t = doc.add_table(rows=0, cols=2)
            t.autofit = False
            for col, w in zip(t.columns, [Inches(2.4), Inches(4.1)]):
                for cell in col.cells:
                    cell.width = w
            return t

        def _add_aligned_row(table, label, hint=''):
            row = table.add_row()
            cell_l, cell_r = row.cells[0], row.cells[1]
            # Label
            p_l = cell_l.paragraphs[0]
            p_l.text = ''
            r_l = p_l.add_run(label)
            _set_font(r_l, size=Pt(11))
            _set_paragraph_format(p_l, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  line_spacing=1.5, space_before_pt=2, space_after_pt=2)
            # Line cell — fixed-width underscores so lines align across rows
            p_r = cell_r.paragraphs[0]
            p_r.text = ''
            r_line = p_r.add_run('_' * 40)
            _set_font(r_line, size=Pt(11))
            if hint:
                r_h = p_r.add_run('  ' + hint)
                _set_font(r_h, size=Pt(9))
            _set_paragraph_format(p_r, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  line_spacing=1.5, space_before_pt=2, space_after_pt=2)

        # Testator + Date table
        sig_t = _build_signing_table()
        _add_aligned_row(sig_t, 'Signature of the Testator:')
        _add_aligned_row(sig_t, 'Date of this Will:', '(dd/mm/yyyy)')

        # Spacer + attestation
        sp = doc.add_paragraph()
        _set_paragraph_format(sp, line_spacing=1.0, space_after_pt=6)
        p_att = doc.add_paragraph()
        r_att = p_att.add_run(
            'This Last Will and Testament was signed by the Testator in the '
            'presence of us both and attested by us in the presence of both '
            'Testator and of each other:'
        )
        _set_font(r_att, size=Pt(11))
        _set_paragraph_format(p_att, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              space_before_pt=8, space_after_pt=12, line_spacing=1.5)

        # Two witness blocks — each a fresh aligned table
        for prefix in ['First', 'Second']:
            wt = _build_signing_table()
            _add_aligned_row(wt, f'Signature of {prefix} Witness:')
            _add_aligned_row(wt, 'Full Name:')
            _add_aligned_row(wt, 'NRIC / Passport No.:')
            _add_aligned_row(wt, 'Address:')
            _add_aligned_row(wt, '')  # 2nd address line — empty label, line still aligned
            _add_aligned_row(wt, '')  # 3rd address line
            _add_aligned_row(wt, 'Contact No.:')
            sp2 = doc.add_paragraph()
            _set_paragraph_format(sp2, line_spacing=1.0, space_after_pt=6)

        # End of Document marker — normal weight, no italic (per user spec).
        # Bold per user feedback would be too heavy; user said "bold, normal,
        # not italic" so render as bold + normal style + no italic.
        p_end = doc.add_paragraph()
        r_end = p_end.add_run('- End of Document -')
        _set_font(r_end, size=Pt(10), bold=True, italic=False)
        _set_paragraph_format(p_end, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before_pt=12, space_after_pt=0, line_spacing=1.2)

    # ── PREPARED BY — final page with law firm details (short centered lines) ──
    if firm_info and (firm_info.get('firm_name') or firm_info.get('firm_address')):
        from docx.enum.text import WD_BREAK
        from docx.oxml.ns import qn as _qn
        p_break = doc.add_paragraph()
        rb = p_break.add_run()
        rb.add_break(WD_BREAK.PAGE)

        # Vertical spacer to push firm block to mid-page
        for _ in range(6):
            p = doc.add_paragraph()
            _set_paragraph_format(p, line_spacing=1.0)

        def _hr_para(space_before=0, space_after=14):
            """Short horizontal line (~40% page width) centered — uses a
            single-cell table with a top border, narrower than full page."""
            t = doc.add_table(rows=1, cols=1)
            t.autofit = False
            # Set width via XML (Word ignores Inches().width on table directly)
            t.alignment = 1  # center
            cell = t.rows[0].cells[0]
            cell.width = Inches(2.6)
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(_qn('w:val'), 'single')
            top.set(_qn('w:sz'), '6')
            top.set(_qn('w:color'), '000000')
            tcBorders.append(top)
            tcPr.append(tcBorders)
            # Center the table itself
            tblPr = t._element.find(_qn('w:tblPr'))
            if tblPr is not None:
                jc = OxmlElement('w:jc')
                jc.set(_qn('w:val'), 'center')
                tblPr.append(jc)
                tblW = OxmlElement('w:tblW')
                tblW.set(_qn('w:w'), '3744')   # ~2.6 inches in dxa (1 inch = 1440)
                tblW.set(_qn('w:type'), 'dxa')
                tblPr.append(tblW)
            # Empty paragraph in the cell (just to draw the line)
            cell.paragraphs[0].text = ''
            # Outer paragraph spacing
            _set_paragraph_format(cell.paragraphs[0], line_spacing=1.0)

        _hr_para()
        p_h = doc.add_paragraph()
        r_h = p_h.add_run('PREPARED BY:')
        _set_font(r_h, size=Pt(14), bold=True, underline=True)
        _set_paragraph_format(p_h, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before_pt=0, space_after_pt=14, line_spacing=1.4)
        _hr_para()

        if firm_info.get('firm_name'):
            p = doc.add_paragraph()
            r = p.add_run(firm_info['firm_name'].upper())
            _set_font(r, size=Pt(14), bold=True)
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_before_pt=0, space_after_pt=4, line_spacing=1.3)

        p_sub = doc.add_paragraph()
        r_sub = p_sub.add_run('ADVOCATES & SOLICITORS')
        _set_font(r_sub, size=Pt(12), bold=True)
        _set_paragraph_format(p_sub, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before_pt=0, space_after_pt=18, line_spacing=1.3)

        for line in [
            firm_info.get('firm_address', ''),
            f"TEL NO: {firm_info['firm_phone']}" if firm_info.get('firm_phone') else '',
            f"EMAIL: {firm_info['firm_email']}" if firm_info.get('firm_email') else '',
        ]:
            if not line:
                continue
            p = doc.add_paragraph()
            r = p.add_run(line)
            _set_font(r, size=Pt(11))
            _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_before_pt=0, space_after_pt=4, line_spacing=1.5)

        _hr_para(space_before=14, space_after=0)

    # ── Save ──
    if output_path is None:
        tmp_dir = tempfile.mkdtemp()
        output_path = os.path.join(tmp_dir, 'will.docx')
    doc.save(output_path)
    return output_path
