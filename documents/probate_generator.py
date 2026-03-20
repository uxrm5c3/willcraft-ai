"""Probate form generator — fills .docx templates with placeholder replacement."""
import copy
import json
import os
import re
import uuid
import zipfile
from datetime import datetime
from docx import Document


# Per-form field mapping: which fields each form requires and where the data comes from.
# Used to show users what information will be filled in each form.
FORM_FIELDS = {
    'doc01': {
        'name': 'Originating Summons (Saman Pemula)',
        'fields': [
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Deceased address', 'Will — Testator'),
            ('Date of death', 'Step 1 — Death Certificate'),
            ('Place of death', 'Step 1 — Death Certificate'),
            ('Applicant (Executor) name & NRIC', 'Will — Executor'),
            ('Applicant address', 'Will — Executor'),
            ('Court location', 'Step 2 — Court Info'),
            ('Case number', 'Step 2 — Court Info'),
            ('Firm name & address', 'Step 2 — Law Firm'),
            ('Firm reference', 'Step 2 — Law Firm'),
        ],
    },
    'doc02': {
        'name': 'Affidavit under Probate Act',
        'fields': [
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Deceased address', 'Will — Testator'),
            ('Date of death', 'Step 1 — Death Certificate'),
            ('Time of death', 'Step 1 — Death Certificate'),
            ('Place of death', 'Step 1 — Death Certificate'),
            ('Death certificate number', 'Step 1 — Death Certificate'),
            ('Applicant name & NRIC', 'Will — Executor'),
            ('Applicant relationship', 'Will — Executor'),
            ('Estate value', 'Step 1 — Estate Value'),
            ('Exhibit references', 'Auto-generated'),
        ],
    },
    'doc03': {
        'name': 'Oath of Administration',
        'fields': [
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Date of death', 'Step 1 — Death Certificate'),
            ('Place of death', 'Step 1 — Death Certificate'),
            ('Applicant name & NRIC', 'Will — Executor'),
            ('Applicant address', 'Will — Executor'),
            ('Court location & case number', 'Step 2 — Court Info'),
        ],
    },
    'doc04': {
        'name': 'Witness 1 Affidavit',
        'fields': [
            ('Witness 1 name & NRIC', 'Step 3 — Witnesses'),
            ('Witness 1 address', 'Step 3 — Witnesses'),
            ('Deceased name', 'Will — Testator'),
            ('Date of death', 'Step 1 — Death Certificate'),
            ('Court location & case number', 'Step 2 — Court Info'),
        ],
    },
    'doc05': {
        'name': 'Witness 2 Affidavit',
        'fields': [
            ('Witness 2 name & NRIC', 'Step 3 — Witnesses'),
            ('Witness 2 address', 'Step 3 — Witnesses'),
            ('Deceased name', 'Will — Testator'),
            ('Date of death', 'Step 1 — Death Certificate'),
            ('Court location & case number', 'Step 2 — Court Info'),
        ],
    },
    'doc06': {
        'name': 'Assets & Liabilities Schedule',
        'fields': [
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Properties (title, lot, mukim, address)', 'Step 5 — Assets'),
            ('Bank accounts (bank, account no., value)', 'Step 5 — Assets'),
            ('Vehicles (desc, reg no., engine, chassis)', 'Step 5 — Assets'),
            ('Other assets (description, value)', 'Step 5 — Assets'),
            ('Liabilities (description, value)', 'Step 5 — Assets'),
        ],
    },
    'doc07': {
        'name': 'Beneficiary List',
        'fields': [
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Beneficiary names & NRIC', 'Step 4 — Beneficiaries'),
            ('Beneficiary relationships', 'Step 4 — Beneficiaries'),
            ('Court location & case number', 'Step 2 — Court Info'),
        ],
    },
    'doc08': {
        'name': 'Notice of Solicitor Appointment',
        'fields': [
            ('Deceased name', 'Will — Testator'),
            ('Applicant (Executor) name & NRIC', 'Will — Executor'),
            ('Court location & case number', 'Step 2 — Court Info'),
            ('Firm name & address', 'Step 2 — Law Firm'),
            ('Firm phone & fax', 'Step 2 — Law Firm'),
            ('Firm reference', 'Step 2 — Law Firm'),
        ],
    },
    'form14a': {
        'name': 'Land Transfer (Form 14A)',
        'fields': [
            ('Property title number', 'Step 5 — Assets (Property)'),
            ('Property lot number', 'Step 5 — Assets (Property)'),
            ('Property mukim', 'Step 5 — Assets (Property)'),
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Applicant name & NRIC', 'Will — Executor'),
            ('Court location & case number', 'Step 2 — Court Info'),
        ],
    },
    'form346': {
        'name': 'Personal Representative (Form 346)',
        'fields': [
            ('Property title number', 'Step 5 — Assets (Property)'),
            ('Property lot number', 'Step 5 — Assets (Property)'),
            ('Deceased name & NRIC', 'Will — Testator'),
            ('Applicant name & NRIC', 'Will — Executor'),
            ('Court case number', 'Step 2 — Court Info'),
        ],
    },
}


def _estimate_line_len(text):
    """Estimate the visual length of text in approximate character widths.

    Tabs count as 8 chars (default tab stop ~0.5 inch ≈ 5-6 chars at 11pt).
    """
    length = 0
    for ch in text:
        if ch == '\t':
            length += 8
        else:
            length += 1
    return length


def _trim_leading_whitespace(text, max_len=60):
    """If text is too long for one line, reduce leading tabs/spaces to fit.

    Uses conservative max_len=60 to account for bold text and tab stops.
    Preserves at least one tab for indented lines. Returns trimmed text.
    """
    est = _estimate_line_len(text)
    if est <= max_len:
        return text

    # Find where the actual content starts (after leading whitespace)
    content_start = 0
    for j, ch in enumerate(text):
        if ch not in (' ', '\t'):
            content_start = j
            break

    if content_start == 0:
        return text  # No leading whitespace to trim

    leading = text[:content_start]
    content = text[content_start:]
    content_len = _estimate_line_len(content)

    # Calculate how much leading space we can afford
    available = max_len - content_len
    if available < 0:
        available = 0

    # Rebuild leading whitespace: use tabs (each ~8 chars) to fill available space
    if available >= 8:
        new_leading = '\t' * (available // 8)
    elif available >= 1:
        new_leading = ' ' * available
    else:
        new_leading = ''

    return new_leading + content


def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph, handling split runs.

    After replacement, trims leading whitespace if the line would be too
    long (e.g. tabs pushing a name off the right edge).
    """
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text or '{{' not in full_text:
        return False

    new_text = full_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value) if value else '')

    if new_text != full_text and paragraph.runs:
        # Smart adjustment: trim leading whitespace if line is too long
        new_text = _trim_leading_whitespace(new_text)

        # Keep first run's formatting, clear others
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

        return True
    return False


def fill_template(template_path, replacements, output_path):
    """Open a .docx template, replace all {{PLACEHOLDER}} markers, save.

    Args:
        template_path: Path to the .docx template file
        replacements: Dict mapping {{PLACEHOLDER}} -> value
        output_path: Path to save the generated .docx file

    Returns:
        output_path on success
    """
    doc = Document(template_path)

    # Process all paragraphs in main body
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    # Process tables (many court forms have tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, replacements)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_in_paragraph(para, replacements)
        if section.footer:
            for para in section.footer.paragraphs:
                replace_in_paragraph(para, replacements)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_in_paragraph(para, replacements)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def build_replacements(probate_app, will_record):
    """Build the full replacement dict from probate application + will data.

    Args:
        probate_app: ProbateApplication model instance
        will_record: Will model instance (can be None for LA)

    Returns:
        Dict of {{PLACEHOLDER}} -> value
    """
    is_la = probate_app.application_type == 'la'

    if is_la or not will_record:
        # LA: deceased/applicant from probate fields
        testator = {
            'full_name': probate_app.deceased_name or '',
            'nric_passport': probate_app.deceased_nric or '',
            'residential_address': probate_app.deceased_address or '',
        }
        primary_exec = {
            'full_name': probate_app.applicant_name or '',
            'nric_passport': probate_app.applicant_nric or '',
            'address': probate_app.applicant_address or '',
            'relationship': probate_app.applicant_relationship or '',
        }
        beneficiaries = json.loads(probate_app.beneficiaries_data or '[]')
        gifts = json.loads(probate_app.assets_data or '[]')
    else:
        # Probate: from will data
        testator = json.loads(will_record.step1_data or '{}')
        step2 = json.loads(will_record.step2_data or '{}')
        executors = step2.get('executors', []) if isinstance(step2, dict) else step2
        primary_exec = executors[0] if executors else {}
        # Prefer probate.beneficiaries_data (populated in step 4), fall back to will data
        _probate_bens = json.loads(probate_app.beneficiaries_data or '[]') if probate_app.beneficiaries_data and probate_app.beneficiaries_data != '[]' else []
        beneficiaries = _probate_bens if _probate_bens else json.loads(will_record.step4_data or '[]')
        gifts = json.loads(will_record.step5_data or '[]')

    # Build applicant initials for exhibit references (e.g., TLL from TAN LI LI)
    applicant_name = primary_exec.get('full_name', probate_app.witness1_name or '')
    initials = ''.join(word[0] for word in applicant_name.split() if word) if applicant_name else 'APP'

    replacements = {
        # Deceased (testator/si mati)
        '{{DECEASED_NAME}}': testator.get('full_name', ''),
        '{{DECEASED_NRIC}}': testator.get('nric_passport', ''),
        '{{DECEASED_ADDRESS}}': testator.get('residential_address', ''),

        # Death details
        '{{DATE_OF_DEATH}}': probate_app.date_of_death or '',
        '{{TIME_OF_DEATH}}': probate_app.time_of_death or '',
        '{{PLACE_OF_DEATH}}': probate_app.place_of_death or '',
        '{{DEATH_CERT_NO}}': probate_app.death_cert_number or '',

        # Applicant (primary executor)
        '{{APPLICANT_NAME}}': primary_exec.get('full_name', ''),
        '{{APPLICANT_NRIC}}': primary_exec.get('nric_passport', ''),
        '{{APPLICANT_ADDRESS}}': primary_exec.get('address', ''),
        '{{APPLICANT_RELATIONSHIP}}': primary_exec.get('relationship', ''),
        '{{APPLICANT_CAPACITY}}': 'waris kadimnya',

        # Court details
        '{{COURT_LOCATION}}': probate_app.court_location or '',
        '{{COURT_STATE}}': probate_app.court_state or '',
        '{{CASE_NUMBER}}': probate_app.case_number or '',
        '{{FILING_YEAR}}': probate_app.filing_year or str(datetime.now().year),

        # Law firm
        '{{FIRM_NAME}}': probate_app.firm_name or '',
        '{{FIRM_NAME_UPPER}}': (probate_app.firm_name or '').upper(),
        '{{FIRM_ADDRESS}}': probate_app.firm_address or '',
        '{{FIRM_PHONE}}': probate_app.firm_phone or '',
        '{{FIRM_FAX}}': probate_app.firm_fax or '',
        '{{FIRM_REFERENCE}}': probate_app.firm_reference or '',

        # Lawyer details (for Form 14A/346 attestation)
        '{{LAWYER_NAME}}': probate_app.lawyer_name or primary_exec.get('full_name', ''),
        '{{LAWYER_NRIC}}': probate_app.lawyer_nric or primary_exec.get('nric_passport', ''),
        '{{LAWYER_BAR_NO}}': probate_app.lawyer_bar_number or '',
        '{{BAR_COUNCIL_NO}}': probate_app.lawyer_bar_number or '',

        # Witnesses
        '{{WITNESS1_NAME}}': probate_app.witness1_name or '',
        '{{WITNESS1_NRIC}}': probate_app.witness1_nric or '',
        '{{WITNESS1_ADDRESS}}': probate_app.witness1_address or '',
        '{{WITNESS2_NAME}}': probate_app.witness2_name or '',
        '{{WITNESS2_NRIC}}': probate_app.witness2_nric or '',
        '{{WITNESS2_ADDRESS}}': probate_app.witness2_address or '',

        # Estate value
        '{{ESTATE_VALUE}}': probate_app.estate_value_estimate or '',
        '{{ESTATE_VALUE_WORDS}}': probate_app.estate_value_estimate or '',

        # Exhibit references
        '{{EXHIBIT_1}}': f'{initials}-1',
        '{{EXHIBIT_2}}': f'{initials}-2',
        '{{EXHIBIT_3}}': f'{initials}-3',
        '{{EXHIBIT_4}}': f'{initials}-4',

        # Beneficiary list (for Doc 07)
        '{{BENEFICIARY1_RELATIONSHIP}}': '',

        # Date
        '{{TODAY_DATE}}': datetime.now().strftime('%d-%m-%Y'),
    }

    # Build firm address lines for forms that use multi-line format
    if probate_app.firm_address:
        # Handle both real newlines and literal \n from form input
        addr_text = probate_app.firm_address.replace('\\n', '\n')
        addr_lines = [l.strip() for l in addr_text.split('\n') if l.strip()]
        for i, line in enumerate(addr_lines[:4], 1):
            replacements[f'{{{{FIRM_ADDRESS_LINE{i}}}}}'] = line
        # Also update the full address with proper formatting
        replacements['{{FIRM_ADDRESS}}'] = ', '.join(addr_lines)
    # Ensure all FIRM_ADDRESS_LINE placeholders have values (empty if not set)
    for i in range(1, 5):
        key = f'{{{{FIRM_ADDRESS_LINE{i}}}}}'
        if key not in replacements:
            replacements[key] = ''

    # Build beneficiary info
    if beneficiaries:
        ben_lines = []
        for i, b in enumerate(beneficiaries):
            name = b.get('beneficiary_name', b.get('full_name', ''))
            nric = b.get('nric_passport_birthcert', b.get('nric_passport', ''))
            rel = b.get('relationship', '')
            ben_lines.append(f'{name} (No. K/P: {nric}), {rel}.')
        replacements['{{BENEFICIARY1_RELATIONSHIP}}'] = ben_lines[0].split('), ')[-1].rstrip('.') if ben_lines else ''

    # Property placeholders (for Form 14A/346)
    # Check both will gifts and probate assets for property data
    properties = [g for g in gifts if g.get('gift_type') == 'property']
    # Also check probate assets_data
    all_assets = json.loads(probate_app.assets_data or '[]')
    prop_assets = [a for a in all_assets if a.get('asset_type') == 'property']
    # Prefer probate assets (entered in step 5), fall back to will gifts
    prop_source = prop_assets[0] if prop_assets else (properties[0] if properties else None)
    if prop_source:
        # Handle both nested (will) and flat (probate) formats
        details = prop_source.get('property_details', prop_source)
        prop_title = details.get('title_number', '')
        prop_lot = details.get('lot_number', '')
        prop_mukim = details.get('mukim', '')
        prop_addr = details.get('address', details.get('description', ''))
        replacements.update({
            '{{PROPERTY1_TITLE}}': prop_title,
            '{{PROPERTY1_LOT}}': prop_lot,
            '{{PROPERTY1_MUKIM}}': prop_mukim,
            '{{PROPERTY1_ADDRESS}}': prop_addr,
            '{{PROPERTY2_TITLE}}': prop_title,
            '{{PROPERTY2_LOT}}': prop_lot,
            '{{PROPERTY2_MUKIM}}': prop_mukim,
            # Aliases used by form346
            '{{PROPERTY_TITLE_NO}}': prop_title,
            '{{PROPERTY_LOT}}': prop_lot,
            '{{PROPERTY_MUKIM}}': prop_mukim,
        })

    # Populate from assets_data (Step 4 schedule)
    all_assets = json.loads(probate_app.assets_data or '[]')

    # Bank accounts
    banks = [a for a in all_assets if a.get('asset_type') == 'bank']
    for i, b in enumerate(banks[:5], 1):
        replacements[f'{{{{BANK{i}_NAME}}}}'] = b.get('bank_name', '')
        replacements[f'{{{{BANK{i}_ACCNO}}}}'] = b.get('account_number', '')
        replacements[f'{{{{BANK{i}_VALUE}}}}'] = b.get('value', '')
    # Ensure empty placeholders for unused slots
    for i in range(1, 6):
        for suffix in ('_NAME', '_ACCNO', '_VALUE'):
            key = f'{{{{BANK{i}{suffix}}}}}'
            if key not in replacements:
                replacements[key] = ''

    # Vehicles
    vehicles = [a for a in all_assets if a.get('asset_type') == 'vehicle']
    for i, v in enumerate(vehicles[:3], 1):
        replacements[f'{{{{VEHICLE{i}_DESC}}}}'] = v.get('description', '')
        replacements[f'{{{{VEHICLE{i}_REGNO}}}}'] = v.get('reg_number', '')
        replacements[f'{{{{VEHICLE{i}_ENGINE}}}}'] = v.get('engine_number', '')
        replacements[f'{{{{VEHICLE{i}_CHASSIS}}}}'] = v.get('chassis_number', '')
        replacements[f'{{{{VEHICLE{i}_VALUE}}}}'] = v.get('value', '')
    for i in range(1, 4):
        for suffix in ('_DESC', '_REGNO', '_ENGINE', '_CHASSIS', '_VALUE'):
            key = f'{{{{VEHICLE{i}{suffix}}}}}'
            if key not in replacements:
                replacements[key] = ''

    # Other assets (EPF, insurance, shares, etc.)
    others = [a for a in all_assets if a.get('asset_type') == 'other']
    for i, o in enumerate(others[:5], 1):
        replacements[f'{{{{OTHER{i}_DESC}}}}'] = o.get('description', '')
        replacements[f'{{{{OTHER{i}_VALUE}}}}'] = o.get('value', '')
    for i in range(1, 6):
        for suffix in ('_DESC', '_VALUE'):
            key = f'{{{{OTHER{i}{suffix}}}}}'
            if key not in replacements:
                replacements[key] = ''

    # Liabilities
    liabilities = [a for a in all_assets if a.get('asset_type') == 'liability']
    for i, l in enumerate(liabilities[:5], 1):
        replacements[f'{{{{LIABILITY{i}_DESC}}}}'] = l.get('description', '')
        replacements[f'{{{{LIABILITY{i}_VALUE}}}}'] = l.get('value', '')
    for i in range(1, 6):
        for suffix in ('_DESC', '_VALUE'):
            key = f'{{{{LIABILITY{i}{suffix}}}}}'
            if key not in replacements:
                replacements[key] = ''

    # Also update property placeholders from assets_data if not already set from will
    asset_properties = [a for a in all_assets if a.get('asset_type') == 'property']
    if asset_properties and '{{PROPERTY1_TITLE}}' not in replacements:
        p = asset_properties[0]
        replacements.update({
            '{{PROPERTY1_TITLE}}': p.get('title_number', ''),
            '{{PROPERTY1_LOT}}': p.get('lot_number', ''),
            '{{PROPERTY1_MUKIM}}': p.get('mukim', ''),
            '{{PROPERTY1_ADDRESS}}': p.get('description', ''),
        })

    return replacements


def recommend_forms(will_record, probate_app=None):
    """Analyze will/probate data and recommend which forms are needed.

    Returns list of dicts: {form_code, recommended, reason}
    """
    is_la = probate_app and probate_app.application_type == 'la'

    # Check assets_data first (Step 4 schedule), fallback to will gifts
    assets_data = json.loads(probate_app.assets_data or '[]') if probate_app else []
    has_property_from_assets = any(a.get('asset_type') == 'property' for a in assets_data)

    if is_la:
        has_property = has_property_from_assets
    elif will_record:
        gifts = json.loads(will_record.step5_data or '[]')
        has_property = has_property_from_assets or any(g.get('gift_type') == 'property' for g in gifts)
    else:
        has_property = has_property_from_assets

    recommendations = [
        {
            'form_code': 'doc01',
            'recommended': True,
            'reason': 'Required — the main court filing to start the application',
        },
        {
            'form_code': 'doc02',
            'recommended': True,
            'reason': "Required — applicant's sworn statement about the estate",
        },
        {
            'form_code': 'doc03',
            'recommended': True,
            'reason': "Required — applicant's oath to manage the estate honestly",
        },
        {
            'form_code': 'doc04',
            'recommended': not is_la,
            'reason': 'Not needed — no will witnesses for LA' if is_la
                      else 'Recommended — sworn statement from the first will witness',
        },
        {
            'form_code': 'doc05',
            'recommended': not is_la,
            'reason': 'Not needed — no will witnesses for LA' if is_la
                      else 'Recommended — sworn statement from the second will witness',
        },
        {
            'form_code': 'doc06',
            'recommended': True,
            'reason': "Required — list of the deceased's assets and debts",
        },
        {
            'form_code': 'doc07',
            'recommended': True,
            'reason': 'Required — list of people who will inherit',
        },
        {
            'form_code': 'doc08',
            'recommended': True,
            'reason': 'Required — formal notice of lawyer appointment',
        },
        {
            'form_code': 'form14a',
            'recommended': has_property,
            'reason': 'Needed — transfers property from deceased to beneficiary' if has_property
                      else 'Not needed — no property in the estate',
        },
        {
            'form_code': 'form346',
            'recommended': has_property,
            'reason': 'Needed — registers executor at land office for property' if has_property
                      else 'Not needed — no property in the estate',
        },
    ]
    return recommendations


def _fix_signing_alignment(doc):
    """Fix signing name lines that use excessive spaces/tabs for right-alignment.

    After placeholder replacement, lines like:
        '                                    \t    PARAMSOTHY A/P V APPUKUDDY'
    overflow because the spaces+name is too wide. Fix by trimming leading
    whitespace and using right-alignment on the paragraph instead.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in doc.paragraphs:
        if not para.runs:
            continue
        full_text = ''.join(r.text for r in para.runs)
        stripped = full_text.strip()
        if not stripped:
            continue
        leading_ws = len(full_text) - len(full_text.lstrip())
        # Target: lines with 20+ leading spaces that contain a name (no {{ }})
        # These are signing name lines like "                    JOHN DOE"
        if leading_ws >= 20 and '{{' not in stripped and len(stripped) < 100:
            # Check it's a name-like line (mostly uppercase, short)
            words = stripped.replace('[', '').replace(']', '').split()
            if words and not any(kw in stripped.lower() for kw in [
                'mahkamah', 'dalam', 'negeri', 'saman', 'afidavit', 'borang',
                'peguambela', 'difailkan', 'setem', 'pembayaran'
            ]):
                # Right-align the paragraph and remove leading whitespace
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.runs[0].text = stripped
                for run in para.runs[1:]:
                    run.text = ''


def _cleanup_empty_placeholders(doc):
    """Final cleanup pass — handle empty values like a seasoned probate lawyer.

    - Fill blanks with underscores for fields that need to be filled manually
    - Remove lines that are entirely empty after replacement
    - Clean dangling colons/punctuation where values are empty
    - Handle optional fields gracefully (fax, reference, time of death, etc.)
    """
    def _clean_text(text):
        """Apply cleanup rules to a text string."""
        if not text or not text.strip():
            return text
        changed = text
        # NRIC/IC patterns with empty value: "No. K/P: )" -> fill blank
        changed = re.sub(r'(No\.\s*K/?P:?\s*)\)', r'\1__________)', changed)
        changed = re.sub(r'(No\.\s*K/?P:?\s*)([,\s]*(?:yang|adalah|$))', r'\1__________\2', changed)
        # Phone/fax with empty value
        changed = re.sub(r'(Tel\.?:?\s*)([,\s]*Fax|$)', r'\1__________\2', changed)
        changed = re.sub(r'(Fax\.?:?\s*)$', r'\1__________', changed)
        changed = re.sub(r'(Fax\.?:?\s*)(\n)', r'\1__________\2', changed)
        # Vehicle details with empty values
        changed = re.sub(r'(No\. Pendaftaran:?\s*)(\n|$)', r'\1__________\2', changed)
        changed = re.sub(r'(No\. Enjin:?\s*)(\n|$)', r'\1__________\2', changed)
        changed = re.sub(r'(No\. Casis:?\s*)(\n|$)', r'\1__________\2', changed)
        # Kenderaan jenama with empty desc
        changed = re.sub(r'(Kenderaan jenama)\s*\n', r'\1 __________\n', changed)
        # Clean empty "Ruj. Kami:" or "Ruj:" (firm reference)
        changed = re.sub(r'(Ruj\.?\s*(?:Kami)?:?\s*)(\n|$)', r'\1__________\2', changed)
        # Account number patterns
        changed = re.sub(r'(No:?\s*)(\n|$)', r'\1__________\2', changed)
        # Remove lines that are just whitespace between non-empty content
        # (don't remove intentional spacing)
        return changed

    # Process paragraphs
    for para in doc.paragraphs:
        full = ''.join(r.text for r in para.runs)
        changed = _clean_text(full)
        if changed != full and para.runs:
            para.runs[0].text = changed
            for r in para.runs[1:]:
                r.text = ''

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = ''.join(r.text for r in para.runs)
                    changed = _clean_text(full)
                    if changed != full and para.runs:
                        para.runs[0].text = changed
                        for r in para.runs[1:]:
                            r.text = ''

    # Process headers/footers
    for section in doc.sections:
        for part in (section.header, section.footer):
            if not part:
                continue
            for para in part.paragraphs:
                full = ''.join(r.text for r in para.runs)
                changed = _clean_text(full)
                if changed != full and para.runs:
                    para.runs[0].text = changed
                    for r in para.runs[1:]:
                        r.text = ''


def _clean_doc06_empty_rows(doc, probate_app):
    """After placeholder replacement, clean up doc06 tables for empty asset categories.

    - Remove rows where all placeholders resolved to empty (no data for that category)
    - If an entire category is empty, show 'Tiada' (None)
    """
    all_assets = json.loads(probate_app.assets_data or '[]')
    has_property = any(a.get('asset_type') == 'property' for a in all_assets)
    has_vehicle = any(a.get('asset_type') == 'vehicle' for a in all_assets)
    has_bank = any(a.get('asset_type') == 'bank' for a in all_assets)
    has_other = any(a.get('asset_type') == 'other' for a in all_assets)
    has_liability = any(a.get('asset_type') == 'liability' for a in all_assets)

    for table in doc.tables:
        rows_to_remove = []
        for ri, row in enumerate(table.rows):
            cell_text = ' '.join(cell.text.strip() for cell in row.cells)
            # Skip header rows and total rows
            if ri == 0 or 'JUMLAH' in cell_text:
                continue
            # Check if this is a data row that's now empty after replacement
            # (placeholder was replaced with empty string)
            content_cell = row.cells[1] if len(row.cells) > 1 else row.cells[0]
            text = content_cell.text.strip()
            # Row is effectively empty if it only has template text with no actual data
            if not text or text in ('Akan ditaksir', '0'):
                # Check if description column (cell 1) is empty
                if not text:
                    rows_to_remove.append(ri)
            # Check for rows with pattern like "Kenderaan jenama \nNo. Pendaftaran: \n..."
            # where all values are empty
            if 'Kenderaan jenama' in text and not has_vehicle:
                rows_to_remove.append(ri)
            if 'Akaun Simpanan' in text and not has_bank:
                rows_to_remove.append(ri)

        # Remove empty rows from bottom up
        for ri in sorted(set(rows_to_remove), reverse=True):
            try:
                row_elem = table.rows[ri]._tr
                row_elem.getparent().remove(row_elem)
            except Exception:
                pass

    # If Table 1 (immovable property) has no property rows, add "Tiada"
    if len(doc.tables) > 1 and not has_property:
        table1 = doc.tables[1]
        # Clear data rows (keep header R0 and total row)
        for ri in range(len(table1.rows) - 2, 0, -1):
            content = table1.rows[ri].cells[1].text.strip() if len(table1.rows[ri].cells) > 1 else ''
            if not content or 'JUMLAH' not in content:
                for cell in table1.rows[ri].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                if ri == 1 and len(table1.rows[ri].cells) > 1:
                    table1.rows[ri].cells[0].paragraphs[0].add_run('1.')
                    table1.rows[ri].cells[1].paragraphs[0].add_run('Tiada')
                    table1.rows[ri].cells[2].paragraphs[0].add_run('0')

    # If Table 2 (movable property) has no movable assets at all, add "Tiada"
    if len(doc.tables) > 2 and not (has_vehicle or has_bank or has_other):
        table2 = doc.tables[2]
        for ri in range(len(table2.rows) - 2, 0, -1):
            content = table2.rows[ri].cells[1].text.strip() if len(table2.rows[ri].cells) > 1 else ''
            if 'JUMLAH' not in content:
                for cell in table2.rows[ri].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                if ri == 1 and len(table2.rows[ri].cells) > 1:
                    table2.rows[ri].cells[0].paragraphs[0].add_run('1.')
                    table2.rows[ri].cells[1].paragraphs[0].add_run('Tiada')
                    table2.rows[ri].cells[2].paragraphs[0].add_run('0')


def generate_probate_forms(probate_app, will_record, selected_codes, templates_map, output_dir):
    """Generate all selected probate forms.

    Args:
        probate_app: ProbateApplication model instance
        will_record: Will model instance
        selected_codes: List of form_code strings to generate
        templates_map: Dict mapping form_code -> template file path
        output_dir: Directory to save generated files

    Returns:
        List of dicts: {form_code, form_name, file_path}
    """
    replacements = build_replacements(probate_app, will_record)
    results = []

    for code in selected_codes:
        template_path = templates_map.get(code)
        if not template_path or not os.path.exists(template_path):
            continue

        output_name = f'{code}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        output_path = os.path.join(output_dir, output_name)

        try:
            fill_template(template_path, replacements, output_path)
            # Post-processing: clean up empty values and fix alignment
            try:
                doc = Document(output_path)
                _cleanup_empty_placeholders(doc)
                _fix_signing_alignment(doc)
                if code == 'doc06':
                    _clean_doc06_empty_rows(doc, probate_app)
                doc.save(output_path)
            except Exception as e2:
                print(f"Warning: post-processing {code} failed: {e2}")
            results.append({
                'form_code': code,
                'file_path': output_path,
            })
        except Exception as e:
            print(f"Error generating {code}: {e}")
            continue

    return results


def convert_to_pdf(docx_path):
    """Convert a .docx file to PDF using LibreOffice headless.

    Returns the PDF file path, or None on failure.
    """
    import subprocess
    output_dir = os.path.dirname(docx_path)
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            capture_output=True, timeout=60, check=True
        )
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception as e:
        print(f'PDF conversion failed for {docx_path}: {e}')
    return None


def create_zip(form_files, zip_path, as_pdf=False):
    """Create a ZIP file containing all generated forms.

    Args:
        form_files: List of dicts with 'form_code' and 'file_path'
        zip_path: Path for the output ZIP file
        as_pdf: If True, convert .docx files to PDF before zipping

    Returns:
        zip_path on success
    """
    os.makedirs(os.path.dirname(zip_path), exist_ok=True)
    now = datetime.now().timetuple()[:6]  # (year, month, day, hour, minute, second)
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in form_files:
            fpath = f['file_path']
            if not os.path.exists(fpath):
                continue
            # Use proper form name for archive filename
            form_name = f.get('form_name') or f.get('form_code', 'form')
            safe_name = form_name.replace(' ', '_').replace('/', '_').replace('&', 'and')
            if as_pdf and fpath.lower().endswith(('.docx', '.doc')):
                pdf_path = convert_to_pdf(fpath)
                if pdf_path and os.path.exists(pdf_path):
                    data = open(pdf_path, 'rb').read()
                    info = zipfile.ZipInfo(f'{safe_name}.pdf', date_time=now)
                    zf.writestr(info, data)
                    continue
            # Fallback: add original file
            ext = os.path.splitext(fpath)[1] or '.docx'
            data = open(fpath, 'rb').read()
            info = zipfile.ZipInfo(f'{safe_name}{ext}', date_time=now)
            zf.writestr(info, data)
    return zip_path
